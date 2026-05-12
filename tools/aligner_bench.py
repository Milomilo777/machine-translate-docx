"""
Aligner benchmark — run the FA aligner over a set of real Persian-translated
docx files and report per-file split-quality stats so we can A/B different
heuristic passes without re-translating.

Usage (from project root):
    PYTHONPATH=src python tools/aligner_bench.py FILE1.docx FILE2.docx ...

The aligner runs in-place on temp copies (originals are never touched).
Output: a markdown table on stdout + a summary line.
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
import tempfile
from pathlib import Path

from machine_translate_docx.openai_tools.persian_double_lines import (
    FASubtitleAligner,
    MAX_CHARS,
    _display_len,
)


def _row_metrics(rows):
    """Per-doc summary: how many rows fit, how many are tight, how many over."""
    over   = sum(1 for r in rows if _display_len(r) > MAX_CHARS)
    tight  = sum(1 for r in rows if 0 < _display_len(r) <= 8)
    long_lines = sum(1 for r in rows if _display_len(r) > MAX_CHARS - 4)
    return over, tight, long_lines


def bench_one(path: Path, llm_threshold: int = 0) -> dict:
    tmp = Path(tempfile.mkdtemp()) / path.name
    shutil.copy(path, tmp)
    aligner = FASubtitleAligner(llm_threshold=llm_threshold)
    stats = aligner.align(str(tmp), str(tmp))

    # Re-open the output and read its FA column for extra metrics
    from docx import Document
    d = Document(str(tmp))
    fa_rows: list[str] = []
    if d.tables:
        for row in d.tables[0].rows:
            cells = row.cells
            if len(cells) >= 3:
                fa_rows.append(cells[2].text.strip())
    over, tight, long_lines = _row_metrics(fa_rows)

    return {
        "file":           path.name,
        "rows":           stats.get("total_rows", 0),
        "groups":         stats.get("groups", 0),
        "doubles":        stats.get("doubles", 0),
        "triples":        stats.get("triples", 0),
        "over_limit":     stats.get("over_limit", 0),
        "double_ratio":   round(stats.get("doubles", 0) / max(stats.get("groups", 1), 1), 2),
        "long_lines":     long_lines,
        "tight_short":    tight,
        "llm_corrected":  stats.get("llm_corrected", 0),
        "tokens_used":    stats.get("tokens_used", 0),
        "elapsed_s":      stats.get("elapsed_seconds", 0),
    }


def main():
    p = argparse.ArgumentParser()
    p.add_argument("files", nargs="+", help="Persian-translated docx files")
    p.add_argument("--llm-threshold", type=int, default=0,
                   help="0..100; 0 = pure mechanical (default), 40 = LLM rescue for hard groups only")
    args = p.parse_args()

    results = []
    for f in args.files:
        path = Path(f)
        if not path.exists():
            print(f"SKIP (not found): {f}", file=sys.stderr)
            continue
        try:
            r = bench_one(path, llm_threshold=args.llm_threshold)
            results.append(r)
        except Exception as exc:
            print(f"FAIL: {f}: {exc!r}", file=sys.stderr)
            continue

    if not results:
        sys.exit(1)

    cols = ["file", "rows", "groups", "doubles", "triples", "over_limit",
            "double_ratio", "long_lines", "tight_short",
            "llm_corrected", "tokens_used", "elapsed_s"]
    headers = ["File", "Rows", "Grp", "Dbl", "Tri", "Over", "Dbl%",
               "Long", "Tight", "LLM", "Tok", "Sec"]
    widths = [max(len(h), max(len(str(r[c])) for r in results)) for h, c in zip(headers, cols)]

    def _fmt(row, fields):
        return " | ".join(str(row[c]).ljust(widths[i]) for i, c in enumerate(fields))

    print(" | ".join(h.ljust(widths[i]) for i, h in enumerate(headers)))
    print("-+-".join("-" * w for w in widths))
    for r in results:
        # shorten file name to a fixed prefix
        short = r["file"]
        if len(short) > 60:
            short = short[:57] + "..."
        row = {**r, "file": short}
        widths_use = list(widths)
        widths_use[0] = max(len(headers[0]), len(short), widths[0])
        print(_fmt(row, cols))

    # Aggregate line
    n = len(results)
    agg = {
        "rows":         sum(r["rows"]         for r in results),
        "groups":       sum(r["groups"]       for r in results),
        "doubles":      sum(r["doubles"]      for r in results),
        "triples":      sum(r["triples"]      for r in results),
        "over_limit":   sum(r["over_limit"]   for r in results),
        "long_lines":   sum(r["long_lines"]   for r in results),
        "tight_short":  sum(r["tight_short"]  for r in results),
        "llm_corrected": sum(r.get("llm_corrected", 0) for r in results),
        "tokens_used":  sum(r.get("tokens_used", 0) for r in results),
        "elapsed_s":    round(sum(r["elapsed_s"] for r in results), 2),
    }
    agg["double_ratio"] = round(agg["doubles"] / max(agg["groups"], 1), 2)
    print()
    print(f"AGGREGATE ({n} files, threshold={args.llm_threshold}): "
          f"{agg['rows']} rows, {agg['groups']} groups, "
          f"{agg['doubles']} doubles ({int(agg['double_ratio']*100)}%), "
          f"{agg['triples']} triples, "
          f"over_limit={agg['over_limit']}, "
          f"long(>44)={agg['long_lines']}, "
          f"tight(<=8)={agg['tight_short']}, "
          f"llm_fired={agg['llm_corrected']}, "
          f"tokens={agg['tokens_used']}, "
          f"total {agg['elapsed_s']}s")


if __name__ == "__main__":
    main()
