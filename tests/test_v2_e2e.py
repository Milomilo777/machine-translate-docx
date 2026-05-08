"""End-to-end smoke test for the v2 frontend.

Boots local_launcher.py in mock-backend mode on a fixed test port and drives
the page with Playwright. Verifies the upload → progress → download flow
without hitting OpenAI / Selenium.

Marked `@pytest.mark.live` so the default `pytest -q` run skips it. Run
manually with:

    pytest -m live -v

Prerequisites:
    pip install playwright
    playwright install chromium

The first two tests (page load + i18n endpoint) are pure HTTP checks and
should pass deterministically. The Playwright-driven tests (upload flow +
locale toggle) wait for `Alpine.$data(document.body).i18n.en.title` to
populate before interacting — this is the only reliable signal that the
async `init()` in `app.js` (which awaits `fetch('/v2/i18n.json')`) has
fully resolved.
"""
from __future__ import annotations

import socket
import subprocess
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
LAUNCHER = ROOT / "local_launcher.py"

# Use a non-default port so a stray dev launcher on 3000 doesn't collide.
TEST_PORT = 3099
TEST_URL  = f"http://127.0.0.1:{TEST_PORT}"


# ── helpers ──────────────────────────────────────────────────────────────────

def _port_is_listening(port: int, host: str = "127.0.0.1") -> bool:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.settimeout(0.2)
        try:
            s.connect((host, port))
            return True
        except OSError:
            return False


def _wait_for_url(url: str, timeout: float = 30.0) -> None:
    """Poll an HTTP URL until it returns 200 or `timeout` seconds elapse."""
    deadline = time.time() + timeout
    last_err: Exception | None = None
    while time.time() < deadline:
        try:
            with urllib.request.urlopen(url, timeout=1.0) as resp:
                if resp.status == 200:
                    return
        except (urllib.error.URLError, ConnectionError, TimeoutError) as exc:
            last_err = exc
        time.sleep(0.3)
    raise RuntimeError(f"URL {url} did not become reachable within {timeout}s "
                       f"(last error: {last_err})")


def _make_tiny_docx(path: Path) -> Path:
    """Create a minimal valid .docx file using python-docx."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Hello world. This is a smoke-test document.")
    doc.add_paragraph("It contains exactly two paragraphs of English text.")
    doc.save(str(path))
    return path


def _wait_for_alpine_ready(page, timeout_ms: int = 15_000) -> None:
    """Wait until Alpine has loaded i18n and the factory's init() resolved."""
    page.wait_for_function(
        """
        () => {
          if (!window.Alpine || !window.Alpine.$data) return false;
          const data = window.Alpine.$data(document.body);
          return !!(data && data.i18n && data.i18n.en && data.i18n.en.title);
        }
        """,
        timeout=timeout_ms,
    )


# ── fixtures ─────────────────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def launcher_proc():
    """Spawn local_launcher.py in mock mode and tear it down on test exit."""
    if _port_is_listening(TEST_PORT):
        pytest.skip(f"Port {TEST_PORT} is already in use — cannot start launcher.")

    proc = subprocess.Popen(
        [sys.executable, str(LAUNCHER),
         "--backend", "mock", "--no-browser",
         "--port", str(TEST_PORT)],
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    try:
        _wait_for_url(f"{TEST_URL}/v2/", timeout=30.0)
        yield proc
    finally:
        proc.terminate()
        try:
            proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            proc.kill()


@pytest.fixture(scope="module")
def docx_fixture(tmp_path_factory):
    path = tmp_path_factory.mktemp("v2-e2e") / "smoke.docx"
    return _make_tiny_docx(path)


# ── tests ────────────────────────────────────────────────────────────────────

@pytest.mark.live
def test_v2_page_loads(launcher_proc):
    """Sanity: the launcher serves /v2/ and the page contains expected markers."""
    with urllib.request.urlopen(f"{TEST_URL}/v2/", timeout=2.0) as resp:
        assert resp.status == 200
        body = resp.read().decode("utf-8", errors="replace")
    assert "Translate a DOCX"     in body, "hero h1 fallback text missing"
    assert 'x-data="docTranslator()"' in body, "Alpine bootstrap hook missing"
    assert "/v2/tailwind.css"     in body, "compiled Tailwind not linked"
    assert 'class="sr-only"' in body, "newsletter sr-only label missing"


@pytest.mark.live
def test_v2_i18n_endpoint(launcher_proc):
    """The launcher must serve /v2/i18n.json with both locales present."""
    import json
    with urllib.request.urlopen(f"{TEST_URL}/v2/i18n.json", timeout=2.0) as resp:
        assert resp.status == 200
        data = json.loads(resp.read().decode("utf-8"))
    assert "en" in data and "fa" in data
    assert set(data["en"].keys()) == set(data["fa"].keys()), "locale key parity broken"
    assert "title" in data["en"]
    assert "نشانی ایمیل" == data["fa"]["newsletter_email_label"]


@pytest.mark.live
def test_v2_upload_to_download_flow(launcher_proc, docx_fixture):
    """Drive the v2 page with Playwright: upload, wait, see download link."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            context = browser.new_context()
            page = context.new_page()
            page.goto(f"{TEST_URL}/v2/", wait_until="networkidle")
            _wait_for_alpine_ready(page)

            page.set_input_files('input[type="file"]', str(docx_fixture))
            page.wait_for_selector("text=smoke.docx", timeout=5_000)

            translate_btn = page.locator("button", has_text="Translate").first
            translate_btn.click()

            download_link = page.locator('a:has-text("Download")').first
            download_link.wait_for(state="visible", timeout=20_000)

            href = download_link.get_attribute("href")
            assert href and href.startswith("/download/"), f"unexpected href: {href}"

            with urllib.request.urlopen(f"{TEST_URL}{href}", timeout=5.0) as resp:
                body = resp.read()
            assert resp.status == 200
            assert body.startswith(b"PK\x03\x04"), "download is not a valid docx zip"
        finally:
            browser.close()


@pytest.mark.live
def test_v2_locale_toggle(launcher_proc):
    """Clicking the locale toggle flips <html lang> and re-renders strings."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            page = browser.new_context().new_page()
            page.goto(f"{TEST_URL}/v2/", wait_until="networkidle")
            _wait_for_alpine_ready(page)

            # Default locale `en` → toggle button shows the OTHER label "فارسی".
            toggle = page.locator("button", has_text="فارسی").first
            toggle.click()
            page.wait_for_function(
                "document.documentElement.getAttribute('lang') === 'fa'",
                timeout=3_000,
            )
            assert page.evaluate("document.documentElement.getAttribute('dir')") == "rtl"

            # Toggle back — button now reads "English".
            toggle = page.locator("button", has_text="English").first
            toggle.click()
            page.wait_for_function(
                "document.documentElement.getAttribute('lang') === 'en'",
                timeout=3_000,
            )
        finally:
            browser.close()
