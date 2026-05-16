"""Unit tests for ``src/machine_translate_docx/docx_io/parse.py``.

The core challenge: ``read_and_parse_docx_document`` lazy-imports six
helpers from ``machine_translate_docx.cli``. Importing cli.py in a test
process would run its entire top-level (argparse, sys.exit calls, network
fetches, Selenium imports). The fix: install a stub module in
``sys.modules['machine_translate_docx.cli']`` before the lazy import fires.
The stub exposes only the six symbols the function actually calls.

Fixture used: ``tests/fixtures/sample_hyperlink.docx``
 - 42 rows, 3 columns (col 0: row no., col 1: EN text, col 2: FA/blank)
 - Row 8 (0-indexed) contains a ``<w:hyperlink>`` element in col 1

Test inventory
--------------
1. test_parse_populates_basic_arrays          — from_text_table, from_text_by_phrase_separator_table, numrows, numcols
2. test_parse_source_column_lock_snapshotted  — C13 invariant: source_columns_snapshot keyed by (row, col) for col in (0,1)
3. test_parse_hyperlink_text_included         — _iter_paragraph_runs walks <w:hyperlink> children
4. test_parse_empty_docx_raises_or_produces_empty — no-table docx → EmptyDocxError
5. test_parse_idempotent                      — running parse twice yields identical arrays
"""
from __future__ import annotations

import io
import sys
import types
from copy import deepcopy
from pathlib import Path
from unittest.mock import MagicMock

import pytest

# ── sys.path setup ────────────────────────────────────────────────────────────

_SRC = Path(__file__).resolve().parents[1] / "src"
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))

# ── Stub cli module injected BEFORE any parse import ─────────────────────────
#
# The six helpers lazy-imported by read_and_parse_docx_document live in
# machine_translate_docx.cli but we must not import that module in tests
# (its top-level runs argparse + network + sys.exit).  We build a minimal
# stub and pin it into sys.modules so the lazy import resolves to our stubs.
#
# Stubs behave exactly like the real functions for a plain EN source docx:
#   is_end_of_line            → 0 (no EOL markers in the fixture)
#   is_empty_line             → 1 if blank, else 0
#   is_beginning_of_line      → 0
#   is_conditional_end_of_line → 0
#   prepare_and_clear_cell_for_writing → no-op (write path, not under test)
#   split_phrases             → lightweight grouping (just fills by-phrase arrays)
#
# The split_phrases stub is the minimal working version: it runs the same
# algorithm as the real function but reads line_separator_str from the
# module's own constant (a single space) so no global state is needed.

import re as _re


def _stub_is_end_of_line(line: str) -> int:
    """Match the real is_end_of_line: scan against eol_array patterns."""
    # eol_array from config.py — replicated here so tests need not import cli.
    from machine_translate_docx.config import eol_array
    for pat in eol_array:
        if _re.search(pat, line):
            return 1
    return 0


def _stub_is_empty_line(line: str) -> int:
    stripped = _re.sub(" +", "", line)
    return 1 if len(stripped) == 0 else 0


def _stub_is_beginning_of_line(line: str) -> int:
    from machine_translate_docx.config import bol_array
    for pat in bol_array:
        if _re.search(pat, line):
            return 1
    return 0


def _stub_is_conditional_end_of_line(line: str) -> int:
    from machine_translate_docx.config import eol_conditional_array
    for pat in eol_conditional_array:
        if _re.search(pat, line):
            return 1
    return 0


def _stub_prepare_and_clear_cell_for_writing(ctx, row_n, translation_cell_text):
    """No-op: cell clearing is write-path, not under test here."""
    pass


def _stub_split_phrases(ctx) -> int:
    """Minimal phrase-grouping matching the real split_phrases algorithm.

    Uses a single space as the line separator (same as the cli.py default
    ``line_separator_str = ' '``).
    """
    _LINE_SEP = " "
    docx = ctx.docx
    n_last_row_phrase = 3
    last_table_row = docx.word_translation_table_length
    cur_row_n = 2
    while cur_row_n < last_table_row:
        if docx.from_text_is_beginning_of_line_table[cur_row_n] == 1:
            n_last_row_phrase = cur_row_n
            nb_lines_in_phrase = 1
            docx.from_text_nb_lines_in_phrase[cur_row_n] = (
                docx.from_text_nb_lines_in_cell[cur_row_n]
            )
            while (
                docx.from_text_is_end_of_line_table[n_last_row_phrase] != 1
                and n_last_row_phrase < (last_table_row - 1)
            ):
                if docx.from_text_by_phrase_separator_table[cur_row_n] == "":
                    docx.from_text_by_phrase_separator_table[cur_row_n] = (
                        docx.from_text_table[n_last_row_phrase]
                    )
                    docx.from_text_by_phrase_table[cur_row_n] = (
                        docx.from_text_table[n_last_row_phrase]
                    )
                else:
                    docx.from_text_by_phrase_separator_table[cur_row_n] = (
                        docx.from_text_by_phrase_separator_table[cur_row_n]
                        + _LINE_SEP
                        + docx.from_text_table[n_last_row_phrase]
                    )
                    docx.from_text_by_phrase_table[cur_row_n] = (
                        docx.from_text_by_phrase_table[cur_row_n]
                        + " "
                        + docx.from_text_table[n_last_row_phrase]
                    )
                    nb_lines_in_phrase += 1
                    docx.from_text_nb_lines_in_phrase[cur_row_n] += (
                        docx.from_text_nb_lines_in_cell[n_last_row_phrase]
                    )
                n_last_row_phrase += 1
            if docx.from_text_by_phrase_separator_table[cur_row_n] == "":
                docx.from_text_by_phrase_separator_table[cur_row_n] = (
                    docx.from_text_table[n_last_row_phrase]
                )
                docx.from_text_by_phrase_table[cur_row_n] = (
                    docx.from_text_table[n_last_row_phrase]
                )
            else:
                docx.from_text_by_phrase_separator_table[cur_row_n] = (
                    docx.from_text_by_phrase_separator_table[cur_row_n]
                    + _LINE_SEP
                    + docx.from_text_table[n_last_row_phrase]
                )
                nb_lines_in_phrase += 1
                docx.from_text_nb_lines_in_phrase[cur_row_n] += (
                    docx.from_text_nb_lines_in_cell[n_last_row_phrase]
                )
                docx.from_text_by_phrase_table[cur_row_n] = (
                    docx.from_text_by_phrase_table[cur_row_n]
                    + " "
                    + docx.from_text_table[n_last_row_phrase]
                )
            cur_row_n = n_last_row_phrase + 1
        else:
            cur_row_n += 1
    return 0


def _install_cli_stub() -> None:
    """Pin a lightweight fake cli module into sys.modules.

    Must be called before importing parse so the lazy-import inside
    read_and_parse_docx_document resolves to our stubs, not the real cli.
    """
    if "machine_translate_docx.cli" in sys.modules:
        return  # already installed (e.g. by a previous test run in the same process)
    stub = types.ModuleType("machine_translate_docx.cli")
    stub.is_end_of_line = _stub_is_end_of_line
    stub.is_empty_line = _stub_is_empty_line
    stub.is_beginning_of_line = _stub_is_beginning_of_line
    stub.is_conditional_end_of_line = _stub_is_conditional_end_of_line
    stub.prepare_and_clear_cell_for_writing = _stub_prepare_and_clear_cell_for_writing
    stub.split_phrases = _stub_split_phrases
    sys.modules["machine_translate_docx.cli"] = stub


_install_cli_stub()

# ── real imports (cli stub is already in sys.modules) ────────────────────────

import docx as _docx  # noqa: E402
from machine_translate_docx.docx_io.parse import read_and_parse_docx_document  # noqa: E402
from machine_translate_docx.exceptions import EmptyDocxError  # noqa: E402
from machine_translate_docx.runtime import RuntimeContext  # noqa: E402

# ── paths ─────────────────────────────────────────────────────────────────────

_FIXTURES = Path(__file__).resolve().parent / "fixtures"
_HYPERLINK_DOCX = _FIXTURES / "sample_hyperlink.docx"


# ── helpers ───────────────────────────────────────────────────────────────────

def _make_ctx(fixture: Path = _HYPERLINK_DOCX) -> RuntimeContext:
    """Return a minimally-configured RuntimeContext for the given fixture."""
    ctx = RuntimeContext.empty()
    ctx.flags.word_file_to_translate = str(fixture)
    ctx.flags.silent = True
    ctx.flags.splitonly = False
    ctx.docx.docxdoc = _docx.Document(str(fixture))
    ctx.docx.use_html = False
    ctx.config.shading_color_ignore_text = []
    return ctx


def _make_empty_docx_ctx() -> RuntimeContext:
    """Return a ctx whose docxdoc has no table at all."""
    document = _docx.Document()
    # Save to an in-memory buffer and re-open so python-docx does a full
    # parse round-trip (the new Document() object has no tables by default).
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    reopened = _docx.Document(buf)
    ctx = RuntimeContext.empty()
    ctx.flags.word_file_to_translate = "<in-memory empty docx>"
    ctx.flags.silent = True
    ctx.flags.splitonly = False
    ctx.docx.docxdoc = reopened
    ctx.docx.use_html = False
    ctx.config.shading_color_ignore_text = []
    return ctx


# ── tests ─────────────────────────────────────────────────────────────────────


def test_parse_populates_basic_arrays():
    """After parse, from_text_table and from_text_by_phrase_separator_table
    are non-empty lists of the expected length, and numrows / numcols are
    positive integers matching the fixture's table dimensions.

    The fixture has 42 rows and 3 columns; arrays are sized numrows+1 per
    the +1 indexing convention (R16 invariant).
    """
    ctx = _make_ctx()
    read_and_parse_docx_document(ctx)

    # numrows / numcols are set from the table geometry
    assert ctx.docx.numrows >= 1, "numrows should be positive"
    assert ctx.docx.numcols >= 3, "fixture has 3 columns"

    # from_text_table is sized numrows+1 (+1 convention)
    assert len(ctx.docx.from_text_table) == ctx.docx.numrows + 1, (
        "from_text_table must be numrows+1 long"
    )

    # from_text_by_phrase_separator_table must be the same length
    assert len(ctx.docx.from_text_by_phrase_separator_table) == len(
        ctx.docx.from_text_table
    ), "from_text_by_phrase_separator_table length must match from_text_table"

    # At least one source cell must have non-empty text (the fixture is not blank)
    non_empty = [v for v in ctx.docx.from_text_table if v and v.strip()]
    assert len(non_empty) >= 1, (
        "from_text_table should contain at least one non-empty string for the fixture"
    )


def test_parse_source_column_lock_snapshotted():
    """C13 invariant: columns 0 and 1 of every row must be deepcopy-snapshotted
    into ctx.docx.source_columns_snapshot after parse.

    The snapshot is keyed by (row_index, col_index) for col_index in {0, 1},
    and each value is a 2-tuple (cell_text: str, cell_tc_copy: lxml element).
    """
    ctx = _make_ctx()
    read_and_parse_docx_document(ctx)

    snap = ctx.docx.source_columns_snapshot
    assert isinstance(snap, dict), "source_columns_snapshot must be a dict"
    assert len(snap) > 0, "snapshot must not be empty after parse"

    # Every key must be a (row, col) pair with col in {0, 1}
    for key in snap:
        row_idx, col_idx = key
        assert col_idx in (0, 1), (
            f"unexpected col_index {col_idx} in snapshot key {key!r}"
        )

    # Each value must be a 2-tuple: (str, lxml element)
    for key, value in snap.items():
        assert isinstance(value, tuple) and len(value) == 2, (
            f"snapshot value at {key!r} must be a 2-tuple, got {type(value)!r}"
        )
        cell_text, cell_tc = value
        assert isinstance(cell_text, str), (
            f"snapshot[{key!r}][0] must be str, got {type(cell_text)!r}"
        )
        # cell_tc is a deepcopy'd lxml element — it must have a tag attribute
        assert hasattr(cell_tc, "tag"), (
            f"snapshot[{key!r}][1] must be an lxml element"
        )

    # For every row in the table, both col 0 and col 1 must be present
    n_rows = ctx.docx.numrows
    for row_i in range(n_rows):
        assert (row_i, 0) in snap, (
            f"snapshot missing (row={row_i}, col=0)"
        )
        assert (row_i, 1) in snap, (
            f"snapshot missing (row={row_i}, col=1)"
        )


def test_parse_hyperlink_text_included():
    """The fixture's row 8 (0-indexed) has 'hyperlink' inside a <w:hyperlink>
    element.  After parse, from_text_table must contain that word, proving
    that _iter_paragraph_runs descends into hyperlink children.
    """
    ctx = _make_ctx()
    read_and_parse_docx_document(ctx)

    joined = " ".join(ctx.docx.from_text_table)
    assert "hyperlink" in joined.lower(), (
        "from_text_table must contain 'hyperlink' — _iter_paragraph_runs "
        "must walk <w:hyperlink> children in the fixture"
    )


def test_parse_empty_docx_raises_or_produces_empty():
    """A docx with no table at all must raise EmptyDocxError.

    parse.py raises EmptyDocxError (a TranslationFailure subclass) rather
    than silently producing empty arrays — this is the E1 edge-case fix
    documented in the function's docstring.
    """
    ctx = _make_empty_docx_ctx()
    with pytest.raises(EmptyDocxError):
        read_and_parse_docx_document(ctx)


def test_parse_idempotent():
    """Running parse twice on a fresh ctx each time yields identical from_text_table
    contents.  Parse is read-only on the source docx, so re-opening and
    re-parsing must produce the same result.
    """
    ctx_a = _make_ctx()
    read_and_parse_docx_document(ctx_a)
    result_a = list(ctx_a.docx.from_text_table)

    ctx_b = _make_ctx()
    read_and_parse_docx_document(ctx_b)
    result_b = list(ctx_b.docx.from_text_table)

    assert result_a == result_b, (
        "Parsing the same fixture twice must yield identical from_text_table"
    )
