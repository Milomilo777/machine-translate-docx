"""Tests for machine_translate_docx.docx_io.metadata.

Covers both public helpers:
- write_destination_language_in_docx_cell
- set_docx_properties_comment_for_history

All tests use in-memory docx.Document() instances — no disk I/O.
"""
from __future__ import annotations

import datetime

import pytest
from docx import Document

import machine_translate_docx.docx_io.metadata as mod
from machine_translate_docx.docx_io.metadata import (
    set_docx_properties_comment_for_history,
    write_destination_language_in_docx_cell,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _doc_3x3(cell_1_2_text: str = "ORIGINAL") -> Document:
    """Return an in-memory Document with a 3-row × 3-col table."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.cell(1, 2).text = cell_1_2_text
    return doc


def _doc_1x1() -> Document:
    """Return an in-memory Document with a 1-row × 1-col table."""
    doc = Document()
    doc.add_table(rows=1, cols=1)
    return doc


# ---------------------------------------------------------------------------
# write_destination_language_in_docx_cell
# ---------------------------------------------------------------------------

def test_splitonly_true_leaves_cell_unchanged():
    """When splitonly=True the function must not touch cell (1, 2)."""
    doc = _doc_3x3("ORIGINAL")
    write_destination_language_in_docx_cell(
        doc,
        splitonly=True,
        dest_lang_name="Persian",
        dest_lang="fa",
    )
    assert doc.tables[0].cell(1, 2).text == "ORIGINAL"


def test_splitonly_false_writes_dest_lang_name():
    """When splitonly=False and dest_lang_name is valid, cell (1, 2) gets it."""
    doc = _doc_3x3("ORIGINAL")
    write_destination_language_in_docx_cell(
        doc,
        splitonly=False,
        dest_lang_name="Persian",
        dest_lang="fa",
    )
    assert doc.tables[0].cell(1, 2).text == "Persian"


def test_splitonly_false_none_name_falls_back_to_dest_lang():
    """When dest_lang_name is None the function falls back to dest_lang."""
    doc = _doc_3x3("ORIGINAL")
    write_destination_language_in_docx_cell(
        doc,
        splitonly=False,
        dest_lang_name=None,
        dest_lang="fa",
    )
    assert doc.tables[0].cell(1, 2).text == "fa"


def test_splitonly_false_both_none_does_not_crash():
    """When both dest_lang_name and dest_lang are None it must not raise."""
    doc = _doc_3x3("ORIGINAL")
    # Should silently swallow both exceptions
    write_destination_language_in_docx_cell(
        doc,
        splitonly=False,
        dest_lang_name=None,
        dest_lang=None,
    )
    # Cell text may be anything; the important thing is no exception was raised.


def test_no_row1_col2_does_not_crash():
    """A docx with only a 1×1 table must not raise."""
    doc = _doc_1x1()
    write_destination_language_in_docx_cell(
        doc,
        splitonly=False,
        dest_lang_name="Persian",
        dest_lang="fa",
    )
    # Must reach here without exception


# ---------------------------------------------------------------------------
# set_docx_properties_comment_for_history
# ---------------------------------------------------------------------------

_FROZEN_INSTANT = datetime.datetime(2026, 5, 16, 10, 30, 45)


class _FrozenDateTime:
    """Stub that makes datetime.now() deterministic.

    The stub replaces the *class* datetime.datetime inside the module.
    Its now() must return a real datetime.datetime instance so that
    strftime() works — we capture the instance before the monkeypatch
    runs (while the real datetime.datetime is still in scope).
    """

    @classmethod
    def now(cls) -> datetime.datetime:
        return _FROZEN_INSTANT


def test_comment_format_exact(monkeypatch):
    """The comment must match the documented format string exactly."""
    monkeypatch.setattr(mod.datetime, "datetime", _FrozenDateTime)

    doc = Document()
    set_docx_properties_comment_for_history(
        doc,
        program_version="1.2.3",
        engine="chatGPT",
    )

    expected = (
        "Document translated by SMTV Robot version 1.2.3 "
        "using chatGPT engine on 16/05/2026 10:30:45."
    )
    assert doc.core_properties.comments == expected


def test_comment_is_stored_on_docxdoc(monkeypatch):
    """The comment must be readable back via core_properties.comments."""
    monkeypatch.setattr(mod.datetime, "datetime", _FrozenDateTime)

    doc = Document()
    set_docx_properties_comment_for_history(
        doc,
        program_version="0.0.1",
        engine="DeepL",
    )

    assert doc.core_properties.comments == (
        "Document translated by SMTV Robot version 0.0.1 "
        "using DeepL engine on 16/05/2026 10:30:45."
    )


def test_comment_overwrites_previous(monkeypatch):
    """A second call must replace the previous comment, not append."""
    monkeypatch.setattr(mod.datetime, "datetime", _FrozenDateTime)

    doc = Document()
    # First stamp
    set_docx_properties_comment_for_history(
        doc,
        program_version="1.0.0",
        engine="Google",
    )
    first = doc.core_properties.comments

    # Second stamp with different args
    set_docx_properties_comment_for_history(
        doc,
        program_version="2.0.0",
        engine="chatGPT",
    )
    second = doc.core_properties.comments

    assert first != second
    assert "2.0.0" in second
    assert "chatGPT" in second
    # Old version must not linger
    assert "1.0.0" not in second
