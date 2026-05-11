"""Unit tests for ``src/docx_io/cells.py``.

Two functions land their first dedicated test coverage in the
2026-05-10 G2 thread-globals pass:

  * ``_paragraph_shading_color`` / ``_run_shading_color`` — the pair of
    private helpers that walk a parsed paragraph or run XML element
    looking for a ``<w:shd w:fill="…"/>`` colour attribute.
  * ``get_cell_data(ctx, cell, row_n)`` — the per-cell reader that
    walks every paragraph + run, applies the colour-ignore list from
    ``ctx.config.shading_color_ignore_text``, and returns
    ``(cell_text, cell_is_gray, cell_is_red)``.

The shading-helper tests use raw XML strings so they have zero
python-docx dependency. The ``get_cell_data`` test uses a tiny
hand-built docx file in memory via ``docx.Document()`` so the run /
paragraph iteration is exercised end-to-end.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest

_SRC = Path(__file__).resolve().parents[1] / "src"
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))

import docx  # noqa: E402  — sys.path tweak above

from machine_translate_docx.docx_io.cells import (  # noqa: E402
    _paragraph_shading_color,
    _run_shading_color,
    get_cell_data,
)
from machine_translate_docx.runtime import RuntimeContext  # noqa: E402


# ── shading helpers ──────────────────────────────────────────────────────────

_W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def test_paragraph_shading_color_returns_fill_when_present():
    xml = (
        f'<w:p {_W_NS}>'
        f'<w:pPr><w:shd w:val="clear" w:color="auto" w:fill="D9D9D9"/></w:pPr>'
        f'<w:r><w:t>hello</w:t></w:r>'
        f'</w:p>'
    )
    assert _paragraph_shading_color(xml) == "D9D9D9"


def test_paragraph_shading_color_none_when_absent():
    xml = f'<w:p {_W_NS}><w:r><w:t>hello</w:t></w:r></w:p>'
    assert _paragraph_shading_color(xml) is None


def test_run_shading_color_returns_fill_when_present():
    xml = (
        f'<w:r {_W_NS}>'
        f'<w:rPr><w:shd w:val="clear" w:color="auto" w:fill="FF00FF"/></w:rPr>'
        f'<w:t>hello</w:t>'
        f'</w:r>'
    )
    assert _run_shading_color(xml) == "FF00FF"


def test_run_shading_color_none_when_absent():
    xml = f'<w:r {_W_NS}><w:t>hello</w:t></w:r>'
    assert _run_shading_color(xml) is None


# ── get_cell_data ────────────────────────────────────────────────────────────

@pytest.fixture
def ctx_with_room_for_one_row():
    ctx = RuntimeContext.empty()
    # get_cell_data writes to ctx.docx.from_text_nb_lines_in_cell[row_n - 1].
    # For row_n=1 we need an array of length >= 1.
    ctx.docx.from_text_nb_lines_in_cell = [0]
    ctx.config.shading_color_ignore_text = []
    return ctx


def _make_single_cell_doc(text_runs: list[str]):
    """Build a one-cell python-docx doc with the given runs, return the cell."""
    document = docx.Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    for text in text_runs:
        paragraph.add_run(text)
    # Round-trip through bytes so the XML is fresh on every read — mimics
    # what python-docx does when the file is written then re-read.
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    reopened = docx.Document(buf)
    return reopened.tables[0].rows[0].cells[0]


def test_get_cell_data_plain_text(ctx_with_room_for_one_row):
    cell = _make_single_cell_doc(["Hello world"])
    text, is_gray, is_red = get_cell_data(ctx_with_room_for_one_row, cell, 1)
    assert text == "Hello world"
    assert is_gray == 0
    assert is_red == 0
    assert ctx_with_room_for_one_row.docx.from_text_nb_lines_in_cell[0] == 1


def test_get_cell_data_strips_pause_and_enter_markers(ctx_with_room_for_one_row):
    cell = _make_single_cell_doc(["Line 1 <pause> Line 2 <enter> tail"])
    text, _, _ = get_cell_data(ctx_with_room_for_one_row, cell, 1)
    assert "<pause>" not in text.lower()
    assert "<enter>" not in text.lower()
    assert "Line 1" in text
    assert "tail" in text
    # Line count = 1 + nb_pause + nb_enter = 1 + 1 + 1 = 3
    assert ctx_with_room_for_one_row.docx.from_text_nb_lines_in_cell[0] == 3


def test_get_cell_data_collapses_whitespace(ctx_with_room_for_one_row):
    cell = _make_single_cell_doc(["Hello   world\n\nfoo"])
    text, _, _ = get_cell_data(ctx_with_room_for_one_row, cell, 1)
    assert text == "Hello world foo"
