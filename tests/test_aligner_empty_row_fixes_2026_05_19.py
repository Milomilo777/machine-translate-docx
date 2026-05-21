"""Regression tests for the 5 aligner empty-row fixes (2026-05-19).

Five distinct bugs were producing empty FA rows in the post-aligner
output where the pre-aligner input had real translated content. The
combined effect was 24/38/69 rows of lost FA across AJAR / VE / Teaser
test files. After the fixes, 137 rows of FA content are restored.

Each test below pins one of the five fixes so a future refactor that
re-introduces the regression fails CI.

See ``notes/aligner-work-2026-05-19/`` for the full diagnostic trail.
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

_SRC = Path(__file__).resolve().parents[1] / "src"
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))

from machine_translate_docx.openai_tools.persian_double_lines import (
    _has_midline_dot,
    _is_bridge,
    _split_for_n_rows,
    MAX_CHARS,
)


# ── Fix 1: trailing close-quote after `.` should NOT count as midline dot ───

def test_midline_dot_strips_trailing_quote_period():
    """``بکنی."`` ends a quoted sentence — the `.` is at the end, not
    midline. Without the fix, no_doubling fires and content is lost."""
    assert _has_midline_dot('بکنی."') is False


def test_midline_dot_strips_trailing_quote_exclamation():
    """``موسیقی!"`` — same pattern with `!`."""
    assert _has_midline_dot('موسیقی!"') is False


def test_midline_dot_strips_trailing_quote_question():
    """``درسته؟"`` — same pattern with Persian `؟`."""
    assert _has_midline_dot('درسته؟"') is False


def test_midline_dot_strips_trailing_curly_quote():
    """``"بکن."`` with curly close-quote ``”`` — also stripped."""
    assert _has_midline_dot('بکن.”') is False


def test_midline_dot_still_catches_real_multi_sentence():
    """``این. آن دیگر.`` — TWO sentences in one cell. Must fire."""
    assert _has_midline_dot('این. آن دیگر.') is True


def test_midline_dot_catches_multi_sentence_with_quote():
    """``گفت "این کار را نکن." بعد رفت.`` — quoted clause then a real
    second sentence ``بعد رفت.``. Must fire."""
    assert _has_midline_dot('گفت "این کار را نکن." بعد رفت.') is True


# ── Fix 2: ALL-CAPS-label regex must NOT be case-folded ─────────────────────

def test_is_bridge_lowercase_phrase_with_colon():
    """``the rehearsal went:`` is a real subtitle line that ends with
    `:`, not a NARRATOR-style label. Must NOT be a bridge — its FA
    content gets distributed into the alignment otherwise."""
    assert _is_bridge('the rehearsal went:', '', False) is False


def test_is_bridge_lowercase_letters_only_with_colon():
    """``come over here:`` — also not a bridge despite letters+spaces+colon."""
    assert _is_bridge('come over here:', '', False) is False


def test_is_bridge_real_all_caps_label():
    """``NARRATOR:`` IS a bridge — must still fire."""
    assert _is_bridge('NARRATOR:', '', False) is True


def test_is_bridge_real_all_caps_multi_word_label():
    """``VEGGIE ELITE HOST:`` IS a bridge."""
    assert _is_bridge('VEGGIE ELITE HOST:', '', False) is True


# ── Fix 3: ellipsis is a pause, not a midline terminator ────────────────────

def test_midline_dot_ignores_mid_ellipsis():
    """``شود"... آن هم فقط`` — ellipsis is a pause, not a hard stop.
    Without the fix, the three `.` chars trigger no_doubling for an
    actually single-clause continuation."""
    assert _has_midline_dot('خراب نمی‌شود"... آن هم فقط') is False


def test_midline_dot_ignores_unicode_ellipsis():
    """``...`` written as the single-char ``…`` must also not fire."""
    assert _has_midline_dot('خراب نمی‌شود"… آن هم فقط') is False


def test_midline_dot_ignores_trailing_ellipsis():
    """``دا-دا-دا-دا...`` — trailing ellipsis is end-of-clause."""
    assert _has_midline_dot('دا-دا-دا-دا...') is False


# ── Fix 4: short text + n_rows >= 3 must NOT short-circuit to 1 chunk ───────

def test_split_for_n_rows_short_text_2_rows_returns_single_chunk():
    """Short text + 2 rows → single chunk is fine (will be doubled)."""
    chunks = _split_for_n_rows("سلام دنیا چطوری", 2)
    assert len(chunks) == 1


def test_split_for_n_rows_short_text_3_rows_attempts_split():
    """Short text + 3 rows → splitter tries to find ≥2 chunks so the
    distributor can avoid blanking the 3rd row via _enforce_no_triple.

    For text that admits a safe break, the splitter SHOULD return
    multiple chunks. Test with text that has a comma + space break."""
    chunks = _split_for_n_rows("دا-دا-دا-دا... اما میکروفن بالا نمی‌آید.", 3)
    # The splitter should find a break around the ellipsis.
    assert len(chunks) >= 2, (
        f"Short text + 3 rows should produce ≥2 chunks (got {len(chunks)}: {chunks})"
    )


def test_split_for_n_rows_unsplittable_short_text_returns_single():
    """If splitter genuinely can't split the text, single chunk is OK.

    Single-word text has no break points; expect 1 chunk back."""
    chunks = _split_for_n_rows("سلام", 5)
    # Either falls back to single chunk or empty — both acceptable;
    # but content must not be silently dropped.
    assert chunks, "Single-word splitter must return at least one chunk"
    assert "".join(chunks).replace(" ", "").startswith("سلام") or "سلام" in "".join(chunks)


# ── Fix 5: no_doubling path with chunks < target_rows uses orig fa_parts ───

def test_align_group_no_doubling_preserves_input_rows():
    """Group has 4 input rows, each with FA fitting MAX_CHARS, and one
    cell has midline punctuation (triggers no_doubling). The splitter
    will merge them into fewer chunks. The fix-5 fallback to original
    fa_parts must kick in so all 4 output rows are filled."""
    from machine_translate_docx.openai_tools.persian_double_lines import (
        FASubtitleAligner,
    )
    aligner = FASubtitleAligner(model="gpt-5.4-mini", llm_threshold=0)
    group = {
        "row_indices": [10, 11, 12, 13],
        "fa_parts": [
            "نامِ... فرض کنیم - طراحی",        # 24 chars, ellipsis triggers no_doubling
            "لباس: \"خارج از آفریقا\"،",         # 22 chars
            "جنی اسمیت. \"جی. اس.\"",            # 18 chars
            "اگر اسم مستعاری داشته باشد.",        # 26 chars
        ],
        "en_parts":  ["x", "x", "x", "x"],
        "en_per_row": ["x", "x", "x", "x"],
    }
    rows = aligner._align_group(group)
    assert len(rows) == 4
    nonempty = sum(1 for r in rows if r.strip())
    # All 4 rows should be non-empty in the preserved-original path
    assert nonempty == 4, (
        f"no_doubling 4-row group with all-fit cells must keep all 4 rows "
        f"filled (got {nonempty}/4 non-empty: {rows!r})"
    )


# ── End-to-end test: aligning a synthetic docx must not lose content ────────

def test_aligner_no_content_loss_on_simple_doubled_input(tmp_path):
    """Build a tiny bilingual docx with 3-row groups where the FA is
    distributed across all rows. After aligning, every input FA cell
    must map to ≥1 non-empty output row."""
    import docx as _d
    doc = _d.Document()
    table = doc.add_table(rows=0, cols=3)

    # Row pattern: (en, fa) tuples. Add a header row, then 3-row phrase.
    test_rows = [
        # speaker header (bridge)
        ("00:00:01 Speaker(f):", ""),
        # 3-row phrase, each cell fits MAX_CHARS
        ("There's a way to do this", "راهی هست که این کار را"),
        ("when the time is right.",  "وقتی زمانش درست باشد"),
        ("Just follow these steps.", "این مراحل را دنبال کنید."),
        # bridge
        ("00:00:05 Speaker(f):", ""),
    ]
    for en, fa in test_rows:
        row = table.add_row()
        row.cells[0].text = ""        # timestamp col
        row.cells[1].text = en
        row.cells[2].text = fa

    in_path = tmp_path / "input.docx"
    doc.save(str(in_path))

    from machine_translate_docx.openai_tools.persian_double_lines import (
        FASubtitleAligner,
    )
    aligner = FASubtitleAligner(model="gpt-5.4-mini", llm_threshold=0)
    out_path = tmp_path / "output.docx"
    aligner.align(str(in_path), str(out_path))

    # Read output back
    out_doc = _d.Document(str(out_path))
    fa_cells = [row.cells[2].text.strip() for table in out_doc.tables for row in table.rows]
    # The three non-empty content rows should remain non-empty
    content_rows = [fa_cells[1], fa_cells[2], fa_cells[3]]
    nonempty = sum(1 for c in content_rows if c)
    assert nonempty == 3, (
        f"Expected all 3 phrase rows to keep FA after aligning; got "
        f"{nonempty}/3 non-empty: {content_rows!r}"
    )


# ── 2026-05-19 (WAU + AW 3153 analysis): five new bridge patterns ────────────
#
# Row-by-row diff against human-edited references uncovered five EN row
# shapes that the previous _BRIDGE_RE did not catch. Each one was
# receiving spillover content from the preceding semantic group. Tests
# below pin each pattern + its negative counterpart so a future refactor
# that drops one of them fails CI.


# Pattern 1: bare speaker shortcut "(f):" / "(m):" / "(h):" standalone
def test_bridge_bare_female_speaker():
    assert _is_bridge('(f):', '', False) is True


def test_bridge_bare_male_speaker():
    assert _is_bridge('(m):', '', False) is True


def test_bridge_bare_speaker_with_whitespace():
    assert _is_bridge('( f ):', '', False) is True


def test_bridge_bare_speaker_uppercase():
    assert _is_bridge('(F):', '', False) is True


def test_bridge_bare_speaker_no_colon():
    """Optional trailing colon — `(f)` alone is still a bridge."""
    assert _is_bridge('(f)', '', False) is True


def test_not_bridge_speaker_inside_sentence():
    """Lee (f): then real dialogue ⇒ not standalone, NOT a bridge."""
    assert _is_bridge('Lee (f): She started talking.', '', False) is False


def test_not_bridge_parens_mention_mid_sentence():
    """`(m) and (f) markers exist.` — parens in the middle of a real
    English sentence; must NOT be a bridge."""
    assert _is_bridge('(m) and (f) markers exist.', '', False) is False


# Pattern 2: timecode + bare speaker shortcut
def test_bridge_timecode_bare_speaker():
    assert _is_bridge('9:39 (f):', '', False) is True


def test_bridge_timecode_with_parens_bare_speaker():
    assert _is_bridge('(9:20) (f):', '', False) is True


def test_not_bridge_timecode_in_sentence():
    """`9:39 marks the time` — timecode followed by real sentence."""
    assert _is_bridge('9:39 marks the time.', '', False) is False


# Pattern 3: title + name + colon standalone
def test_bridge_dr_name_colon():
    assert _is_bridge('Dr. Jena Questen:', '', False) is True


def test_bridge_master_name_colon():
    assert _is_bridge('Master Ching Hai:', '', False) is True


def test_bridge_mr_name_colon():
    assert _is_bridge('Mr. Smith:', '', False) is True


def test_bridge_mrs_name_colon():
    assert _is_bridge('Mrs. Brown:', '', False) is True


def test_bridge_prof_name_colon():
    assert _is_bridge('Prof. Johnson:', '', False) is True


def test_not_bridge_dr_with_dialogue():
    """`Dr. Smith said the patient was sick.` — title + name + words,
    NOT a label, must NOT match."""
    assert _is_bridge('Dr. Smith said the patient was sick.', '', False) is False


def test_not_bridge_dr_continuation():
    """`Dr. Jena Questen explained that` — name + verb, must NOT match."""
    assert _is_bridge('Dr. Jena Questen explained that', '', False) is False


# Pattern 4: section header with parenthetical descriptor
def test_bridge_outro_with_parens():
    assert _is_bridge('OUTRO (IN ENGLISH):', '', False) is True


def test_bridge_intro_with_parens():
    assert _is_bridge('INTRO (BMD):', '', False) is True


def test_bridge_host_with_parens():
    assert _is_bridge('HOST (CONTINUED):', '', False) is True


def test_not_bridge_outro_sentence():
    """`OUTRO is short for outro music.` — sentence about OUTRO, not a
    header. Must NOT be a bridge."""
    assert _is_bridge('OUTRO is short for outro music.', '', False) is False


# Pattern 5: Priority News section pointer
def test_bridge_priority_news_with_label():
    assert _is_bridge('Priority News: animal-people clip', '', False) is True


def test_bridge_priority_news_bare():
    assert _is_bridge('Priority News:', '', False) is True


def test_not_bridge_priority_news_sentence():
    """`Priority News was discussed in detail.` — sentence with no
    colon after News. Must NOT match (the previous draft of the
    pattern over-matched here)."""
    assert _is_bridge('Priority News was discussed in detail.', '', False) is False


# Pattern 6: template placeholder with 3+ underscores
def test_bridge_template_placeholder():
    """`in _____ with subtitles` — the underscores are an editing
    placeholder for a language name; FA stays blank."""
    assert _is_bridge('in _____ with subtitles', '', False) is True


def test_bridge_template_placeholder_at_end():
    assert _is_bridge('with subtitles in _____', '', False) is True


def test_not_bridge_single_underscore():
    """Single underscore (e.g. variable name) shouldn't trigger."""
    assert _is_bridge('the my_var was set.', '', False) is False


# ── 2026-05-20 (AJAR 3154 + GAT 3154 analysis): three more findings ─────────
#
# After the 2026-05-19 patch eliminated 6 / 15 over-fills in WAU 3153, two
# more test documents (AJAR 3154, GAT 3154) revealed three remaining
# patterns the bridge detector still missed:
#
#   A1) Group / chorus speakers — "Band(all):", "Vocalists(all):"
#   A2) Multi-name speakers with comma + ampersand — "Wendy, Carnie
#       Wilson & Owen Elliot(f):"
#   A3) Bare "(m): real dialogue" — not a bridge but a new turn boundary
#
# A1 + A2 are fixed by broadening _SPEAKER_RE and the timecode+speaker
# bridge pattern to accept multi-name characters and the full speaker
# token list. A3 is fixed by an absorption-loop guard in _parse_groups.


# A1 — Group / chorus speaker tags
def test_bridge_band_all():
    """`Band(all):` — band/chorus speaker tag, FA should stay blank."""
    assert _is_bridge('Band(all):', '', False) is True


def test_bridge_vocalists_all():
    """`Vocalists(all):` — same pattern, different keyword."""
    assert _is_bridge('Vocalists(all):', '', False) is True


def test_bridge_performer_m():
    """`Performer(m):` — broadened from `[mf]` to include any speaker token."""
    assert _is_bridge('Performer(m):', '', False) is True


def test_bridge_timecode_band_all():
    """`02:52:44 Band(all):` — AJAR 3154 row #336/#356 case."""
    assert _is_bridge('02:52:44 Band(all):', '', False) is True


# A2 — Multi-name speakers with comma + ampersand
def test_bridge_multi_name_speaker_with_comma_ampersand():
    """`Wendy, Carnie Wilson & Owen Elliot(f):` — AJAR 3154 row #317/#375
    case. Old _SPEAKER_RE char class `[A-Za-z\s\-]` rejected `,` `&`."""
    assert _is_bridge('Wendy, Carnie Wilson & Owen Elliot(f):', '', False) is True


def test_bridge_multi_name_with_comma_and():
    """`Al, Matthew & Adam Jardine(m):` — AJAR 3154 row #245 case."""
    assert _is_bridge('Al, Matthew & Adam Jardine(m):', '', False) is True


def test_not_bridge_speaker_dialogue_after_colon():
    """`Lee (f): She continued.` — must still NOT match (dialogue, not label)."""
    assert _is_bridge('Lee (f): She continued.', '', False) is False


# A3 — Absorption-loop guard for "(m): real dialogue" rows
def test_not_bridge_bare_speaker_with_content():
    """`(m): Look at you! You pudge!` is NOT a bridge — the content
    after the colon is real dialogue that should be translatable."""
    assert _is_bridge('(m): Look at you! You pudge!', '', False) is False


def test_absorption_guard_blocks_speaker_turn_absorption(tmp_path):
    """The actual GAT 3154 #80-82 case end-to-end.

    Build a 5-row table where:
      row 0: complete FA, sentence-ending
      row 1: EN starts with "(m): ..." and FA is empty (polish stage
              didn't translate it)
      row 2: EN is more dialogue, FA also empty
      row 3: bridge HOST(m):
      row 4: next sentence, complete FA

    Without the absorption guard, rows 1+2 get absorbed as "spare slots"
    of row 0's sentence group; the row 0 chunk gets split into pieces
    and spills into rows 1, 2 (observed bug). With the guard, rows 1+2
    stay blank and the row 0 chunk lives in row 0 only.
    """
    import docx as _d

    doc = _d.Document()
    t = doc.add_table(rows=5, cols=3)
    rows = [
        ('', 'and she is absolutely feisty.', 'و کاملاً پرجنب‌وجوش است.'),
        ('', '(m): Look at you! You pudge!', ''),
        ('', 'Let go! Let go!', ''),
        ('', 'HOST (m):', ''),
        ('', 'Earth-loving viewers,', 'بینندگان دوستدار زمین'),
    ]
    for ri, r in enumerate(rows):
        for ci in range(3):
            t.cell(ri, ci).text = r[ci]
    in_path = tmp_path / "guard.docx"
    doc.save(str(in_path))

    from machine_translate_docx.openai_tools.persian_double_lines import (
        FASubtitleAligner,
    )
    aligner = FASubtitleAligner(model="gpt-5.4-mini", llm_threshold=0)
    out_path = tmp_path / "guard_out.docx"
    aligner.align(str(in_path), str(out_path))

    out_doc = _d.Document(str(out_path))
    fa_cells = [r.cells[2].text.strip()
                for table in out_doc.tables for r in table.rows]

    # Row 0: keep the full FA
    assert fa_cells[0], "Row 0 FA must remain non-empty"
    # Row 1: must stay blank — the absorption guard prevents this
    assert fa_cells[1] == "", (
        f"Row 1 (m): turn-start should stay blank, got {fa_cells[1]!r}"
    )
    # Row 2: also stays blank (no FA in polish input, not absorbed)
    assert fa_cells[2] == "", (
        f"Row 2 should stay blank, got {fa_cells[2]!r}"
    )
