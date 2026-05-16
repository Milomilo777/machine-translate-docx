"""Unit tests for ``src/machine_translate_docx/docx_io/save.py``.

Covers the five public / semi-public surfaces:

  * ``engine_suffix(ctx)``
  * ``_resolve_output_path(ctx)``          — mutates ctx in-place
  * ``_write_minimal_sidecar(ctx)``        — NOTE: takes ctx only (no log_path arg)
  * ``_restore_source_column(ctx)``
  * ``save_docx_file(ctx, docxdoc, *, silent, write_translation_log_fn)``

Deviation from the task description:
  ``_write_minimal_sidecar`` takes only ``ctx`` (not ctx + log_path).
  It resolves the log path internally via ``log_paths.resolve_log_path``.
  Tests patch ``resolve_log_path`` to redirect output into tmp_path.

  ``_resolve_output_path`` takes only ``ctx`` (no positional path args).
  It reads/writes ``ctx.flags.word_file_to_translate`` and
  ``ctx.flags.word_file_to_translate_save_as_path``.
"""
from __future__ import annotations

import io
import json
import os
import sys
from copy import deepcopy
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

_SRC = Path(__file__).resolve().parents[1] / "src"
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))

import docx as _docx  # python-docx  # noqa: E402

from machine_translate_docx.runtime import RuntimeContext  # noqa: E402
from machine_translate_docx.docx_io.save import (  # noqa: E402
    engine_suffix,
    save_docx_file,
    _resolve_output_path,
    _write_minimal_sidecar,
    _restore_source_column,
)


# ── helpers ──────────────────────────────────────────────────────────────────

def _make_ctx(
    *,
    engine: str = "chatgpt",
    method: str = "api",
    with_polish: bool = False,
    dest_lang: str = "fa",
    input_path: str = "/tmp/test.docx",
) -> RuntimeContext:
    ctx = RuntimeContext.empty()
    ctx.engine.engine = engine
    ctx.engine.method = method
    ctx.flags.with_polish = with_polish
    ctx.language.dest_lang = dest_lang
    ctx.language.src_lang = "en"
    ctx.flags.word_file_to_translate = input_path
    ctx.flags.word_file_to_translate_save_as_path = None
    return ctx


def _make_docx_in_memory() -> io.BytesIO:
    """Return a BytesIO of a minimal one-table docx (3 rows x 3 cols)."""
    document = _docx.Document()
    table = document.add_table(rows=3, cols=3)
    for ri in range(3):
        for ci in range(3):
            table.rows[ri].cells[ci].text = f"orig_r{ri}_c{ci}"
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


def _open_docx_from_bytes(buf: io.BytesIO) -> _docx.Document:
    buf.seek(0)
    return _docx.Document(buf)


# ── engine_suffix tests ──────────────────────────────────────────────────────

def test_engine_suffix_polish():
    """chatgpt + api + with_polish=True → '_Polish'."""
    ctx = _make_ctx(engine="chatgpt", method="api", with_polish=True)
    assert engine_suffix(ctx) == "_Polish"


def test_engine_suffix_chatgpt_api_no_polish():
    """chatgpt + api + with_polish=False → '_chatGPT'."""
    ctx = _make_ctx(engine="chatgpt", method="api", with_polish=False)
    assert engine_suffix(ctx) == "_chatGPT"


def test_engine_suffix_google():
    """google engine → '_Google' regardless of method or polish."""
    ctx = _make_ctx(engine="google", method="phrasesblock", with_polish=False)
    assert engine_suffix(ctx) == "_Google"


def test_engine_suffix_deepl():
    """deepl engine → '_Deepl' regardless of method or polish."""
    ctx = _make_ctx(engine="deepl", method="phrasesblock", with_polish=False)
    assert engine_suffix(ctx) == "_Deepl"


# ── _resolve_output_path tests ────────────────────────────────────────────────

def test_resolve_output_path_no_collision(tmp_path):
    """No collision: resolved path gets _{LANG}{engine_tag}.docx suffix."""
    input_file = str(tmp_path / "episode.docx")
    # Create a stub input file so the function can reference it.
    Path(input_file).write_bytes(b"PK stub")

    ctx = _make_ctx(engine="chatgpt", method="api", with_polish=True,
                    dest_lang="fa", input_path=input_file)
    _resolve_output_path(ctx)

    out = ctx.flags.word_file_to_translate_save_as_path
    assert out is not None
    # Should contain lang tag PER and engine tag _Polish.
    assert "_PER" in out
    assert "_Polish" in out
    assert out.lower().endswith(".docx")
    # The output file must NOT exist yet (no collision suffix needed).
    assert not os.path.exists(out)


def test_resolve_output_path_with_collision(tmp_path):
    """Collision avoidance: _1, _2 suffixes prevent overwriting."""
    input_file = str(tmp_path / "episode.docx")
    Path(input_file).write_bytes(b"PK stub")

    ctx = _make_ctx(engine="chatgpt", method="api", with_polish=True,
                    dest_lang="fa", input_path=input_file)

    # First call — no collision.
    _resolve_output_path(ctx)
    first_out = ctx.flags.word_file_to_translate_save_as_path
    assert first_out is not None

    # Simulate the first file existing on disk.
    Path(first_out).write_bytes(b"PK translated v1")

    # Second call — collision: expect _1 suffix.
    ctx2 = _make_ctx(engine="chatgpt", method="api", with_polish=True,
                     dest_lang="fa", input_path=input_file)
    _resolve_output_path(ctx2)
    second_out = ctx2.flags.word_file_to_translate_save_as_path
    assert second_out != first_out
    assert second_out.endswith("_1.docx")

    # Simulate the _1 file also existing.
    Path(second_out).write_bytes(b"PK translated v2")

    # Third call — collision on _1: expect _2 suffix.
    ctx3 = _make_ctx(engine="chatgpt", method="api", with_polish=True,
                     dest_lang="fa", input_path=input_file)
    _resolve_output_path(ctx3)
    third_out = ctx3.flags.word_file_to_translate_save_as_path
    assert third_out.endswith("_2.docx")


# ── _write_minimal_sidecar tests ──────────────────────────────────────────────

def test_write_minimal_sidecar_shape(tmp_path):
    """Minimal sidecar JSON has the expected top-level keys and summary sub-keys.

    _write_minimal_sidecar(ctx) resolves the log path internally via
    log_paths.resolve_log_path(out_path). We patch that resolver to
    redirect output into tmp_path so the test is hermetic.
    """
    out_docx = str(tmp_path / "episode_PER_Google.docx")
    log_file = str(tmp_path / "episode_PER_Google_log.json")

    ctx = _make_ctx(engine="google", method="phrasesblock", dest_lang="fa",
                    input_path=str(tmp_path / "episode.docx"))
    ctx.flags.word_file_to_translate_save_as_path = out_docx
    ctx.language.src_lang = "en"
    ctx.language.dest_lang = "fa"
    ctx.docx.from_text_table = ["Hello", "World", ""]
    ctx.docx.to_text_by_phrase_separator_table = ["سلام", "جهان", ""]

    with patch(
        "machine_translate_docx.log_paths.resolve_log_path",
        return_value=log_file,
    ):
        _write_minimal_sidecar(ctx)

    assert os.path.exists(log_file), "sidecar file was not created"
    with open(log_file, encoding="utf-8") as fh:
        payload = json.load(fh)

    # Top-level keys
    assert "run_info" in payload
    assert "blocks" in payload
    assert "summary" in payload

    # blocks must be a list (empty for non-OpenAI engines)
    assert isinstance(payload["blocks"], list)

    # summary sub-keys
    summary = payload["summary"]
    for key in ("total_blocks", "total_tokens", "total_cost_usd",
                 "elapsed_total_seconds", "source_rows_nonempty",
                 "target_rows_nonempty", "row_count"):
        assert key in summary, f"summary missing key: {key!r}"

    # Spot-check computed values.
    assert summary["source_rows_nonempty"] == 2   # "Hello" + "World"
    assert summary["target_rows_nonempty"] == 2   # "سلام" + "جهان"
    assert summary["row_count"] == 3              # max(len(src), len(tgt))


# ── _restore_source_column tests ──────────────────────────────────────────────

def test_restore_source_column_no_drift(tmp_path):
    """Drifted cols 0+1 are restored to their parse-time snapshot.

    Verification strategy: round-trip through a BytesIO so python-docx
    re-parses the XML after the lxml-level replace, avoiding any stale
    in-memory wrapper references. The saved docx carries the replaced
    <w:tc> elements, so re-reading it gives the definitive result.
    """
    buf = _make_docx_in_memory()
    document = _open_docx_from_bytes(buf)
    table = document.tables[0]

    ctx = RuntimeContext.empty()
    ctx.docx.table = table

    # Build the snapshot: capture col 0 and col 1 for every row.
    snapshot: dict = {}
    orig_texts: dict = {}
    for ri, row in enumerate(table.rows):
        for ci in (0, 1):
            if ci < len(row.cells):
                cell = row.cells[ci]
                orig_text = cell.text
                orig_texts[(ri, ci)] = orig_text
                orig_tc = deepcopy(cell._tc)
                snapshot[(ri, ci)] = (orig_text, orig_tc)

    ctx.docx.source_columns_snapshot = snapshot

    # Drift: overwrite col 0 text in every row (mutate the XML in place).
    for row in table.rows:
        if row.cells:
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.text = "DRIFTED"

    # Verify drift happened before restore.
    assert "DRIFTED" in table.rows[0].cells[0].text

    # Restore.
    _restore_source_column(ctx)

    # Round-trip: save the document to bytes and re-open so python-docx
    # re-parses the XML tree, giving fresh Cell objects backed by the
    # replaced <w:tc> elements.
    out_buf = io.BytesIO()
    document.save(out_buf)
    out_buf.seek(0)
    reopened = _docx.Document(out_buf)
    restored_table = reopened.tables[0]

    for ri in range(len(restored_table.rows)):
        for ci in (0, 1):
            expected_text = orig_texts[(ri, ci)]
            actual_text = restored_table.rows[ri].cells[ci].text
            assert actual_text == expected_text, (
                f"row {ri} col {ci}: expected {expected_text!r}, got {actual_text!r}"
            )


# ── save_docx_file callback test ──────────────────────────────────────────────

def test_save_docx_file_invokes_write_log_callback(tmp_path):
    """write_translation_log_fn is called once with the resolved log path
    when with_polish=True and ctx.openai.translation_log has 'blocks'."""
    input_file = str(tmp_path / "episode.docx")
    Path(input_file).write_bytes(b"PK stub")

    ctx = _make_ctx(engine="chatgpt", method="api", with_polish=True,
                    dest_lang="fa", input_path=input_file)
    # Populate translation_log so the with_polish branch is taken.
    ctx.openai.translation_log = {"blocks": [{"block": 1}]}

    # Build a real (minimal) python-docx document so docxdoc.save() works.
    buf = _make_docx_in_memory()
    docxdoc = _open_docx_from_bytes(buf)

    mock_write_log = MagicMock()

    # Patch resolve_log_path so we control the log path handed to the callback.
    expected_log_path = str(tmp_path / "episode_PER_Polish_log.json")
    with patch(
        "machine_translate_docx.log_paths.resolve_log_path",
        return_value=expected_log_path,
    ), patch.dict(os.environ, {"MTD_DISABLE_SIDECAR": ""}):
        save_docx_file(
            ctx,
            docxdoc,
            silent=True,
            write_translation_log_fn=mock_write_log,
        )

    mock_write_log.assert_called_once()
    actual_log_path = mock_write_log.call_args[0][0]
    assert actual_log_path == expected_log_path
