"""Unit tests for ``src/machine_translate_docx/docx_io/runs.py``.

Covers ``_iter_paragraph_runs`` — the iterator that walks every ``<w:r>``
element below a paragraph, including those nested inside ``<w:hyperlink>``.

The hyperlink test uses ``tests/fixtures/sample_hyperlink.docx``: row 8 of
its single table contains the paragraph "Here is a hyperlink with alt text."
where the word "hyperlink" sits inside a ``<w:hyperlink>`` element.
Native ``paragraph.runs`` silently drops that run; ``_iter_paragraph_runs``
must include it.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest

_SRC = Path(__file__).resolve().parents[1] / "src"
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))

import docx  # noqa: E402

from machine_translate_docx.docx_io.runs import _iter_paragraph_runs  # noqa: E402

# ── helpers ───────────────────────────────────────────────────────────────────

_FIXTURE = Path(__file__).resolve().parent / "fixtures" / "sample_hyperlink.docx"


def _make_paragraph(text_runs: list[str]):
    """Return a fresh python-docx paragraph (inside a one-cell table) with the
    given runs added in order. Round-tripped through bytes so the XML reflects
    what python-docx actually stores on disk.

    Note: do NOT call ``para.text = ""`` before adding runs — that assignment
    inserts an empty ``<w:r/>`` stub into the XML which would inflate run counts
    and confuse the assertions.
    """
    document = docx.Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    para = cell.paragraphs[0]
    for t in text_runs:
        para.add_run(t)
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    reopened = docx.Document(buf)
    return reopened.tables[0].rows[0].cells[0].paragraphs[0]


# ── tests ─────────────────────────────────────────────────────────────────────

def test_iter_paragraph_runs_plain_text():
    """Single plain-text run yields exactly one Run with the expected text."""
    para = _make_paragraph(["Hello world"])
    runs = list(_iter_paragraph_runs(para))
    assert len(runs) == 1
    assert runs[0].text == "Hello world"


def test_iter_paragraph_runs_multiple_runs():
    """Three runs are yielded in document order; concatenation matches."""
    para = _make_paragraph(["Hello ", "world", "!"])
    runs = list(_iter_paragraph_runs(para))
    assert len(runs) == 3
    assert runs[0].text == "Hello "
    assert runs[1].text == "world"
    assert runs[2].text == "!"
    assert "".join(r.text for r in runs) == "Hello world!"


def test_iter_paragraph_runs_hyperlink_included():
    """Run nested inside ``<w:hyperlink>`` is included in the iteration.

    The fixture paragraph at table row 8, cell 1 reads:
        "Here is a hyperlink with alt text."
    Native ``paragraph.runs`` returns only 2 runs (skipping "hyperlink").
    ``_iter_paragraph_runs`` must return 3 runs and the concatenation must
    contain the word "hyperlink".
    """
    doc = docx.Document(str(_FIXTURE))
    para = doc.tables[0].rows[8].cells[1].paragraphs[0]

    # Sanity-check that this is the right paragraph.
    assert "hyperlink" in para.text

    native_run_count = len(para.runs)
    all_runs = list(_iter_paragraph_runs(para))
    concatenated = "".join(r.text for r in all_runs)

    # The iterator must expose MORE runs than the native accessor because
    # native drops the hyperlink-nested run.
    assert len(all_runs) > native_run_count

    # The word "hyperlink" must appear in the concatenated text of all runs.
    assert "hyperlink" in concatenated

    # Full round-trip text should match paragraph.text.
    assert concatenated == para.text


def test_iter_paragraph_runs_empty_paragraph():
    """Empty paragraph produces an empty iterator (no runs yielded)."""
    para = _make_paragraph([])
    runs = list(_iter_paragraph_runs(para))
    assert runs == []
