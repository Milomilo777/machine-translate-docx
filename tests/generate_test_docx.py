"""
Generates a minimal test DOCX file for automated translation tests.
Usage: python tests/generate_test_docx.py
Output: tests/fixtures/test_input.docx
/ تولید فایل DOCX آزمایشی برای تست‌های خودکار
"""
import os
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "fixtures")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "test_input.docx")


def add_colored_run(paragraph, text: str, rgb: tuple):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(*rgb)
    return run


def generate():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    # Row 1: normal translatable text
    doc.add_paragraph("You can try translating this sample file.")

    # Row 2: normal text with line break
    doc.add_paragraph("The program supports multiple translation engines.")

    # Row 3: dark blue text — should be BYPASSED (not translated)
    p = doc.add_paragraph()
    add_colored_run(p, "Dark Blue text is ignored by the translator.",
                   (0, 0, 139))

    # Row 4: more normal text
    doc.add_paragraph("It does not require Microsoft Office.")

    # Row 5: normal text — translatable
    doc.add_paragraph("God Bless!")

    doc.save(OUTPUT_FILE)
    print(f"[OK] Test DOCX generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    generate()
