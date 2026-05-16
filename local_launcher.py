from __future__ import annotations

import argparse
import hashlib
import json
import mimetypes
import os
import shutil
import tempfile
import threading
import time
import uuid
import webbrowser
import sys
from dataclasses import dataclass, asdict
from email.parser import BytesParser
from email.policy import default as email_policy
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
import subprocess
from urllib.parse import unquote, urlparse


import re as _re

# Windows console (cp1252) crashes when print() emits ▶ ✓ ✗ — that the
# job-progress logs use. Force stdout/stderr to UTF-8 so the _process_job
# thread does not die on the first decoration. (Audit finding F-013.)
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except (AttributeError, OSError):
    pass

# W-2 (2026-05-11): also export PYTHONUTF8 / PYTHONIOENCODING into the
# parent process environment so anything we shell out to that does NOT
# go through `subprocess.Popen(env=...)` (e.g. an aligner helper that
# spawns its own subprocess, a developer running scripts from the
# launcher's REPL) inherits a UTF-8 IO mode by default. Subprocess.Popen
# calls below override the env explicitly anyway — this is a
# belt-and-suspenders default.
os.environ.setdefault("PYTHONUTF8", "1")
os.environ.setdefault("PYTHONIOENCODING", "utf-8")

ROOT = Path(__file__).resolve().parent
INDEX_FILE = ROOT / "index.ejs"
WEB_V2_DIR = ROOT / "web" / "v2"
WEB_STATIC_DIR = ROOT / "web" / "static"
SUBSCRIBERS_FILE = ROOT / "subscribers.txt"

# 5-day cache for API-translated files (sha256(payload) + lang + engine
# + ai_model). Cache key + paths held in memory; payloads themselves stay on
# disk under runtime_dir/cache/<hash>/. Pruned by start_cleanup_thread().
# 2026-05-15: bumped from 36 hours to 5 days. With the new "raw cache +
# post-process split" architecture, a single cache entry serves both
# Basic and Double Lines requests for the same source file, so longer
# retention pays off — a returning user who tries a different split
# method on a 4-day-old file still skips translation + polish.
CACHE_TTL_SEC = 5 * 24 * 60 * 60
_API_ENGINES = frozenset({"chatgpt", "chatgpt-polish"})  # only API engines cache

# ISO 639-2/B codes matching what machine_translate_docx.cli produces via langcodes
_LANG_ALPHA3B = {
    'fa': 'PER', 'ar': 'ARA', 'de': 'GER', 'fr': 'FRE',
    'zh-hans': 'CHI', 'zh-hant': 'CHI', 'zh-cn': 'CHI', 'zh-tw': 'CHI',
    'ko': 'KOR', 'ja': 'JPN', 'ru': 'RUS', 'es': 'SPA',
    'it': 'ITA', 'pt': 'POR', 'nl': 'DUT', 'pl': 'POL',
    'tr': 'TUR', 'sv': 'SWE', 'no': 'NOR', 'da': 'DAN',
    'fi': 'FIN', 'he': 'HEB', 'uk': 'UKR', 'cs': 'CZE',
    'ro': 'RUM', 'hu': 'HUN', 'vi': 'VIE', 'th': 'THA',
}


def _lang_suffix(target_language: str) -> str:
    return _LANG_ALPHA3B.get(target_language.lower(), target_language.replace('-', '').upper())


def _double_lines_output_path(base_path: Path) -> Path:
    """Return the Persian Double Lines variant of a translated docx path.

    Phase 6 contract: the suffix is `_Double_Lines` and is appended after
    the engine suffix, just before `.docx`. Examples:

        sample_PER_Polish.docx  → sample_PER_Polish_Double_Lines.docx
        sample_PER_chatGPT.docx → sample_PER_chatGPT_Double_Lines.docx
    """
    stem = _re.sub(r"(?i)\.docx$", "", base_path.name)
    return base_path.with_name(f"{stem}_Double_Lines.docx")


def _engine_suffix_for(translation_engine: str | None) -> str:
    """Filename suffix appended after the lang code, per engine.

    Mirrors the table in `save_docx_file._engine_suffix(ctx)` on the
    backend side. Used by `_fallback_output_path` when the launcher
    has to guess the output filename because the subprocess never
    printed `Saved file name:`.
    """
    # cleanup pass — Cloudflare gating made them never-reach-prod.
    return {
        'google':         '_Google',
        'deepl':          '_Deepl',
        'chatgpt':        '_chatGPT',
        'chatgpt-polish': '_Polish',
    }.get((translation_engine or '').lower().strip(), '')


def _sanitize_filename(name: str) -> str:
    """Reduce a user-supplied filename to a safe basename.

    R-5 (2026-05-11 audit): caps the result to 200 characters,
    preserving the extension. Without the cap, a 1000-character UTF-8
    filename hits Windows' 255-byte path limit and python-docx fails
    when saving with a confusing OSError. The cap also bounds anything
    we later splice into archive folder names.
    """
    name = Path(name).name
    name = name.replace("\x00", "")
    name = name.replace("（", "(").replace("）", ")")
    name = name.replace("&", "")
    name = name.strip() or "upload.docx"
    if len(name) > 200:
        # Preserve the suffix (last `.ext`, max 10 chars) so MIME / docx
        # detection still works on the truncated name.
        suffix = ""
        dot = name.rfind(".")
        if 0 <= dot >= len(name) - 11:
            suffix = name[dot:]
        head_budget = 200 - len(suffix)
        name = name[:head_budget] + suffix
    return name


# ── 5-day cache for API-translated outputs (CACHE_TTL_SEC above) ─────────────

def _cache_key(payload: bytes, target_lang: str, engine: str,
               ai_model: str | None, split_engine: str | None = None) -> str:
    """SHA-256 over payload + lang + engine + ai_model + split_engine.

    Two requests collide ONLY when the uploaded bytes are byte-identical AND
    the language/engine/model/split_engine quadruple matches. A one-byte
    difference in the docx zip (e.g. different metadata, different
    timestamp) yields a different key — by design.

    2026-05-15 — `split_engine` added to the key. The launcher routes
    Persian Double Lines and Basic through different CLI flag combinations
    (B1-guard forces `splitTranslate=false` for persian_double_lines, true
    for basic), so the cached ``main_path`` file has a different row
    shape depending on which split method ran first. Without
    `split_engine` in the key, a Persian-Double-Lines-first run would
    cache a raw "everything in row 1" docx, and a subsequent Basic run
    would replay that wrong-shape cache (bug observed on MOS 3148,
    2026-05-15). Including `split_engine` forces a cache miss on switch,
    which re-runs the engine with the correct flags for the new method.
    """
    h = hashlib.sha256()
    h.update(payload)
    h.update(b"\x00")
    h.update(target_lang.encode("utf-8", errors="replace"))
    h.update(b"\x00")
    h.update(engine.encode("utf-8", errors="replace"))
    h.update(b"\x00")
    h.update((ai_model or "").encode("utf-8", errors="replace"))
    h.update(b"\x00")
    h.update((split_engine or "").encode("utf-8", errors="replace"))
    return h.hexdigest()


# ── Telegram alert helpers ───────────────────────────────────────────────────


def _parse_telegram_chat_ids(raw: str) -> list[str]:
    """Split a comma- / semicolon- / whitespace-separated list of chat ids.

    Each id is left as a string because Telegram accepts either a
    numeric DM id (``987654321``), a numeric group id (``-987654321``),
    or a public channel handle (``@my_channel``). Validating the shape
    here would block the legitimate channel-handle case.

    Empty / whitespace-only entries are dropped. Duplicates are
    preserved on purpose — operators sometimes intentionally repeat a
    chat id with a different separator while debugging.
    """
    if not raw:
        return []
    parts: list[str] = []
    for piece in raw.replace(";", ",").split(","):
        for sub in piece.split():        # collapse internal whitespace
            sub = sub.strip()
            if sub:
                parts.append(sub)
    return parts


def _telegram_escape(s: str) -> str:
    """Escape Markdown-special chars inside Telegram alert strings.

    Telegram's `Markdown` parse mode (the legacy one we use, not
    `MarkdownV2`) treats ``*``, ``_``, ``\\``` and ``[`` as formatting
    characters. A user-supplied filename like ``my_doc.docx`` would
    otherwise turn the next word italic. The helper neutralises the
    five characters that can break a code span / italic / bold /
    link, leaving everything else alone — Telegram silently passes
    through anything Markdown does not recognise.
    """
    if s is None:
        return ""
    return (
        str(s)
        .replace("\\", "\\\\")
        .replace("`",  "\\`")
        .replace("*",  "\\*")
        .replace("_",  "\\_")
        .replace("[",  "\\[")
    )


# ── Newsletter subscriber list ───────────────────────────────────────────────

_RE_EMAIL = _re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def _append_subscriber(email: str) -> tuple[bool, str]:
    """Append `email` to subscribers.txt if it validates and is new.

    Returns (ok, message). The file is line-delimited UTF-8; existing
    duplicates are silently ignored so re-submission is idempotent.
    """
    email = (email or "").strip().lower()
    if not _RE_EMAIL.match(email):
        return False, "invalid email format"
    if len(email) > 254:
        return False, "email too long"

    try:
        existing: set[str] = set()
        if SUBSCRIBERS_FILE.exists():
            with SUBSCRIBERS_FILE.open("r", encoding="utf-8") as f:
                existing = {line.strip().lower() for line in f if line.strip()}
        if email in existing:
            return True, "already subscribed"
        with SUBSCRIBERS_FILE.open("a", encoding="utf-8") as f:
            f.write(email + "\n")
        return True, "subscribed"
    except Exception as exc:
        # R-4 (2026-05-11 audit): do not echo the raw exception text into
        # the user-visible JSON response — a PermissionError leaks the
        # absolute filesystem path; a UnicodeDecodeError leaks the
        # offending bytes. Log the full traceback to stderr for the
        # operator and return a generic message.
        import traceback as _tb
        print(
            f"[subscribe] server error: {exc!r}\n{_tb.format_exc()}",
            file=sys.stderr,
            flush=True,
        )
        return False, "server error"


# Server-side upload limits
_MAX_DOCX_UNCOMPRESSED = 50 * 1024 * 1024   # 50 MB total uncompressed entries
_MAX_DOCX_COMPRESSED   = 20 * 1024 * 1024   # 20 MB max raw POST body
_DOCX_MAGIC_PK         = b"PK\x03\x04"      # ZIP local file header (DOCX is a ZIP)
# Required entries inside a real DOCX. Validation rejects ZIPs missing them.
_DOCX_REQUIRED_PARTS   = ("[Content_Types].xml", "word/document.xml")

# Concurrency cap on real-backend subprocesses. Each subprocess loads
# python-docx + openai client + tiktoken (≈250-500 MB). Two slots is a
# safe default for a workstation; increase via MTD_MAX_CONCURRENT_JOBS env var.
_MAX_CONCURRENT_JOBS = int(os.environ.get("MTD_MAX_CONCURRENT_JOBS", "2"))
_job_semaphore       = threading.Semaphore(_MAX_CONCURRENT_JOBS)


def _validate_docx_payload(payload: bytes) -> str | None:
    """Return None when payload is a safe DOCX, otherwise an error message.

    Three layers:
      1. Magic bytes — first four bytes must be ZIP local-file header (PK\\x03\\x04).
      2. Decompressed-size cap — sum of all entries' file_size must stay under
         _MAX_DOCX_UNCOMPRESSED. Defends against zip-bomb DOCX uploads.
      3. DOCX shape — the archive must contain the required parts so a plain
         ZIP can't masquerade as a Word document.
    """
    if not payload or not payload.startswith(_DOCX_MAGIC_PK):
        return "File does not look like a DOCX (missing ZIP header)."

    import io
    import zipfile
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as zf:
            total = 0
            names = set(zf.namelist())
            for zi in zf.infolist():
                total += zi.file_size
                if total > _MAX_DOCX_UNCOMPRESSED:
                    return (
                        f"DOCX uncompressed size exceeds the "
                        f"{_MAX_DOCX_UNCOMPRESSED // (1024 * 1024)} MB limit."
                    )
            for required in _DOCX_REQUIRED_PARTS:
                if required not in names:
                    return f"DOCX is missing required part: {required}"
    except zipfile.BadZipFile:
        return "DOCX file is corrupted or not a valid ZIP archive."
    return None


def _load_index_html() -> str:
    return INDEX_FILE.read_text(encoding="utf-8", errors="replace")


def _inject_client_patch(html: str) -> str:
    patch = """
<script>
(function () {
  const originalFetch = window.fetch.bind(window);
  window.fetch = async function (input, init) {
    try {
      const url = typeof input === 'string' ? input : (input && input.url) ? input.url : '';
      const body = init && init.body;
      if (url.includes('/upload') && body instanceof FormData) {
        const aiModel = document.getElementById('aiModel');
        const enableSound = document.getElementById('enableSound');
        const soundSelect = document.getElementById('soundSelect');

        if (aiModel && !body.has('aiModel')) body.append('aiModel', aiModel.value);
        if (enableSound && !body.has('enableSound')) body.append('enableSound', enableSound.checked ? 'on' : 'off');
        if (soundSelect && !body.has('soundSelect')) body.append('soundSelect', soundSelect.value);
      }
    } catch (err) {
      console.warn('Local launcher patch failed:', err);
    }
    return originalFetch(input, init);
  };
})();
</script>
"""
    if "</body>" in html:
        return html.replace("</body>", patch + "\n</body>", 1)
    return html + patch


def _read_int(path: Path, default: int = 0) -> int:
    try:
        return int(path.read_text(encoding="utf-8").strip())
    except Exception:
        return default


def _write_int(path: Path, value: int) -> None:
    path.write_text(str(int(value)), encoding="utf-8")


def _build_placeholder_docx(
    output_path: Path,
    *,
    job_id: str,
    original_name: str,
    source_file: str,
    source_language: str,
    target_language: str,
    translation_engine: str,
    split_translate: bool,
    split_engine: str | None,
    ai_model: str | None,
    enable_sound: str | None,
    sound_select: str | None,
) -> None:
    # python-docx is available in this environment, so we can generate a
    # visible DOCX instead of returning a byte-for-byte copy.
    from docx import Document

    doc = Document()
    doc.add_heading("Local Test Output", level=0)
    doc.add_paragraph(
        "This document was generated by local_launcher.py as a local UI test placeholder."
    )
    doc.add_paragraph(f"Job ID: {job_id}")
    doc.add_paragraph(f"Original upload name: {original_name}")
    doc.add_paragraph(f"Stored source file: {source_file}")
    doc.add_paragraph(f"Source language: {source_language}")
    doc.add_paragraph(f"Target language: {target_language}")
    doc.add_paragraph(f"Translation engine: {translation_engine}")
    doc.add_paragraph(f"Split translation: {'yes' if split_translate else 'no'}")
    doc.add_paragraph(f"Split engine: {split_engine or '(not sent)'}")
    doc.add_paragraph(f"AI model: {ai_model or '(not sent)'}")
    doc.add_paragraph(f"Sound enabled: {enable_sound or '(not sent)'}")
    doc.add_paragraph(f"Sound selection: {sound_select or '(not sent)'}")
    doc.add_paragraph(
        "If you want a real end-to-end translation run, the current machine still needs the missing Python dependencies used by src/machine_translate_docx.py."
    )
    doc.save(output_path)


@dataclass
class Job:
    status: str
    filename: str | None
    error: str | None
    created_at: float
    progress: int = 0             # 0-100; updated by PROGRESS:N markers from backend


class LocalState:
    def __init__(self, runtime_dir: Path, backend_mode: str, python_exe: Path, script_path: Path):
        self.runtime_dir = runtime_dir
        self.uploads_dir = runtime_dir / "uploads"
        self.logs_dir = runtime_dir / "logs"
        self.views_dir = runtime_dir / "views"
        self.ssl_dir = runtime_dir / "ssl"
        self.cache_dir = runtime_dir / "cache"
        self.count_file = runtime_dir / "count.txt"
        self.index_file = self.views_dir / "index.ejs"
        self.backend_mode = backend_mode
        self.python_exe = python_exe
        self.script_path = script_path
        self.lock = threading.Lock()
        self.jobs: dict[str, Job] = {}
        # Subprocess handles, keyed by job_id. Held while a real-backend
        # job runs so POST /cancel/<id> can terminate it. Mock jobs never
        # populate this. Cleared when the subprocess exits.
        self.job_procs: dict[str, "subprocess.Popen"] = {}
        self.total_uploads = 0
        # cache_key → (timestamp, payload_dict). The dict carries:
        #   main_path:                Path  — engine output, no splitter
        #   source_path:              Path  — original upload bytes
        #   translation_array:        list[str] — for line-aware splitters
        #   phrase_separator_table:   list[str]
        #   engine, ai_model, src_lang, dest_lang: str | None
        # The cache key intentionally excludes the splitter so a re-upload
        # with a different Split Method reuses the cached translation and
        # only re-runs the splitter (phase 4 of the persian-double-lines
        # roadmap). Pre-phase-4 (timestamp, list-of-tuples) entries are
        # evicted on access.
        self.cache: dict[str, tuple[float, dict]] = {}

    def boot(self) -> None:
        for directory in [
            self.runtime_dir,
            self.uploads_dir,
            self.logs_dir,
            self.views_dir,
            self.ssl_dir,
            self.cache_dir,
        ]:
            directory.mkdir(parents=True, exist_ok=True)

        self.index_file.write_text(_inject_client_patch(_load_index_html()), encoding="utf-8")
        _write_int(self.count_file, 0)

        # The original server expects these files to exist during startup.
        # Their contents are not used in mock mode, so placeholders are fine.
        (self.ssl_dir / "private.key").write_text("LOCAL-MOCK-KEY", encoding="utf-8")
        (self.ssl_dir / "certificate.crt").write_text("LOCAL-MOCK-CERT", encoding="utf-8")
        (self.ssl_dir / "ca_bundle.crt").write_text("LOCAL-MOCK-CA", encoding="utf-8")

    def register_job(self) -> str:
        job_id = uuid.uuid4().hex
        with self.lock:
            self.jobs[job_id] = Job(
                status="pending",
                filename=None,
                error=None,
                created_at=time.time(),
            )
            self.total_uploads += 1
        return job_id

    def update_job(self, job_id: str, **changes) -> None:
        # C3 (internal audit 2026-05-13): the cancellation path can pop
        # a job mid-flight; a late `update_job` from the stdout reader
        # thread then KeyError'd and killed that thread. Guard with
        # `.get` so a stale update is a silent no-op.
        with self.lock:
            job = self.jobs.get(job_id)
            if job is None:
                return
            self.jobs[job_id] = Job(**{**asdict(job), **changes})

    def get_job(self, job_id: str) -> Job | None:
        with self.lock:
            return self.jobs.get(job_id)

    def job_snapshot(self) -> dict[str, Job]:
        with self.lock:
            return dict(self.jobs)

    def cancel_job(self, job_id: str) -> tuple[bool, str]:
        """Terminate a running real-backend job.

        Returns (ok, message). The subprocess is killed (SIGTERM/TerminateProcess);
        the running stdout-reader thread observes the broken pipe, marks the
        job as ``status='cancelled'``, and exits. Idempotent — calling on an
        already-finished job returns (False, message).

        R-6 (2026-05-11): the status check + status flip happen under the
        same lock so two concurrent /cancel calls don't both observe
        ``pending`` and both call ``proc.kill()``. We still release the
        lock for the actual ``proc.kill()`` syscall because that can
        block briefly on Windows (TerminateProcess + handle settle), and
        we don't want to hold the launcher's only lock across a syscall.
        """
        with self.lock:
            job = self.jobs.get(job_id)
            if job is None:
                return False, "no such job"
            if job.status != "pending":
                return False, f"job already {job.status}"
            proc = self.job_procs.get(job_id)
            # Reserve the cancel by flipping status immediately, under
            # the lock, so a second concurrent caller sees
            # ``status == 'cancelled'`` and bails out at the line above.
            job.status = "cancelled"
            job.error  = "cancelled by user"

        if proc is None:
            # Mock-backend or pre-subprocess; status already flipped.
            return True, "cancelled (no subprocess to kill)"
        try:
            proc.kill()
        except Exception as exc:
            # Status is already 'cancelled' — log but don't roll back;
            # the subprocess may already have exited on its own.
            return False, f"kill failed: {exc}"
        return True, "cancelled"

    def cleanup_old_jobs(self, max_age_sec: int = 3600) -> int:
        """Remove finished jobs older than `max_age_sec`. Returns count removed."""
        now = time.time()
        removed = 0
        with self.lock:
            for jid in [
                j for j, job in self.jobs.items()
                if job.status in ("done", "error", "cancelled") and (now - job.created_at) > max_age_sec
            ]:
                del self.jobs[jid]
                self.job_procs.pop(jid, None)
                removed += 1
        return removed

    # ── 36-hour cache ────────────────────────────────────────────────────────

    def cache_lookup(self, key: str) -> dict | None:
        """Return the cache payload dict if a fresh entry exists, else None.

        Stale entries (older than CACHE_TTL_SEC) are evicted on access, as
        are entries whose `main_path` has disappeared from disk. Pre-phase-4
        entries (the legacy (timestamp, list-of-tuples) shape) are also
        evicted on access so the next upload re-runs the engine cleanly.
        """
        with self.lock:
            entry = self.cache.get(key)
        if not entry:
            return None
        ts, data = entry
        if not isinstance(data, dict):
            self._evict(key)
            return None
        if time.time() - ts > CACHE_TTL_SEC:
            self._evict(key)
            return None
        main_path = data.get("main_path")
        if not main_path or not Path(main_path).exists():
            self._evict(key)
            return None
        return data

    def cache_store(
        self,
        key: str,
        *,
        main_path: Path,
        source_file: Path | None = None,
        translation_array: list[str] | None = None,
        phrase_separator_table: list[str] | None = None,
        engine: str | None = None,
        ai_model: str | None = None,
        src_lang: str | None = None,
        dest_lang: str | None = None,
    ) -> None:
        """Persist a cached translation under `runtime/cache/<key>/`.

        Stores the engine's main output docx and a copy of the source
        upload (so a future request with a different splitter can apply
        it without needing the original upload to still be on disk).
        Optional translation arrays are kept for line-aware splitters and
        the line-count reconciler; either may be empty in early phases.
        """
        if not main_path.exists():
            return
        try:
            entry_dir = self.cache_dir / key
            entry_dir.mkdir(parents=True, exist_ok=True)
            dst_main = entry_dir / main_path.name
            if not dst_main.exists() or dst_main.stat().st_size != main_path.stat().st_size:
                shutil.copy2(main_path, dst_main)
            # 2026-05-13: also cache the JSON sidecar so a cache-hit
            # delivers the run log too. Old behaviour cached only the
            # docx, so re-using a cached translation gave the user a
            # docx without its accompanying _log.json.
            dst_log: Path | None = None
            sidecar_src = main_path.with_name(
                _re.sub(r"(?i)\.docx$", "_log.json", main_path.name)
            )
            if sidecar_src.exists():
                dst_log = entry_dir / sidecar_src.name
                if not dst_log.exists() or dst_log.stat().st_size != sidecar_src.stat().st_size:
                    shutil.copy2(sidecar_src, dst_log)
            dst_src: Path | None = None
            if source_file and source_file.exists():
                dst_src = entry_dir / "_source.docx"
                if not dst_src.exists() or dst_src.stat().st_size != source_file.stat().st_size:
                    shutil.copy2(source_file, dst_src)
            data: dict = {
                "main_path": dst_main,
                "log_path":  dst_log,
                "source_path": dst_src,
                "translation_array": list(translation_array or []),
                "phrase_separator_table": list(phrase_separator_table or []),
                "engine": engine,
                "ai_model": ai_model,
                "src_lang": src_lang,
                "dest_lang": dest_lang,
            }
            with self.lock:
                self.cache[key] = (time.time(), data)
        except Exception as exc:
            print(f"[cache] store failed for {key[:12]}…: {exc}")

    def _evict(self, key: str) -> None:
        """Remove cache entry + its on-disk copies. Quiet on missing files."""
        with self.lock:
            entry = self.cache.pop(key, None)
        if not entry:
            return
        entry_dir = self.cache_dir / key
        try:
            if entry_dir.exists():
                shutil.rmtree(entry_dir, ignore_errors=True)
        except Exception:
            pass

    def cleanup_stale_cache(self, max_age_sec: int = CACHE_TTL_SEC) -> int:
        """Evict cache entries older than `max_age_sec`. Returns count evicted."""
        now = time.time()
        with self.lock:
            stale = [k for k, (ts, _) in self.cache.items()
                     if now - ts > max_age_sec]
        for k in stale:
            self._evict(k)
        return len(stale)

    def start_cleanup_thread(self, interval_sec: int = 600, max_age_sec: int = 3600) -> None:
        """Spawn a daemon thread that periodically prunes finished jobs and
        stale cache entries (independent TTLs)."""
        def _loop():
            while True:
                time.sleep(interval_sec)
                try:
                    n = self.cleanup_old_jobs(max_age_sec=max_age_sec)
                    if n:
                        print(f"[cleanup] pruned {n} finished job(s) older than {max_age_sec}s")
                    c = self.cleanup_stale_cache()
                    if c:
                        print(f"[cleanup] evicted {c} cache entr(ies) older than {CACHE_TTL_SEC}s")
                except Exception as exc:
                    print(f"[cleanup] error: {exc}")
        threading.Thread(target=_loop, daemon=True, name="job-cleanup").start()

    # ── Weekly newsletter export to Telegram (2026-05-11) ────────────────
    #
    # Every Saturday at 12:00 in the operator's chosen timezone (default
    # `Europe/Paris`, overridable via MTD_SCHEDULER_TZ), the launcher
    # uploads `subscribers.txt` as a Telegram document to the same chat
    # ids configured for failure alerts (`MTD_TELEGRAM_CHAT_ID`,
    # comma-separated). Empty file → silent skip. Failure → set a
    # `pending_warning` flag in
    # `runtime_dir/subscribers_report_state.json`; the next launcher
    # boot reads that flag and emits a "[warn] subscribers report
    # failed last week (N emails pending)" line, then clears the flag.
    #
    # No Telegram token configured? The whole scheduler stays dormant —
    # the boot-time check just prints once and skips.

    def boot_subscribers_report_check(self) -> None:
        """Read the state file and print a warning if last week's report failed."""
        state_path = self.runtime_dir / "subscribers_report_state.json"
        try:
            if not state_path.is_file():
                return
            data = json.loads(state_path.read_text(encoding="utf-8"))
        except Exception:
            return
        if data.get("pending_warning"):
            count = data.get("last_run_count", 0)
            ts    = data.get("last_run_ts", "?")
            print(
                f"[subscribers] WARNING: last attempt at {ts} failed "
                f"({count} email(s) pending). Reason: {data.get('last_run_reason', 'unknown')}",
                file=sys.stderr, flush=True,
            )
            # Clear the flag so we don't nag every boot. The scheduler
            # will re-set it if the next Saturday attempt also fails.
            data["pending_warning"] = False
            try:
                state_path.write_text(
                    json.dumps(data, indent=2), encoding="utf-8"
                )
            except Exception:
                pass

    def start_subscribers_report_thread(self) -> None:
        """Daemon thread that fires `_run_subscribers_report` on the
        Saturday-noon-Europe schedule. The poll interval is 60 s; the
        thread sleeps 24 h after a successful (or empty-and-skipped)
        run to avoid re-firing in the same Saturday window.
        """
        token = os.environ.get("MTD_TELEGRAM_TOKEN", "").strip()
        if not token:
            # No Telegram configured — scheduler is a no-op. (The boot
            # check still runs if a state file exists from a previous
            # session when Telegram WAS configured.)
            print("[subscribers] Telegram not configured — weekly report disabled")
            return

        tz_name = os.environ.get("MTD_SCHEDULER_TZ", "Europe/Paris").strip() or "Europe/Paris"
        # `zoneinfo` is stdlib on Python 3.9+. If for some reason the
        # tz database is missing (very stripped-down installs), fall
        # back to UTC + warn.
        try:
            from zoneinfo import ZoneInfo
            tz = ZoneInfo(tz_name)
        except Exception as exc:
            print(f"[subscribers] zoneinfo({tz_name}) unavailable ({exc}); falling back to UTC")
            tz = None

        def _loop():
            import datetime as _dt
            last_fired_on: str = ""   # YYYY-MM-DD of last fire (in TZ)
            while True:
                try:
                    now = _dt.datetime.now(tz) if tz else _dt.datetime.utcnow()
                    # weekday: Mon=0 … Sat=5, Sun=6.
                    is_window = (
                        now.weekday() == 5 and
                        now.hour == 12 and
                        now.minute < 5   # 12:00..12:04 trigger window
                    )
                    today_key = now.strftime("%Y-%m-%d")
                    if is_window and last_fired_on != today_key:
                        last_fired_on = today_key
                        self._run_subscribers_report(now_iso=now.isoformat(timespec="seconds"))
                except Exception as exc:
                    print(f"[subscribers] scheduler tick error: {exc!r}",
                          file=sys.stderr, flush=True)
                time.sleep(60)

        threading.Thread(target=_loop, daemon=True, name="subscribers-report").start()
        print(f"[subscribers] scheduler armed (Sat 12:00 {tz_name})")

    def _run_subscribers_report(self, *, now_iso: str) -> None:
        """One attempt at the weekly Telegram upload.

        Empty `subscribers.txt` is treated as success-by-skipping (no
        notification, no error). Non-empty + Telegram success → state
        marked OK. Telegram failure → state.pending_warning = True so
        the next launcher boot surfaces it.
        """
        state_path = self.runtime_dir / "subscribers_report_state.json"
        # Read the subscribers file.
        if not SUBSCRIBERS_FILE.is_file():
            return
        try:
            content = SUBSCRIBERS_FILE.read_text(encoding="utf-8")
        except Exception as exc:
            print(f"[subscribers] could not read subscribers.txt: {exc!r}",
                  file=sys.stderr, flush=True)
            return
        emails = [ln.strip() for ln in content.splitlines() if ln.strip()]
        if not emails:
            # Silent skip per user instruction — empty list never alerts.
            return

        token   = os.environ.get("MTD_TELEGRAM_TOKEN", "").strip()
        chat_raw = os.environ.get("MTD_TELEGRAM_CHAT_ID", "").strip()
        if not (token and chat_raw):
            return
        chat_ids = _parse_telegram_chat_ids(chat_raw)
        if not chat_ids:
            return

        # Compose the message + send subscribers.txt as a document.
        text = (
            f"📬 *Newsletter snapshot*\n"
            f"• subscribers: *{len(emails)}*\n"
            f"• timestamp:   `{_telegram_escape(now_iso)}`\n"
            f"\n(File attached.)"
        )
        # Reuse the existing helpers via a tiny per-recipient loop.
        any_success = False
        last_error: str | None = None
        for chat_id in chat_ids:
            try:
                self._send_telegram_text(
                    token=token,
                    chat_id=chat_id,
                    text=text,
                )
                # Send the file separately so a 50-MB cap or a flaky
                # network on the attachment doesn't lose the text.
                self._telegram_send_document(
                    token=token,
                    chat_id=chat_id,
                    file_path=SUBSCRIBERS_FILE,
                    caption=f"subscribers.txt ({len(emails)} addresses)",
                )
                any_success = True
            except Exception as exc:
                last_error = repr(exc)
                print(f"[subscribers] send to {chat_id} failed: {exc!r}",
                      file=sys.stderr, flush=True)

        # Persist outcome to state file. The boot-time check reads
        # `pending_warning` and surfaces it on next launch.
        try:
            state_data = {
                "last_run_ts":     now_iso,
                "last_run_count":  len(emails),
                "last_run_result": "ok" if any_success else "failed",
                "last_run_reason": "" if any_success else (last_error or "unknown"),
                "pending_warning": not any_success,
            }
            state_path.write_text(
                json.dumps(state_data, indent=2), encoding="utf-8"
            )
            if any_success:
                print(f"[subscribers] weekly report sent ({len(emails)} email(s))")
            else:
                print(
                    f"[subscribers] weekly report FAILED ({len(emails)} email(s) pending)",
                    file=sys.stderr, flush=True,
                )
        except Exception as exc:
            print(f"[subscribers] could not write state file: {exc!r}",
                  file=sys.stderr, flush=True)

    def _send_telegram_text(
        self, *, token: str, chat_id: str, text: str,
    ) -> None:
        """Minimal text-only sendMessage helper used by the weekly
        report. Raises on any non-OK response so the per-recipient
        loop above can record `last_error`.
        """
        import urllib.request as _ur
        url = f"https://api.telegram.org/bot{token}/sendMessage"
        payload = {
            "chat_id":    chat_id,
            "text":       text,
            "parse_mode": "Markdown",
            "disable_web_page_preview": True,
        }
        req = _ur.Request(
            url,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with _ur.urlopen(req, timeout=10) as resp:
            data = resp.read()
            if b'"ok":true' not in data:
                raise RuntimeError(
                    f"telegram rejected sendMessage: {data[:200].decode(errors='replace')}"
                )


def _parse_multipart(headers, body: bytes) -> tuple[dict[str, str], dict[str, tuple[str, bytes]]]:
    content_type = headers.get("Content-Type", "")
    if "multipart/form-data" not in content_type:
        return {}, {}

    raw = (
        f"Content-Type: {content_type}\r\n"
        f"MIME-Version: 1.0\r\n\r\n"
    ).encode("utf-8") + body
    message = BytesParser(policy=email_policy).parsebytes(raw)

    fields: dict[str, str] = {}
    files: dict[str, tuple[str, bytes]] = {}

    for part in message.iter_parts():
        name = part.get_param("name", header="content-disposition")
        if not name:
            continue
        filename = part.get_filename()
        payload = part.get_payload(decode=True) or b""
        if filename:
            files[name] = (filename, payload)
        else:
            charset = part.get_content_charset() or "utf-8"
            fields[name] = payload.decode(charset, errors="replace")

    return fields, files


def _verify_password(candidate: str, stored_hash: str) -> bool:
    """Constant-time-ish password check for HTTP Basic auth.

    Supports bcrypt (``$2b$...``) and PBKDF2-SHA256 hashes written by
    ``scripts/setup_wizard.py``. Returns False on any error.
    """
    if not stored_hash or not candidate:
        return False
    try:
        if stored_hash.startswith("$2"):
            try:
                import bcrypt  # type: ignore
                return bool(bcrypt.checkpw(
                    candidate.encode("utf-8"),
                    stored_hash.encode("ascii"),
                ))
            except ImportError:
                return False
        if stored_hash.startswith("pbkdf2_sha256$"):
            import hashlib
            import hmac as _hmac
            _, iters_s, salt_hex, digest_hex = stored_hash.split("$", 3)
            iters = int(iters_s)
            salt = bytes.fromhex(salt_hex)
            expected = bytes.fromhex(digest_hex)
            actual = hashlib.pbkdf2_hmac(
                "sha256", candidate.encode("utf-8"), salt, iters
            )
            return _hmac.compare_digest(expected, actual)
    except Exception:
        return False
    return False


class MockTranslatorHandler(BaseHTTPRequestHandler):
    server_version = "LocalDocxTranslator/1.0"

    # feat/server-deploy (2026-05-14): public paths bypass HTTP Basic.
    _PUBLIC_PATHS = frozenset({"/health", "/favicon.ico"})
    _PUBLIC_PREFIXES = ("/static/",)

    @property
    def state(self) -> LocalState:
        return self.server.state  # type: ignore[attr-defined]

    @property
    def index_html(self) -> str:
        return self.server.index_html  # type: ignore[attr-defined]

    def log_message(self, fmt, *args):
        # Keep the console readable. The launcher prints its own status lines.
        print(f"[http] {self.address_string()} - {fmt % args}")

    def _is_public_path(self, path: str) -> bool:
        if path in self._PUBLIC_PATHS:
            return True
        return any(path.startswith(p) for p in self._PUBLIC_PREFIXES)

    def _check_auth(self, path: str) -> bool:
        """Return True if request may proceed; False if a 401 was sent."""
        if self._is_public_path(path):
            return True
        creds = getattr(self.server, "_mtd_auth", None) or {}
        if not creds.get("password_hash"):
            return True  # workstation mode — no auth configured
        import base64
        import hmac
        header = self.headers.get("Authorization", "")
        if not header.lower().startswith("basic "):
            self._send_auth_challenge()
            return False
        try:
            raw = base64.b64decode(header[6:].strip()).decode("utf-8")
            user, _sep, pwd = raw.partition(":")
        except Exception:
            self._send_auth_challenge()
            return False
        if (hmac.compare_digest(user, creds.get("username", ""))
                and _verify_password(pwd, creds.get("password_hash", ""))):
            return True
        self._send_auth_challenge()
        return False

    def _send_auth_challenge(self) -> None:
        self.send_response(HTTPStatus.UNAUTHORIZED)
        self.send_header("WWW-Authenticate", 'Basic realm="mtd"')
        self.send_header("Content-Type", "text/plain; charset=utf-8")
        self.end_headers()
        try:
            self.wfile.write(b"401 - authentication required\n")
        except Exception:
            pass

    def _handle_health(self) -> None:
        """Liveness probe — no auth. Used by Caddy / UptimeRobot."""
        import time as _time
        boot_ts = getattr(self.server, "_mtd_boot_time", _time.time())
        self._send_json({
            "status":  "ok",
            "version": getattr(self.server, "_mtd_version", "unknown"),
            "uptime":  int(_time.time() - boot_ts),
        })

    def _send_security_headers(self) -> None:
        """B14 (audit 2026-05-13): defence-in-depth response headers.

        These reduce the blast radius of a future XSS or supply-chain
        compromise. They do not fix any current vulnerability; they
        just prevent a hypothetical malicious script from exfiltrating
        upload names, framing the launcher inside an iframe, or
        guessing content-type to bypass MIME sniffing.
        """
        # The v2 SPA loads its CSS inline (Tailwind) and its JS from /v2/app.js
        # on the same origin. CDN audio files (Pixabay) play via the existing
        # frontend audio element — `media-src` opens that exact case. The
        # WebSocket-free architecture means no `connect-src` exemption needed
        # beyond same-origin.
        # B14 (audit 2026-05-13, revised 2026-05-13 evening):
        # the legacy `index.ejs` loads Tailwind CSS from `cdn.jsdelivr.net`
        # and audio cues from `cdn.pixabay.com`. The original CSP locked
        # style-src to 'self' and blocked Tailwind, which rendered the page
        # without any styles. Explicitly allow these two CDNs for the
        # respective resource types. Everything else (scripts, fetch,
        # frame embedding) stays locked down to 'self' / DENY.
        # Legacy index.ejs uses inline <script> blocks for the whole
        # form / engine / language wiring (~500 lines). The original
        # B14 CSP locked script-src to 'self' which silently disabled
        # every inline script — the language dropdowns rendered empty
        # because populateLanguage() never ran. Allow 'unsafe-inline'
        # for scripts (same as styles, also inline-heavy via Tailwind
        # utility tags). The remaining frame-ancestors / X-Frame-Options
        # / nosniff protections still apply.
        self.send_header(
            "Content-Security-Policy",
            "default-src 'self'; "
            "script-src 'self' 'unsafe-inline'; "
            "style-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net; "
            "img-src 'self' data:; "
            "media-src 'self' https://cdn.pixabay.com; "
            "font-src 'self' https://cdn.jsdelivr.net; "
            "connect-src 'self'; "
            "frame-ancestors 'none'",
        )
        self.send_header("X-Frame-Options", "DENY")
        self.send_header("X-Content-Type-Options", "nosniff")
        self.send_header("Referrer-Policy", "no-referrer")

    def _send_json(self, payload: dict, status: int = 200) -> None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        self._send_security_headers()
        self.end_headers()
        self.wfile.write(data)

    def _send_text(self, text: str, status: int = 200, content_type: str = "text/plain; charset=utf-8") -> None:
        data = text.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        self._send_security_headers()
        self.end_headers()
        self.wfile.write(data)

    def _send_file(self, file_path: Path, download_name: str) -> None:
        # A1 (2026-05-12): confine download to uploads_dir. Reject paths
        # that escape via .., absolute roots, or symlinks pointing outside
        # uploads_dir. Without this, /download/<name> would happily serve
        # any file the launcher process can read.
        try:
            uploads_root = self.state.uploads_dir.resolve()
            resolved = file_path.resolve()
            resolved.relative_to(uploads_root)
        except (ValueError, OSError):
            self._send_text("Not found", HTTPStatus.NOT_FOUND)
            return
        if not resolved.exists() or not resolved.is_file():
            self._send_text("Not found", HTTPStatus.NOT_FOUND)
            return

        file_path = resolved
        data = file_path.read_bytes()
        # 2026-05-13: Content-Type by extension. Previously hard-coded
        # to docx, which made the JSON sidecar (also served through
        # this endpoint) arrive with a Word MIME and tripped some
        # downloaders. Map the common cases; default to octet-stream.
        ext = file_path.suffix.lower()
        if ext == ".docx":
            ctype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext == ".json":
            ctype = "application/json; charset=utf-8"
        else:
            ctype = "application/octet-stream"
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", ctype)
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Content-Disposition", f'attachment; filename="{download_name}"')
        self.send_header("Cache-Control", "no-store")
        self._send_security_headers()
        self.end_headers()
        self.wfile.write(data)

    def _send_zip_for_job(self, job_id: str) -> None:
        """[RETIRED] Multi-file ZIP packaging.

        Phase 7 of the persian-double-lines roadmap collapsed the output
        from three files (TranslatePolish + Classic + Double) to one
        single docx per job, so there is nothing left to bundle. The
        /download-zip/ route is kept to avoid 404s for any client that
        still has the URL cached, but always responds 410 GONE.
        """
        self._send_text("ZIP download is no longer active.", HTTPStatus.GONE)

    def do_GET(self):
        parsed = urlparse(self.path)
        path = parsed.path

        # feat/server-deploy (2026-05-14): public health probe first,
        # auth gate next.
        if path == "/health":
            self._handle_health()
            return
        if not self._check_auth(path):
            return

        if path == "/":
            # 2026-05-13: re-read index.ejs from disk on every request
            # instead of relying on the in-memory copy captured at boot.
            # Otherwise a fix to index.ejs only takes effect after a
            # full launcher restart, which surprised the user when the
            # CDN→/static migration didn't appear in the browser.
            try:
                fresh_html = _inject_client_patch(_load_index_html())
            except Exception as exc:
                print(f"[WARN] /: re-reading index.ejs failed ({exc!r}); falling back to boot-cached copy.")
                fresh_html = self.index_html
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Cache-Control", "no-store, no-cache, must-revalidate, proxy-revalidate")
            self.send_header("Pragma", "no-cache")
            self.send_header("Expires", "0")
            self._send_security_headers()
            self.end_headers()
            self.wfile.write(fresh_html.encode("utf-8"))
            return

        if path == "/count":
            self._send_json({"count": str(_read_int(self.state.count_file, 0))})
            return

        if path == "/pricing":
            # U-2 (2026-05-11 audit): expose the OpenAI per-1M-token
            # pricing table so the v2 frontend can show a pre-flight
            # cost estimate. We surface only the models in
            # ``config.VALID_AI_MODELS`` to avoid advertising stale ids.
            #
            # Schema:
            #   {
            #     "models": [
            #       {"id": "gpt-5.5", "input": 5.00, "cached": 0.50, "output": 30.00},
            #       ...
            #     ],
            #     "currency": "USD",
            #     "unit": "per_1M_tokens"
            #   }
            #
            # The numbers come from the polisher's PRICES dict (it's the
            # most complete one and covers both whitelisted models). If
            # the polisher import fails for any reason, the launcher
            # falls back to an empty list — the frontend then degrades
            # to "estimate unavailable".
            try:
                _src_dir = ROOT / "src"
                if str(_src_dir) not in sys.path:
                    sys.path.insert(0, str(_src_dir))
                from machine_translate_docx.config import VALID_AI_MODELS as _VALID
                # Inline the same table the polisher uses — the polisher
                # method is per-instance, so we duplicate it as a
                # module-level constant here. If you bump the polisher
                # numbers, mirror them in this dict.
                _PRICES = {
                    "gpt-5.5":      {"input": 5.00, "cached": 0.50,  "output": 30.00},
                    "gpt-5.4-mini": {"input": 0.75, "cached": 0.075, "output": 4.50},
                }
                models = [
                    {"id": m, **_PRICES[m]}
                    for m in _VALID
                    if m in _PRICES
                ]
            except Exception as _exc:
                print(f"[/pricing] table unavailable: {_exc!r}", file=sys.stderr)
                models = []
            self._send_json({
                "models":   models,
                "currency": "USD",
                "unit":     "per_1M_tokens",
            })
            return

        if path == "/robotscount":
            jobs = self.state.job_snapshot()
            active = sum(1 for job in jobs.values() if job.status == "pending")
            self._send_json({"count": {"all": len(jobs), "user": active}})
            return

        if path.startswith("/status/"):
            job_id = path.removeprefix("/status/")
            job = self.state.get_job(job_id)
            if not job:
                self._send_json({"ok": False, "status": "not_found"}, HTTPStatus.NOT_FOUND)
                return
            payload: dict = {
                "ok": True,
                "status": job.status,
                "filename": job.filename,
                "error": job.error,
                "progress": job.progress,
            }
            self._send_json(payload)
            return

        if path.startswith("/download/"):
            file_name = unquote(path.removeprefix("/download/"))
            self._send_file(self.state.uploads_dir / file_name, file_name)
            return

        # 2026-05-13: /log/<filename> returns the JSON sidecar so the
        # legacy frontend can render a run-summary card after a job
        # finishes. Same path-traversal guard as /download.
        if path.startswith("/log/"):
            file_name = unquote(path.removeprefix("/log/"))
            try:
                uploads_root = self.state.uploads_dir.resolve()
                # Map the output docx name to its _log.json sibling.
                base_no_ext = _re.sub(r"(?i)\.docx$", "", file_name)
                log_path = (self.state.uploads_dir / f"{base_no_ext}_log.json").resolve()
                log_path.relative_to(uploads_root)
            except (ValueError, OSError):
                self._send_json({"ok": False, "error": "not found"}, HTTPStatus.NOT_FOUND)
                return
            if not log_path.exists():
                self._send_json({"ok": False, "error": "no sidecar"}, HTTPStatus.NOT_FOUND)
                return
            try:
                data = log_path.read_text(encoding="utf-8")
                self._send_json({"ok": True, "log": json.loads(data)})
            except Exception as exc:
                self._send_json({"ok": False, "error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR)
            return

        if path.startswith("/download-zip/"):
            # ZIP download is disabled — see _send_zip_for_job() for explanation.
            job_id = unquote(path.removeprefix("/download-zip/"))
            self._send_zip_for_job(job_id)
            return

        if path == "/robots.txt":
            self._send_text("User-agent: *\nDisallow:\n")
            return

        # ── /static/* — shared assets (e.g. Tailwind served locally) ─────────
        # 2026-05-13: legacy frontend used to pull Tailwind from
        # cdn.jsdelivr.net but some users see the CDN load fail
        # (Brave's blocker, corporate proxy, slow DNS). Ship a local
        # copy under web/static/ and route /static/* to it. Path
        # traversal blocked by resolve()+relative_to.
        if path.startswith("/static/"):
            relative = path.removeprefix("/static/").lstrip("/")
            if ".." in relative.split("/") or relative.startswith("/"):
                self._send_text("Not found", HTTPStatus.NOT_FOUND)
                return
            asset = (WEB_STATIC_DIR / relative).resolve()
            try:
                asset.relative_to(WEB_STATIC_DIR.resolve())
            except ValueError:
                self._send_text("Not found", HTTPStatus.NOT_FOUND)
                return
            if asset.is_file():
                ctype, _ = mimetypes.guess_type(str(asset))
                self.send_response(HTTPStatus.OK)
                self.send_header("Content-Type", ctype or "application/octet-stream")
                self.send_header("Content-Length", str(asset.stat().st_size))
                self.send_header("Cache-Control", "public, max-age=86400")
                self._send_security_headers()
                self.end_headers()
                self.wfile.write(asset.read_bytes())
                return
            self._send_text("Not found", HTTPStatus.NOT_FOUND)
            return

        # ── v2 frontend (Claude-inspired UI) ─────────────────────────────────
        # The legacy index.ejs continues to live at "/" untouched. v2 lives
        # alongside it under /v2 and shares the same backend endpoints.
        if path == "/v2" or path == "/v2/":
            html_file = WEB_V2_DIR / "index.html"
            if html_file.exists():
                data = html_file.read_text(encoding="utf-8")
                self.send_response(HTTPStatus.OK)
                self.send_header("Content-Type", "text/html; charset=utf-8")
                self.send_header("Cache-Control", "no-store")
                self._send_security_headers()
                self.end_headers()
                self.wfile.write(data.encode("utf-8"))
            else:
                self._send_text("v2 frontend not deployed", HTTPStatus.NOT_FOUND)
            return

        if path.startswith("/v2/"):
            relative = path.removeprefix("/v2/").lstrip("/")
            # Reject path traversal and only serve files under web/v2/.
            if ".." in relative.split("/") or relative.startswith("/"):
                self._send_text("Not found", HTTPStatus.NOT_FOUND)
                return
            asset = (WEB_V2_DIR / relative).resolve()
            try:
                asset.relative_to(WEB_V2_DIR.resolve())
            except ValueError:
                self._send_text("Not found", HTTPStatus.NOT_FOUND)
                return
            if asset.is_file():
                ctype, _ = mimetypes.guess_type(str(asset))
                self.send_response(HTTPStatus.OK)
                self.send_header("Content-Type", ctype or "application/octet-stream")
                self.send_header("Content-Length", str(asset.stat().st_size))
                self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
                self._send_security_headers()
                self.end_headers()
                self.wfile.write(asset.read_bytes())
                return
            self._send_text("Not found", HTTPStatus.NOT_FOUND)
            return

        self._send_text("Not found", HTTPStatus.NOT_FOUND)

    def do_POST(self):
        parsed = urlparse(self.path)

        # feat/server-deploy (2026-05-14): all POSTs require auth.
        if not self._check_auth(parsed.path):
            return

        # Cancel a running job. The frontend Cancel button hits this.
        # We kill the subprocess; the job thread observes the broken
        # pipe and exits without overwriting the cancelled status.
        if parsed.path.startswith("/cancel/"):
            job_id = parsed.path.removeprefix("/cancel/").strip("/")
            ok, msg = self.state.cancel_job(job_id)
            self._send_json(
                {"ok": ok, "message": msg},
                HTTPStatus.OK if ok else HTTPStatus.BAD_REQUEST,
            )
            return

        # Newsletter subscribe endpoint — accepts JSON {"email": "..."} OR
        # multipart/form-data with field name "email". Used by the v2 UI
        # but also accessible from any client.
        if parsed.path == "/subscribe":
            content_length = int(self.headers.get("Content-Length", "0"))
            body = self.rfile.read(content_length)
            ctype = (self.headers.get("Content-Type") or "").lower()
            email = ""
            if "application/json" in ctype:
                try:
                    email = (json.loads(body or b"{}").get("email") or "").strip()
                except Exception:
                    email = ""
            elif "multipart/form-data" in ctype:
                fields, _ = _parse_multipart(self.headers, body)
                email = fields.get("email", "").strip()
            else:
                # urlencoded
                from urllib.parse import parse_qs
                qs = parse_qs(body.decode("utf-8", errors="replace"))
                email = (qs.get("email", [""])[0] or "").strip()
            ok, msg = _append_subscriber(email)
            self._send_json(
                {"ok": ok, "message": msg},
                HTTPStatus.OK if ok else HTTPStatus.BAD_REQUEST,
            )
            return

        if parsed.path != "/upload":
            self._send_text("Not found", HTTPStatus.NOT_FOUND)
            return

        # A7 (2026-05-12): reject oversize uploads *before* reading the
        # body into memory. The DOCX inner-validation already catches
        # zip-bombs (uncompressed cap), but without a compressed-size cap
        # we still read N MB into RAM on every request, which is a cheap
        # DoS surface. The cap matches the smaller of the two limits.
        try:
            content_length = int(self.headers.get("Content-Length", "0"))
        except (TypeError, ValueError):
            content_length = -1
        if content_length < 0:
            self._send_json(
                {"ok": False, "comment": "Missing or invalid Content-Length."},
                HTTPStatus.LENGTH_REQUIRED,
            )
            return
        if content_length > _MAX_DOCX_COMPRESSED:
            self._send_json(
                {"ok": False, "comment": (
                    f"Upload exceeds the "
                    f"{_MAX_DOCX_COMPRESSED // (1024 * 1024)} MB limit."
                )},
                HTTPStatus.REQUEST_ENTITY_TOO_LARGE,
            )
            return
        body = self.rfile.read(content_length)
        fields, files = _parse_multipart(self.headers, body)

        uploaded = files.get("file")
        if not uploaded:
            self._send_json({"ok": False, "comment": "Please upload a DOCX file."}, HTTPStatus.BAD_REQUEST)
            return

        original_name, payload = uploaded

        # Server-side validation runs *before* writing to disk.
        # Magic bytes + zip-bomb cap protect against malicious / malformed uploads
        # even when the client-side check in index.ejs is bypassed.
        validation_error = _validate_docx_payload(payload)
        if validation_error:
            self._send_json(
                {"ok": False, "comment": validation_error},
                HTTPStatus.BAD_REQUEST,
            )
            return

        safe_name = _sanitize_filename(original_name)
        saved_name = f"{int(time.time() * 1000)}-{safe_name}"
        upload_path = self.state.uploads_dir / saved_name
        upload_path.write_bytes(payload)

        # 2026-05-13: optional XLSX search-and-replace file. When the
        # legacy frontend sends a second file under the name "xlsxFile",
        # save it next to the docx and pass --xlsxreplacefile to the
        # CLI. Missing → CLI falls back to the bundled default.
        xlsx_uploaded = files.get("xlsxFile")
        xlsx_path: Path | None = None
        if xlsx_uploaded:
            xlsx_orig_name, xlsx_payload = xlsx_uploaded
            if xlsx_payload and xlsx_orig_name.lower().endswith(".xlsx"):
                xlsx_safe = _sanitize_filename(xlsx_orig_name)
                xlsx_saved = f"{int(time.time() * 1000)}-{xlsx_safe}"
                xlsx_path = self.state.uploads_dir / xlsx_saved
                xlsx_path.write_bytes(xlsx_payload)

        job_id = self.state.register_job()

        source_language = fields.get("sourceLanguage", "Auto")
        target_language = fields.get("targetLanguage", "en")
        translation_engine = fields.get("translationEngine", "google")
        split_engine = fields.get("splitEngine")
        ai_model = fields.get("aiModel")
        # Aligner LLM threshold (legacy frontend slider, 0..100). Default 0
        # = mechanical-only (matches today's aligner behaviour). The CLI
        # passes this through to FASubtitleAligner where it is currently
        # a no-op; it becomes meaningful when a hybrid aligner is wired.
        try:
            aligner_llm_threshold = int(fields.get("alignerLlmThreshold", "0"))
        except (TypeError, ValueError):
            aligner_llm_threshold = 0
        aligner_llm_threshold = max(0, min(100, aligner_llm_threshold))

        # 2026-05-15 — Persian Double Lines MAX_CHARS override. Default 48
        # (the historical broadcast-CPL constant). UI exposes 24..70 with
        # a stored preference in localStorage. Clamp here too so a hand-
        # crafted request cannot drive the aligner outside the safe band.
        try:
            aligner_max_chars = int(fields.get("alignerMaxChars", "48"))
        except (TypeError, ValueError):
            aligner_max_chars = 48
        aligner_max_chars = max(24, min(70, aligner_max_chars))
        enable_sound = fields.get("enableSound")
        sound_select = fields.get("soundSelect")
        split_translate = fields.get("splitTranslate", "false").lower() in {"true", "1", "on", "yes"}

        # ── 36-hour cache short-circuit ──────────────────────────────────────
        # Only API engines cache (chatgpt, chatgpt-polish). Selenium engines
        # are stateful (cookie consent, login) and not worth caching.
        cache_key: str | None = None
        if translation_engine.lower() in _API_ENGINES:
            cache_key = _cache_key(
                payload, target_language, translation_engine, ai_model,
                split_engine,
            )
            cached = self.state.cache_lookup(cache_key)
            if cached:
                # Materialise the cached engine output back into uploads/
                # and apply the requested Split Method on top. A different
                # splitter than was used last time is the whole point of
                # the phase-4 dict-shape cache.
                served = self._materialise_cached_output(
                    cached,
                    split_engine=split_engine,
                    target_language=target_language,
                    aligner_max_chars=aligner_max_chars,
                )
                if served:
                    splitter_only = served.name.endswith("_Double_Lines.docx")
                    self.state.update_job(
                        job_id, status="done",
                        filename=served.name,
                        progress=100, error=None,
                    )
                    print(
                        f"[cache hit] {cache_key[:12]}… reused; "
                        f"split={split_engine or 'none'} splitter_only={splitter_only}"
                    )
                    self._send_json({
                        "ok": True, "jobId": job_id,
                        "cacheHit": True, "splitterOnly": splitter_only,
                    })
                    return

        print(f"[job {job_id}] upload={saved_name}")
        print(f"[job {job_id}] source={source_language} target={target_language} engine={translation_engine} split={split_translate}")
        print(f"[job {job_id}] splitEngine={split_engine or '(not sent)'} aiModel={ai_model or '(not sent)'} enableSound={enable_sound or '(not sent)'} soundSelect={sound_select or '(not sent)'}")

        thread = threading.Thread(
            target=self._process_job,
            args=(
                job_id,
                upload_path,
                safe_name,
                source_language,
                target_language,
                translation_engine,
                split_translate,
                split_engine,
                ai_model,
                enable_sound,
                sound_select,
                cache_key,
                aligner_llm_threshold,
                xlsx_path,
                aligner_max_chars,
            ),
            daemon=True,
        )
        thread.start()

        self._send_json({"ok": True, "jobId": job_id, "cacheHit": False})

    def _process_job(
        self,
        job_id: str,
        source_file: Path,
        original_name: str,
        source_language: str,
        target_language: str,
        translation_engine: str,
        split_translate: bool,
        split_engine: str | None,
        ai_model: str | None,
        enable_sound: str | None,
        sound_select: str | None,
        cache_key: str | None = None,
        aligner_llm_threshold: int = 0,
        xlsx_path: Path | None = None,
        aligner_max_chars: int = 48,
    ) -> None:
        _job_t0 = time.time()
        print(f"[job {job_id}] ▶ start — file: {original_name} | lang: {target_language} | engine: {translation_engine}")
        self.state.update_job(job_id, progress=5)
        # Limit concurrent backend subprocesses to keep memory bounded.
        # When the cap is reached, a job waits here (status remains 'pending'
        # so the frontend keeps polling) until a slot frees up.
        _job_semaphore.acquire()
        self.state.update_job(job_id, progress=10)
        try:
            if self.state.backend_mode == "mock":
                time.sleep(1.2)

                stem = Path(original_name).stem
                suffix = _lang_suffix(target_language) or "OUT"
                output_name = f"{stem}_{suffix}.docx"
                output_path = self.state.uploads_dir / output_name
                _build_placeholder_docx(
                    output_path,
                    job_id=job_id,
                    original_name=original_name,
                    source_file=source_file.name,
                    source_language=source_language,
                    target_language=target_language,
                    translation_engine=translation_engine,
                    split_translate=split_translate,
                    split_engine=split_engine,
                    ai_model=ai_model,
                    enable_sound=enable_sound,
                    sound_select=sound_select,
                )

                count = _read_int(self.state.count_file, 0) + 1
                _write_int(self.state.count_file, count)

                self.state.update_job(job_id, status="done", filename=output_name, error=None)
                print(f"[job {job_id}] done -> {output_name} (placeholder DOCX)")
                return

            output_path = self._run_real_backend(
                job_id=job_id,
                source_file=source_file,
                target_language=target_language,
                translation_engine=translation_engine,
                split_translate=split_translate,
                split_engine=split_engine,
                source_language=source_language,
                ai_model=ai_model,
                aligner_llm_threshold=aligner_llm_threshold,
                xlsx_path=xlsx_path,
            )

            count = _read_int(self.state.count_file, 0) + 1
            _write_int(self.state.count_file, count)

            # Apply the requested Split Method on top of the engine's
            # raw translated output. For Persian Double Lines this runs
            # the FA mechanical aligner in-process; for any other splitter
            # (or none) the engine output is served unchanged.
            served_path = self._apply_splitter(
                output_path,
                split_engine=split_engine,
                target_language=target_language,
                aligner_max_chars=aligner_max_chars,
            )

            self.state.update_job(
                job_id,
                status="done",
                filename=served_path.name,
                error=None,
            )

            # Cache outputs for 36 hours (API engines only — `cache_key` is
            # only set for them in do_POST). The cache stores the engine's
            # *raw* translated docx (no splitter applied) so a re-upload
            # with a different Split Method can reuse it.
            if cache_key:
                self.state.cache_store(
                    cache_key,
                    main_path=output_path,
                    source_file=source_file,
                    engine=translation_engine,
                    ai_model=ai_model,
                    src_lang=source_language,
                    dest_lang=target_language,
                )

            _job_elapsed = time.time() - _job_t0
            print(f"[job {job_id}] ✓ done in {_job_elapsed:.0f}s -> {served_path.name}")
        except Exception as exc:
            _job_elapsed = time.time() - _job_t0
            # If the user already cancelled this job, the subprocess kill
            # raised inside us as a non-zero exit. Don't overwrite the
            # 'cancelled' status with 'error' in that case.
            cur = self.state.get_job(job_id)
            if cur and cur.status == "cancelled":
                print(f"[job {job_id}] ✗ cancelled by user after {_job_elapsed:.0f}s")
            else:
                self.state.update_job(job_id, status="error", filename=None, error=str(exc))
                print(f"[job {job_id}] ✗ error after {_job_elapsed:.0f}s: {exc}")
        finally:
            _job_semaphore.release()

    def _map_engine(self, translation_engine: str) -> tuple[str, list[str]]:
        engine = translation_engine.lower().strip()
        extra: list[str] = []

        if engine == "chatgpt-polish":
            engine = "chatgpt"
            extra.extend(["--enginemethod", "api", "--with-polish"])
        elif engine == "chatgpt":
            extra.extend(["--enginemethod", "api"])

        return engine, extra

    def _run_real_backend(
        self,
        *,
        job_id: str,
        source_file: Path,
        target_language: str,
        translation_engine: str,
        split_translate: bool,
        split_engine: str | None,
        source_language: str,
        ai_model: str | None,
        aligner_llm_threshold: int = 0,
        xlsx_path: Path | None = None,
    ) -> Path:
        engine, extra_flags = self._map_engine(translation_engine)

        # B1-guard (2026-05-12 revised): only the persian_double_lines
        # split path runs the FA aligner, and that path does its OWN
        # row distribution on the saved docx. For the basic split path
        # the legacy line-by-line splitter (document_split_phrases) is
        # the only thing that fans the single-call FA output back across
        # the cell column — without `--split` the entire polished
        # translation lands in row 1 and every following row stays
        # empty (bug observed on News Scroll NS 3145, 2026-05-12).
        # So: force off ONLY for persian_double_lines, leave basic alone.
        if (
            translation_engine == "chatgpt-polish"
            and target_language.lower().startswith("fa")
            and split_engine == "persian_double_lines"
        ):
            split_translate = False

        # 2026-05-11 src-layout migration: invoke the CLI as a *module*
        # (`python -m machine_translate_docx.cli`) instead of a script
        # path. The new package uses relative imports — running it as a
        # script (`python /…/cli.py`) makes Python treat cli.py as the
        # top-level module and every `from .config import …` fails. The
        # `-m` form gives Python the right package context.
        # PYTHONPATH below points at ``src/`` so the package is findable
        # without `pip install -e .`.
        is_pkg_layout = self.state.script_path.parent.name == "machine_translate_docx"
        if is_pkg_layout:
            cmd = [
                str(self.state.python_exe),
                "-m", "machine_translate_docx.cli",
                "--docxfile", str(source_file),
                "--destlang", target_language,
                "--silent", "--exitonsuccess",
                "--engine", engine,
            ]
        else:
            # Legacy fallback for pre-migration checkouts.
            cmd = [
                str(self.state.python_exe),
                str(self.state.script_path),
                "--docxfile", str(source_file),
                "--destlang", target_language,
                "--silent", "--exitonsuccess",
                "--engine", engine,
            ]

        if ai_model:
            cmd.extend(["--aimodel", ai_model])
        else:
            cmd.extend(["--aimodel", "gpt-5.5"])

        if source_language and source_language.lower() != "auto":
            cmd.extend(["--srclang", source_language])
        if split_translate:
            cmd.append("--split")
        if split_engine in ("openai", "persian_double_lines"):
            cmd.extend(["--splitengine", split_engine])
        if split_engine == "persian_double_lines" and isinstance(aligner_llm_threshold, int):
            cmd.extend(["--alignerllmthreshold", str(aligner_llm_threshold)])
        if xlsx_path is not None and xlsx_path.exists():
            cmd.extend(["--xlsxreplacefile", str(xlsx_path)])
        if engine == "google":
            cmd.append("--showbrowser")

        cmd.extend(extra_flags)

        print(f"[job {job_id}] running real backend via: {self.state.python_exe}")
        print(f"[job {job_id}] command: {' '.join(cmd)}")

        # bufsize=1 → line-buffered. Without this, Python pipes default to
        # full-buffering (~8 KB), so PROGRESS:N markers emitted by the
        # backend can be held back for many seconds before the launcher
        # sees them — making the UI bar jump from 10 % straight to 100 %.
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
            encoding="utf-8",
            errors="replace",
            cwd=str(ROOT),
            env={
                **{k: v for k, v in os.environ.items() if k.upper() not in {
                    "HTTP_PROXY",
                    "HTTPS_PROXY",
                    "ALL_PROXY",
                    "NO_PROXY",
                    "http_proxy",
                    "https_proxy",
                    "all_proxy",
                    "no_proxy",
                }},
                # Force UTF-8 stdout/stderr in the subprocess so aligner
                # print() calls containing non-ASCII chars (e.g. arrows, dashes)
                # don't crash with UnicodeEncodeError on Windows CP1252 consoles.
                "PYTHONIOENCODING": "utf-8",
                "PYTHONUTF8": "1",
                # Point `python -m machine_translate_docx.cli` at the
                # in-repo package without requiring `pip install -e .`.
                "PYTHONPATH": str(ROOT / "src"),
            },
        )

        # Register the subprocess so POST /cancel/<id> can kill it.
        with self.state.lock:
            self.state.job_procs[job_id] = proc

        saved_filename: str | None = None
        # B-001 (2026-05-10): backend now emits a structured
        # `[FAIL] reason=<token> message=<text>` line when it detects an
        # empty docx or an empty engine return. Capture them here so we
        # can attach a human-readable reason to the job error and so
        # the B-002 archive hook can stamp the same tokens into the
        # failure folder's meta.json.
        captured_log_buf: list[str] = []
        captured_fail: tuple[str, str] | None = None
        # R-2 (2026-05-11): wrap the reader in try/finally so that if an
        # exception escapes the loop body (e.g. a sudden disk-full while
        # appending to captured_log_buf), we still drain the rest of
        # stdout and close the pipe — otherwise the child blocks on a
        # full pipe and stays alive after the parent has moved on.
        try:
            if proc.stdout is not None:
                for line in proc.stdout:
                    stripped = line.rstrip()
                    # Parse PROGRESS:N lines into job.progress without printing
                    # them as noise — they are tiny status pings, not log content.
                    if stripped.startswith("PROGRESS:"):
                        try:
                            pct = int(stripped.split(":", 1)[1].strip())
                            pct = max(0, min(100, pct))
                            self.state.update_job(job_id, progress=pct)
                        except ValueError:
                            pass
                        continue
                    if stripped:
                        print(f"[job {job_id}] {stripped}")
                        # Keep a bounded copy for B-002 archival.
                        captured_log_buf.append(stripped)
                        if len(captured_log_buf) > 5000:
                            # Trim to most-recent 5000 lines so we cap memory
                            # for runaway logs while preserving the tail.
                            del captured_log_buf[:1000]
                    if "Saved file name:" in stripped:
                        _, _, part = stripped.partition("Saved file name:")
                        candidate = part.strip()
                        if candidate:
                            saved_filename = candidate
                    if stripped.startswith("[FAIL] reason="):
                        # Format: [FAIL] reason=<token> message=<text>
                        body = stripped[len("[FAIL] "):]
                        reason_token = ""
                        message = ""
                        for kv in body.split(" message=", 1):
                            if kv.startswith("reason="):
                                reason_token = kv[len("reason="):].strip()
                                break
                        if " message=" in body:
                            message = body.split(" message=", 1)[1].strip()
                        captured_fail = (reason_token or "translation_failure", message or stripped)
        finally:
            # Drain anything left in the pipe so the child can exit, then
            # close it. Without this, a parse exception mid-loop would
            # leak the file descriptor + leave the child blocked.
            if proc.stdout is not None:
                try:
                    for _ in proc.stdout:
                        pass
                except Exception:
                    pass
                try:
                    proc.stdout.close()
                except Exception:
                    pass

        code = proc.wait()
        # R-1 (2026-05-11): subprocess is finished — drop the entry from
        # `job_procs` so the dict doesn't grow unbounded across long
        # sessions. The cleanup thread still prunes finished jobs at
        # 1-hour granularity, but doing it here means the OS file handle
        # held by the Popen object is freed immediately on completion.
        with self.state.lock:
            self.state.job_procs.pop(job_id, None)
        if captured_fail is not None:
            reason_token, message = captured_fail
            self._archive_failed_job(
                job_id=job_id,
                source_file=source_file,
                reason=reason_token,
                message=message,
                stdout_lines=captured_log_buf,
                target_language=target_language,
                translation_engine=translation_engine,
            )
            raise RuntimeError(f"{reason_token}: {message}")
        if code != 0:
            self._archive_failed_job(
                job_id=job_id,
                source_file=source_file,
                reason="backend_nonzero_exit",
                message=f"Backend exited with code {code}",
                stdout_lines=captured_log_buf,
                target_language=target_language,
                translation_engine=translation_engine,
            )
            raise RuntimeError(f"Backend exited with code {code}")

        output_path = Path(saved_filename) if saved_filename else self._fallback_output_path(source_file, target_language, translation_engine)
        deadline = time.time() + 120
        while time.time() < deadline:
            if output_path.exists():
                output_path = self._strip_timestamp(output_path)
                return output_path
            time.sleep(0.5)

        raise FileNotFoundError(f"Output file not found: {output_path}")

    # ── B-002 (2026-05-10) — failure archive + alerting ──────────────────────
    #
    # When a job ends in error we keep an on-disk post-mortem so the
    # operator can triage what went wrong without scrolling stdout. The
    # archive lives at:
    #
    #     runtime_dir/failures/<job_id>__<UTC iso ts>/
    #         input.docx     — the original upload
    #         stdout.log     — stdout / stderr captured during the run
    #         meta.json      — job_id, lang, engine, model, reason, ts
    #         UNREVIEWED.txt — sentinel (deleted when an operator opens
    #                          the folder; lets `ls` show pending issues)
    #
    # Two optional alerting hooks fire if the corresponding env var is
    # set; both are explicitly cheap-and-free:
    #
    #     MTD_FAILURE_EMAIL=op@example.com
    #         Sends a plain-text email via smtplib using
    #         MTD_SMTP_HOST / MTD_SMTP_PORT / MTD_SMTP_USER /
    #         MTD_SMTP_PASS / MTD_SMTP_FROM. No third-party dep.
    #
    #     MTD_FAILURE_WEBHOOK=https://discord.com/api/webhooks/...
    #         POSTs JSON to a Discord/Slack/Mattermost incoming
    #         webhook. Works with any service that takes
    #         {"content": "<text>"} (Discord shape) — others can map
    #         in a tiny shim if they prefer different keys.
    #
    # Failures are *never* fatal — a flaky email server should never
    # block the launcher from continuing to serve other jobs.

    def _archive_failed_job(
        self,
        *,
        job_id: str,
        source_file: Path,
        reason: str,
        message: str,
        stdout_lines: list[str],
        target_language: str,
        translation_engine: str,
    ) -> None:
        """Copy input + stdout + meta into runtime_dir/failures/<id>__<ts>/.

        Best-effort. Errors are swallowed and printed so a misconfigured
        runtime_dir cannot kill the launcher.

        R-3 (2026-05-11 audit): if the primary location is unwritable
        (read-only volume, no permissions, disk full), fall back to a
        temp dir under the OS-default tempdir so the operator at least
        has SOMETHING to triage. The traceback is printed in both cases
        so the failure is not invisible.
        """
        import datetime as _dt
        import tempfile as _tf
        import traceback as _tb

        ts = _dt.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
        primary = self.state.runtime_dir / "failures" / f"{job_id}__{ts}"
        base = None
        try:
            primary.mkdir(parents=True, exist_ok=True)
            base = primary
        except Exception as _e:
            print(
                f"[B-002] primary failures dir {primary} unwritable "
                f"({_e!r}); falling back to system tempdir",
                file=sys.stderr,
                flush=True,
            )
            try:
                fallback_root = Path(_tf.mkdtemp(prefix="mtd-failure-"))
                base = fallback_root / f"{job_id}__{ts}"
                base.mkdir(parents=True, exist_ok=True)
            except Exception as _e2:
                print(
                    f"[B-002] fallback tempdir also unwritable "
                    f"({_e2!r}); archive skipped\n{_tb.format_exc()}",
                    file=sys.stderr,
                    flush=True,
                )
                return

        try:

            # Input file
            try:
                shutil.copy2(source_file, base / "input.docx")
            except Exception as _e:
                print(f"[B-002] failed to archive input.docx: {_e}")

            # stdout/stderr capture
            try:
                (base / "stdout.log").write_text(
                    "\n".join(stdout_lines), encoding="utf-8"
                )
            except Exception as _e:
                print(f"[B-002] failed to write stdout.log: {_e}")

            # meta.json
            try:
                meta = {
                    "job_id":            job_id,
                    "timestamp_utc":     ts,
                    "reason":            reason,
                    "message":           message,
                    "target_language":   target_language,
                    "translation_engine": translation_engine,
                    "input_filename":    source_file.name,
                }
                (base / "meta.json").write_text(
                    json.dumps(meta, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
            except Exception as _e:
                print(f"[B-002] failed to write meta.json: {_e}")

            # Sentinel — deletable by an operator after triage so `ls`
            # shows what is still unreviewed.
            try:
                (base / "UNREVIEWED.txt").write_text(
                    "Delete this file when the failure has been triaged.\n",
                    encoding="utf-8",
                )
            except Exception:
                pass

            print(f"[B-002] archived failed job to {base}")

            # Best-effort alerting.
            self._alert_failure(
                job_id=job_id,
                reason=reason,
                message=message,
                meta=meta,
                archive_dir=base,
            )
        except Exception as exc:
            # R-3: surface the full traceback so the operator can see
            # exactly which write step failed. The single-line print was
            # too cryptic — a permission error inside meta.json's open()
            # would surface as just `archive skipped: [Errno 13]` with
            # no path or call site.
            import traceback as _tb
            print(
                f"[B-002] archive skipped (unexpected error): {exc!r}\n"
                f"{_tb.format_exc()}",
                file=sys.stderr,
                flush=True,
            )

    def _alert_failure(
        self,
        *,
        job_id: str,
        reason: str,
        message: str,
        meta: dict,
        archive_dir: Path,
    ) -> None:
        """Fire optional email + webhook alerts. Both are env-gated."""
        # ── email (smtplib) ───────────────────────────────────────────
        email_to = os.environ.get("MTD_FAILURE_EMAIL", "").strip()
        if email_to:
            try:
                import smtplib
                from email.message import EmailMessage

                host = os.environ.get("MTD_SMTP_HOST", "localhost")
                port = int(os.environ.get("MTD_SMTP_PORT", "25"))
                user = os.environ.get("MTD_SMTP_USER", "")
                pwd  = os.environ.get("MTD_SMTP_PASS", "")
                sender = os.environ.get(
                    "MTD_SMTP_FROM",
                    f"machine-translate-docx@{host}",
                )

                em = EmailMessage()
                em["From"] = sender
                em["To"]   = email_to
                em["Subject"] = (
                    f"[machine-translate-docx] FAILED job {job_id} ({reason})"
                )
                body_lines = [
                    f"Job:    {job_id}",
                    f"Reason: {reason}",
                    f"Lang:   {meta.get('target_language')}",
                    f"Engine: {meta.get('translation_engine')}",
                    f"Input:  {meta.get('input_filename')}",
                    f"Folder: {archive_dir}",
                    "",
                    f"Message: {message}",
                    "",
                    "(See the archive folder for input.docx, stdout.log, meta.json.)",
                ]
                em.set_content("\n".join(body_lines))

                with smtplib.SMTP(host, port, timeout=10) as smtp:
                    if user and pwd:
                        try:
                            smtp.starttls()
                        except Exception:
                            pass
                        smtp.login(user, pwd)
                    smtp.send_message(em)
                print(f"[B-002] alert email sent to {email_to}")
            except Exception as exc:
                print(f"[B-002] alert email skipped: {exc}")

        # ── webhook (Discord/Slack/Mattermost shape) ──────────────────
        webhook_url = os.environ.get("MTD_FAILURE_WEBHOOK", "").strip()
        if webhook_url:
            try:
                import urllib.request
                payload = {
                    "content": (
                        f":x: machine-translate-docx FAILED job `{job_id}` "
                        f"({reason}) — engine `{meta.get('translation_engine')}` "
                        f"lang `{meta.get('target_language')}` — "
                        f"`{archive_dir}`"
                    ),
                }
                req = urllib.request.Request(
                    webhook_url,
                    data=json.dumps(payload).encode("utf-8"),
                    headers={"Content-Type": "application/json"},
                    method="POST",
                )
                with urllib.request.urlopen(req, timeout=10):
                    pass
                print("[B-002] alert webhook posted")
            except Exception as exc:
                print(f"[B-002] alert webhook skipped: {exc}")

        # ── Telegram bot (env-gated) ──────────────────────────────────
        # Two env vars required:
        #   MTD_TELEGRAM_TOKEN     — bot token from @BotFather
        #   MTD_TELEGRAM_CHAT_ID   — one OR MORE chat ids, comma- /
        #                            whitespace-separated. Each id can
        #                            be a numeric DM id (e.g.
        #                            `987654321`), a group id (e.g.
        #                            `-987654321`), or a public channel
        #                            handle (e.g. `@my_alerts_channel`).
        #
        # See docs/telegram-alerts-setup.md for the @BotFather walkthrough,
        # the @userinfobot shortcut for chat_ids, and the multi-recipient
        # options (multi-DM / private group / channel).
        #
        # Best-effort: any exception per recipient is logged and swallowed
        # so a flaky network or one bad chat_id cannot block the rest of
        # the alert fan-out or the failure-archive path.
        tg_token     = os.environ.get("MTD_TELEGRAM_TOKEN", "").strip()
        tg_chat_raw  = os.environ.get("MTD_TELEGRAM_CHAT_ID", "").strip()
        if tg_token and tg_chat_raw:
            chat_ids = _parse_telegram_chat_ids(tg_chat_raw)
            for chat_id in chat_ids:
                try:
                    self._send_telegram_alert(
                        token=tg_token,
                        chat_id=chat_id,
                        job_id=job_id,
                        reason=reason,
                        message=message,
                        meta=meta,
                        archive_dir=archive_dir,
                    )
                except Exception as exc:
                    # Should not happen — _send_telegram_alert already
                    # swallows internally. Belt-and-suspenders so a
                    # malformed chat_id can't kill the loop.
                    print(f"[telegram] recipient {chat_id} skipped: {exc!r}",
                          file=sys.stderr, flush=True)

    def _send_telegram_alert(
        self,
        *,
        token: str,
        chat_id: str,
        job_id: str,
        reason: str,
        message: str,
        meta: dict,
        archive_dir: Path,
    ) -> None:
        """POST a failure alert to Telegram (text + optional docx).

        The text alert always fires. The docx attachment is sent
        separately so a too-large or missing file does NOT prevent the
        text from getting through. Set ``MTD_TELEGRAM_NO_ATTACHMENT=1``
        to suppress the docx upload entirely (useful when the docs
        contain sensitive content and you only want a heads-up).

        Token + chat_id come from the caller, not from environment, so
        the helper is unit-testable.
        """
        import urllib.request as _ur
        import urllib.error   as _ue

        text = (
            "❌ *machine-translate-docx* — job failed\n"
            f"• reason: `{_telegram_escape(reason)}`\n"
            f"• job id: `{_telegram_escape(job_id)}`\n"
            f"• engine: `{_telegram_escape(str(meta.get('translation_engine')))}`\n"
            f"• lang:   `{_telegram_escape(str(meta.get('target_language')))}`\n"
            f"• file:   `{_telegram_escape(str(meta.get('input_filename')))}`\n"
            f"\n_{_telegram_escape(message)[:500]}_"
        )

        # ── 1. Text alert (always) ─────────────────────────────────────
        try:
            url = f"https://api.telegram.org/bot{token}/sendMessage"
            payload = {
                "chat_id":    chat_id,
                "text":       text,
                "parse_mode": "Markdown",
                # Silently fail on link-preview attempts so Telegram does
                # not try to expand any URL in the message.
                "disable_web_page_preview": True,
            }
            req = _ur.Request(
                url,
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            with _ur.urlopen(req, timeout=10) as resp:
                body = resp.read()
                if b'"ok":true' not in body:
                    print(f"[telegram] message rejected: {body[:200].decode(errors='replace')}",
                          file=sys.stderr, flush=True)
                    return
            print(f"[telegram] alert text sent to chat {chat_id[:6]}…")
        except _ue.HTTPError as exc:
            print(f"[telegram] message HTTP error: {exc.code} {exc.reason}",
                  file=sys.stderr, flush=True)
            return
        except Exception as exc:
            print(f"[telegram] message skipped: {exc!r}",
                  file=sys.stderr, flush=True)
            return

        # ── 2. Optional input.docx attachment ─────────────────────────
        if os.environ.get("MTD_TELEGRAM_NO_ATTACHMENT", "").strip():
            return
        archived_input = archive_dir / "input.docx"
        if not archived_input.is_file():
            return
        size = archived_input.stat().st_size
        # Telegram Bot API caps uploads at 50 MB on the cloud endpoint;
        # we cap at 20 MB to stay comfortably below limits across
        # mirrors and keep the alert path fast.
        if size > 20 * 1024 * 1024:
            print(f"[telegram] attachment skipped (size {size} B > 20 MB cap)")
            return
        try:
            self._telegram_send_document(
                token=token,
                chat_id=chat_id,
                file_path=archived_input,
                caption=f"input.docx for failed job {job_id} ({reason})",
            )
            print(f"[telegram] attachment sent ({size} B)")
        except Exception as exc:
            print(f"[telegram] attachment skipped: {exc!r}",
                  file=sys.stderr, flush=True)

    @staticmethod
    def _telegram_send_document(
        *,
        token: str,
        chat_id: str,
        file_path: Path,
        caption: str,
    ) -> None:
        """Upload a single file via Telegram's ``sendDocument`` endpoint.

        Pure stdlib multipart construction — no third-party SDK. The
        boundary is randomised so the same payload cannot collide
        with file content. The full body is sent in one POST; for a
        20 MB cap this fits comfortably in memory.
        """
        import urllib.request as _ur
        import uuid as _uuid

        url      = f"https://api.telegram.org/bot{token}/sendDocument"
        boundary = "----mtd-" + _uuid.uuid4().hex
        crlf     = b"\r\n"

        # multipart/form-data body
        def _field(name: str, value: str) -> bytes:
            return (
                b'--' + boundary.encode() + crlf +
                f'Content-Disposition: form-data; name="{name}"'.encode() +
                crlf + crlf +
                value.encode("utf-8") + crlf
            )

        with open(file_path, "rb") as fh:
            file_bytes = fh.read()

        body = b"".join([
            _field("chat_id", chat_id),
            _field("caption", caption),
            b'--' + boundary.encode() + crlf,
            (f'Content-Disposition: form-data; name="document"; '
             f'filename="{file_path.name}"').encode() + crlf,
            b'Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            + crlf + crlf,
            file_bytes, crlf,
            b'--' + boundary.encode() + b'--' + crlf,
        ])

        req = _ur.Request(
            url,
            data=body,
            headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
            method="POST",
        )
        with _ur.urlopen(req, timeout=30) as resp:
            data = resp.read()
            if b'"ok":true' not in data:
                raise RuntimeError(
                    f"telegram rejected sendDocument: {data[:200].decode(errors='replace')}"
                )

    def _strip_timestamp(self, path: Path) -> Path:
        """Rename file to remove leading timestamp prefix (e.g. 1778036666789-).

        W-8 (2026-05-11): if a `_log.json` sidecar exists next to the
        docx, rename it the same way and rewrite its
        ``run_info.output_file`` field to the post-rename name. Without
        this, the sidecar survives under the old timestamped name and
        its ``output_file`` points at a docx that no longer exists.
        """
        clean = _re.sub(r'^\d{10,}-', '', path.name)
        if clean == path.name:
            return path
        clean_path = path.with_name(clean)
        if not clean_path.exists():
            path.rename(clean_path)
        else:
            # A2 (2026-05-12): the previous code dropped the newly produced
            # file when a clean-named version already existed and served the
            # OLD file. That violates the collision-safety rule and could
            # silently hand the user back yesterday's translation. Preserve
            # the new file by walking the same `_1`, `_2`, … suffix the CLI
            # uses on collision.
            stem  = _re.sub(r'(?i)\.docx$', '', clean)
            ext   = clean[len(stem):]   # preserves original case (.docx / .DOCX)
            idx   = 1
            candidate = clean_path.with_name(f"{stem}_{idx}{ext}")
            while candidate.exists():
                idx += 1
                candidate = clean_path.with_name(f"{stem}_{idx}{ext}")
            path.rename(candidate)
            clean_path = candidate

        # Mirror the rename onto the JSON sidecar if one exists.
        sidecar_old = path.with_name(_re.sub(r"(?i)\.docx$", "_log.json", path.name))
        sidecar_new = clean_path.with_name(_re.sub(r"(?i)\.docx$", "_log.json", clean_path.name))
        if sidecar_old.exists() and sidecar_old != sidecar_new:
            try:
                if sidecar_new.exists():
                    sidecar_new.unlink()
                sidecar_old.rename(sidecar_new)
                # Rewrite output_file inside the sidecar to match.
                try:
                    data = json.loads(sidecar_new.read_text(encoding="utf-8"))
                    if isinstance(data.get("run_info"), dict):
                        data["run_info"]["output_file"] = clean_path.name
                        sidecar_new.write_text(
                            json.dumps(data, ensure_ascii=False, indent=2),
                            encoding="utf-8",
                        )
                except Exception as _e:
                    print(f"[W-8] sidecar payload rewrite skipped: {_e}")
            except Exception as _e:
                print(f"[W-8] sidecar rename skipped: {_e}")
        return clean_path

    def _apply_splitter(
        self,
        base_path: Path,
        *,
        split_engine: str | None,
        target_language: str,
        aligner_max_chars: int = 48,
    ) -> Path:
        """Apply the requested Split Method to a translated docx.

        For ``persian_double_lines`` (FA target only): run the FA mechanical
        aligner in-process — no API call, no extra subprocess — and emit
        ``{stem}_Double_Lines.docx`` next to the input. For any other
        splitter or target language, return ``base_path`` unchanged.

        Falls back to the input path on any aligner error so the user
        always receives at least the engine's translated docx.

        ``aligner_max_chars`` (added 2026-05-15) controls the broadcast
        CPL ceiling that the FA aligner enforces per chunk. UI exposes
        24..70 with default 48; we clamp to the same band here so a
        hand-crafted request cannot push the aligner out of range.
        """
        if split_engine != "persian_double_lines":
            return base_path
        if not (target_language or "").lower().startswith("fa"):
            return base_path
        # 2026-05-13 (News Scroll NS 3146 fix): if the file already
        # carries _Double_Lines in its name, refuse to re-align it.
        # Safety net for any future code path that produces an
        # already-aligned file (kept after F7c rollback in case
        # someone calls the aligner directly from a CLI script).
        if "_Double_Lines" in base_path.stem:
            return base_path
        try:
            out_path = _double_lines_output_path(base_path)
            if out_path.exists():
                return out_path
            src_dir = str(ROOT / "src")
            if src_dir not in sys.path:
                sys.path.insert(0, src_dir)
            # Lazy import — keeps the launcher's start-up cheap and avoids
            # loading python-docx until a Persian Double Lines job arrives.
            from machine_translate_docx.openai_tools.persian_double_lines import FASubtitleAligner
            # Clamp again at the boundary in case a caller passed an
            # unvalidated value (e.g. cache replay path).
            mc = max(24, min(70, int(aligner_max_chars)))
            print(f"[splitter] persian_double_lines: MAX_CHARS={mc}")
            aligner = FASubtitleAligner(
                model="gpt-5.4-mini",   # aligner is hardcoded mini (C1)
                llm_threshold=0,        # purely mechanical; no LLM call
                token_budget=0,
                max_chars=mc,
            )
            aligner.align(str(base_path), str(out_path))
            return out_path
        except Exception as exc:
            print(f"[splitter] persian_double_lines failed: {exc}")
            return base_path

    def _materialise_cached_output(
        self,
        cached: dict,
        *,
        split_engine: str | None,
        target_language: str,
        aligner_max_chars: int = 48,
    ) -> Path | None:
        """Copy the cached engine output back into uploads/, applying the
        requested splitter (Phase 4: Persian Double Lines re-runs the FA
        aligner against the cached translation in <2 s — no engine call).
        """
        main_src = cached.get("main_path")
        if not main_src:
            return None
        main_src = Path(main_src)
        if not main_src.exists():
            return None
        uploads_dir = self.state.uploads_dir
        base_dst = uploads_dir / main_src.name
        if not base_dst.exists():
            try:
                shutil.copy2(main_src, base_dst)
            except Exception as exc:
                print(f"[cache] copy-back failed: {exc}")
                return None
        # 2026-05-13: copy the sidecar too so /download/<…>_log.json
        # works on a cache-hit job.
        log_src = cached.get("log_path")
        if log_src:
            log_src = Path(log_src)
            if log_src.exists():
                log_dst = uploads_dir / log_src.name
                try:
                    if not log_dst.exists():
                        shutil.copy2(log_src, log_dst)
                except Exception as exc:
                    print(f"[cache] log copy-back failed: {exc}")
        return self._apply_splitter(
            base_dst,
            split_engine=split_engine,
            target_language=target_language,
            aligner_max_chars=aligner_max_chars,
        )

    def _fallback_output_path(
        self,
        source_file: Path,
        target_language: str,
        translation_engine: str | None = None,
    ) -> Path:
        suffix     = _lang_suffix(target_language) or "OUT"
        engine_tag = _engine_suffix_for(translation_engine)
        stem       = _re.sub(r'^\d{10,}-', '', source_file.stem)
        return source_file.with_name(f"{stem}_{suffix}{engine_tag}.docx")


def _find_free_port(start_port: int) -> int:
    import socket

    port = start_port
    while port < start_port + 100:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
            sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            try:
                sock.bind(("127.0.0.1", port))
                return port
            except OSError:
                port += 1
    raise RuntimeError("No free port found.")


def main() -> int:
    # feat/server-deploy (2026-05-14): load config.toml first so its
    # `[server]` section can override the built-in defaults below.
    # Ensure src/ is on sys.path BEFORE the package import — the
    # launcher is usually invoked via `python local_launcher.py` from
    # the repo root, where the package is at `src/machine_translate_docx`.
    _src_for_cfg = str(ROOT / "src")
    if _src_for_cfg not in sys.path:
        sys.path.insert(0, _src_for_cfg)
    try:
        from machine_translate_docx.server_config import (
            bootstrap as _bootstrap_cfg,
            get_auth as _get_auth,
            get_server as _get_server,
        )
        _cfg = _bootstrap_cfg()
        _srv = _get_server(_cfg)
        _auth = _get_auth(_cfg)
    except Exception as _exc:
        print(f"[server_config] WARNING: bootstrap failed ({_exc!r}); using defaults.")
        _cfg, _srv, _auth = {}, {"host": "127.0.0.1", "port": 3000, "max_concurrent_jobs": 1}, {}

    parser = argparse.ArgumentParser(description="Run a local browser-ready simulator for the DOCX translator UI.")
    parser.add_argument("--port", type=int, default=_srv["port"], help="Preferred port to use. Falls back to the next free port if busy.")
    parser.add_argument("--no-browser", action="store_true", help="Do not open the browser automatically.")
    parser.add_argument("--host", default=_srv["host"], help="Host interface to bind.")
    parser.add_argument("--backend", choices=["real", "mock"], default="real", help="Use the real Python backend or the local placeholder mode.")
    parser.add_argument("--python-exe", default="", help="Python interpreter to use for the real backend. Defaults to the current interpreter.")
    parser.add_argument("--setup", action="store_true", help="Run the interactive setup wizard and exit.")
    args = parser.parse_args()

    if args.setup:
        # Delegate to scripts/setup_wizard.py
        import runpy
        runpy.run_path(str(ROOT / "scripts" / "setup_wizard.py"), run_name="__main__")
        return 0

    runtime_dir = Path(tempfile.gettempdir()) / "machine_translate_docx_local"
    if runtime_dir.exists():
        shutil.rmtree(runtime_dir, ignore_errors=True)

    python_exe = Path(args.python_exe) if args.python_exe else Path(sys.executable)
    if args.backend == "real" and not python_exe.exists():
        raise FileNotFoundError(f"Python executable not found: {python_exe}")

    # 2026-05-11: backend script moved from `src/machine_translate_docx.py`
    # into the new `src/machine_translate_docx/cli.py` package layout.
    # The launcher invokes it as a regular .py path (NOT `python -m`)
    # because the CLI carries argparse + lots of import-time side effects
    # that historically rely on being run as a top-level script, not as
    # a sub-module. The old name is still recognised as a fallback so
    # mid-migration checkouts work.
    script_path = ROOT / "src" / "machine_translate_docx" / "cli.py"
    if not script_path.exists():
        # Fallback for old checkouts where the file hasn't moved yet.
        legacy = ROOT / "src" / "machine_translate_docx.py"
        if legacy.exists():
            script_path = legacy
        else:
            raise FileNotFoundError(f"Backend script not found: {script_path}")

    state = LocalState(runtime_dir, args.backend, python_exe, script_path)
    state.boot()

    # Periodically prune finished jobs older than 1 h so the in-memory job
    # store does not grow unbounded across long-running sessions.
    state.start_cleanup_thread(interval_sec=600, max_age_sec=3600)

    # Weekly Telegram export of subscribers.txt (Saturday noon Europe).
    # Boot-time check first — if last attempt failed, log a one-line
    # warning and clear the flag. Then arm the scheduler (no-op if
    # Telegram env vars are not configured).
    state.boot_subscribers_report_check()
    state.start_subscribers_report_thread()

    port = _find_free_port(args.port)
    server = ThreadingHTTPServer((args.host, port), MockTranslatorHandler)
    server.state = state  # type: ignore[attr-defined]
    server.index_html = _inject_client_patch(_load_index_html())  # type: ignore[attr-defined]
    # feat/server-deploy: stash auth creds + boot time + version so the
    # handler can serve /health without auth and gate other routes.
    import time as _time_now
    server._mtd_auth = _auth                          # type: ignore[attr-defined]
    server._mtd_boot_time = _time_now.time()          # type: ignore[attr-defined]
    try:
        from machine_translate_docx import __version__ as _mtd_ver  # type: ignore
    except Exception:
        _mtd_ver = "dev"
    server._mtd_version = _mtd_ver                    # type: ignore[attr-defined]
    if _auth.get("password_hash"):
        print(f"[auth] HTTP Basic enabled for user '{_auth.get('username', '')}'.")
    else:
        print("[auth] disabled (no [auth] section in config.toml — workstation mode).")

    url = f"http://{args.host}:{port}/"
    print(f"Local launcher ready: {url}")
    print(f"Runtime dir: {runtime_dir}")
    print("Press Ctrl+C to stop.")

    if not args.no_browser:
        threading.Timer(1.0, lambda: webbrowser.open(url, new=2)).start()

    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopping...")
    finally:
        server.server_close()

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
