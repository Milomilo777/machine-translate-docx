#!/usr/bin/python3
# - *- coding: utf- 8 - *-
PROGRAM_VERSION="2026-03-29"

import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pkg_resources")
# Suppress deprecation warning from pyobjc using pkg_resources
warnings.filterwarnings(
    "ignore",
    category=UserWarning,
    module="objc._bridgesupport"
)

import sys
import io

# If all these flags appear anywhere on the command line, exit quietly.
UNWANTED_FLAGS = {"-B", "-S", "-E", "-s", "-c"}

if UNWANTED_FLAGS.issubset(set(sys.argv[1:])):
    # Optional extra safety: ensure there's something after -c (the inline code)
    try:
        c_index = sys.argv.index("-c")
        if c_index + 1 < len(sys.argv):
            # There *is* an argument after -c; likely the resource_tracker helper
            sys.exit(0)
        # If -c is the last token (rare/invalid), still exit if you want:
        # sys.exit(0)
    except ValueError:
        # -c not found (shouldn't happen when issubset passed), but just in case:
        sys.exit(0)
    
# For bidirectional text display right to left and left to right
from bidi.algorithm import get_display

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
sys.stdout.reconfigure(encoding='utf-8')
import platform

json_configuration_url='https://raw.githubusercontent.com/translation-robot/machine-translate-docx/main/src/configuration/configuration.json'
# Day 0 is October 3rd 2017


print("*********************************************************")
print("*  machine-translate-docx program version : %s" % (PROGRAM_VERSION))
print("*********************************************************")

print("Python programming language %s\n" % (platform.python_version()))

import gc
import pprint
from pprint import pprint
import traceback
import shlex
import subprocess
import os
#from googletrans import Translator
import re
import time
import codecs
import urllib
import urllib.request
import requests
import json
import json5 # json 5 with the ability to have comments


from inspect import currentframe, getframeinfo
import chardet
import getpass
import datetime

import zipfile
import xml.dom.minidom
# used to get elements in XML, shading in docx for example
from lxml import etree

# This library automatically downloads chrome driver
# pyderman was replaced with webdriver_manager
# then selenium 4.11.2 managed downloading the drivers
#import pyderman
# For selenium 3

# When translation engine is deepl or chatgpt : use undetected_chromedriver
# Else, use standard selenium webdriver
# drivers are loaded after getting the engine value

from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.remote.remote_connection import LOGGER
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import StaleElementReferenceException, TimeoutException

from screeninfo import get_monitors


#from selenium.webdriver import Firefox, FirefoxOptions
from time import sleep
import argparse
import clipboard
#import pyperclip

import psutil

#import winsound

import docx
from docx import Document
from docx import oxml
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT,WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

# For japanese
#import tinysegmenter

#from tinysegmenter import TinySegmenter
# For Thai
#from thai_tokenizer import Tokenizer as thai_tokenizer_tokenizer
#from thai_tokenizer import data
#from thai_tokenizer.data import bpe_merges
#import thai_tokenizer 
#import thai-segmenter

import timeit
import datetime
import progressbar


from timeit import default_timer as timer

import re
import inspect

from xlsx_translation_memory import xlsx_translation_memory

# Module-level translation-memory handle. Pre-refactor code assumed this
# existed as a module-level global, but no top-level binding was ever
# committed — `initialize_translation_memory_xlsx` only created a local
# `xtm`, so every later read raised NameError once the code path ran
# live (silent in tests because the F-012 mid-layer wasn't exercised).
# Initialise as None here so `if xtm is not None` reads correctly when
# `--xlsxreplacefile` is not provided.
xtm = None

import html

from urllib.parse import urlencode, quote_plus

from openpyxl import load_workbook
from openpyxl import Workbook

from bs4 import BeautifulSoup

# pip install pycryptodome
# used for passwords (deepl, etc)

# Load configuration from a json file on internet (github for example)

import json

from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.text.run import Run as _DocxRun

import glob
from langcodes import *
import math
import shutil
import signal
import atexit

import random

from config import (
    DefaultJsonConfiguration,
    validate_json_string,
    get_nested_value_from_json_array,
    google_translate_lang_codes,
    deepl_translate_lang_codes,
    office_language_tags,
    right_to_left_languages_list,
    eol_array,
    eol_conditional_array,
    bol_array,
    MAX_LINE_SIZE,
    COUNTRY_QUERY_HTTP_TIMEOUT,
)

from runtime import RuntimeContext
from selenium_utils import (
    safe_click,
    browser_fill_form_field_value,
    set_chrome_window_2_3_screen,
    create_webdriver,
    minimize_browser,
    clean_up_previous_chrome_selenium_drivers,
    cleanup_selenium_chrome_temp_folders,
)
from engines.google import (
    selenium_chrome_google_translate,
    selenium_chrome_google_click_cookies_consent_button,
)
from engines.deepl import (
    selenium_chrome_deepl_log_in,
    selenium_chrome_deepl_log_off,
    deepl_close_messages,
    selenium_chrome_deepl_translate,
)
# Stale ``engines._prompts`` shim was removed in C3 of the 2026-05-10
# cleanup. The only callers of ``build_translation_prompt`` were the
# now-deleted chatgpt-web / perplexity-web engines.
from runner import selenium_chrome_translate_maxchar_blocks as _runner_translate_maxchar_blocks

# Module-level RuntimeContext singleton — Phase F1 transition shim.
# Lazily built by `_get_ctx()` on first call from any function that has
# already been threaded with a `ctx` parameter. Phase F1.6 collapses this
# helper into an explicit `ctx = RuntimeContext.empty()` at the top of
# main(); until then, threaded functions can be called from un-threaded
# callers without a signature cascade.
_ctx: RuntimeContext | None = None


def _get_ctx() -> RuntimeContext:
    """Return the singleton RuntimeContext, snapshotting current module-level
    globals on first call. Threaded callers should propagate their own
    parameter where possible — this helper exists only to bridge un-threaded
    callers to threaded callees during the F1 transition."""
    global _ctx
    if _ctx is None:
        _ctx = RuntimeContext.empty()
        # Snapshot whatever module-level globals are already populated by
        # import time. The dataclass defaults cover anything that is not.
        # Each F1.x batch may extend this snapshot as it threads more
        # globals.
        try:
            _ctx.browser.driver = driver
        except NameError:
            pass
        try:
            _ctx.language.src_lang_name  = src_lang_name
            _ctx.language.dest_lang_name = dest_lang_name
        except NameError:
            pass
        # F1.2 — configuration snapshot
        try:
            _ctx.config.json_configuration_array = json_configuration_array
        except NameError:
            pass
        try:
            _ctx.config.max_translation_block_size = MAX_TRANSLATION_BLOCK_SIZE
        except NameError:
            pass
        # F1.4 — engine + flags snapshot
        try:
            _ctx.engine.engine = translation_engine
        except NameError:
            pass
        try:
            _ctx.engine.method = engine_method
        except NameError:
            pass
        try:
            _ctx.flags.splitonly = splitonly
        except NameError:
            pass
        try:
            _ctx.docx.translation_errors_count = translation_errors_count
        except NameError:
            pass
        try:
            _ctx.browser.deepl_sleep_wait_translation_seconds = deepl_sleep_wait_translation_seconds
        except NameError:
            pass
        # F1.5 — DeepL session flags + OpenAI handles
        try:
            _ctx.browser.closed_cookies_accept_message_bool = closed_cookies_accept_message_bool
        except NameError:
            pass
        try:
            _ctx.browser.close_install_extension_message_bool = close_install_extension_message_bool
        except NameError:
            pass
        try:
            _ctx.browser.deepl_nb_clear_cached_times = deepl_nb_clear_cached_times
        except NameError:
            pass
        try:
            _ctx.browser.logged_into_deepl = logged_into_deepl
        except NameError:
            pass
        try:
            _ctx.openai.translator = oai_translator
        except NameError:
            pass
        try:
            _ctx.openai.polisher = oai_polisher
        except NameError:
            pass
        try:
            _ctx.openai.translation_log = translation_log
        except NameError:
            pass
        # F1.5 — language + paths + translation_array
        try:
            _ctx.language.src_lang = src_lang
        except NameError:
            pass
        try:
            _ctx.language.dest_lang = dest_lang
        except NameError:
            pass
        try:
            _ctx.flags.word_file_to_translate = word_file_to_translate
        except NameError:
            pass
        try:
            _ctx.flags.use_api = use_api
        except NameError:
            pass
        try:
            _ctx.docx.translation_array = translation_array
        except NameError:
            pass
        try:
            _ctx.docx.blocks_nchar_max_to_translate_array = blocks_nchar_max_to_translate_array
        except NameError:
            pass
        # G1 — webdriver module + paths + remaining flags
        try:
            _ctx.browser.webdriver_module = webdriver
        except NameError:
            pass
        try:
            _ctx.browser.chrome_options = chrome_options
        except NameError:
            pass
        try:
            _ctx.flags.exitonsuccess = exitonsuccess
        except NameError:
            pass
        try:
            _ctx.flags.silent = silent
        except NameError:
            pass
        try:
            _ctx.flags.viewdocx = viewdocx
        except NameError:
            pass
        try:
            _ctx.flags.xlsxreplacefile = xlsxreplacefile
        except NameError:
            pass
        # G4 — aimodel + with_polish snapshots
        try:
            _ctx.flags.aimodel = args.aimodel
        except (NameError, AttributeError):
            pass
        try:
            _ctx.flags.with_polish = args.with_polish
        except (NameError, AttributeError):
            pass
    return _ctx


def _atexit_cleanup_driver() -> None:
    """Best-effort browser shutdown on interpreter exit.

    The happy-path quit lives at the bottom of ``main()``; if anything
    above it raises, the child Chrome process gets orphaned. Registering
    this with ``atexit`` makes sure the driver is closed on any normal
    termination — including crashes — so the launcher's job pool doesn't
    accumulate zombie Chrome processes between failed jobs.
    """
    try:
        if _ctx is not None and _ctx.browser.driver is not None:
            try:
                _ctx.browser.driver.quit()
            except Exception:
                pass
    except Exception:
        pass


import atexit as _atexit
_atexit.register(_atexit_cleanup_driver)


def _sync_globals_from_ctx(ctx: RuntimeContext) -> None:
    """Mirror every populated attribute of ``ctx.docx`` (and selected
    ``ctx.flags`` / ``ctx.language`` fields) onto this module's namespace
    so legacy functions that still read by bare name see the populated
    state.

    Phase H bridge: rather than threading ~40 remaining functions one
    at a time, we re-export the same objects under their historical
    module-level names. Lists and dicts on ``ctx.docx`` are referenced
    by identity, so any in-place mutation by a downstream helper is
    visible through both names. After a full replacement (e.g.
    ``ctx.docx.translation_array = new_list``), call this helper again
    to refresh the module-level alias.

    Called from ``main()`` at pipeline boundaries: after
    ``read_and_parse_docx_document``, after ``translate_docx``, and
    after the polish/aligner step. Cheap (one ``setattr`` per attr).
    """
    import sys as _sys
    _mod = _sys.modules[__name__]
    # Mirror every public dataclass field on ctx.docx.
    for _name, _value in vars(ctx.docx).items():
        if _name.startswith("_"):
            continue
        setattr(_mod, _name, _value)
    # A few non-docx ctx fields are also read by bare name in legacy code.
    if getattr(ctx.language, "dest_lang", None) is not None:
        setattr(_mod, "dest_lang", ctx.language.dest_lang)
    if getattr(ctx.language, "src_lang", None) is not None:
        setattr(_mod, "src_lang", ctx.language.src_lang)
    # Browser handle: many helpers (download utilities, Perplexity login,
    # the recovery branches inside the Selenium engines) read `driver` as
    # a bare name. Mirror the active handle so they reach the live session.
    if getattr(ctx.browser, "driver", None) is not None:
        setattr(_mod, "driver", ctx.browser.driver)
    # OpenAI handles — read by bare name in legacy translate / polish helpers.
    if getattr(ctx.openai, "translator", None) is not None:
        setattr(_mod, "oai_translator", ctx.openai.translator)
    if getattr(ctx.openai, "polisher", None) is not None:
        setattr(_mod, "oai_polisher", ctx.openai.polisher)


# Track the child processes
def kill_child_process():
    import time
    parent = psutil.Process(os.getpid())
    children = parent.children(recursive=True)

    #print(f"[CLEANUP] Found {len(children)} child process(es).")

    for child in children:
        try:
            #print(f"[CLEANUP] Terminating PID {child.pid} ({' '.join(child.cmdline())})")
            child.terminate()
        except psutil.NoSuchProcess:
            print(f"[CLEANUP] PID {child.pid} already gone.")
        except Exception as e:
            print(f"[CLEANUP] Could not terminate PID {child.pid}: {e}")

    # Give processes some time to exit gracefully
    _, alive = psutil.wait_procs(children, timeout=3)

    for child in alive:
        try:
            print(f"[CLEANUP] Forcing kill on PID {child.pid} ({' '.join(child.cmdline())})")
            child.kill()
        except psutil.NoSuchProcess:
            print(f"[CLEANUP] PID {child.pid} disappeared before kill().")
        except Exception as e:
            print(f"[CLEANUP] Could not kill PID {child.pid}: {e}")

    # Re-check if any children remain
    still_alive = parent.children(recursive=True)
    if still_alive:
        print(f"[CLEANUP] WARNING: Still alive: {[p.pid for p in still_alive]}")
    else:
        #print("[CLEANUP] All child processes terminated.")
        pass

# Run cleanup on normal exit
atexit.register(kill_child_process)

# Trap termination signals
def handle_signal(signum, frame):
    kill_child_process()
    raise SystemExit(0)

signal.signal(signal.SIGINT, handle_signal)
signal.signal(signal.SIGTERM, handle_signal)

#from parsivar import Normalizer
my_hazm_normalizer = None

#from hazm import Normalizer
#import hazm

# validate_json_string + get_nested_value_from_json_array + DefaultJsonConfiguration
# now live in src/config.py.


def test_internet(host="8.8.8.8", port=53, timeout=3):
    """
    Host: 8.8.8.8 (google-public-dns-a.google.com)
    OpenPort: 53/tcp
    Service: domain (DNS/TCP)
    """
    try:
        socket.setdefaulttimeout(timeout)
        socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect((host, port))
        return True
    except socket.error as ex:
        print(ex)
        return False

try:
    json_online_configuration = requests.get(json_configuration_url).content
except Exception:
    print("Warning, unable to get configuration from internet at {json_configuration_url}")
    if not test_internet():
        print("Warning, internet connection seems to be down, google name servers don't respond")
        time.sleep(1)
        
    json_online_configuration = "{}"


# Find default configuration file name from other configuration files
local_configuration_json_path_key = ["local_configuration", "json_filename_path"]
local_configuration_json_path = get_nested_value_from_json_array([json_online_configuration,DefaultJsonConfiguration], local_configuration_json_path_key)

# determine if application is a script file or frozen exe
if getattr(sys, 'frozen', False):
    application_path = os.path.dirname(sys.executable)
elif __file__:
    application_path = os.path.dirname(__file__)

configuration_file_full_path = os.path.join(application_path, local_configuration_json_path)

try:
    if os.path.isfile(configuration_file_full_path):
        with open(configuration_file_full_path) as configuration_file:
          local_json_contents = configuration_file.read()
    else:
        #print(f"Optional local json configuration file not found at {configuration_file_full_path}, ignoring")
        local_json_contents = None
except Exception:
    local_json_contents = None
          
json_configuration_array = [local_json_contents,json_online_configuration,DefaultJsonConfiguration]

# Find default sleep time for update message
version_checker_sleep_seconds_on_update_key = ["version_checker", "sleep_seconds_on_update"]
version_checker_sleep_seconds_on_update = get_nested_value_from_json_array(json_configuration_array,
    version_checker_sleep_seconds_on_update_key, default_when_none=30)
# We assume the program does not need update. Value of 1 is for update needed.
str_needs_update = "0"


# Colors : grey and pink backgroud to ignore
# https://learn.microsoft.com/en-us/office/vba/api/word.wdcolor
shading_color_ignore_text_update_key = ["document", "shading_color_ignore_text"]
shading_color_ignore_text = get_nested_value_from_json_array(json_configuration_array,
    shading_color_ignore_text_update_key)

process_platform = platform.system()
if platform.system() == 'Windows':
    from win32com.client import Dispatch

tried_login_in_deepl = False
logged_into_deepl = False

from_text_table = [''] *1
from_text_is_greyed_table = [0] *1
from_text_is_red_color_table = [0] *1
from_text_is_end_of_line_table = [0] *1
from_text_is_beginning_of_line_table = [0] *1
from_text_is_empty_line_table = [0] *1
from_text_is_conditional_end_of_line_table = [0] *1
from_text_by_phrase_separator_table = [''] *1
from_text_by_phrase_table = [''] *1
#number of lines in per phrase
from_text_nb_lines_in_phrase = [0] *1
from_text_nb_lines_in_cell = [0] *1
#
to_text_by_phrase_separator_table = [''] *1
to_text_by_phrase_separator_removed_table = [''] *1
to_text_splited_table1 = [''] *1
to_text_by_phrase_table = [''] *1
to_text_table = [''] *1
to_raw_translated_table = [''] *1
to_text_removed_line_separator = [''] *1
translation_result_using_separator = [''] *1
translation_result_phrase_array = [[]] *1
translation_result = [''] *1
from_text_is_read = [0] *1
word_translation_table_length = 0
table = None

table_cells = [['' for i in range(1)] for j in range(1)]

docxfile_table_number_of_phrases = 0

selenium_chrome_machine_translate_once = None

numerrors_deepl = 0

# We have found zero phrase up to now
docxfile_table_number_of_characters = 0
docxfile_table_number_of_phrases = 0
docxfile_table_number_of_words = 0

numrows = 0
numcols = 0

E_mail_str = 'sm' + 'tv' + '.' + 'bot' + '@g' + 'mail' + '.' + 'c' + 'o' + 'm'

#import pandas as pd
#import multiprocessing

cf = currentframe()
filename = getframeinfo(cf).filename

start_time = datetime.datetime.now()

for m in get_monitors():
    #print(str(m))
    break

parser = argparse.ArgumentParser()

#parser = argparse.ArgumentParser(description = "Translate everything!")
#parser.add_argument('--source-language', required = True, choices = Languages, help="Specify the source language!")
parser.add_argument('--srclang', '-sl', required = False, help="Specify the default source language, en is default (hi,ja,ru,de,ru,hi,ja,in, etc)", default='en')
parser.add_argument('--destlang', '--dl', required = False, help="Specify the destination language with 2 letter code (hi,ja,ru,de,ru,hi,ja,in, etc)")
parser.add_argument('--engine', '-e', required = False, help="Specify the translation engine (google, deepl, chatgpt, perplexity)")
parser.add_argument('--enginemethod', '-m', required = False, help="Specify the method (javascript, phrasesblock, singlephrase, xlsxfile, textfile )")
parser.add_argument('--aimodel', '-am', required = False, help="Specify the ai model when applicable")
parser.add_argument('--docxfile', '-d', required = False, help="Input file name")
parser.add_argument('--xlsxreplacefile', '-x', required = False, help="Excel xlsx search and replace file")
parser.add_argument('--destfont', '-f', required = False, help="Destination font name")
parser.add_argument('--useapi', '-a', required = False, help="Use api to get translation, lower quality but faster", action='store_true')
parser.add_argument('--split', '-s', required = False, help="Split web translation into cells", action='store_true')
parser.add_argument('--splitengine', '-p', required = False, help="Specify split engine (openai | persian_double_lines)")
parser.add_argument('--splitonly', required = False, help="Split translation into lines only, do not translate.", action='store_true')
parser.add_argument('--showbrowser', '-b', required = False, help="Show browser", action='store_true')
parser.add_argument('--exitonsuccess', '-t', required = False, help="Exit progream on success", action='store_true')
parser.add_argument('--viewdocx', '-l', required = False, help="Open the docx file with the default application after completion.", action='store_true')
parser.add_argument('--silent', '-q', required = False, help="Silent, do not ask question and exit silently", action='store_true')
parser.add_argument("--verbose", '-v', help="increase output verbosity", action="store_true")
parser.add_argument("--clientip", '-i', help="Client IP for statistics")
#parser.add_argument('--destination-file', required = True, help="Output file name")
#args = parser.parse_args()
parser.add_argument('--version', required = False, help="Show program version", action='store_true')
parser.add_argument('--with-polish', required = False, help="Run a Persian polish pass after translation and before splitting (API only)", action='store_true', dest='with_polish')

try:
    args = parser.parse_args()
except Exception:  #print("Waiting for the input_element...")
    var = traceback.format_exc()
    print(var)
    #input ("Type enter to continue")

show_version = args.version
silent = args.silent
if show_version:

    print("\nDeveloper: %s\n" %(E_mail_str))
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not silent:
        input("\nEnter to close program")
    sys.exit(0)

if args.docxfile is None:
    parser.print_help()
    print("\nDeveloper: smtv.bot@gmail.com\n")
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not silent:
        input("\nEnter to close program")
    else:
        print("Program ended with errors")
    sys.exit(1)

use_html = False


script_folder = os.path.dirname(os.path.realpath(__file__))
os_path = os.environ["PATH"]
new_os_path = "%s;%s" %(script_folder, os_path)

#print("\nAdding %s to path.\n" % (script_folder))
os.environ["PATH"] = new_os_path

use_translation_api = False


#line_separator_str = ' () '
line_separator_str = ' '
#line_separator_nospace_str = '()'
line_separator_nospace_str = '()'
line_separator_regex_str = ' ?\(\) ?'

#pp = pprint.PrettyPrinter(indent=4)


html_file_path = ''

nb_character_total = 0

# Maximum 5000 characters on the free version
# but only 1500 if not logged on to deepl with free version
deepl_max_char_bloc_size_key = ['deepl', 'no_account','maximum_character_block']
deepl_maximum_character_block = get_nested_value_from_json_array(json_configuration_array, deepl_max_char_bloc_size_key)

deepl_sleep_wait_translation_seconds = 0.1
translation_errors_count = 0

word_file_to_translate = args.docxfile

viewdocx = args.viewdocx
client_ip = args.clientip

xlsxreplacefile = args.xlsxreplacefile
dest_font = args.destfont
split_translation = args.split

split_engine = args.splitengine
if split_engine is not None:
    split_engine = split_engine.lower()
    if split_engine not in ('openai', 'persian_double_lines'):
        print("Unknown split engine, accepted values : openai, persian_double_lines. Defaulting to non AI line splitting.")
        split_engine = None

use_api = args.useapi
#use_browser = args.useapi

showbrowser = args.showbrowser
exitonsuccess = args.exitonsuccess
splitonly = args.splitonly
with_polish = args.with_polish
if splitonly:
    split_translation = True

driver = None

src_lang = args.srclang
dest_lang = args.destlang
if dest_lang is not None:
    dest_lang = dest_lang.lower()
else:
    dest_lang = ""
    if not splitonly:
        dest_lang = input ("Please enter language translation code (fr,de,ru,hi,etc.)")

# cjk_segmenter = None 
# if dest_lang == 'zh':
# import jieba
# jieba.enable_paddle()# 启动paddle模式。 0.40版之后开始支持，早期版本不支持
# if dest_lang == 'zh-cn':
# dest_lang = 'zh-CN'
# import jieba
# jieba.enable_paddle()# 启动paddle模式。 0.40版之后开始支持，早期版本不支持
# if dest_lang == 'zh-tw':
# dest_lang = 'zh-TW'
# import jieba
# jieba.enable_paddle()# 启动paddle模式。 0.40版之后开始支持，早期版本不支持
# if dest_lang == 'th':
# from newmm_tokenizer.tokenizer import word_tokenize
# if dest_lang == 'ko':
# from soynlp.word import WordExtractor
# if dest_lang == 'ja':
# from tinysegmenter import TinySegmenter
# cjk_segmenter = TinySegmenter()
# if dest_lang == 'fa':
# from hazm import Normalizer
# my_hazm_normalizer = Normalizer()


cjk_segmenter = None 
if dest_lang == 'zh-cn':
    dest_lang = 'zh-CN'
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang == 'zh-tw':
    dest_lang = 'zh-TW'
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang.lower() == 'zh-hant' or dest_lang == 'zh-hans':
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
    
if dest_lang == 'th':
    from newmm_tokenizer.tokenizer import word_tokenize
if dest_lang == 'zh' or dest_lang == 'ja' or dest_lang == 'ko':
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang == 'fa':
    from hazm import Normalizer
    my_hazm_normalizer = Normalizer()



valid_online_json = validate_json_string(json_online_configuration)
if not valid_online_json == True:
    print(f"json_online_configuration={json_online_configuration}")
    print(f"Warning: Json file at {json_configuration_url} is not valid. Ignoring this configuration file.")
else:
    #print(f"Using JSON configuration file at {json_configuration_url} : OK")
    pass
    
valid_local_json = validate_json_string(local_json_contents)
if os.path.isfile(configuration_file_full_path):
    if not valid_local_json == True:
        print(f"Warning: Json file at {configuration_file_full_path} is not valid. Ignoring this configuration file.")
    else:
        print(f"Using JSON configuration file at {configuration_file_full_path}")

print("")


src_lang_name = (google_translate_lang_codes.get(src_lang))
if src_lang_name is None:
    src_lang_name = ""
    if not splitonly:
        print("Source language name for %s not found. Continuing as it is." % (dest_lang))
else:
    print("Source language name for '%s' : %s" % (src_lang, src_lang_name))

dest_lang_name = (google_translate_lang_codes.get(dest_lang))

if dest_lang_name is None:
    dest_lang_name = deepl_translate_lang_codes.get(dest_lang)
    if not splitonly and dest_lang_name is None:
        print("Target language name for %s not found. Continuing as it is." % (dest_lang))
else:
    print("Target language name for '%s' : %s" % (dest_lang, dest_lang_name))

dest_lang_tag = ""
try:
    dest_lang_tag = office_language_tags[dest_lang]
except Exception:
    pass

translation_engine = args.engine

if translation_engine is not None:
    translation_engine = translation_engine.lower()
else:
    translation_engine = ""

if translation_engine in ['perplexity', 'chatgpt', 'deepl']:
    showbrowser = True
elif translation_engine in ['deepl', 'chatgpt']:
    pass  # keep the value as is
else:
    translation_engine = 'google'

if use_api and translation_engine != 'chatgpt':
    use_api = False 
    
perplexity_max_char_bloc_size_key = ['perplexity', 'account','maximum_character_block']
perplexity_maximum_character_block = get_nested_value_from_json_array(json_configuration_array, perplexity_max_char_bloc_size_key)

engine_method = args.enginemethod
engine_method = "%s" % engine_method
engine_method = engine_method.strip().lower()

translation_array = []

if splitonly:
    engine_method = ''
elif translation_engine == 'google':
    if engine_method == 'api' or use_api == True:
        engine_method = 'api'
    elif engine_method  == 'singlephrase':
        engine_method = 'singlephrase'
    elif engine_method  == 'phrasesblock':
        engine_method = 'phrasesblock'
    elif engine_method =='xlsxfile':
        engine_method = 'xlsxfile'
        # There is a bug on xlsxfile method, show browser for debugging purpose
        showbrowser = True
    elif engine_method  == 'textfile':
        engine_method = 'textfile'
        # There is a bug on textfile method, show browser for debugging purpose
        showbrowser = True
    elif engine_method  == 'javascript':
        engine_method = 'javascript'
    else:
        # Default for `--engine google` was historically `javascript`,
        # which uploads a local HTML file to translate.google.com. That
        # path stopped working when modern Chrome started blocking the
        # widget on file:// URLs (~2022); it now fails fast with an
        # error message but produces an empty docx. Switch the default
        # to `phrasesblock`, which uses the textarea URL — fast and
        # actually working. Users who genuinely want the old file-mode
        # path can still pass `--enginemethod javascript` explicitly.
        engine_method = 'phrasesblock'
elif translation_engine == 'deepl':
    if engine_method == 'singlephrase' or use_api == True:
        engine_method = 'singlephrase'
    elif engine_method  == 'phrasesblock':
        engine_method = 'phrasesblock'
    else:
        engine_method = 'phrasesblock'
elif translation_engine == 'chatgpt':
    # chatgpt-web was removed in the 2026-05-10 cleanup; only the API
    # path remains for chatgpt.
    if engine_method == 'api' or use_api == True:
        engine_method = 'api'
        showbrowser = False
    else:
        engine_method = 'api'
        showbrowser = False

elif translation_engine == 'perplexity':
    # perplexity-web was removed in the 2026-05-10 cleanup; only the
    # webservice path remains.
    if engine_method == 'webservice':
        engine_method = 'webservice'
    else:
        engine_method = 'webservice'
else:
    engine_method = "web"

if engine_method == 'webservice':
    showbrowser = False

if translation_engine == 'chatgpt' and engine_method == 'api':
    from openai_tools import OpenAITranslator, OpenAIPolisher
    # FASubtitleAligner is no longer imported here (Phase 1): the aligner is
    # decoupled from chatgpt-polish and reached only via the Persian Double
    # Lines Split Method, which performs its own local import on demand.
    chatgpt_max_char_bloc_size_key = ['chatgpt', 'api','maximum_character_block']
else:
    chatgpt_max_char_bloc_size_key = ['chatgpt', 'no_account','maximum_character_block']
chatgpt_maximum_character_block = get_nested_value_from_json_array(json_configuration_array, chatgpt_max_char_bloc_size_key)

# Load openai line splitting package
from openai_tools import OpenAISubtitleSplitter
    
if translation_engine == 'perplexity':
    MAX_TRANSLATION_BLOCK_SIZE = perplexity_maximum_character_block
elif translation_engine == 'chatgpt':
    MAX_TRANSLATION_BLOCK_SIZE = chatgpt_maximum_character_block
else:
    MAX_TRANSLATION_BLOCK_SIZE = deepl_maximum_character_block
# Override MAX_TRANSLATION_BLOCK_SIZE value after logging on Deepl

# When translation engine is deepl or chatgpt : use undetected_chromedriver
# Else, use standard selenium webdriver

if translation_engine in ['perplexity', 'chatgpt'] and engine_method != "webservice":
    import undetected_chromedriver as webdriver
else:
    from selenium import webdriver  # regular selenium webdriver
    
def lineno():
    """Returns the current line number in our program."""
    return inspect.currentframe().f_back.f_lineno
    
def linux_distribution():
    try:
        return platform.linux_distribution()
    except Exception:
        return "N/A"

oai_translator = None
oai_polisher = None
translation_log = {"run_info": {}, "blocks": [], "summary": {}}


def write_translation_log(log_path: str):
    """Write translation_log as formatted JSON next to the output DOCX."""
    import json as _json
    import datetime as _dt

    blocks = translation_log.get("blocks", [])

    total_prompt = 0
    total_completion = 0
    total_total = 0
    total_cached = 0
    total_cost = 0.0
    total_elapsed = 0.0

    for b in blocks:
        for key in ("translation", "polish"):
            call = b.get(key) or {}
            tok = call.get("tokens") or {}
            total_prompt     += tok.get("prompt", 0)
            total_completion += tok.get("completion", 0)
            total_total      += tok.get("total", 0)
            total_cached     += tok.get("cached", 0)
            total_cost       += call.get("cost_usd", 0.0)
            total_elapsed    += call.get("elapsed_seconds", 0.0)

    translation_log["run_info"]["output_file"] = log_path.replace("_log.json", ".docx")
    translation_log["summary"] = {
        "total_blocks":        len(blocks),
        "total_tokens": {
            "prompt":          total_prompt,
            "completion":      total_completion,
            "total":           total_total,
            "cached":          total_cached,
        },
        "total_cost_usd":      round(total_cost, 6),
        "elapsed_total_seconds": round(total_elapsed, 3),
    }

    with open(log_path, "w", encoding="utf-8") as fh:
        _json.dump(translation_log, fh, ensure_ascii=False, indent=2)

    print(f"[INFO] Translation log saved → {log_path}")


def print_os_info():

    print("""Python version: %s
    dist: %s
    linux_distribution: %s
    system: %s
    machine: %s
    platform: %s
    uname: %s
    version: %s
    mac_ver: %s
    """ % (
    sys.version.split('\n'),
    str(platform.dist()),
    linux_distribution(),
    platform.system(),
    platform.machine(),
    platform.platform(),
    platform.uname(),
    platform.version(),
    platform.mac_ver(),
    ))

if not os.path.exists(word_file_to_translate) :
    print("ERROR: File not found: %s" % (word_file_to_translate))
    sys.exit(1)

splitted_filename = os.path.splitext(os.path.basename(word_file_to_translate))

# number of segment separated by dot in the docx filename
splitted_filename_size = len(splitted_filename)

docx_file_name =  "%s%s" % (splitted_filename[splitted_filename_size-2], splitted_filename[splitted_filename_size-1])

if splitted_filename_size > 1:
    word_file_to_translate_extension = splitted_filename[splitted_filename_size-1].lower()

if word_file_to_translate_extension == ".docx":
    try:
        docxdoc = docx.Document(word_file_to_translate)
    except Exception:
        print(f"Error, file {word_file_to_translate} does not appear to be a valid Microsoft Word docx file.")
        print("Please check that the file is a valid document and rerun on a valid Microsoft docx document.\n")
        
        print("\nDeveloper: %s" % (E_mail_str))
        print("Program version: %s\n" % (PROGRAM_VERSION))
        if not silent:
            input("Enter to close program")
        else:
            print("Program ended with errors")
        sys.exit(2)
    styles = docxdoc.styles
    
    if dest_lang_tag != '':
        styles_element = docxdoc.styles.element
        try:
            # Some office suite like WPS does not handle language tag in a document, ignore it
            rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
            lang_default = rpr_default.xpath('w:lang')[0]
            lang_default.set(docx.oxml.shared.qn('w:val'),dest_lang_tag)
        except Exception:
            # Ignore the language tag of the document, it is not supported by some office suites
            pass

    # Create Right to Left Style if it is not found
    try:
        rtlstyle = styles['rtlstyle']
    except Exception:
        rtlstyle = docxdoc.styles.add_style('rtlstyle', WD_STYLE_TYPE.CHARACTER)
    if dest_lang == "" or dest_lang is None:
        dest_lang_name_from_cell = docxdoc.tables[0].cell(1, 2).text
        print("Lang cell: %s" % (dest_lang_name_from_cell))
        for lang_code, lang_name in google_translate_lang_codes.items():
            #print("%s : %s" % (lang_code, lang_name))
            if dest_lang_name_from_cell == lang_name:
                dest_lang = lang_code

print("File: %s" %(args.docxfile))
print("Language code: %s" %(dest_lang))
print("Language name: %s" %(dest_lang_name))

print("Destination font: %s" %(dest_font))
print("Split: %s" %(split_translation))
print("Splitonly: %s" %(splitonly))


tmx_file_path = "%s\%s_%s.tmx" % (os.path.dirname(word_file_to_translate), os.path.splitext(os.path.basename(word_file_to_translate))[0],dest_lang)
#print(tmx_file_path)


print("Extension: %s" % (word_file_to_translate_extension))

if not os.path.exists(word_file_to_translate):
    print("ERROR: docxfile '%s' not found, exiting." % (word_file_to_translate))


if word_file_to_translate_extension != ".docx":
    print("ERROR: not a docx file. Please convert to docx first then run on docx file. Exiting.")
    if not silent:
        input("Enter to close program")
    else:
        print("Program ended with errors")
    os._exit(3)

print("")


location_primary_country_checker_url_key = ["local_configuration", "json_filename_path"]
location_primary_country_checker_url_key = ["location", "primary_country_checker_url"]
location_primary_country_checker_url = get_nested_value_from_json_array(json_configuration_array, location_primary_country_checker_url_key)

location_secondary_country_checker_url_key = ["location", "secondary_country_checker_url"]
location_secondary_country_checker_url = get_nested_value_from_json_array(json_configuration_array, location_secondary_country_checker_url_key)

location_http_query_timeout_key = ["location", "http_query_timeout"]
location_http_query_timeout = get_nested_value_from_json_array(json_configuration_array, location_http_query_timeout_key)

# Check if location_http_query_timeout is not an integer > 0
if not isinstance(location_http_query_timeout, int) or location_http_query_timeout <= 0:
    location_http_query_timeout = COUNTRY_QUERY_HTTP_TIMEOUT  # Set to 3 if the condition is not met

chrome_driver_restricted_countries_key = ["chrome_driver", "restricted_countries"]
chrome_driver_restricted_countries = get_nested_value_from_json_array(json_configuration_array, chrome_driver_restricted_countries_key)

chrome_driver_mirror_url_key = ["chrome_driver", "mirror_url"]
chrome_driver_mirror_url = get_nested_value_from_json_array(json_configuration_array, chrome_driver_mirror_url_key)

#print(f"location_primary_country_checker_url = {location_primary_country_checker_url}")
#print(f"location_secondary_country_checker_url = {location_secondary_country_checker_url}")
#print(f"chrome_driver_restricted_countries = {chrome_driver_restricted_countries}")
#print(f"chrome_driver_mirror_url = {chrome_driver_mirror_url}")

def fetch_country_data(url):
    """Fetch country data from the specified URL."""
    try:
        response = requests.get(url, timeout=location_http_query_timeout)
        response.raise_for_status()  # Check if the request was successful (status code 200)
        
        # Parse the JSON response
        data = response.json()
        
        # Check if the status is success and return the country name
        if data.get("status") == "success":
            return data.get('country')
        else:
            print(f"Failed to retrieve IP information: {data.get('message')}")
            return None
            
    except requests.exceptions.RequestException as e:
        print(f"HTTP request failed: {e}")
    except json.JSONDecodeError:
        print("Failed to parse the JSON response.")
    return None

def check_mirror_url(url):
    """Check if the mirror URL responds with HTTP 200 or 400 status codes."""
    try:
        response = requests.get(url, timeout=location_http_query_timeout)
        return response.status_code in [200, 400]
    except requests.exceptions.RequestException as e:
        print(f"Mirror URL check failed: {e}")
        return False

def set_SE_DRIVER_MIRROR_URL_if_needed(country_name, mirror_url):
    """Set the SE_DRIVER_MIRROR_URL environment variable if the country is restricted and mirror URL is valid."""
    if country_name in chrome_driver_restricted_countries:
        print(f"The host country ({country_name}) is restricted from downloading Google Chrome Driver, using proxy to bypass restrictions...")
        
        # Check the mirror URL and set environment variable if it responds with HTTP 200 or 400
        if check_mirror_url(mirror_url):
            os.environ['SE_DRIVER_MIRROR_URL'] = mirror_url
            print(f"SE_DRIVER_MIRROR_URL set to: {os.environ['SE_DRIVER_MIRROR_URL']}")
        else:
            print(f"Mirror URL ({mirror_url}) did not respond with HTTP 200 or 400.")
    else:
        print(f"Using Google Chrome Driver from {country_name}...")


# Set chrome driver download proxy URL for restricted countries
country_name = fetch_country_data(location_primary_country_checker_url)

# If primary URL fails or does not return a valid country name, fallback to the secondary URL
if not country_name:
    print("Falling back to secondary URL...")
    country_name = fetch_country_data(location_secondary_country_checker_url)

# Set environment variable if needed
set_SE_DRIVER_MIRROR_URL_if_needed(country_name, chrome_driver_mirror_url)

# Set up Chrome options
# Set the user-data-dir to the parent of the profiles

#chrome_options.add_argument(f"--user-data-dir={user_data_dir}") 
#chrome_options.add_argument(r'--profile-directory=Default')


user_data_dir = fr"C:\Temp\Chrome"
# Set the user-data-dir to the parent of the profiles



chrome_options = Options()
chrome_options.add_argument("--disable-web-security")
chrome_options.add_argument("--disable-xss-auditor")
chrome_options.add_argument("--lang=en-GB")
#chrome_options.add_argument("--verbose")
chrome_options.add_argument("--log-level=3")  # fatal
chrome_options.add_argument("--password-store=basic")


if  translation_engine.lower() == "chatgpt" and False:
    print(f"Using Chrome profile")
    print(f"Using user data dir: {user_data_dir}")
    chrome_options.add_argument(f"--user-data-dir={user_data_dir}") 
    chrome_options.add_argument(r'--profile-directory=Default')

#chrome_options.add_argument("load-extension=C:\\Users\Patriot\\AppData\\Local\\Google\\Chrome\\User Data\\Default\\Extensions\\mooikfkahbdckldjjndioackbalphokd\\3.17.0_0")

if not showbrowser :
    chrome_options.add_argument("--headless")
    if platform.system() == "Linux":  # Linux
        chrome_options.add_argument("--disable-gpu")         # remove GPU fallback flutters
        chrome_options.add_argument("--disable-software-rasterizer")
        chrome_options.add_argument("--enable-features=UseOzonePlatform")
        chrome_options.add_argument("--use-system-clipboard")
        chrome_options.add_argument("--disable-features=site-per-process")
        chrome_options.add_argument("--allow-running-insecure-content")
        chrome_options.add_argument("--disable-extensions")
        chrome_options.add_argument("--remote-allow-origins=*")
        chrome_options.add_argument("--window-size=1920,1080")
elif platform.system() == "Linux":  # Linux    
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-backgrounding-occluded-windows")
    chrome_options.add_argument("--disable-features=OptimizationGuideModelDownloading")
    chrome_options.add_argument("--disable-infobars")

if translation_engine.lower() == "chatgpt" and False:
    # Get the Windows username
    try:
        username = os.getlogin()
        home_dir = os.path.expanduser("~")

        # Construct the Chrome user data directory path
        user_data_dir = os.path.join(home_dir, "AppData", "Local", "Google", "Chrome", "User Data")
        
        user_data_dir = fr"C:\Temp\Chrome"
        print(f"Using Chrome user data directory: {user_data_dir}")

        # Set up ChromeOptions
        print(f"Using Chrome user data directory: {user_data_dir}")
        chrome_options.add_argument(f"--user-data-dir={user_data_dir}")  # Path to the user data directory
        chrome_options.add_argument("--disable-blink-features=AutomationControlled")
        chrome_options.add_argument("--profile-directory=Default")  # Use the "Default" profile

        print(f"Using Chrome user data directory: {user_data_dir}")
    #word_file_to_translate = r'X:\travail\smtv-hindi\NWN 584 sf2 - table fix1.doc'
    except Exception:
        var = traceback.format_exc()
        print(var)
        print("Failed to add chrome options")


# Used to tokenize thai
#thai_segmenter = thai_tokenizer_tokenizer()
#word_tokenize(text)

#translator = Translator(service_urls=['translate.google.com'], user_agent='Mozilla/5.0 (Windows NT 6.1; Win64; x64; rv:47.0) Gecko/20100101 Firefox/47.0')
#user_agent='Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36')

#driver = webdriver.Chrome()
#driver.set_window_position(0, 350)
#driver.set_window_size(400, 650)

#driver.get("https://translate.google.com/#%s/%s/" % (src_lang,dest_lang))
#driver = uc.Firefox()

def get_translated_cells_content(lineno, to_translate):
    print("get_translation for line %d" % (lineno))
    print("from_text_nb_lines_in_phrase %d" % (from_text_nb_lines_in_phrase[lineno]))
    translation = ""

    if dest_lang.lower() == 'ja' or dest_lang.lower() == 'zh-cn' or dest_lang.lower() == 'zh-tw' or dest_lang.lower() == 'ko':
        cell_space = ''
    else:
        cell_space = ' '

    last_row_n = from_text_nb_lines_in_phrase[lineno] + lineno
    for row_n in range(lineno, last_row_n):
        #cell_text = docxdoc.tables[0].cell(row_n, 2).text
        cell_text = table_cells[row_n][2].text
        cell_text = cell_text.strip()
        if cell_text != "":
            print("adding cell %s " % (cell_text))
            if row_n == lineno:
                translation = cell_text
            else:
                translation = translation + cell_space + cell_text
    return translation

found_google_cookies_consent_button = False
google_translate_first_page_loaded = False

def selenium_chrome_translate_get_from_text_array(to_translate, index):
    # Guard: pre-refactor code raised IndexError here whenever the block
    # loop produced fewer translated lines than the document has phrases
    # (off-by-one or alignment mismatch). Returning '' instead lets the
    # caller log "Error translating='…'" and continue, rather than
    # killing the whole job.
    if 0 < index <= len(translation_array):
        return translation_array[index - 1]
    print(f"[WARN] translation_array index out of range: index={index}, len={len(translation_array)}")
    return ""
    
def selenium_chrome_google_translate_text_file(ctx: RuntimeContext, text_file_path):
    driver = ctx.browser.driver  # Phase H: bind active browser handle
    try:

        if not ctx.browser.google_translate_first_page_loaded:
            selenium_chrome_google_click_cookies_consent_button(ctx)

        driver.get("https://translate.google.com/?sl=%s&tl=%s&op=docs" % (ctx.language.src_lang,ctx.language.dest_lang))
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        
        browse_file_element_xpath = "//label[contains(.,'Browse your computer')]"
        
        #browse_file_element = WebDriverWait(driver, 25).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))))
        
        #safe_click(driver, browse_file_element)
        
        # Waiting for URL : https://translate.googleusercontent.com/translate_f
        
        print("Selecting file %s for uploading..." % (text_file_path))
        text_file_element_xpath = "//input[@name='file']"
        text_file_element_xpath = "//input[@id='i37']"
        text_file_element_xpath = "//div[3]/input"
        text_file_element = WebDriverWait(driver, 925).until(EC.presence_of_element_located((By.XPATH, text_file_element_xpath)))
        
        text_file_element.send_keys(text_file_path)

        #text_file_translate_button_xpath = "//div[2]/div[2]/button/span"
        text_file_translate_button_xpath = "//div[2]/div/button/span"
        
        
        
        text_file_translate_button = WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.XPATH, text_file_translate_button_xpath)))
        
        print("Clicking on Translate button...")
        safe_click(driver, text_file_translate_button)
        
        # Wait for result text translation page to be loaded
        loop_wait_translation_count = 200
        
        print("Waiting for translation result...")
        while(driver.current_url != 'https://translate.googleusercontent.com/translate_f' and loop_wait_translation_count > 0):
            time.sleep(0.1)
            loop_wait_translation_count = loop_wait_translation_count - 1
        
        #Wait for page status loaded to be complete
        WebDriverWait(driver, 15).until(lambda driver: driver.execute_script('return document.readyState') == 'complete')
        
        print("Translation received.")
        
        #print("loop_wait_translation_count: %d" % (200 - loop_wait_translation_count))
        #print("URL: %s" % (driver.current_url))
        
        html_translation = driver.page_source
        text_translated_document_str = html_translation.replace('<html><head><meta charset="UTF-8"></head><body><pre>', '')
        text_translated_document_str = text_translated_document_str.replace('</pre></body></html>', '')
        text_translated_document_str = html.unescape(text_translated_document_str)
        
        ctx.docx.translation_array = text_translated_document_str.split('\n')
        
        text_translated_document_str_nb_lines = len(ctx.docx.translation_array)
        
        #print ("text_translated_document_str_nb_linestext_translated_document_str_nb_lines: %s" % text_translated_document_str_nb_lines)
        print ("docxfile_table_number_of_phrases: %s" % ctx.docx.docxfile_table_number_of_phrases)
        
        if ctx.docx.docxfile_table_number_of_phrases == text_translated_document_str_nb_lines:
            #print("OK, we got the right number of translated lines !")
            pass
        else:
            print("oups ! we got %s translated lines out of %s" % (text_translated_document_str_nb_lines, ctx.docx.docxfile_table_number_of_phrases))
            translation_succeded = False
        
        #print("text_translated_document_str:")
        #print(text_translated_document_str)
        
    except Exception:
        print("Error getting google translation from text file.")
        var = traceback.format_exc()
        print(var)
        sys.exit(7)
    return ctx.docx.translation_array
    
    
def selenium_chrome_google_translate_html_javascript_file(ctx: RuntimeContext, html_file_path):
    # Phase H: bind the active browser driver into a local name. The
    # function reassigns `driver` later (line ~1485, recovery branch),
    # which makes Python treat `driver` as local for the entire body —
    # without this seed read, every earlier `driver.get(...)` raised
    # UnboundLocalError. After a successful reassign we mirror the new
    # handle back to ctx.browser.driver so downstream calls see it.
    driver = ctx.browser.driver

    html_file_path_escaped = html_file_path.replace('#','%23')
    file_url = 'file://' + html_file_path_escaped

    nb_retry = 3
    
    while nb_retry > 0:
        nb_retry = nb_retry -1
        try:
            driver.get(file_url)
            
            print("Reading translation")

            try:
                scrollHeight = driver.execute_script("return document.body.scrollHeight")
                innerHeight = driver.execute_script("return window.innerHeight")
                bar = progressbar.ProgressBar(max_value=scrollHeight)
                bar.update(0)
            except Exception:
                var = traceback.format_exc()
                print(var)

            paragraphs = driver.find_elements(by=By.XPATH, value='//p[@class="translation"]')

            # PROGRESS milestones for the Google-javascript path. The block
            # loop in runner.py emits 25 / 50 / 75 by block, but this code
            # path never goes through the runner — without these emits the
            # UI would jump from "10" (backend started) straight to "100".
            _gj_total = max(1, len(paragraphs))
            _gj_progress_emitted: set = set()

            try:

                # How to detect a paragraph is translated is that it has the string below
                #translated_substring_old = '<font style="vertical-align: inherit;">'
                #translated_substring_new = '<font dir="auto" style="vertical-align: inherit;"><font dir="auto" style="vertical-align: inherit;">'
                re_translated_substring = re.compile('^[ \t\r\n]{0,}<font ')
                scroll_offset_paragraph = 60

                for index, paragraph in enumerate(paragraphs, start=1):
                    # Emit PROGRESS markers proportional to paragraph progress.
                    _pct = int((index / _gj_total) * 100)
                    for _m in (15, 30, 50, 75, 90):
                        if _pct >= _m and _m not in _gj_progress_emitted:
                            print(f"PROGRESS:{_m}", flush=True)
                            _gj_progress_emitted.add(_m)
                
                    #input("Time out here next line ")
                    viewport_top = driver.execute_script("return window.pageYOffset;")
                    viewport_bottom = viewport_top + driver.execute_script("return window.innerHeight;")

                    # Get the coordinates of the element
                    element_top = paragraph.location['y']
                    element_bottom = element_top + paragraph.size['height']
                    
                    paragraph_html = paragraph.get_attribute('innerHTML')
                    #print(f"{paragraph_html}")
                    
                    location = paragraph.location
                    x = location['x']
                    y = location['y']
                    scroll_position = location['y'] - scroll_offset_paragraph
                    
                    wait_translation_sleep_sec = 0.05
                    
                    # or viewport_top <= element_bottom <= viewport_bottom
                    if viewport_top <= element_top <= viewport_bottom:
                        #print("The element is displayed at the current scroll position.")
                        pass
                    else:
                        #print("The element is not displayed at the current scroll position.")
                        try:
                            driver.execute_script(f"window.scrollTo(0, {scroll_position});")
                            time.sleep(wait_translation_sleep_sec)
                            #print("The element should NOW be displayed at the current scroll position.")
                        except Exception as e:
                            #Ignore and continue if there is an error
                            #print(f"Error scrolling paragraph {index}: {str(e)}")
                            pass
                    
                    # Wait until the translation is available        
                    
                    wait_translation_max_sleep_sec = 30
                    loop_wait_translation_count = wait_translation_max_sleep_sec / wait_translation_sleep_sec
                    match_translated_tag = re_translated_substring.match(paragraph_html)
                    while not match_translated_tag and loop_wait_translation_count > 0:
                        #print(f"Sleeping in Paragraph {index}")
                        time.sleep(wait_translation_sleep_sec)
                        paragraph_html = paragraph.get_attribute('innerHTML')
                        match_translated_tag = re_translated_substring.match(paragraph_html)
                        #print(f"\n{paragraph_html}\n")
                        loop_wait_translation_count = loop_wait_translation_count - 1
                    
                    try:
                        bar.update(scroll_position + scroll_offset_paragraph)    
                    except Exception:  # Ignore progressbar errors at the end
                        pass
            except Exception:
                var = traceback.format_exc()
                print(var)
            
            bar.update(scrollHeight)
            progressbar.streams.flush()
            bar.finish()
            
            # Read translation from HTML
            html_translation = driver.page_source
            #soup = BeautifulSoup(html_translation)
            soup = BeautifulSoup(html_translation, features="lxml")
            pTags = soup.find_all('p', {'class':"translation"})
            ctx.docx.translation_array = []
            for pTranstlation in pTags:
                pData = pTranstlation.text
                if ctx.language.dest_lang.lower() == 'fa':
                   pData =  my_hazm_normalizer.normalize(text=pData)
                ctx.docx.translation_array.append(pData)

            return (ctx.docx.translation_array)
  
        except Exception:
            var = traceback.format_exc()
            print(var)
                        
            
            print("Here do something exit with session failed ")
                
            ctx.browser.chrome_options = Options()
            ctx.browser.chrome_options.add_argument("--disable-web-security")
            ctx.browser.chrome_options.add_argument("--disable-xss-auditor")
            ctx.browser.chrome_options.add_argument("--log-level=3")  # fatal
            ctx.browser.chrome_options.add_argument("--lang=en-GB")
            ctx.browser.chrome_options.add_argument("--password-store=basic")
            
            if not ctx.flags.showbrowser:
                ctx.browser.chrome_options.add_argument("--headless")
            
            docxfile_table_number_of_lines = numrows
            if ctx.flags.use_api or ctx.flags.splitonly:
                print("\nCreating a new browser for stats")
                                                      
                service = Service()
                driver = webdriver.Chrome(service=service, options=ctx.browser.chrome_options)
                ctx.browser.driver = driver  # mirror new handle back into ctx





    
def getDownLoadedFileNameFirefox(waitTime):
    driver.execute_script("window.open()")
    WebDriverWait(driver,10).until(EC.new_window_is_opened)
    driver.switch_to.window(driver.window_handles[-1])
    driver.get("about:downloads")

    endTime = time.time()+waitTime
    while True:
        try:
            fileName = driver.execute_script("return document.querySelector('#contentAreaDownloadsView .downloadMainArea .downloadContainer description:nth-of-type(1)').value")
            if fileName:
                return fileName
        except Exception:
            pass
        time.sleep(2)
        if time.time() > endTime:
            break


# method to get the downloaded file name
def getDownLoadedFileNameChrome(waitTime):
    driver.execute_script("window.open()")
    # switch to new tab
    driver.switch_to.window(driver.window_handles[-1])
    # navigate to chrome downloads
    driver.get('chrome://downloads')
    # define the endTime
    endTime = time.time()+waitTime
    while True:
        try:
            # get downloaded percentage
            downloadPercentage = driver.execute_script(
                "return document.querySelector('downloads-manager').shadowRoot.querySelector('#downloadsList downloads-item').shadowRoot.querySelector('#progress').value")
            # check if downloadPercentage is 100 (otherwise the script will keep waiting)
            if downloadPercentage == 100:
                # return the file name once the download is completed
                return driver.execute_script("return document.querySelector('downloads-manager').shadowRoot.querySelector('#downloadsList downloads-item').shadowRoot.querySelector('div#content  #file-link').text")
        except Exception:
            pass
        time.sleep(1)
        if time.time() > endTime:
            break


# function to wait for download to finish and then rename the latest downloaded file
def get_last_downloaded_file_path():
    # function to wait for all chrome downloads to finish
    def chrome_downloads(drv):
        if not "chrome://downloads" in drv.current_url: # if 'chrome downloads' is not current tab
            drv.execute_script("window.open('');") # open a new tab
            drv.switch_to.window(driver.window_handles[1]) # switch to the new tab
            drv.get("chrome://downloads/") # navigate to chrome downloads
        dld_file_paths = drv.execute_script("""
            return document.querySelector('downloads-manager')
            .shadowRoot.querySelector('#downloadsList')
            .items.filter(e => e.state === 'COMPLETE')
            .map(e => e.filePath || e.file_path || e.fileUrl || e.file_url);
            """)
        print("dld_file_paths=%s" % (dld_file_paths))
        #input("dld_file_paths press enter")
        return dld_file_paths
    # wait for all the downloads to be completed
    dld_file_paths = []
    while len(dld_file_paths) == 0:
        dld_file_paths = chrome_downloads(driver)
        print("len dld_file_paths=%d" % (len(dld_file_paths)))
        print("res dld_file_paths=%s" % (dld_file_paths))
        #input("Opened download status page")
        #WebDriverWait(driver, 120, 1).until(chrome_downloads) # returns list of downloaded file paths)
    # Close the current tab (chrome downloads)
    if "chrome://downloads" in driver.current_url:
        driver.close()
    # Switch back to original tab
    driver.switch_to.window(driver.window_handles[0]) 
    # get latest downloaded file name and path
    dlFilename = dld_file_paths[0] # latest downloaded file from the list
    # wait till downloaded file appears in download directory
    time_to_wait = 20 # adjust timeout as per your needs
    time_counter = 0
    while not os.path.isfile(dlFilename):
        print ("We have dlFilename=%s" % (dlFilename))
        time.sleep(5)
        time_counter += 1
        if time_counter > time_to_wait:
            break
    # rename the downloaded file
    print("dlFilename=%s" %(dlFilename))
    return dlFilename
    #shutil.move(dlFilename, os.path.join(download_dir,newFilename))
    return

def selenium_chrome_google_translate_xlsx_file(ctx: RuntimeContext, xlsx_file_path):
    driver = ctx.browser.driver  # Phase H: bind active browser handle
    try:

        if not ctx.browser.google_translate_first_page_loaded:
            selenium_chrome_google_click_cookies_consent_button(ctx)

        driver.get("https://translate.google.com/?sl=%s&tl=%s&op=docs" % (ctx.language.src_lang,ctx.language.dest_lang))
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        
        browse_file_element_xpath = "//label[contains(.,'Browse your computer')]"
        
        #browse_file_element = WebDriverWait(driver, 25).until(EC.presence_of_element_located((By.XPATH, browse_file_element_xpath)))))
        
        #safe_click(driver, browse_file_element)
        
        # Waiting for URL : https://translate.googleusercontent.com/translate_f
        
        print("Selecting file %s for uploading..." % (xlsx_file_path))
        xlsx_file_element_xpath = "//input[@name='file']"
        xlsx_file_element_xpath = "//input[@id='i37']"
        #xlsx_file_element_xpath = "//span[contains(.,'Browse your computer')]"
        #xlsx_file_element_xpath = '//button[normalize-space()="Browse your computer"]'
        xlsx_file_element_xpath = "//div[3]/input"
        xlsx_file_element = WebDriverWait(driver, 925).until(EC.presence_of_element_located((By.XPATH, xlsx_file_element_xpath)))
        
        xlsx_file_element.send_keys(xlsx_file_path)
        #input(" HERE : xlsx_file_path")

        #xlsx_file_translate_button_xpath = "//div[2]/div[2]/button/span"
        xlsx_file_translate_button_xpath = "//div[2]/div/button/span"
        
        xlsx_file_translate_button = WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.XPATH, xlsx_file_translate_button_xpath)))
        
        print("Clicking on Translate button...")
        #input("BEFORE : xlsx_file_translate_button")
        safe_click(driver, xlsx_file_translate_button)
        #input(" HERE : xlsx_file_translate_button")
        
        # Wait for result text translation page to be loaded
        loop_wait_translation_count = 200
        
         
        res_downloaded_xlsx_translation  = False
        print("Waiting for translation result...")
        while(driver.current_url != 'https://translate.googleusercontent.com/translate_f' and loop_wait_translation_count > 0 and res_downloaded_xlsx_translation == False):
            time.sleep(0.1)
            loop_wait_translation_count = loop_wait_translation_count - 1
                
            if ("https://www.google.com/sorry" in driver.current_url):
                print("We found a CAPTCHA window")
                if not silent:
                    input("Press enter after solving CAPTCHA")
                else:
                    # In silent mode (e.g. spawned by local_launcher.py)
                    # there's no user to solve the CAPTCHA — fail fast
                    # instead of hanging the launcher subprocess pipe.
                    raise RuntimeError("Google CAPTCHA encountered in silent mode — cannot proceed without user interaction")
                
            download_button_xpath = '//button[normalize-space()="Download translation"]'
            download_button_xpath = '//div[2]/div/button/span'
            download_button_xpath = '//div/button[2]/span[2]'
            
            
            try:
                download_button = WebDriverWait(driver, 0.1).until(EC.presence_of_element_located((By.XPATH, download_button_xpath)))
                safe_click(driver, download_button)
                print("We found a download button")
                print("Waiting for download to finish")
                #xlsx_tranlated_path = getDownLoadedFileNameChrome(15)
                downloaded_xlsx_translation_path = get_last_downloaded_file_path()
                if len(downloaded_xlsx_translation_path) > 0:
                    print("downloaded_xlsx_translation_path=%s" %(downloaded_xlsx_translation_path))
                    res_downloaded_xlsx_translation = True
                
            except Exception:
                pass
        
        if res_downloaded_xlsx_translation:
            print("Translation xlsx file downloaded at : %s" %(downloaded_xlsx_translation_path))
            #input("Press enter")
        
        #input("After download button")
        #Wait for page status loaded to be complete
        WebDriverWait(driver, 10).until(EC.presence_of_element_located(driver.execute_script('return document.readyState') == 'complete'))
        
        print("Translation received.")
        #input("HERE AGAIN .")
        
        #print("loop_wait_translation_count: %d" % (200 - loop_wait_translation_count))
        #print("URL: %s" % (driver.current_url))
        
        html_translation = driver.page_source
        print("html_translation")
        print(html_translation)
        print("\n________________\nhtml_to_text")
        print("BeautifulSoup to text:")
        soup = BeautifulSoup(html_translation, features="lxml")
        soup = BeautifulSoup(html_translation)
        tdTag = soup.find_all("td")
        ctx.docx.translation_array = []
        for td in tdTag:
            pData = td.text
            ctx.docx.translation_array.append(td.text)
            #res = soup.get_text()
            print(pData)
        #input("after pData")
        print(ctx.docx.translation_array)
        print("__________________________")
        
        # text_translated_document_str = html_translation.replace('<html><head></head><body><pre>', '')
        # text_translated_document_str = text_translated_document_str.replace('</pre></body></html>', '')
        # text_translated_document_str = html.unescape(text_translated_document_str)
        
        #ctx.docx.translation_array = text_translated_document_str.split('\n')
        
        text_translated_document_str_nb_lines = len(ctx.docx.translation_array)
        
        #print ("text_translated_document_str_nb_linestext_translated_document_str_nb_lines: %s" % text_translated_document_str_nb_lines)
        print ("docxfile_table_number_of_phrases: %s" % ctx.docx.docxfile_table_number_of_phrases)
        
        if ctx.docx.docxfile_table_number_of_phrases == text_translated_document_str_nb_lines:
            #print("OK, we got the right number of translated lines !")
            pass
        else:
            print("oups ! we got %s translated lines out of %s" % (text_translated_document_str_nb_lines, ctx.docx.docxfile_table_number_of_phrases))
            translation_succeded = False
        
        #print("text_translated_document_str:")
        #print(text_translated_document_str)
        
    except Exception:
        print("Error getting google translation from text file.")
        var = traceback.format_exc()
        print(var)
        sys.exit(8)
    return ctx.docx.translation_array



def remove_span_tag(text):
    search_opening_html_span_tag = r'(?i)<span class="[a-zA-Z]+">'
    search_replace_opening_span = re.compile(search_opening_html_span_tag)
                
    subn_result = search_replace_opening_span.subn("", text)
    subn_count = subn_result[1]
    if subn_count > 0:
        #print ("Replaced '%s' by '%s' %d times." % (search_opening_html_span_tag, "", subn_count))
        text = subn_result[0]
        #if subn_count > 0:
        #    print ("Replaced span %d times" % (subn_count))
            
    search_closing_html_span_tag = r'(?i)</span>'
    search_replace_closing_span = re.compile(search_closing_html_span_tag)
                
    subn_result = search_replace_closing_span.subn("", text)
    subn_count = subn_result[1]
    if subn_count > 0:
        #print ("Replaced '%s' by '%s' %d times." % (search_opening_html_span_tag, "", subn_count))
        text = subn_result[0]
        #if subn_count > 0:
        #    input ("Replaced span %d times" % (subn_count))
                
    return text



def selenium_chrome_perplexity_wait_log_in():
    """Phase F1.2: drops vestigial global declarations — neither
    json_configuration_array nor MAX_TRANSLATION_BLOCK_SIZE is read or
    written in this function body. (The only caller is currently
    commented-out in main().)"""
    driver.set_window_size(600, 600)
    #driver.maximize_window()

    loop_count = 200
    sleep_wait_sec = 5

    while True:
        try:
            driver.get("https://www.perplexity.ai/")
            
            # Wait up to 10 seconds for the signed-in avatar to appear
            WebDriverWait(driver, 3).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, '[data-testid="sidebar-popover-trigger-signed-in"]'))
            )
            print("✅ User is logged in perplexity.")
            return True

        except Exception:
            var = traceback.format_exc()
            print(var)
            
    return False



# build_translation_prompt was deleted in C3 of the 2026-05-10
# cleanup pass — its only callers were the now-deleted chatgpt-web /
# perplexity-web engines.


# selenium_webservice_perplexity_translate was extracted to
# ``src/engines/perplexity_webservice.py`` in C3.1 of the 2026-05-10
# architecture cleanup. Imported here so the runner-injection call
# site (and any historical caller) keeps resolving.
from engines.perplexity_webservice import selenium_webservice_perplexity_translate


# set_translation_function was extracted to ``src/dispatch.py`` in the
# 2026-05-10 architecture cleanup. The dispatch module is the single
# source of truth for engine routing — both ``set_translation_function``
# AND the ``use_phrasesblock`` predicate live there, so future drift
# between them is impossible.
from dispatch import set_translation_function, use_phrasesblock as _dispatch_use_phrasesblock
import dispatch as _dispatch_module
_dispatch_module.set_array_dispatcher(selenium_chrome_translate_get_from_text_array)


def selenium_chrome_machine_translate(ctx: RuntimeContext, to_translate, index):
    """Per-call wrapper around ``ctx.engine.dispatcher`` with retry logic.

    Threaded in Phase F1.4: reads the dispatcher pointer, the active
    engine + method, the retry counter, and the inter-retry sleep
    factor from ``ctx`` instead of module globals.
    """
    translation = ""
    translation_try_count = 1
    max_try_count = 15
    try:
        while translation_try_count < max_try_count and translation == "":
            if translation_try_count > 1:
                print("Retrying to translate again (%d)..." % (translation_try_count))
                ctx.docx.translation_errors_count += 1
                ctx.browser.deepl_sleep_wait_translation_seconds *= 1.1
                print("%d translation retry so far..." % (ctx.docx.translation_errors_count))
            if ctx.engine.engine == 'deepl':
                if ctx.engine.method == 'phrasesblock':
                    translation = ctx.engine.dispatcher(to_translate, index)
                else:
                    translation = ctx.engine.dispatcher(to_translate, translation_try_count - 1)
            elif ctx.engine.method in (
                'textfile', 'xlsxfile', 'phrasesblock',
                'javascript', 'webservice', 'api',
            ):
                translation = ctx.engine.dispatcher(to_translate, index)
            else:
                translation = ctx.engine.dispatcher(to_translate)
            translation_try_count += 1
    except Exception:
        print("Error in selenium_chrome_machine_translate function.")
    return translation
    
def initialize_translation_memory_xlsx(ctx: RuntimeContext):
    global xtm
    # If --xlsxreplacefile was provided in the command line
    if xlsxreplacefile is not None:
        print("xlsxreplacefile: %s" % (xlsxreplacefile))
        xtm = xlsx_translation_memory.xlsx_translation_memory(xlsxreplacefile)
        print("")
    else:
        xtm = xlsx_translation_memory.xlsx_translation_memory(None)


def is_end_of_line(line):
    for eol in eol_array:
        #print("Testing is_end_of_line '%s' on string '%s'" %(eol, line))
        if re.search(eol, line):
            #print("Found is_end_of_line '%s' on string '%s'" %(eol, line))
            return 1
    return 0


def is_conditional_end_of_line(line):
    for ceol in eol_conditional_array:
        #print("Testing is_conditional_end_of_line '%s' on string '%s'" %(ceol, line))
        if re.search(ceol, line):
            return 1
    return 0

def is_beginning_of_line(line):
    for bol in bol_array:
        if re.search(bol, line):
            return 1
    return 0

def is_empty_line(line):
    line_trimmed = re.sub(' +', '', line)
    length = len(line_trimmed)
    if length == 0:
        return 1
    return 0

def get_paragraph_shading_color(xml_paragraph_str):
    paragraph_xml = etree.fromstring(xml_paragraph_str)
    attrib_fill = None
    
    namespaces = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    try:
        namespaces = {paragraph_xml.prefix : paragraph_xml.nsmap[paragraph_xml.prefix]}
    except Exception:  #print("Could not determine namespace")
        pass
    attrib_fill = None
    
    for e in paragraph_xml.findall('.//w:pPr/w:shd', namespaces):
        #print("e:", etree.tostring(e, pretty_print=True))
        try:
            attrib_val = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            attrib_color = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
            attrib_fill = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
            #print(f"attrib_color : {attrib_color}")
            #print(f"attrib_fill : {attrib_fill}")
            #print(f"attrib_val : {attrib_val}")
        except Exception:
            pass
    return attrib_fill


def get_run_shading_color(xml_run_str):
    run_xml = etree.fromstring(xml_run_str)
    
    namespaces = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    try:
        namespaces = {run_xml.prefix : run_xml.nsmap[run_xml.prefix]}
    except Exception:  #print("Could not determine namespace")
        pass
    attrib_fill = None
    
    for e in run_xml.findall('.//w:rPr/w:shd', namespaces):
        #print("e:", etree.tostring(e, pretty_print=True))
        try:
            attrib_val = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            attrib_color = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
            attrib_fill = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        except Exception:
            pass
    return attrib_fill

# Return cell_non_greyed_text (string), cell_is_gray (integer for boolean)
# _iter_paragraph_runs was extracted to src/docx_io/runs.py in the
# 2026-05-10 docx_io extraction pass. Imported here so callers in
# this file (the get_cell_data path) keep resolving the name.
from docx_io import _iter_paragraph_runs


def get_cell_data(ctx: RuntimeContext, cell,row_n):
    cell_is_gray = None
    cell_is_red = None
    cell_non_greyed_text = ''

    re_enter = re.compile('enter')
    re_newline = re.compile('\n')

    n_paragraph = 0
    n_cell_lines = 1


    for paragraph in cell.paragraphs:
        paragraphs_text = ""
        n_paragraph = n_paragraph + 1

        #print("paragraph:", paragraph._p.xml)

        root = etree.fromstring(paragraph._p.xml)
        p_shading_color = get_paragraph_shading_color(paragraph._p.xml)

        # Materialise the run list once so both p_text (used for
        # <pause> / <enter> counting) and the run-by-run loop below
        # see the same source of truth — including hyperlinked text.
        # Relying on paragraph.text would tie us to a python-docx
        # implementation detail that has flipped behaviour across
        # versions on whether <w:hyperlink> contents are included.
        paragraph_runs = list(_iter_paragraph_runs(paragraph))
        p_text = ''.join(r.text for r in paragraph_runs)
        nb_pause = len(re.findall('(?i)(<pause>)', p_text))
        nb_enter = len(re.findall('(?i)(<enter>)', p_text))

        n_cell_lines = n_cell_lines + nb_pause + nb_enter

        if p_shading_color is not None:
            #print(paragraph.text)
            #input("Found a shaded paragraph")
            if p_shading_color in shading_color_ignore_text:
                continue

        #if n_paragraph > 1:
        #    print("paragraph %d" % (n_paragraph))
        previous_run_text = ""
        # Walk every <w:r> below the paragraph, including those nested
        # inside <w:hyperlink>. Using paragraph.runs alone drops the
        # text of every clickable link in the document.
        for run in paragraph_runs:
            current_run_text = run.text
            
            #print("cell row %d has %d runs," % (row_n, len(paragraph.runs) ))
            #print(f"current_run_text : '{current_run_text
            
            root = etree.fromstring(run.element.xml)
            run_shading_color = get_run_shading_color(run.element.xml)
            
            if run_shading_color is not None:
                #print(f"run.element.xml : {run.element.xml}")
                #print(f"current_run_text : {current_run_text}")
                #input(f"Found a shaded run {run_shading_color}")
                if run_shading_color in shading_color_ignore_text:
                    #print(f"Color {run_shading_color} in the list of colors to ignore text")
                    pass
            
            # if re_enter.match(current_run_text):
                # print("found enter")

            if str(run.font.color.rgb) == "FF0000":
                if cell_is_red == None:
                    cell_is_red = 1
            else:
                if current_run_text != "":
                    if cell_is_red == None:
                        cell_is_red = 0
                    else:
                        cell_is_red = cell_is_red * 0
                
            if run.font.highlight_color == WD_COLOR_INDEX.RED :
                pass

            if run.font.highlight_color == WD_COLOR_INDEX.GRAY_25 or run.font.highlight_color == WD_COLOR_INDEX.GRAY_50 or run.font.strike or run.font.double_strike or run.font.highlight_color == WD_COLOR_INDEX.PINK or run.font.highlight_color == WD_COLOR_INDEX.RED or run_shading_color in shading_color_ignore_text:
                #print("Found GRAY_25")
                cell_non_greyed_text = cell_non_greyed_text + ' '
                if cell_is_gray == None:
                    cell_is_gray = 1
                
            else:
                #print("Not gray")
                if current_run_text != "":
                    cell_non_greyed_text = cell_non_greyed_text + current_run_text
                    if cell_is_gray == None:
                        cell_is_gray = 0
                    else:
                        cell_is_gray = cell_is_gray * 0
                    #return cell_is_gray
            previous_run_text = current_run_text
        #if (paragraphs_text.upper() == '<ENTER>' or paragraphs_text.upper() == '<PAUSE>'):
        #    print("Found <ENTER> or <PAUSE>")
        #    #input("press enter")
        
    
    ctx.docx.from_text_nb_lines_in_cell[row_n-1] = n_cell_lines
    #if n_cell_lines > 1:
    #    print("%d lines" % (n_cell_lines))
    #    #input("here")

    cell_non_greyed_text = cell_non_greyed_text.replace('’', "'")
    cell_non_greyed_text = cell_non_greyed_text.replace("\n", " ")
    cell_non_greyed_text = cell_non_greyed_text.replace("\r", " ")
    cell_non_greyed_text = re.sub(r'[\r\n\u2028\u2029]+', ' ', cell_non_greyed_text)
    
    cell_non_greyed_text = re.sub("(?i)<pause>", "", cell_non_greyed_text) #'remove <pause> case insensitive
    cell_non_greyed_text = re.sub("(?i)<enter>", "", cell_non_greyed_text) #'remove <pause> case insensitive
    
    cell_non_greyed_text = re.sub(' +', ' ', cell_non_greyed_text)
    cell_non_greyed_text = cell_non_greyed_text.strip()
    

    #if cell_is_gray == 1:
    #    print("FOUND A GRAY CELL")
    #time.sleep(4)
    return cell_non_greyed_text, cell_is_gray, cell_is_red



def change_cell_font(cell):

    #print("cell has %d runs," % (len(paragraph[0].runs) ))
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = dest_font
    return

def join_from_lines(line_start, line_end, separator_str):
    joined_str = ""
    row_n = line_start
    joined_str = from_text_table[row_n]
    row_n = row_n + 1
    while row_n <= line_end:
        joined_str += from_text_table[row_n]
        row_n += 1
    #print "joined_str (%d, %d)=%s<br>" % (line_start, line_end, joined_str)
    return joined_str

def tokenize_text_to_array(text, lang_code):
    lang_code = lang_code + ""
    lang_code = lang_code.lower()

    words = []
    # In japanese tokenize words
    if lang_code == 'ja' or lang_code== 'zh-cn' or lang_code == 'zh' or lang_code == 'zh-tw' or lang_code == 'ko' or lang_code== 'zh-hans' or lang_code == 'zh-hant':
        words = cjk_segmenter.tokenize(text)
    # In other languages, just use spaces
    elif lang_code == 'th':
        #words = thai_segmenter(text)
        words = word_tokenize(text)
    # In other languages, just use spaces
    else:
        #xtm.tokenize_phrase(text, dest_lang)

        # search do not split here
        #xtm.pprint_translation_memory_list()

        # Old simple split method replaced by tokenize_phrase method having do not split
        # words = text.split()

        words = xtm.tokenize_phrase(text, lang_code)
        #input("Wait, remove tokenize_phrase here..")


    return words

def divide_array(words_array, dest_lang, width):
    dest_lang = dest_lang.lower()

    print("Divide into max %d size lines" % (width))
    count = len(words_array)
    offsets = [0]
    for w in words_array:
        offsets.append(offsets[-1] + len(w))

    minima = [0] + [10 ** 20] * count
    breaks = [0] * (count + 1)

    def cost(i, j):
        w = offsets[j] - offsets[i] + j - i - 1
        if w > width:
            return 10 ** 10
        return minima[i] + (width - w) ** 2

    def search(i0, j0, i1, j1):
        stack = [(i0, j0, i1, j1)]
        while stack:
            i0, j0, i1, j1 = stack.pop()
            if j0 < j1:
                j = (j0 + j1) // 2
                for i in range(i0, i1):
                    c = cost(i, j)
                    if c <= minima[j]:
                        minima[j] = c
                        breaks[j] = i
                stack.append((breaks[j], j+1, i1, j1))
                stack.append((i0, j0, breaks[j]+1, j))

    n = count + 1
    i = 0
    offset = 0
    while True:
        r = min(n, 2 ** (i + 1))
        edge = 2 ** i + offset
        search(0 + offset, edge, edge, r + offset)
        x = minima[r - 1 + offset]
        for j in range(2 ** i, r - 1):
            y = cost(j + offset, r - 1 + offset)
            if y <= x:
                n -= j
                i = 0
                offset += j
                break
        else:
            if r == n:
                break
            i = i + 1

    lines = []
    j = count
    while j > 0:
        i = breaks[j]
        # In japanese just join words_array without adding any spaces
        if dest_lang == 'ja' or dest_lang == 'zh-cn' or dest_lang == 'zh-tw' \
            or dest_lang.lower() == 'zh-hans' or dest_lang.lower() == 'zh-hant' or dest_lang == 'ko' \
                or dest_lang == 'th':
            lines.append(''.join(words_array[i:j]))
        # In other languages, join words_array using a space
        else:
            lines.append(' '.join(words_array[i:j]))

        j = i
    lines.reverse()
    return lines


def split_phrases(ctx: RuntimeContext):
    """Group consecutive cells into a single phrase keyed by the
    first cell of the phrase.

    Phase H bridge: the historical version wrote to module-level
    globals (`from_text_by_phrase_separator_table`,
    `from_text_by_phrase_table`, `from_text_nb_lines_in_phrase`)
    that were sized `[''] * 1`. Every row beyond index 0 raised
    IndexError, which the enclosing `try/except` in
    `read_and_parse_docx_document` swallowed silently — leaving the
    arrays empty so every translation block downstream became
    empty too. All reads + writes now go through ``ctx.docx`` which
    is properly sized at parse time.
    """
    docx = ctx.docx
    n_last_row_phrase = 3
    last_table_row = docx.word_translation_table_length
    cur_row_n = 2
    while cur_row_n < (last_table_row):
        if docx.from_text_nb_lines_in_cell[cur_row_n] > 1:
            pass
        if docx.from_text_is_beginning_of_line_table[cur_row_n] == 1:
            n_last_row_phrase = cur_row_n
            nb_lines_in_phrase = 1
            docx.from_text_nb_lines_in_phrase[cur_row_n] = docx.from_text_nb_lines_in_cell[cur_row_n]
            while docx.from_text_is_end_of_line_table[n_last_row_phrase] != 1 \
                and n_last_row_phrase < (last_table_row - 1):
                if docx.from_text_by_phrase_separator_table[cur_row_n] == "":
                    docx.from_text_by_phrase_separator_table[cur_row_n] = docx.from_text_table[n_last_row_phrase]
                    docx.from_text_by_phrase_table[cur_row_n] = docx.from_text_table[n_last_row_phrase]
                else:
                    docx.from_text_by_phrase_separator_table[cur_row_n] = docx.from_text_by_phrase_separator_table[cur_row_n] + line_separator_str + docx.from_text_table[n_last_row_phrase]
                    docx.from_text_by_phrase_table[cur_row_n] = docx.from_text_by_phrase_table[cur_row_n] + ' ' + docx.from_text_table[n_last_row_phrase]
                    nb_lines_in_phrase += 1
                    if docx.from_text_nb_lines_in_cell[cur_row_n] > 1:
                        pass
                    docx.from_text_nb_lines_in_phrase[cur_row_n] += docx.from_text_nb_lines_in_cell[n_last_row_phrase]
                n_last_row_phrase += 1
            if docx.from_text_by_phrase_separator_table[cur_row_n] == "":
                docx.from_text_by_phrase_separator_table[cur_row_n] = docx.from_text_table[n_last_row_phrase]
                docx.from_text_by_phrase_table[cur_row_n] = docx.from_text_table[n_last_row_phrase]
            else:
                docx.from_text_by_phrase_separator_table[cur_row_n] = docx.from_text_by_phrase_separator_table[cur_row_n] + line_separator_str + docx.from_text_table[n_last_row_phrase]
                nb_lines_in_phrase += 1
                docx.from_text_nb_lines_in_phrase[cur_row_n] += docx.from_text_nb_lines_in_cell[n_last_row_phrase]
                docx.from_text_by_phrase_table[cur_row_n] = docx.from_text_by_phrase_table[cur_row_n] + ' ' + docx.from_text_table[n_last_row_phrase]
            if use_html:
                print("(%d)from_text_by_phrase_table[%d]=%s<br>" % (n_last_row_phrase, cur_row_n, docx.from_text_by_phrase_table[cur_row_n]))
            nb_lines_in_phrase_str = "[%s]" % (nb_lines_in_phrase)

            cur_row_n = n_last_row_phrase + 1
        else:
            cur_row_n += 1

    return 0

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def generate_tmx_file():
    print("In generate_tmx_file")

    try:
        f = open(tmx_file_path, 'w', encoding='utf-8')

        # Writing TMX Header
        username = getpass.getuser()
        datenow = datetime.datetime.now()
        creation_date = "%s%0.2d%0.2dT%0.2d%0.2d%0.2dZ" % (datenow.year, datenow.month, datenow.day, datenow.hour, datenow.minute,
 datenow.second)
        header = """<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE tmx PUBLIC "-//LISA OSCAR:1998//DTD for Translation Memory eXchange//EN" "tmx14.dtd" >
<tmx version="1.4">
<header
	creationtool="SMTV translation robot"
	creationtoolversion="1.0"
	srclang="%s"
	adminlang=%s
	datatype="unknown"
	o-tmf="unknown"
	segtype="sentence"
	creationid="%s"
	creationdate="%s">
</header>
<body>\n""" % (src_lang, src_lang, username,creation_date)
        f.write(header)

        for i, line in enumerate(from_text_table):
            item = from_text_by_phrase_separator_table[i]
            item.strip()
            from_language = src_lang
            phrase_separator_removed_str = ''

            p_remove_separator = re.compile(line_separator_regex_str)
            p_remove_double_spaces = re.compile(' +')
            p_remove_parenthesis_spaces = re.compile('\( +')

            item = from_text_by_phrase_table[i]
            item_escaped = from_text_by_phrase_table[i].replace("&", "&amp;")
            item_escaped = item_escaped.replace("<", "&lt;")
            item_escaped = item_escaped.replace(">", "&gt;")

            item_translation = to_text_by_phrase_separator_table[i].replace("&", "&amp;")
            item_translation = item_translation.replace("<", "&lt;")
            item_translation = item_translation.replace(">", "&gt;")
            if item_escaped.strip() != "":
                segment = """<tu changeid="french user 1" changedate="%s" creationid="Black Mamba RS7" creationdate="%s" creationtool="SMTV translation robot" creationtoolversion="1.0.0">
<tuv xml:lang="en-US"><seg>%s</seg></tuv>
<tuv xml:lang="%s"><seg>%s</seg></tuv>
</tu>""" % (creation_date, creation_date, item_escaped, dest_lang, item_translation)
                f.write(segment)
                f.write("\n")

        # Writing TMX Footer
        footer = """</body>
</tmx>\n"""
        f.write(footer)
    except Exception:
        var = traceback.format_exc()
        print(var)


def prepare_and_clear_cell_for_writing(ctx: RuntimeContext, row_n, translation_cell_text):
    """Clear and re-init a target-language cell.

    Threaded in Phase F1.3: reads ``ctx.docx.table_cells`` in place of the
    historical ``table_cells`` global. ``dest_lang`` and ``rtlstyle`` are
    not yet on ctx; they remain as ambient module reads (closed over by
    Python's name lookup) and will be threaded in a later sub-phase.
    """
    paragraph_no = 0
    # Skip rows that don't have a third cell (footer / single-column notes
    # at the end of subtitle DOCX files). Without this guard the legacy
    # code would IndexError on every short row and spam the log.
    if row_n >= len(ctx.docx.table_cells) or len(ctx.docx.table_cells[row_n]) < 3:
        return
    current_cell = ctx.docx.table_cells[row_n][2]

    current_cell._element.clear_content()
    
    # Clear paragraphs in the cell
    for paragraph in current_cell.paragraphs:
        if paragraph_no != 0:
            delete_paragraph(paragraph)
        else:
            paragraph.text = ''
        paragraph_no += 1

    # Ensure there's at least one paragraph in the cell
    if len(current_cell.paragraphs) == 0:
        cell_paragraph = current_cell.add_paragraph("")
    else:
        cell_paragraph = current_cell.paragraphs[0]

    # Add orientation for Right-to-Left (RTL) languages
    if dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(translation_cell_text)
        run.style = rtlstyle  # Ensure `rtlstyle` exists in the document
        font = run.font
        font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = translation_cell_text
        # Check if the 'Normal' style exists before applying
        try:
            cell_paragraph.style = 'Normal'
        except KeyError:
            #print("Warning: 'Normal' style not found. Falling back to 'Default Paragraph Font'.")
            try:
                cell_paragraph.style = 'Default Paragraph Font'
            except KeyError:
                #print("Error: No usable default style found. Proceeding without style assignment.")
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Apply font changes if necessary
    if dest_font != "":
        change_cell_font(current_cell)

    ctx.docx.table_cells[row_n][2] = current_cell


def cell_set_1st_paragraph(ctx: RuntimeContext, row_n, paragraph_text):
    paragraph_no = 0
    current_cell = ctx.docx.table_cells[row_n][2]
    
    #print("cell_add_paragraph")
    #print("paragraph[%d]: %s" % (row_n,paragraph_text))
    cell_paragraph = cell_paragraph = current_cell.paragraphs[0]

    # Add orientation from Right To Left (RTL) for specific languages
    if dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(paragraph_text,style = "rtlstyle")
        run.style = rtlstyle
        font = run.font
        font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = paragraph_text
        try:
            cell_paragraph.style = 'Normal'
        except KeyError:
            #print("Warning: 'Normal' style not found. Falling back to 'Default Paragraph Font'.")
            try:
                cell_paragraph.style = 'Default Paragraph Font'
            except KeyError:
                #print("Error: No usable default style found. Proceeding without style assignment.")
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if dest_font != "":
        change_cell_font (current_cell)

    ctx.docx.table_cells[row_n][2] = current_cell


def cell_add_paragraph(ctx: RuntimeContext, row_n, paragraph_text):
    paragraph_no = 0
    current_cell = ctx.docx.table_cells[row_n][2]
    
    #print("cell_add_paragraph")
    #print("paragraph[%d]: %s" % (row_n,paragraph_text))
    cell_paragraph = current_cell.add_paragraph("")

    # Add orientation from Right To Left (RTL) for specific languages
    if dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(paragraph_text,style = "rtlstyle")
        run.style = rtlstyle
        font = run.font
        font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = paragraph_text
        try:
            cell_paragraph.style = 'Normal'
        except KeyError:
            #print("Warning: 'Normal' style not found. Falling back to 'Default Paragraph Font'.")
            try:
                cell_paragraph.style = 'Default Paragraph Font'
            except KeyError:
                #print("Error: No usable default style found. Proceeding without style assignment.")
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if dest_font != "":
        change_cell_font (current_cell)

    ctx.docx.table_cells[row_n][2] = current_cell


def read_and_parse_docx_document(ctx: RuntimeContext):
    """Parse the input DOCX into the parallel arrays on ``ctx.docx``.

    Threaded in Phase F1.3: every parallel array, table_cells, the
    table reference, and the row/column geometry move from module
    globals into ``ctx.docx``. The +1 indexing convention
    (arrays sized ``numrows + 1``, accessed at ``[i + 1]``) is
    preserved exactly. R16's structural test
    (``test_docx_arrays_plus_one_indexing``) pins this contract.
    """
    docx = ctx.docx


    start = timeit.timeit()

    if use_html:
        print("Content-Type: text/html\n")

    docx.word_translation_table_length = len(docxdoc.tables[0].rows)

    nb_tables = len(docxdoc.tables)

    nb_character_total = 0

    if use_html:
        print(
            "<!doctype html><head><meta http-equiv=""Content-Type"" content=""text/html"" charset=utf-8 /><title>Winword in python</title></head><h2>tables</h2><span style=""font-family:monospace,monospace;"">")

    # Number of tables</h2>nb_tables=", nb_tables

    numerrors = 0
    # print("docx.word_translation_table_length=%d" %(docx.word_translation_table_length))
    # print("docx_translation_table_length=%d" %(docx_translation_table_length))

    try:
        docx.table = docxdoc.tables[0]
    except Exception:
        print(f"Error: document {docxfile} does not have a table. Exiting.")
        exit(14)
    docx.table_cells = [['' for i in range(len(docx.table.columns))] for j in range(len(docx.table.rows))]

    docx.numrows = len(docx.table.rows)
    docx.numcols = len(docx.table.columns)

    if docx.numcols <= 2:
        print("ERROR : The table has %s column but expected 3" % (docx.numcols))
        print("Exiting\n")

        print("\nDeveloper: %s" %(E_mail_str))
        print("Program version: %s\n" % (PROGRAM_VERSION))
        if not silent:
            input("Enter to close program")
        else:
            print("Program ended with errors")
        sys.exit(11)

    rownum = 0

    docx.from_text_table = [''] * (docx.numrows + 1)
    docx.from_text_is_greyed_table = [0] * (docx.numrows + 1)
    docx.from_text_is_red_color_table = [0] * (docx.numrows + 1)
    docx.from_text_is_end_of_line_table = [0] * (docx.numrows + 1)
    docx.from_text_is_beginning_of_line_table = [0] * (docx.numrows + 1)
    docx.from_text_is_empty_line_table = [0] * (docx.numrows + 1)
    docx.from_text_is_conditional_end_of_line_table = [0] * (docx.numrows + 1)
    docx.from_text_by_phrase_separator_table = [''] * (docx.numrows + 1)
    docx.from_text_by_phrase_table = [''] * (docx.numrows + 1)
    #number of lines in per phrase
    docx.from_text_nb_lines_in_phrase = [0] * (docx.numrows + 1)
    docx.from_text_nb_lines_in_cell = [0] * (docx.numrows + 1)
   #input(docx.numrows)
    #
    docx.to_text_by_phrase_separator_table = [''] * (docx.numrows + 1)
    docx.to_text_by_phrase_separator_removed_table = [''] * (docx.numrows + 1)
    docx.to_text_splited_table1 = [''] * (docx.numrows + 1)
    docx.to_text_by_phrase_table = [''] * (docx.numrows + 1)
    docx.to_text_table = [''] * (docx.numrows + 1)
    docx.to_raw_translated_table = [''] * (docx.numrows + 1)
    docx.to_text_removed_line_separator = [''] * (docx.numrows + 1)
    docx.translation_result_using_separator = [''] * (docx.numrows + 1)
    # `[[]] * n` would have every slot pointing at the same shared list,
    # so any future `array[i].append(...)` would silently mutate every
    # other slot. List-comprehension gives each slot a distinct list.
    docx.translation_result_phrase_array = [[] for _ in range(docx.numrows + 1)]
    docx.translation_result = [''] * (docx.numrows + 1)
    docx.from_text_is_read = [0] * (docx.numrows + 1)

    if use_html :
        print("<br>%s rows.<br>%d colums.<br>" % (docx.numrows, docx.numcols))

    for i, row in enumerate(docx.table.rows):
        col_no = 1
        row_n = i + 1
        
        p_remove_pause = re.compile('(?i)<pause>')
        p_remove_double_spaces = re.compile(' +')
        p_remove_parenthesis_spaces = re.compile('\( +')
        
        try:
            for j, cell in enumerate(row.cells):
                #if cell.text:
                #    df[i][j] = cell.text
                docx.table_cells[i][j] = cell
                # XML is ._tc
                #df[i][j] = cell._tc
                # Defensive lock: snapshot every source-side cell (columns 0
                # and 1 — line-number and EN text) so save_docx_file can
                # restore them before writing the docx to disk. Guarantees
                # the source language column is never altered by any engine
                # or helper, even via a future leak we haven't audited.
                # Store both the visible text and the deepcopy'd XML so the
                # save-time check can prefer text comparison (immune to
                # python-docx's XML re-serialisation noise) and only fall
                # back to XML restore on actual content drift.
                if j in (0, 1):
                    from copy import deepcopy as _dc
                    docx.source_columns_snapshot[(i, j)] = (cell.text, _dc(cell._tc))
                if col_no == 2:
                
                    #docx.from_text_is_greyed_table[row_n] = is_greyed_line(cell)
                    #cellvalue = cell.text.replace('’', "'").strip()
                    #print(docx.from_text_is_greyed_table)
                    #print(docx.from_text_is_red_color_table)
                    #print("row_n=%d" % (row_n))
                    cellvalue, docx.from_text_is_greyed_table[i], docx.from_text_is_red_color_table[i] = get_cell_data(ctx, cell,row_n)
                    p_remove_pause
                    cellvalue = p_remove_pause.sub(' ', cellvalue)
                    cellvalue = p_remove_double_spaces.sub(' ', cellvalue)
                    cellvalue = p_remove_parenthesis_spaces.sub('(', cellvalue)
                    length = len(cellvalue)

                    try:
                        print("%d : %s" % (i, cellvalue), flush=True)
                    except Exception:
                        try:
                            print("%d : %s" % (i, cellvalue.encode("utf-8")))
                        except Exception:
                            print("%d : (unable to print content to screen)" )

                    docx.from_text_is_end_of_line_table[i] = is_end_of_line(cellvalue) or docx.from_text_is_red_color_table[i]
                    docx.from_text_is_empty_line_table[i] = is_empty_line(cellvalue)
                    docx.from_text_is_beginning_of_line_table[i] = is_beginning_of_line(cellvalue)
                    docx.from_text_is_conditional_end_of_line_table[i] = is_conditional_end_of_line(cellvalue)

                    if docx.from_text_is_greyed_table[i] == 1:
                        docx.from_text_is_beginning_of_line_table[i] = 0
                        docx.from_text_is_end_of_line_table[i] = 0
                        
                    if i == 2 and len(cellvalue) > 0:
                        docx.from_text_is_beginning_of_line_table[i] = 1

                    if i > 1:
                        #Test conditionel de fin de ligne
                        if docx.from_text_is_conditional_end_of_line_table[i - 1] == 1 \
                            and docx.from_text_is_beginning_of_line_table[i] == 1:
                            docx.from_text_is_end_of_line_table [i - 1] = 1
                            docx.from_text_is_beginning_of_line_table [i] = 1

                        # Verifier debut de ligne special
                        # Si ligne precedente est vide ou grisee:
                        #    Si ligne courante est non vide et non grisee
                        #        ligne courante est debut de ligne
                        if (docx.from_text_is_empty_line_table[i - 1] == 1 \
                            or docx.from_text_is_greyed_table[i - 1] == 1):
                            if (docx.from_text_is_empty_line_table[i] == 1 \
                                and docx.from_text_is_greyed_table[i] == 1):
                                docx.from_text_is_beginning_of_line_table[i] = 1

                        # Verifier la ligne precedente est fin de ligne
                        # Si ligne precedente est non vide et non grisee
                        #    Si ligne courante est vide ou grisee
                        #        la ligne precedente est fin de ligne
                        if (docx.from_text_is_empty_line_table[i - 1] == 0 \
                            and docx.from_text_is_greyed_table[i - 1] == 0):
                            if (docx.from_text_is_empty_line_table[i] == 1 \
                                or docx.from_text_is_greyed_table[i] == 1):
                                docx.from_text_is_end_of_line_table[i - 1] = 1


                        # Verifier que c'est vraiment un debut de ligne suivant une fin de ligne
                        # Si ligne precedente n'est pas fin de ligne
                        #    et ligne oourante est debut de ligne
                        #        la ligne courante n'est pas un debut de ligne
                        if docx.from_text_is_beginning_of_line_table[i] == 1 and \
                            docx.from_text_is_end_of_line_table[i - 1] == 0 \
                            and docx.from_text_is_greyed_table[i - 1] == 0 \
                            and i > 2:
                            docx.from_text_is_beginning_of_line_table[i] = 0


                        # Verifier qu'on a pas loupe un debut de ligne
                        # Si ligne precedente est fin de ligne
                        #    et ligne oourante n'est pas grisee et pas debut de ligne
                        #        la ligne courante est un debut de ligne
                        if docx.from_text_is_end_of_line_table[i - 1] == 1 \
                            and docx.from_text_is_greyed_table[i] == 0 \
                            and docx.from_text_is_beginning_of_line_table[i] == 0:
                            docx.from_text_is_beginning_of_line_table[i] = 1

                        if (docx.from_text_is_empty_line_table[i - 1] == 1 \
                            or docx.from_text_is_greyed_table[i - 1] == 1) \
                            and (docx.from_text_is_empty_line_table[i] == 0 \
                            and docx.from_text_is_greyed_table[i] == 0):
                            docx.from_text_is_beginning_of_line_table[i] = 1

                        if docx.from_text_is_empty_line_table[i - 1] == 1:
                            docx.from_text_is_beginning_of_line_table[i - 1] = 0

                        if i == docx.numrows:
                            docx.from_text_is_end_of_line_table[i - 1] = 1

                    docx.from_text_table[i] = cellvalue
                col_no = col_no + 1
            
            if not splitonly and i > 1:
                prepare_and_clear_cell_for_writing(ctx, i, '')
            docx.from_text_is_read[i] = 1
        except Exception:
            var = traceback.format_exc()
            print(var)
            numerrors = numerrors + 1

    if docx.from_text_is_greyed_table[docx.numrows] == 0 \
        and docx.from_text_is_empty_line_table[docx.numrows] == 0:
        docx.from_text_is_end_of_line_table[docx.numrows] = 1

    split_phrases(ctx)

    if use_html :
        print("<table border=1 width=800>")

    for row_n in range(1, len(docx.from_text_table)):
        try:
            if use_html :
                print("<tr>")
                print("<td width=50>", row_n)
                print("<td width=250>")

            if docx.from_text_is_beginning_of_line_table[row_n] == 1:
                if use_html :
                    print("<hr style=\"height:5px;border:none;color:#ffff00;background-color:#ffff00;\" />")
            
            if docx.from_text_is_greyed_table[row_n] == 1:
                if use_html :
                    print("'<span style=\"background-color: #DCDCDC\">%s</span>' (%s)" % (docx.from_text_table[row_n], len(docx.from_text_table[row_n])))
                    print("<hr style=\"height:5px;border-top: dotted 2px;color:##DCDCDC;background-color:#DCDCDC;\" />")
            else:

                if use_html :
                    print("'%s' (%s)" % (docx.from_text_table[row_n], len(docx.from_text_table[row_n])))

            if docx.from_text_is_end_of_line_table[row_n] == 1:
                
                if use_html :
                    print("<hr style=\"height:5px;border:none;color:#333;background-color:#333;\" />")
            
            if docx.from_text_is_empty_line_table[row_n] == 1:

                if use_html :
                    print("<hr style=\"height:5px;border-top: dotted 2px;color:##DCDCDC;background-color:#DCDCDC;\" />")
                    print("<td>is_greyed=%s<br>is_end_of_line=%s<br>is_empty_line=%s<br>is_beginning_of_line=%s<br>is_conditional_end_of_line=%s" %(
                docx.from_text_is_greyed_table[row_n], \
                docx.from_text_is_end_of_line_table[row_n], \
                docx.from_text_is_empty_line_table[row_n], \
                docx.from_text_is_beginning_of_line_table[row_n], \
                docx.from_text_is_conditional_end_of_line_table[row_n]))

            if use_html :
                print("<td>'%s' (%d)<td>'%s' (%d)" % (docx.from_text_by_phrase_table[row_n], len(docx.from_text_by_phrase_table[row_n]), \
                                                  docx.from_text_by_phrase_separator_table[row_n], len(docx.from_text_by_phrase_separator_table[row_n])))
        except Exception:
            var = traceback.format_exc()
            print(var)
            numerrors = numerrors + 1

def reverse_string(s):
    return s[::-1]


def generate_html_file_from_phrases_for_google_translate_javascript(ctx: RuntimeContext):
    #input("Here")
    print("Generating html page.")

    ctx.docx.docxfile_table_number_of_phrases = 0
    html_to_translate = '''<html lang=%s >
<head>
  <meta charset="UTF-8">
  <title>machine-translate-docx - %s - %s</title>
</head> 
<body>
<h3 translate="no">%s - %s</h3>
<p lang=%s translate="yes" id="%s">%s</p>
<table border=1 CELLPADDING=5 CELLSPACING=0>
<tr><td translate="no">Line No</td><td translate="no">%s<td translate="no">%s</td></tr>
''' % (src_lang, docx_file_name, dest_lang_name, docx_file_name, dest_lang_name, src_lang_name, dest_lang_name, dest_lang_name, src_lang_name, dest_lang_name)
    
    for i, line in enumerate(from_text_table):
        item = from_text_by_phrase_separator_table[i]
        item.strip()
        
        item_searched_and_replaced_before = item
        
        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if xtm.wb is not None:
                if xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item, count=False)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        item_html_escaped = html.escape(item_searched_and_replaced_before.strip())  
        
        if item_searched_and_replaced_before != '':
            ctx.docx.docxfile_table_number_of_phrases += 1
            html_to_translate = html_to_translate + '''
<tr>
    <td>%s</td>
    <td><p id="phrase_%s_%s" lang="%s" translate="no" class="source">%s</p></td>
    <td id="%s"><p id="phrase_%s_%s" lang=%s class="translation">%s</td>
</tr>
''' % (i+1, i, src_lang, src_lang, item_html_escaped, i, i, dest_lang, src_lang, item_html_escaped)
    html_to_translate = html_to_translate + '''
</table>
<div id="google_translate_element"></div>

<script type="text/javascript">
function googleTranslateElementInit() {
  new google.translate.TranslateElement({pageLanguage: '%s'}, 'google_translate_element');
}
</script>
<script>
    function googleTranslateElementInit() {
        new google.translate.TranslateElement({pageLanguage: '%s', includedLanguages: '%s', autoDisplay: true}, 'google_translate_element');
        var a = document.querySelector("#google_translate_element select");
        a.selectedIndex=1;
        a.dispatchEvent(new Event('change'));
    }
</script>

<script type = "text/javascript">  
window.onload = function(){  
        var a = document.querySelector("#google_translate_element select");
        a.selectedIndex.value = '%s';
		//document.querySelector('#google_translate_element select').value = '%s';
        a.dispatchEvent(new Event('change')); 
}  
</script>

<script type="text/javascript" src="https://translate.google.com/translate_a/element.js?cb=googleTranslateElementInit"></script>
</body>
''' % (src_lang, src_lang, dest_lang,dest_lang,dest_lang)
    #print (html_to_translate)
    try:
        if(platform.system() == "Darwin"):
            # Write to TMPDIR or /tmp folder
            try:
                tmpdir = os.environ['TMPDIR']
                if(tmpdir is None or tmpdir == ""):
                    tmpdir = '/tmp/'
            except Exception:
                tmpdir = '/tmp/'
            html_file_path = tmpdir + docx_file_name + '.' + str(os.getpid()) + '.' + dest_lang + '.html'
        else:
            # Windows, write to file at the same location
            html_file_path = os.path.abspath(os.path.expanduser(os.path.expandvars(word_file_to_translate))) + '.' + str(os.getpid()) + '.' + dest_lang + '.html'
        
        print(f"Writing temporary html file to : {html_file_path}")
        html_file_to_translate = open(html_file_path, 'w', encoding='utf-8')
        html_file_to_translate.write(html_to_translate)
        html_file_to_translate.close()
        #print("HTML FILE WRITTEN !")
    except Exception:
        var = traceback.format_exc()
        print(var)


def generate_text_file_from_phrases(ctx: RuntimeContext, text_file_path):
    ctx.docx.docxfile_table_number_of_phrases = 0
    print("Generating text file for google translation...")
    #if xtm.wb is not None:
    if xtm is not None:
        print("Replacing text before using excel file...\n")
    text_to_translate = ''
    text_to_translate_array = []
    
    for i, line in enumerate(from_text_table):
        item = ctx.docx.from_text_by_phrase_separator_table[i]
        item = item.strip()
        
        item_searched_and_replaced_before = item
        
        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if xtm.wb is not None:
                if xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        if item_searched_and_replaced_before != '':
            #text_to_translate = text_to_translate + '''%s
#''' % (item)
            text_to_translate_array.append(item_searched_and_replaced_before)
            ctx.docx.docxfile_table_number_of_phrases = ctx.docx.docxfile_table_number_of_phrases + 1
    #print (text_to_translate)
    #print (text_to_translate_array)
    
    len_text_to_translate_array = len(text_to_translate_array)
    #print("len(text_to_translate_array)=%d" % (len(text_to_translate_array)))
    
    for index in range(len(text_to_translate_array)):
        #print("%d : '%s'" % (index, text_to_translate_array[index]))
        if index == (len_text_to_translate_array - 1):
            text_to_translate = text_to_translate + '%s' % (text_to_translate_array[index])
        else:
            text_to_translate = text_to_translate + '%s\n' % (text_to_translate_array[index])

    #print("text_to_translate=\n%s" % (text_to_translate))
    try:
        #text_file_path = docx_file_name + '.txt'
        text_file_to_translate = open(text_file_path, 'w', encoding='utf-8')
        text_file_to_translate.write(text_to_translate)
        text_file_to_translate.close()
    except Exception:
        var = traceback.format_exc()
        print(var)
        

def generate_xlsx_file_from_phrases(ctx: RuntimeContext, xlsx_file_path):
    ctx.docx.docxfile_table_number_of_phrases = 0
    print("Generating xlsx file for google translation...")
    #if xtm.wb is not None:
    if xtm is not None:
        print("Replacing xlsx before using excel file...\n")
    text_to_translate = ''
    text_to_translate_array = []
    
    for i, line in enumerate(from_text_table):
        item = ctx.docx.from_text_by_phrase_separator_table[i]
        item = item.strip()
        
        item_searched_and_replaced_before = item
        
        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if xtm.wb is not None:
                if xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        if item_searched_and_replaced_before != '':
            #text_to_translate = text_to_translate + '''%s
            #''' % (item)
            text_to_translate_array.append(item_searched_and_replaced_before)
            ctx.docx.docxfile_table_number_of_phrases = ctx.docx.docxfile_table_number_of_phrases + 1
    #print (text_to_translate)
    #print (text_to_translate_array)
    
    len_text_to_translate_array = len(text_to_translate_array)
    #print("len(text_to_translate_array)=%d" % (len(text_to_translate_array)))
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "English"
    except Exception:
        print ("Error creating empty xlsx workbook")
        var = traceback.format_exc()
        print("ERROR: %s" % (var))
        self.wb = None
        self.ws = None
        if not silent:
            input("Enter to close program")
        else:
            print("Program ended with errors")
        sys.exit(13)
    
    index_current_row = 1
    max_col_length = 0
    for index in range(len(text_to_translate_array)):
        ws.cell(row=index_current_row, column=1, value=text_to_translate_array[index])
        if len(text_to_translate_array[index]) > max_col_length:
            max_col_length = len(text_to_translate_array[index])
        index_current_row = index_current_row + 1
    
    ws.column_dimensions['A'].width = max_col_length + 1000
    
    file_saved = 0
    while file_saved == 0:
        try:
            wb.save(xlsx_file_path)
            print ("Excel XLSX english text to translate file \"%s\" saved..." %(xlsx_file_path))
            file_saved=1
        except Exception:
            var = traceback.format_exc()
            print(var)
            if not silent:
                txt_readline = input("\n\nERROR: File saving failed. Please close microsoft excel or other program and press enter to save the xlsx document.\n")
            else:
                # No user to dismiss the prompt; back off briefly and
                # retry the loop instead of hanging the launcher pipe.
                time.sleep(2)
        

def deepl_double_linefeed_between_phrases(dest_lang):
    single_linefeed_phrase_separator_langs = ('ar', 'bg', 'cs', 'da', 'de', 'el', 'en', 'en-us', 'en-gb',
                       'es', 'et', 'fi', 'fr', 'he', 'hu', 'id', 'it', 'ja', 'ko',
                       'lt', 'lv', 'nb', 'nl', 'pl', 'pt', 'pt-br', 'pt-pt',
                       'ro', 'ru', 'sk', 'sl', 'sv', 'tr', 'uk', 'vi', 'zh-hant', 'zh-hans')
    return dest_lang not in single_linefeed_phrase_separator_langs

def generate_char_blocks_array_from_phrases(ctx: RuntimeContext, text_file_path):
    ctx.docx.docxfile_table_number_of_phrases = 0
    print("Generating %d character blocks for translation..." % (ctx.config.max_translation_block_size))
    #if xtm.wb is not None:
    if xtm is not None:
        print("Replacing text before using excel file...\n")
    text_to_translate = ''
    text_to_translate_array = []
    ctx.docx.blocks_nchar_max_to_translate_array = []
    
    double_linefeed_between_phrases = False
    phrase_separator = "\n"
    phrase_separator_len = 1
    if ctx.engine.engine == 'deepl':
        if deepl_double_linefeed_between_phrases(ctx.language.dest_lang):
            double_linefeed_between_phrases = False
            phrase_separator = "\n\n"
            phrase_separator_len = 2
    
    for i, line in enumerate(from_text_table):
        item = ctx.docx.from_text_by_phrase_separator_table[i]
        item = item.strip()
        
        item_searched_and_replaced_before = item
        
        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if xtm.wb is not None:
                if xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        if item_searched_and_replaced_before != '':
            #text_to_translate = text_to_translate + '''%s
#''' % (item)
            text_to_translate_array.append(item_searched_and_replaced_before)
            ctx.docx.docxfile_table_number_of_phrases = ctx.docx.docxfile_table_number_of_phrases + 1
                
    #print (text_to_translate)
    #print (text_to_translate_array)
    
    len_text_to_translate_array = len(text_to_translate_array)
    #print("len(text_to_translate_array)=%d" % (len(text_to_translate_array)))
    
    
    current_text_block = ""
    current_text_block_len = len(current_text_block)
    
    
    #print(text_to_translate_array)
    #print(len_text_to_translate_array)
    #input("len_text_to_translate_array")
    #print("Before current_text_block generation")
    for index in range(0, len(text_to_translate_array)):
        #print("%d : '%s'" % (index, text_to_translate_array[index]))
        current_phrase_str = text_to_translate_array[index]
        
        if current_text_block_len + len(current_phrase_str) <= ctx.config.max_translation_block_size + phrase_separator_len:
            if len(current_text_block) == 0:
                current_text_block = current_phrase_str
                current_text_block_len = len(current_text_block)
                #print(current_phrase_str)
                #input("adding first phrase")
            else:
               
                current_text_block = current_text_block + phrase_separator + current_phrase_str
                current_text_block_len = current_text_block_len + len(current_phrase_str) + phrase_separator_len
                #print(current_phrase_str)
                #print("current_text_block")
                #print(current_text_block)
        else:
            ctx.docx.blocks_nchar_max_to_translate_array.append(current_text_block)
            #print("Current block of %d characters:\n-------------------------------------------------" % (ctx.config.max_translation_block_size))
            #print("current_text_block")
            #print(current_text_block)
            #print("end")
            #input("OK Here")
            #input("adding more phrase")
            #print(current_text_block.split("\n"))
            #print("------------------------------------------\nBlock size : %d" % (current_text_block_len))
            
            current_text_block = current_phrase_str
            current_text_block_len = len(current_phrase_str)
            #input("------------------------------------------\nType enter to continue")
        
        if index == (len(text_to_translate_array) - 1):
            ctx.docx.blocks_nchar_max_to_translate_array.append(current_text_block)
            #print("Current block of %d characters:\n-------------------------------------------------"  % (ctx.config.max_translation_block_size))
            #print(current_text_block.split("\n"))
            #print("------------------------------------------\nBlock size : %d" % (current_text_block_len))
            #input("------------------------------------------\nType enter to continue")
    
    #print("blocks_nchar_max_to_translate_array:")
    #print("\n******************************************\n".join(ctx.docx.blocks_nchar_max_to_translate_array))
    #print ("len(text_to_translate_array) = %d " % (len(text_to_translate_array) - 1)) 

    #print("text_to_translate=\n%s" % (text_to_translate))
    try:
        #text_file_path = docx_file_name + '.txt'
        #text_file_to_translate = open(text_file_path, 'w', encoding='utf-8')
        #text_file_to_translate.write(text_to_translate)
        #text_file_to_translate.close()
        pass
    except Exception:
        var = traceback.format_exc()
        print(var)


def google_translate_from_text_file(ctx: RuntimeContext):
    #ctx.flags.word_file_to_translate
    text_file_path = docx_file_name + '.txt'
    text_file_full_path = os.path.realpath(text_file_path)
    #print("text_file_full_path=%s" % text_file_full_path)
    generate_text_file_from_phrases(ctx, text_file_full_path)
    #input("There")
    #input("Here, press enter:")
    print("Starting translation in google using text file...")
    ctx.docx.translation_array = selenium_chrome_google_translate_text_file(ctx, text_file_full_path)
    try:
        os.remove(text_file_path)
        pass
    except Exception:
        pass

def google_translate_from_html_javascript(ctx: RuntimeContext):
    #input("There")
    #input("Here, press enter:")
    print("Starting translation in google using html file...")
    
    generate_html_file_from_phrases_for_google_translate_javascript(ctx)
    
    ctx.docx.translation_array = selenium_chrome_google_translate_html_javascript_file(ctx, html_file_path)
    
    try:
        #input("before remove html file")
        os.remove(html_file_path)
        pass
    except Exception:
        pass

    return ctx.docx.translation_array

def google_translate_from_html_xlsxfile(ctx: RuntimeContext):
    xlsx_file_path = ctx.flags.word_file_to_translate + '.xlsx'
    xlsx_file_full_path = os.path.realpath(xlsx_file_path)
    #print("text_file_full_path=%s" % text_file_full_path)
    generate_xlsx_file_from_phrases(ctx, xlsx_file_full_path)
    #input("There")
    #input("Here, press enter:")
    print("Starting translation in google using text file...")
    ctx.docx.translation_array = selenium_chrome_google_translate_xlsx_file(ctx, xlsx_file_full_path)
    try:
        os.remove(xlsx_file_full_path)
        pass
    except Exception:
        pass

def translate_from_phrasesblock(ctx: RuntimeContext):
    text_file_path = docx_file_name + '.txt'
    text_file_full_path = os.path.realpath(text_file_path)
    #print("text_file_full_path=%s" % text_file_full_path)
    #generate_text_file_from_phrases(ctx, text_file_full_path)
    generate_char_blocks_array_from_phrases(ctx, text_file_full_path)

    translation_succeded = True

    #input("phrasesblock")
    print("Starting translation in %s using phrase blocks of %d characters..." % (ctx.engine.engine, ctx.config.max_translation_block_size))

    translation_succeded, ctx.docx.translation_array = _runner_translate_maxchar_blocks(
        ctx,
        selenium_webservice_perplexity_translate=selenium_webservice_perplexity_translate,
    )
    try:
        os.remove(text_file_path)
        pass
    except Exception:
        pass
    return translation_succeded

def translate_docx(ctx: RuntimeContext):
    translation_succeded = True

    # ------------------------------------------------------------------
    # Engine-method specific translators
    # ------------------------------------------------------------------
    if engine_method == "textfile":
        google_translate_from_text_file(ctx)
        return translation_succeded

    if engine_method == "javascript":
        google_translate_from_html_javascript(ctx)
        # Fail loudly when Google's web widget refuses to translate the
        # local file (a known modern-Chrome limitation, since ~2022).
        # Silent failure here used to cascade into ~14 retries per
        # phrase × N phrases of `translation_array index out of range`
        # warnings before producing an empty docx — completely
        # unreadable failure mode for the user.
        if not ctx.docx.translation_array:
            print("[ERROR] Google web translate returned 0 lines.")
            print("[ERROR] Modern Chrome blocks the Google translate widget on")
            print("[ERROR] file:// URLs (CORS / sandboxing). This engine path")
            print("[ERROR] cannot complete in current Chrome versions.")
            print("[INFO] Use the OpenAI API engine (chatgpt) or DeepL instead.")
        return translation_succeded

    if engine_method == "xlsxfile":
        google_translate_from_html_xlsxfile(ctx)
        return translation_succeded

    # ------------------------------------------------------------------
    # Phrase-block logic — predicate lives in src/dispatch.py so it
    # can never drift from set_translation_function (both share the
    # same engine ↔ method matrix). See dispatch.use_phrasesblock for
    # the per-engine policy.
    # ------------------------------------------------------------------
    if _dispatch_use_phrasesblock(translation_engine, engine_method):
        translation_succeded = translate_from_phrasesblock(ctx)

    return translation_succeded



def get_translation_and_replace_after(ctx: RuntimeContext):
    # Phase H: seed local `driver` from ctx so the later reassignment
    # branches don't trigger UnboundLocalError on prior reads.
    driver = ctx.browser.driver

    phrase_no = 0

    # ── DIAGNOSTIC: confirm we enter here with polished ctx.docx.translation_array ─
    if ctx.openai.polisher is not None:
        print(f"[DIAG] get_translation_and_replace_after: ctx.docx.translation_array has "
              f"{len(ctx.docx.translation_array)} lines — distributing into ctx.docx.to_text_by_phrase_separator_table")
    # ─────────────────────────────────────────────────────────────────────

    p_remove_pause = re.compile('(?i)<pause>')
    p_remove_double_spaces = re.compile(' +')
    p_remove_parenthesis_spaces = re.compile('\( +')

    for i, line in enumerate(from_text_table):
        item = ctx.docx.from_text_by_phrase_separator_table[i]
        item = item.strip()
        from_language = ctx.language.src_lang
        phrase_separator_removed_str = ''

        p_remove_separator = re.compile(line_separator_regex_str)
        p_remove_double_spaces = re.compile(' +')

        # Avec separateurs ()

        try:
            web_translation_separators = ''
            if item.strip() != '':
                phrase_no = phrase_no + 1
                print("\n%d/%d" % (i, word_translation_table_length))
                print("Phrase to translate :'%s'\n" % (item.strip()))
                item = item.strip()

                item_searched_and_replaced_before = item
                if xlsxreplacefile is not None:
                    if xtm.wb is not None:
                        item_searched_and_replaced_before, nb_searched_and_replaced_before = xtm.search_and_replace_text('before', item)
                        if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                            continue
                if ctx.flags.splitonly:
                    web_translation_separators = get_translated_cells_content (i, item_searched_and_replaced_before)
                elif ctx.flags.use_api:
                    try:
                        web_translation_separators = ""
                        #if ctx.flags.use_api:
                        #    translation = ctx.openai.translator.translate(item_searched_and_replaced_before, src=ctx.language.src_lang, dest=ctx.language.dest_lang)
                        #    web_translation_separators = translation.text
                        if not len(web_translation_separators) > 0:
                            ctx.flags.use_api = False
                            # Faster google Chrome translate failed, using Selenium as backup

                            if driver is None:
                                print(f"[Line {inspect.currentframe().f_lineno}] Starting Chrome browser\n")
                                
                                service = Service()                                
                                driver = webdriver.Chrome(service=service, options=ctx.browser.chrome_options)
                                
                                driver.set_window_position(100, 100)
                                driver.set_window_size(800, 700)
                                #driver.set_window_size(400, 650)

                            print("phrase_no=%d" % phrase_no)
                            web_translation_separators = selenium_chrome_machine_translate(ctx, item_searched_and_replaced_before, phrase_no)
                    except Exception:
                        ctx.flags.use_api = False
                        # Faster google Chrome translate failed, using Selenium as backup

                        if driver is not None:
                            print(f"Starting Chrome browser\n")
                            
                            service = Service()                                
                            driver = webdriver.Chrome(service=service, options=ctx.browser.chrome_options)

                        if ctx.engine.engine == 'google' and driver is not None:
                            driver.set_window_position(100, 100)
                            driver.set_window_size(800, 700)

                        print("phrase_no = %d" % phrase_no)
                        web_translation_separators = selenium_chrome_machine_translate(ctx, item_searched_and_replaced_before, phrase_no)
                else:
                    if ctx.engine.method == "singlephrase" and ctx.engine.engine == 'deepl':
                        translation_succeded, web_translation_separators  = selenium_chrome_machine_translate(ctx, item_searched_and_replaced_before, phrase_no)
                    else:
                        web_translation_separators = selenium_chrome_machine_translate(ctx, item_searched_and_replaced_before, phrase_no)
                        
                #web_translation_separators = translation.text
                phrase_separator_removed_str = p_remove_double_spaces.sub(' ', web_translation_separators)

                #print("Google translation='%s'" % (phrase_separator_removed_str.encode('utf8')))
                if xlsxreplacefile is not None:
                    nb_searched_and_replaced = 0
                    web_translation_separators_searched_and_replaced, nb_searched_and_replaced = xtm.search_and_replace_text('after', phrase_separator_removed_str)
                    if nb_searched_and_replaced > 0:
                        #print("\nPhrase %d replacements :\n'%s'" % (nb_searched_and_replaced, web_translation_separators))
                        #print("Replaced phrase :\n'%s'" % (web_translation_separators_searched_and_replaced))
                        phrase_separator_removed_str = web_translation_separators_searched_and_replaced

                if ctx.language.dest_lang in right_to_left_languages_list.keys():
                    #phrase_separator_removed_aligned_str = reverse_string (phrase_separator_removed_str)
                    #phrase_separator_removed_aligned_str = "\u202B" + phrase_separator_removed_str + "\u202C"
                    phrase_separator_removed_aligned_str = get_display(phrase_separator_removed_str)
                else:
                    phrase_separator_removed_aligned_str = phrase_separator_removed_str
                try:
                    if ctx.flags.splitonly:
                        print("Translated text :'%s'\n" % (phrase_separator_removed_aligned_str))
                    else:
                        print("%s translation (%s):'%s'" % (ctx.engine.engine.title() ,ctx.language.dest_lang_name, phrase_separator_removed_aligned_str))
                except Exception:
                    print("")
                    print("Google translation='%s'" % (phrase_separator_removed_str.encode('utf8').decode('utf8')))
                if web_translation_separators.strip() == '' and not ctx.flags.splitonly:
                    print("Error translating='%s'" % (item))
                ctx.docx.to_text_by_phrase_separator_table[i] = phrase_separator_removed_str
                phrase_separator_removed_str = p_remove_separator.sub(' ', phrase_separator_removed_str)
                phrase_separator_removed_str.strip()
                to_text_by_phrase_separator_removed_table[i] = phrase_separator_removed_str
        except Exception:
            var = traceback.format_exc()
            ctx.browser.numerrors_deepl = ctx.browser.numerrors_deepl + 1
            web_translation_separators = var
            print("ERROR:%s" % (var))

        item = from_text_by_phrase_table[i]
        try:
            web_translation_no_separators = ''
            if item.strip() != '':
                #google_translation_res = translator.translate(item, src=ctx.language.src_lang, dest='fr')
                #time.sleep(5)
                #web_translation_no_separators = pydeepl.translate(item, to_language)
                phrase_separator_removed_str = p_remove_double_spaces.sub(' ', web_translation_no_separators)
                phrase_separator_removed_str = p_remove_parenthesis_spaces.sub('(', phrase_separator_removed_str)
                to_text_by_phrase_table[i] = phrase_separator_removed_str
        except Exception:
            var = traceback.format_exc()
            numerrors_googletranslate = numerrors_googletranslate + 1
            web_translation_no_separators = var
        Identical_with_without_separators = 'DIFFERENT<BR>'
        if phrase_separator_removed_str == web_translation_no_separators:
            Identical_with_without_separators = 'SAME<BR>'


def document_split_phrases(ctx: RuntimeContext):
    # Split phrases into multiple lines to match source language number of lines

    oai_sub_splitter = None
    if ctx.flags.split_engine == "openai":
        if ctx.openai.translator is not None:
            oai_sub_splitter = OpenAISubtitleSplitter(filename=ctx.flags.word_file_to_translate, doc_id=ctx.openai.translator.get_doc_id())
        else:
            oai_sub_splitter = OpenAISubtitleSplitter()
    
    for i, line in enumerate(from_text_table):
        if ctx.docx.to_text_by_phrase_separator_table[i] != '':
            #ctx.docx.docxfile_table_number_of_phrases = ctx.docx.docxfile_table_number_of_phrases + 1
            ctx.docx.docxfile_table_number_of_characters = ctx.docx.docxfile_table_number_of_characters + len(ctx.docx.from_text_by_phrase_separator_table[i])
            ctx.docx.phrase_number_of_words = len(ctx.docx.from_text_by_phrase_separator_table[i].strip().split(" "))
            #print("Phrase to split: %s" % (ctx.docx.from_text_by_phrase_separator_table[i]))
            #print("number of words: %d" % (ctx.docx.phrase_number_of_words))
            ctx.docx.docxfile_table_number_of_words = ctx.docx.docxfile_table_number_of_words + ctx.docx.phrase_number_of_words
            input_phrase_lines = ""
            #translation_result_using_separator[i] = []
            #translation_result_phrase_array[i] = []
            
            try:
                current_line = ctx.docx.to_text_by_phrase_separator_table[i]
                # Using () as separator for splitting phrases, not used anymore
                #lines = current_line.split(line_separator_nospace_str)
                str_translation_len = len(current_line)

                try:
                    if str_translation_len <= 0:
                        str_phrase_stats = ""
                    else:
                        str_nb_lines = from_text_nb_lines_in_phrase[i]
                        if str_nb_lines > 0:
                            str_line_average = str_translation_len / str_nb_lines
                            str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                        else:
                            str_line_average = 0
                            str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                        #print("str_phrase_stats=%s" % (str_phrase_stats))
                except Exception:
                    var = traceback.format_exc()
                    print("  ERROR:%s<br>" % (var))
                
                
                # --- Step 1: Build input safely ---
                if str_nb_lines > 1:
                    input_phrase_lines = "\n".join(
                        from_text_table[i + idx] for idx in range(str_nb_lines)
                    )
                else:
                    input_phrase_lines = current_line


                # --- Step 2: Split functions ---
                def split_with_openai():
                    def expected_line_count(text):
                        if not text:
                            return 0
                        return text.count("\n") + 1

                    def normalize(lines):
                        # If string → split
                        if isinstance(lines, str):
                            lines = lines.strip().split("\n")

                        # Ensure list of strings
                        if isinstance(lines, list):
                            return [str(l).strip() for l in lines if str(l).strip() != ""]

                        return []

                    def is_valid(lines, expected):
                        return isinstance(lines, list) and len(lines) == expected

                    expected = expected_line_count(input_phrase_lines)

                    oai_sub_splitter.set_model('gpt-5.4-mini')
                    # --- Try 2 times with default model ---
                    for attempt in range(2):
                        response, lines = oai_sub_splitter.split_phrase(
                            ctx.language.src_lang_name, ctx.language.dest_lang_name, input_phrase_lines, current_line
                        )

                        lines = normalize(lines)

                        if is_valid(lines, expected):
                            return lines

                        print(f"[Retry {attempt+1}] Invalid OpenAI split: got {len(lines)} expected {expected}")

                    # --- Switch model ---
                    print("[Switching model to gpt-5.5]")
                    oai_sub_splitter.set_model('gpt-5.5')

                    response, lines = oai_sub_splitter.split_phrase(
                        ctx.language.src_lang_name, ctx.language.dest_lang_name, input_phrase_lines, current_line
                    )

                    # --- Final handling ---
                    if isinstance(lines, list):
                        lines = [str(l).strip() for l in lines if str(l).strip() != ""]
                        if lines:
                            return lines
                        else:
                            return []

                    elif isinstance(lines, str):
                        lines = lines.strip().split("\n")
                        lines = [l.strip() for l in lines if l.strip() != ""]
                        if lines:
                            return [lines[0]]  # return first line only
                        else:
                            return []

                    # fallback safety
                    return []


                def split_with_algorithm():
                    local_avg = str_line_average

                    if local_avg > MAX_LINE_SIZE:
                        local_avg = math.ceil(local_avg)

                    tokens = tokenize_text_to_array(current_line, ctx.language.dest_lang)
                    lines = divide_array(tokens, ctx.language.dest_lang, local_avg + 4)

                    divide_max_try = MAX_LINE_SIZE
                    while (len(lines) > str_nb_lines) and (divide_max_try > 0):
                        local_avg += 1
                        lines = divide_array(tokens, ctx.language.dest_lang, local_avg + 4)
                        divide_max_try -= 1

                    divide_max_try = MAX_LINE_SIZE
                    local_avg = str_translation_len / str_nb_lines
                    local_avg = math.ceil(local_avg)

                    while (len(lines) <= str_nb_lines) and (len(lines) >= 1) and (divide_max_try > 0):
                        local_avg -= 1
                        attempt_lines = divide_array(tokens, ctx.language.dest_lang, local_avg + 4)

                        if len(attempt_lines) <= str_nb_lines:
                            lines = attempt_lines

                        divide_max_try -= 1

                    return lines


                # --- Only 1 line, no splitting ---
                if str_nb_lines == 1:
                    lines_divided = [current_line]
                if str_nb_lines > 1:
                    if ctx.flags.split_engine == "openai":
                        lines_divided = split_with_openai()
                        
                        number_lines = len(lines_divided)
                        if number_lines != str_nb_lines:
                            print('Fallback to algorythm line splitting')
                            lines_divided = split_with_algorithm()
                    else:
                        lines_divided = split_with_algorithm()

                    # --- Step 4: Normalize ---
                    if not lines_divided:
                        lines_divided = [current_line]

                    lines_divided = [str(line).strip() for line in lines_divided]

                    # --- Step 5: Enforce exact number of lines ---
                    number_lines = len(lines_divided)
                    print(number_lines)
                    print(lines_divided)

                    if number_lines != str_nb_lines:
                        print(f"Adjusting output: {number_lines} -> {str_nb_lines}")

                        if number_lines > str_nb_lines:
                            lines_divided = lines_divided[:str_nb_lines - 1] + [
                                " ".join(lines_divided[str_nb_lines - 1:])
                            ]
                        else:
                            lines_divided += [""] * (str_nb_lines - number_lines)

                        number_lines = len(lines_divided)


                # --- Step 6: Store results ---
                translation_result_phrase_array[i] = lines_divided

                number_lines = len(lines_divided)
                for line_no in range(number_lines):
                    try:
                        translation_result_using_separator[line_no + i] = lines_divided[line_no]
                        print(f"{line_no} -> {lines_divided[line_no]}")
                    except Exception:
                        translation_result_using_separator[line_no + i] = "Error"


                # --- Step 7: Logging ---
                try:
                    print(
                        "Splitting phrase : %s (%d) = %d lines"
                        % (ctx.docx.to_text_by_phrase_separator_table[i], i, str_nb_lines)
                    )
                except Exception:
                    try:
                        print(
                            "%s (%d): %d "
                            % (ctx.docx.to_text_by_phrase_separator_table[i].encode("utf-8"), i, str_nb_lines)
                        )
                    except Exception:
                        print("(unable to print content to screen) (%d): %d" % (i, str_nb_lines))


                # --- Final check ---
                if number_lines != str_nb_lines:
                    print("Error in number of line %d, expected %d." % (number_lines, str_nb_lines))
                
            except Exception:
                var = traceback.format_exc()
                print("  ERROR:%s<br>" % (var))


MAX_CHARS = 750

def create_translation_split_prompts():
    """
    Groups source phrases into blocks of max 750 characters (without breaking phrases)
    and prints AI prompts in the requested subtitle format.
    """

    current_block = []
    current_length = 0
    block_index = 1

    for phrase in from_text_table:
        phrase = phrase.strip()
        if not phrase:
            continue

        phrase_length = len(phrase)

        # If adding this phrase would exceed the max size, finalize current block
        if current_length + phrase_length > MAX_CHARS and current_block:
            print_prompt_block(block_index, current_block)
            block_index += 1
            current_block = []
            current_length = 0

        current_block.append(phrase)
        current_length += phrase_length

    # Print last block if any
    if current_block:
        print_prompt_block(block_index, current_block)


def print_prompt_block(block_index, block_phrases):
    """
    Prints a single AI prompt for one block of phrases in the desired format.
    """

    # Prepare source text lines with numbering
    source_lines = "\n".join([f"Input {i+1}:{line}" for i, line in enumerate(block_phrases)])
    num_lines = len(block_phrases)

    prompt = f"""
You are given subtitle text in a source language and its translation in a destination language.

Source Text ({num_lines} lines):
{source_lines}

Destination Text (translation):
# Insert your translation here, line by line

Task:

Reformat the translated text so that it has exactly the same number of lines as the source text, preserving the line structure of the source.

Rules:

- Each line in the source corresponds in order to the translated text.
- If a source sentence is split across multiple lines, the translation must also be split naturally across the same number of lines.
- Do not change any words or punctuation in the translation.
- Each phrase ending with a full stop in the source should preserve its line count in the translation.
- In case the target language grammar is different from the input language grammar, the lines do not need to match the source, but the phrase full stop should determine the number of lines to be split for a phrase(s) from the input.
- Output only the translated text, line by line, with no numbers or labels.

Example:

Source sample (English):

I’ve always had a terrible  
aversion to bullfighting.

Translation sample (French):

J'ai toujours eu une aversion terrible
pour la corrida.
"""

    print(f"\n{'=' * 80}")
    print(f"PROMPT BLOCK #{block_index}")
    print(f"{'=' * 80}")
    print(prompt.strip())


    
def print_html_program_result():
    if use_html :
        print("<table border=1 bgcolor=""#EEEEEE"">")

    for i, line in enumerate(from_text_table):
        Identical_with_without_separators = 'DIFFERENT<BR>'
        if to_text_by_phrase_separator_removed_table[i] == to_text_by_phrase_table[i]:
            Identical_with_without_separators = 'SAME<BR>'
        #print "<tr><td>%s<td>%s<td>%s<td>%s<td>%s%s" % (i, from_text_table[i], from_text_by_phrase_separator_table[i].encode('utf8'), to_text_by_phrase_separator_table[i].encode('utf8'), Identical_with_without_separators.encode('utf8'), to_text_by_phrase_separator_removed_table[i].encode('utf8') )
        if len(from_text_by_phrase_separator_table[i]) == 0:
            Identical_with_without_separators = ''
        if use_html :
            print("<tr><td>%d<td>'%s'<td>%s<td>%s<td>%s<td>%s%s" % (i, from_text_table[i], translation_result_using_separator[i].encode('utf8'), to_text_by_phrase_separator_table[i].encode('utf8'), to_text_by_phrase_table[i].encode('utf8'), Identical_with_without_separators.encode('utf8'), to_text_by_phrase_table[i].encode('utf8') ))
        #sys.exit(0)

    if use_html :
        print("</table><br>elapsedtime = ", elapsedtime)
        print("</span>")

def write_destination_language_in_docx_cell():
    if not splitonly:
        try:
            docxdoc.tables[0].cell(1, 2).text = dest_lang_name
        except Exception:
            try:
                docxdoc.tables[0].cell(1, 2).text = dest_lang
            except Exception:
                pass


def print_console_docx_file_translated(ctx: RuntimeContext):
    print("\nTranslated text:\n")
    numrows = len(table.rows)
    numcols = len(table.columns)
    current_cell_row = 2
    for row_n in range(2, (numrows)):

        str_translation_len = len(ctx.docx.translation_result_using_separator[row_n])
        translation_phrase_lines_len = len(ctx.docx.translation_result_phrase_array[row_n])
        if translation_phrase_lines_len == 0 and current_cell_row < row_n:
            print("%d :" % row_n)
        #print("row_n = %d" %  row_n)

        # Non-split path: write the translation regardless of whether
        # document_split_phrases populated translation_result_phrase_array.
        # The original guard `if translation_phrase_lines_len >= 1` was a
        # silent failure mode \u2014 if the split helper ever skipped a row
        # (e.g. its own to_text guard), the translation would never reach
        # the cell. By keying on to_text_by_phrase_separator_table directly
        # we write whatever the translation engine returned.
        if not split_translation:
            translation_cell_text = ctx.docx.to_text_by_phrase_separator_table[row_n]
            if translation_cell_text:
                prepare_and_clear_cell_for_writing(ctx, row_n, translation_cell_text)
                if dest_lang in right_to_left_languages_list.keys():
                    translation_cell_aligned_text = get_display(translation_cell_text)
                else:
                    translation_cell_aligned_text = translation_cell_text
                print("%d : %s" % (row_n, translation_cell_aligned_text))
            continue

        if translation_phrase_lines_len >= 1 :
            #print("%d : %s" % (row_n,' '.join(ctx.docx.translation_result_phrase_array[row_n])))

            if False:  # legacy non-split branch above already handled by `if not split_translation`
                pass
            else:
                #translation_cell_text = ctx.docx.translation_result_using_separator[row_n]
                #print("len array: %d" % (translation_phrase_lines_len))
                #print("translation_result_phrase_array[%d] : %s" % (row_n,'\n'.join(ctx.docx.translation_result_phrase_array[row_n])))

                translation_phrase_line_pos = 0
                translation_phrase_cell_pos = 0

                while translation_phrase_line_pos < translation_phrase_lines_len:
                    current_cell_row = row_n + translation_phrase_cell_pos
                    cell_lines_len = ctx.docx.from_text_nb_lines_in_cell[row_n + translation_phrase_cell_pos]
                    cell_line_pos = 0
                    current_cell = ctx.docx.table_cells[current_cell_row][2]
                    while cell_line_pos < cell_lines_len \
                        and translation_phrase_line_pos < translation_phrase_lines_len:

                        translation_phrase_line_str = ctx.docx.translation_result_phrase_array[row_n][translation_phrase_line_pos]
                        if dest_lang in right_to_left_languages_list.keys():
                            #translation_cell_aligned_text = reverse_string (translation_phrase_line_str)
                            #translation_cell_aligned_text = "\u202B" + translation_phrase_line_str + "\u202C"
                            translation_cell_aligned_text = get_display(translation_phrase_line_str)
                        else:
                            translation_cell_aligned_text = translation_phrase_line_str
                        if cell_lines_len > 1:
                            print("%d-%d : %s" % (current_cell_row, cell_line_pos + 1, translation_cell_aligned_text))
                        else:
                            print("%d : %s" % (current_cell_row, translation_cell_aligned_text))
                        if cell_line_pos == 0:
                            #print("cell_line_pos=%d" % cell_line_pos)
                            if splitonly:
                                prepare_and_clear_cell_for_writing(ctx, current_cell_row, translation_phrase_line_str)
                            else:
                            #prepare_and_clear_cell_for_writing(current_cell_row, translation_phrase_line_str)
                                cell_set_1st_paragraph(ctx, current_cell_row, translation_phrase_line_str)
                            # Not needed
                            #current_cell.paragraphs[0].text = translation_phrase_line_str
                        else:
                            # Add empty paragraph between translation lines
                            cell_add_paragraph(ctx, current_cell_row, "")
                            # Add the translation line
                            cell_add_paragraph(ctx, current_cell_row, translation_phrase_line_str)
                        cell_line_pos = cell_line_pos + 1
                        translation_phrase_line_pos = translation_phrase_line_pos + 1
                        #input("press enter")
                    translation_phrase_cell_pos = translation_phrase_cell_pos + 1

        try:
            if str_translation_len <= 0:
                str_phrase_stats = ""
            else:
                str_translation_len = len(translation_result_using_separator[row_n])
                str_nb_lines = from_text_nb_lines_in_phrase[row_n]
                if str_nb_lines > 0:
                    str_line_average = str_translation_len / str_nb_lines
                    str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                else:
                    str_line_average = 0
                    str_phrase_stats = ""

        except Exception:
            var = traceback.format_exc()
            print("  ERROR:%s<br>" % (var))


#print("Generating TMX file for translation comparison")
#generate_tmx_file ()
#word.Application.ActiveWindow.Close()
#word.Application.Quit()

def set_docx_properties_comment_for_history(ctx: RuntimeContext):
    now = datetime.datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    docxdoc.core_properties.comments = "Document translated by SMTV Robot version %s using %s engine on %s." % (PROGRAM_VERSION, ctx.engine.engine, dt_string)



def local_time_offset(t=None):
    """Return offset of local zone from GMT, either at present or at time t."""
    localtimezone = 0
    # python2.3 localtime() can't take None
    if t is None:
        t = time.time()
    localtimezone = -time.altzone / 3600
    if (localtimezone - int(localtimezone)) == 0:
        localtimezone = int(localtimezone)
    if time.localtime(t).tm_isdst == False or time.daylight != 1:
        localtimezone = -localtimezone
    return localtimezone


def run_statistics(ctx: RuntimeContext):
    # Phase H: seed local `driver` so subsequent `driver.get(...)` reads
    # don't UnboundLocalError when the reassign branch is skipped.
    driver = ctx.browser.driver

    _orig_run_statistics_body_marker = None  # placeholder kept for the editor diff
    
    statistics_html_statistics_form_url_key = ['statistics', 'html_statistics_form_url']
    statistics_html_statistics_form_url = get_nested_value_from_json_array(ctx.config.json_configuration_array, statistics_html_statistics_form_url_key)
    
    bool_print_stats = False
    
    try:
        if ctx.flags.splitonly:
            action = "splitonly"
        else:
            action = "translate"
        
        docxfile_size = os.path.getsize(ctx.flags.word_file_to_translate)
        if ctx.flags.use_api == True:
            ctx.engine.method = "api"
        elif ctx.engine.method is None or ctx.engine.method == "":
            ctx.engine.method = "web"
        
        if xlsxreplacefile is not None:
            xlsxreplacefile_splitted = os.path.splitext(os.path.basename(xlsxreplacefile))
            xlsxreplacefile_filename_size = len(xlsxreplacefile_splitted)
            xlsxreplacefile_name = "%s%s" % (xlsxreplacefile_splitted[xlsxreplacefile_filename_size-2], xlsxreplacefile_splitted[xlsxreplacefile_filename_size-1])
        else:
            xlsxreplacefile_name = ""
        
        if xlsxreplacefile_name != "":
            replacebeforelistsize = xtm.get_sheet_number_lines('before')
            replacebeforelistreplaced = xtm.get_sheet_number_of_replacements('before')
            replaceafterlistsize = xtm.get_sheet_number_lines('after')
            replaceafterlistreplaced = xtm.get_sheet_number_of_replacements('after')
            donotsplitlistsize = xtm.get_sheet_number_lines('keep_on_same_line')
            donotsplitfound = xtm.get_sheet_number_of_do_not_split_match('keep_on_same_line')
        else:
            replacebeforelistsize = ""
            replacebeforelistreplaced = ""
            replaceafterlistsize = ""
            replaceafterlistreplaced = ""
            donotsplitlistsize = ""
            donotsplitfound = ""
        
        platform_uname = platform.uname()
        platform_system = platform.system()
        platform_node = platform.node()
        platform_release = platform.release()
        platform_version = platform.version()
        platform_machine = platform.machine()
        platform_processor = platform.processor()
        
        cpu_count = psutil.cpu_count()
        mem_total = psutil.virtual_memory().total
        
        cost_google_translate = 20 * ctx.docx.docxfile_table_number_of_characters / 1000000
        
        local_time_offset_str = local_time_offset()
        
        docxfile_page_count = None
        try:
            document = zipfile.ZipFile(ctx.flags.word_file_to_translate)
            dxml = document.read('docProps/app.xml')
            uglyXml = xml.dom.minidom.parseString(dxml)
            docxfile_page_count = uglyXml.getElementsByTagName('Pages')[0].childNodes[0].nodeValue
        except Exception:
            if bool_print_stats:
                print("Unable to get number of pages from document. You can ignore this.")
        
        try:
            archive = zipfile.ZipFile("myDocxOrPptxFile.docx", "r")
            ms_data = archive.read("docProps/app.xml")
            archive.close()
            app_xml = ms_data.decode("utf-8")
            
            regex = r"<(Pages|Slides)>(\d)</(Pages|Slides)>"
            
            matches = re.findall(regex, app_xml, re.MULTILINE)
            match = matches[0] if matches[0:] else [0, 0]
            page_count = match[1]
        except Exception:
            if bool_print_stats:
                print("Unable to get number of pages from document. You can ignore this.")
        
        if bool_print_stats:
            print("Statistics:")
            print("program_version: %s" % (PROGRAM_VERSION))
            print("client_ip: %s" % (ctx.flags.client_ip))
            #https://stackoverflow.com/questions/1695183/how-to-percent-encode-url-parameters-in-python
            
            print("docxfile: %s" % (ctx.flags.word_file_to_translate))
            print("action: %s" % (action))
            print("destlang_code: %s" % (ctx.language.dest_lang))
            print("destlang_name: %s" % (ctx.language.dest_lang_name))
            print("docxfile: %s" % (docx_file_name))
            print("docxfile_page_count: %s" % docxfile_page_count)
            print("docxfile_size: %s" % (docxfile_size))
            print("docxfile_table_number_of_lines: %s" % (numrows))
            print("docxfile_table_number_of_phrases: %s" % (ctx.docx.docxfile_table_number_of_phrases))
            print("docxfile_table_number_of_words: %s" % (ctx.docx.docxfile_table_number_of_words))
            print("docxfile_table_number_of_characters: %s" % (ctx.docx.docxfile_table_number_of_characters))
            print("engine: %s" % (ctx.engine.engine))
            print("xlsxreplacefile: %s" % (xlsxreplacefile_name))
            print("destfont: %s" % (dest_font))
            print("splitonly: %s" % (ctx.flags.splitonly))
            print("split_translation: %s" % (split_translation))
            print("showbrowser: %s" % (ctx.flags.showbrowser))
            print("start_time: %s" % (start_time))
            print("end_time: %s" % (end_time))
            print("elapsed_time: %s" % ((elapsed_time)))
            
            
            if xlsxreplacefile_name != "":
                print("replacebeforelistsize: %s" % (replacebeforelistsize))
                print("replacebeforelistreplaced: %s" % (replacebeforelistreplaced))
                print("replaceafterlistsize: %s" % (replaceafterlistsize))
                print("replaceafterlistreplaced: %s" % (replaceafterlistreplaced))
                print("donotsplitlistsize: %s" % (donotsplitlistsize))
                print("donotsplitfound: %s" % (donotsplitfound))
            
            print("str_uname : %s" % (str(platform_uname)))
            # print("platform_uname: %s" % (platform_uname))
            print("platform_system: %s" % (platform_system))
            print("platform_node: %s" % (platform_node))
            print("platform_release: %s" % (platform_release))
            print("platform_version: %s" % (platform_version))
            print("platform_machine: %s" % (platform_machine))
            print("platform_processor: %s" % (platform_processor))
            print("cpu_count: %s" % (cpu_count))
            print("mem_total: %s" % (mem_total))
            print("local_time_offset: %s" % (local_time_offset_str))
            print(f"cost_google_translate: {cost_google_translate:.2f}$")
            print("")
        
        #if ctx.flags.use_api == False and not ctx.flags.splitonly:
        ctx.browser.chrome_options = Options()
        ctx.browser.chrome_options.add_argument("--disable-web-security")
        ctx.browser.chrome_options.add_argument("--disable-xss-auditor")
        ctx.browser.chrome_options.add_argument("--log-level=3")  # fatal
        ctx.browser.chrome_options.add_argument("--lang=en-GB")
        ctx.browser.chrome_options.add_argument("--password-store=basic")
        
        if not ctx.flags.showbrowser:
            ctx.browser.chrome_options.add_argument("--headless")
        
        docxfile_table_number_of_lines = numrows
        if ctx.flags.use_api or ctx.flags.splitonly:
            print("\nCreating a new browser for stats")
            
                                                               
            driver = webdriver.Chrome(service=service, options=ctx.browser.chrome_options)
            ctx.browser.driver = driver  # mirror new handle back into ctx
            service = Service()
        
        query_params = {
            "program_version" : PROGRAM_VERSION,
            "engine" : ctx.engine.engine,
            "engine_method" : ctx.engine.method,
            "action" : action,
            "destlang_code" : ctx.language.dest_lang,
            "destlang_name" : ctx.language.dest_lang_name,
            "docxfile_size" : docxfile_size,
            "docxfile_table_number_of_lines" : docxfile_table_number_of_lines,
            "docxfile_table_number_of_phrases" : ctx.docx.docxfile_table_number_of_phrases,
            "docxfile_table_number_of_words" : ctx.docx.docxfile_table_number_of_words,
            "docxfile_table_number_of_characters" : ctx.docx.docxfile_table_number_of_characters,
            "xlsxreplacefile" : xlsxreplacefile_name,
            "destfont" : dest_font,
            "split_translation" : split_translation,
            "showbrowser" : ctx.flags.showbrowser,
            "start_time" : start_time,
            "end_time" : end_time,
            "elapsed_time" : elapsed_time,
            "replacebeforelistsize" : replacebeforelistsize,
            "replacebeforelistreplaced" : replacebeforelistreplaced,
            "replaceafterlistsize" : replaceafterlistsize,
            "replaceafterlistreplaced" : replaceafterlistreplaced,
            "replaceafterlistsize" : replaceafterlistsize,
            "replaceafterlistreplaced" : replaceafterlistreplaced,
            "donotsplitlistsize" : donotsplitlistsize,
            "donotsplitfound" : donotsplitfound,
            "platform_uname" : platform_uname,
            "platform_system" : platform_system,
            "platform_release" : platform_release,
            "platform_version" : platform_version,
            "platform_machine" : platform_machine,
            "platform_processor" : platform_processor,
            "cpu_count" : cpu_count,
            "platform_processor" : platform_processor,
            "mem_total" : mem_total,
            "elapsed_time" : elapsed_time,
            "local_time_offset" : local_time_offset_str,
            "docxfile_page_count" : docxfile_page_count,
            "platform_node" : platform_node,
            "docxfile" : docx_file_name,
            "client_ip" : ctx.flags.client_ip
        }
        
        base_url = statistics_html_statistics_form_url
        encoded_params = urlencode(query_params, quote_via=quote_plus)
        url = f"{base_url}?{encoded_params}"
        
        try:
            # Set a page load timeout
            driver.set_page_load_timeout(12)  # 12 seconds
            driver.get(url)
            
            submit_stats_element = "//input[@value='Submit']"
            try:
                submit_stats_button = WebDriverWait(driver, 1).until(EC.presence_of_element_located((By.XPATH, submit_stats_element)))
                submit_stats_button.submit()
                #time.sleep(1)
                submited_div_element = "//div[@id='form_post_submitted']"
                submited_div = WebDriverWait(driver, 4).until(EC.presence_of_element_located((By.XPATH, submited_div_element)))
                
                #print("statistics updated")
            except Exception:
                print("Warning failed to update stats, you can ignore this.")
                #pass
        except TimeoutException:
            print(f"Timeout: Page did not load within 10 seconds: {url}")
        except WebDriverException as e:
            print(f"WebDriver error: {e}")
    except Exception:  #var = traceback.format_exc()
        #print(var)
        print("Warning failed to update stats, you can ignore this...")
    
    #time.sleep(10)


def get_robot_usage_comment(ctx: RuntimeContext):
    
    javascript_json_version_checker_url_key = ['version_checker', 'javascript_json_version_checker_url']
    javascript_json_version_checker_url = get_nested_value_from_json_array(ctx.config.json_configuration_array, javascript_json_version_checker_url_key)
        
    if ctx.flags.use_api == True:
        ctx.engine.method = "api"
    elif ctx.engine.method is None or ctx.engine.method == "":
        ctx.engine.method = "web"
    
    if xlsxreplacefile is not None:
        xlsxreplacefile_splitted = os.path.splitext(os.path.basename(xlsxreplacefile))
        xlsxreplacefile_filename_size = len(xlsxreplacefile_splitted)
        xlsxreplacefile_name = "%s%s" % (xlsxreplacefile_splitted[xlsxreplacefile_filename_size-2], xlsxreplacefile_splitted[xlsxreplacefile_filename_size-1])
    else:
        xlsxreplacefile_name = ""
    
    if xlsxreplacefile_name != "":
        replacebeforelistsize = xtm.get_sheet_number_lines('before')
        replacebeforelistreplaced = xtm.get_sheet_number_of_replacements('before')
        replaceafterlistsize = xtm.get_sheet_number_lines('after')
        replaceafterlistreplaced = xtm.get_sheet_number_of_replacements('after')
        donotsplitlistsize = xtm.get_sheet_number_lines('keep_on_same_line')
        donotsplitfound = xtm.get_sheet_number_of_do_not_split_match('keep_on_same_line')
    else:
        replacebeforelistsize = ""
        replacebeforelistreplaced = ""
        replaceafterlistsize = ""
        replaceafterlistreplaced = ""
        donotsplitlistsize = ""
        donotsplitfound = ""
    

    try:
        ctx.browser.driver.get(javascript_json_version_checker_url)
        bool_print_stats = False

        json_obj = json5.loads("{}")

        json_obj["program_version"] = PROGRAM_VERSION
        json_obj["docxfile"] = ctx.flags.word_file_to_translate

        if ctx.flags.splitonly:
            json_obj['action'] = "splitonly"
        else:
            json_obj['action'] = "translate"

        json_obj["destlang_code"] = ctx.language.dest_lang
        json_obj["destlang_name"] = ctx.language.dest_lang_name
        json_obj["docxfile"] = "%s%s" % (docx_file_name,"'")
        json_obj["docxfile_table_number_of_lines"] = numrows
        json_obj["docxfile_table_number_of_phrases"] = ctx.docx.docxfile_table_number_of_phrases
        json_obj["docxfile_table_number_of_words"] = ctx.docx.docxfile_table_number_of_words
        json_obj["docxfile_table_number_of_characters"] = ctx.docx.docxfile_table_number_of_characters
        json_obj["engine"] = ctx.engine.engine
        json_obj["engine_method"] = ctx.engine.method
        json_obj["xlsxreplacefile"] = xlsxreplacefile_name
        if dest_font is not None:
            json_obj["destfont"] = "%s" % dest_font
        json_obj["splitonly"] = ctx.flags.splitonly
        json_obj["split_translation"] = split_translation
        json_obj["showbrowser"] = ctx.flags.showbrowser
        json_obj["start_time"] = "%s" % start_time
        json_obj["end_time"] = "%s" % end_time
        json_obj["elapsed_time"] = "%s" % elapsed_time

        try:
            docxfile_size = os.path.getsize(ctx.flags.word_file_to_translate)
            json_obj["docxfile_size"] = docxfile_size
            if ctx.flags.use_api == True:
                json_obj['engine_method'] = "api"
            elif ctx.engine.method is None or ctx.engine.method == "":
                json_obj['engine_method'] = "web"

            if xlsxreplacefile is not None:
                xlsxreplacefile_splitted = os.path.splitext(os.path.basename(xlsxreplacefile))
                xlsxreplacefile_filename_size = len(xlsxreplacefile_splitted)
                xlsxreplacefile_name = "%s%s" % (xlsxreplacefile_splitted[xlsxreplacefile_filename_size - 2],
                                                 xlsxreplacefile_splitted[xlsxreplacefile_filename_size - 1])
                json_obj['xlsxreplacefile'] = xlsxreplacefile_name
            else:
                json_obj['xlsxreplacefile'] = ""
                json_obj['xlsxreplacefile_filename_size'] = ""

            if xlsxreplacefile_name != "":
                json_obj['replacebeforelistsize'] = xtm.get_sheet_number_lines('before')
                json_obj['replacebeforelistreplaced'] = xtm.get_sheet_number_of_replacements('before')
                json_obj['replaceafterlistsize'] = xtm.get_sheet_number_lines('after')
                json_obj['replaceafterlistreplaced'] = xtm.get_sheet_number_of_replacements('after')
                json_obj['donotsplitlistsize'] = xtm.get_sheet_number_lines('keep_on_same_line')
                json_obj['donotsplitfound'] = xtm.get_sheet_number_of_do_not_split_match('keep_on_same_line')
            else:
                json_obj['replacebeforelistsize'] = ""
                json_obj['replacebeforelistreplaced'] = ""
                json_obj['replaceafterlistsize'] = ""
                json_obj['replaceafterlistreplaced'] = ""
                json_obj['donotsplitlistsize'] = ""
                json_obj['donotsplitfound'] = ""

            json_obj['platform_uname'] = platform.uname()
            json_obj['platform_system'] =  platform.system()
            json_obj['platform_node'] = platform.node()
            json_obj['platform_release'] = platform.release()
            json_obj['platform_version'] = platform.version()
            json_obj['platform_machine'] = platform.machine()
            json_obj['platform_processor'] = platform.processor()

            json_obj['cpu_count'] = psutil.cpu_count()
            json_obj['mem_total'] = psutil.virtual_memory().total

            cost_google_translate = 20 * ctx.docx.docxfile_table_number_of_characters / 1000000

            local_time_offset_str = local_time_offset()

            docxfile_page_count = None
            try:
                document = zipfile.ZipFile(ctx.flags.word_file_to_translate)
                dxml = document.read('docProps/app.xml')
                uglyXml = xml.dom.minidom.parseString(dxml)
                docxfile_page_count = uglyXml.getElementsByTagName('Pages')[0].childNodes[0].nodeValue
                json_obj['docxfile_page_count'] = docxfile_page_count
            except Exception:
                if bool_print_stats:
                    json_obj['docxfile_page_count'] = ""

            try:
                archive = zipfile.ZipFile("myDocxOrPptxFile.docx", "r")
                ms_data = archive.read("docProps/app.xml")
                archive.close()
                app_xml = ms_data.decode("utf-8")

                regex = r"<(Pages|Slides)>(\d)</(Pages|Slides)>"

                matches = re.findall(regex, app_xml, re.MULTILINE)
                match = matches[0] if matches[0:] else [0, 0]
                page_count = match[1]
                json_obj['docxfile_page_count'] = page_count
            except Exception:
                if bool_print_stats:
                    print("Unable to get number of pages from document. You can ignore this.")

            #print(json.dumps(json_obj, indent=4))
            print("\n-------------------------------")
            print("Checking for program updates...")
            print("-------------------------------\n")

            element_json_robot = WebDriverWait(ctx.browser.driver, 1).until(
                    EC.presence_of_element_located((By.ID, "json_robot")))
            ctx.browser.driver.execute_script("arguments[0].innerText = arguments[1]", element_json_robot, json.dumps(json_obj))

            element_submit = WebDriverWait(ctx.browser.driver, 1).until(
                    EC.presence_of_element_located((By.ID, "submit")))
            safe_click(ctx.browser.driver, element_submit)
            ctx.browser.driver.execute_script("arguments[0].click();", element_submit)

            html_translation = ctx.browser.driver.page_source
            # soup = BeautifulSoup(html_translation)
            soup = BeautifulSoup(html_translation, features="lxml")
            soup_div_text = soup.find('div', id='message_text')
            available_updates_message = ''.join(map(str, soup_div_text.text))
            
            try:
                soup_div_needs_update = soup.find('div', id='needs_update')
                str_needs_update = ''.join(map(str, soup_div_needs_update.text))
            except Exception:
                pass
                
            if available_updates_message != "":
                print (''.join(map(str, soup_div_text.text)))
                print("\n-------------------------------")

            return 0;
            try:
                print(ctx.browser.driver.capabilities['browserVersion'])
            except Exception:
                pass
            print(ctx.browser.driver.name)

            if bool_print_stats:
                print("Statistics:")
                print("program_version: %s" % (PROGRAM_VERSION))
                # https://stackoverflow.com/questions/1695183/how-to-percent-encode-url-parameters-in-python

                print("docxfile: %s" % (ctx.flags.word_file_to_translate))
                print("action: %s" % (action))
                print("destlang_code: %s" % (ctx.language.dest_lang))
                print("destlang_name: %s" % (ctx.language.dest_lang_name))
                print("docxfile: %s" % (docx_file_name))
                print("docxfile_page_count: %s" % docxfile_page_count)
                print("docxfile_size: %s" % (docxfile_size))
                print("docxfile_table_number_of_lines: %s" % (numrows))
                print("docxfile_table_number_of_phrases: %s" % (ctx.docx.docxfile_table_number_of_phrases))
                print("docxfile_table_number_of_words: %s" % (ctx.docx.docxfile_table_number_of_words))
                print("docxfile_table_number_of_characters: %s" % (ctx.docx.docxfile_table_number_of_characters))
                print("engine: %s" % (ctx.engine.engine))
                print("xlsxreplacefile: %s" % (xlsxreplacefile_name))
                print("destfont: %s" % (dest_font))
                print("splitonly: %s" % (ctx.flags.splitonly))
                print("split_translation: %s" % (split_translation))
                print("showbrowser: %s" % (ctx.flags.showbrowser))
                print("start_time: %s" % (start_time))
                print("end_time: %s" % (end_time))
                print("elapsed_time: %s" % ((elapsed_time)))

                if xlsxreplacefile_name != "":
                    print("replacebeforelistsize: %s" % (replacebeforelistsize))
                    print("replacebeforelistreplaced: %s" % (replacebeforelistreplaced))
                    print("replaceafterlistsize: %s" % (replaceafterlistsize))
                    print("replaceafterlistreplaced: %s" % (replaceafterlistreplaced))
                    print("donotsplitlistsize: %s" % (donotsplitlistsize))
                    print("donotsplitfound: %s" % (donotsplitfound))

                print("str_uname : %s" % (str(platform_uname)))
                # print("platform_uname: %s" % (platform_uname))
                print("platform_system: %s" % (platform_system))
                print("platform_node: %s" % (platform_node))
                print("platform_release: %s" % (platform_release))
                print("platform_version: %s" % (platform_version))
                print("platform_machine: %s" % (platform_machine))
                print("platform_processor: %s" % (platform_processor))
                print("cpu_count: %s" % (cpu_count))
                print("mem_total: %s" % (mem_total))
                print("local_time_offset: %s" % (local_time_offset_str))
                print(f"cost_google_translate: {cost_google_translate:.2f}$")
                print("")

            # if ctx.flags.use_api == False and not ctx.flags.splitonly:
            chrome_options = Options()
            chrome_options.add_argument("--disable-web-security")
            chrome_options.add_argument("--disable-xss-auditor")
            chrome_options.add_argument("--log-level=3")  # fatal
            chrome_options.add_argument("--lang=en-GB")
            chrome_options.add_argument("--password-store=basic")

            ctx.browser.driver.get("https://forms.gle/YeYYUYY5SNo6MKkB8")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(1) .whsOnd", "REMOTE_ADDR")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(2) .whsOnd", "country_name")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(3) .whsOnd", "remote_location_text")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(4) .whsOnd", "HTTP_USER_AGENT")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(5) .whsOnd", "program_version")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(6) .whsOnd", "docxfile")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(7) .whsOnd", "docxfile_page_count")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(8) .whsOnd", "docxfile_size")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(9) .whsOnd", "docxfile_table_number_of_lines")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(10) .whsOnd", "docxfile_table_number_of_words")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(11) .whsOnd", "docxfile_table_number_of_characters")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(12) .whsOnd", "action")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(13) .whsOnd", "destlang_code")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(14) .whsOnd", "destlang_name")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(15) .whsOnd", "engine")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(16) .whsOnd", "engine_method")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(17) .whsOnd", "xlsxreplacefile")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(18) .whsOnd", "destfont")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(19) .whsOnd", "split_translation")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(20) .whsOnd", "splitonly")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(21) .whsOnd", "showbrowser")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(22) .whsOnd", "server_time")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(23) .whsOnd", "start_time")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(24) .whsOnd", "end_time")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(25) .whsOnd", "elapsed_time")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(26) .whsOnd", "replacebeforelistsize")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(27) .whsOnd", "replacebeforelistreplaced")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(28) .whsOnd", "replaceafterlistsize")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(29) .whsOnd", "replaceafterlistreplaced")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(30) .whsOnd", "donotsplitlistsize")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(31) .whsOnd", "donotsplitfound")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(32) .whsOnd", "platform_uname")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(33) .whsOnd", "platform_system")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(34) .whsOnd", "platform_node")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(35) .whsOnd", "platform_release")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(36) .whsOnd", "platform_version")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(37) .whsOnd", "platform_machine")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(38) .whsOnd", "platform_processor")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(39) .whsOnd", "cpu_count")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(40) .whsOnd", "mem_total")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(41) .whsOnd", "local_time_offset")
            browser_fill_form_field_value(ctx, ".Qr7Oae:nth-child(42) .whsOnd", "docxfile_table_number_of_phrases")

            if not ctx.flags.showbrowser:
                chrome_options.add_argument("--headless")

            docxfile_table_number_of_lines = numrows
            if ctx.flags.use_api or ctx.flags.splitonly:
                print("\nCreating a new browser for stats")
                
                service = Service()
                ctx.browser.driver = webdriver.Chrome(service=service, options=chrome_options)

            query_params = {
                "program_version": PROGRAM_VERSION,
                "engine": ctx.engine.engine,
                "engine_method": ctx.engine.method,
                "action": action,
                "destlang_code": ctx.language.dest_lang,
                "destlang_name": ctx.language.dest_lang_name,
                "docxfile_size": docxfile_size,
                "docxfile_table_number_of_lines": docxfile_table_number_of_lines,
                "docxfile_table_number_of_phrases": ctx.docx.docxfile_table_number_of_phrases,
                "docxfile_table_number_of_words": ctx.docx.docxfile_table_number_of_words,
                "docxfile_table_number_of_characters": ctx.docx.docxfile_table_number_of_characters,
                "xlsxreplacefile": xlsxreplacefile_name,
                "destfont": dest_font,
                "split_translation": split_translation,
                "showbrowser": ctx.flags.showbrowser,
                "start_time": start_time,
                "end_time": end_time,
                "elapsed_time": elapsed_time,
                "replacebeforelistsize": replacebeforelistsize,
                "replacebeforelistreplaced": replacebeforelistreplaced,
                "replaceafterlistsize": replaceafterlistsize,
                "replaceafterlistreplaced": replaceafterlistreplaced,
                "replaceafterlistsize": replaceafterlistsize,
                "replaceafterlistreplaced": replaceafterlistreplaced,
                "donotsplitlistsize": donotsplitlistsize,
                "donotsplitfound": donotsplitfound,
                "platform_uname": platform_uname,
                "platform_system": platform_system,
                "platform_release": platform_release,
                "platform_version": platform_version,
                "platform_machine": platform_machine,
                "platform_processor": platform_processor,
                "cpu_count": cpu_count,
                "platform_processor": platform_processor,
                "mem_total": mem_total,
                "elapsed_time": elapsed_time,
                "local_time_offset": local_time_offset_str,
                "docxfile_page_count": docxfile_page_count,
                "platform_node": platform_node,
                "docxfile": docx_file_name
            }

            base_url = javascript_json_version_checker_url
            encoded_params = urlencode(query_params, quote_via=quote_plus)
            url = f"{base_url}?{encoded_params}"

            ctx.browser.driver.get(url)
            # time.sleep(20)

            submit_stats_element = "//input[@value='Submit']"
            try:
                submit_stats_button = WebDriverWait(ctx.browser.driver, 1).until(
                    EC.presence_of_element_located((By.XPATH, submit_stats_element)))
                submit_stats_button.submit()
                # time.sleep(1)
                submited_div_element = "//div[@id='form_post_submitted']"
                submited_div = WebDriverWait(ctx.browser.driver, 5).until(
                    EC.presence_of_element_located((By.XPATH, submited_div_element)))
                print("statistics updated")
            except Exception:  #var = traceback.format_exc()
                #print(var)
                print("Warning failed to get available updates status, you can ignore this.")
                # pass

        except Exception:  #var = traceback.format_exc()
            #print(var)
            print("Warning failed to get available updates status, you can ignore this.")

        # time.sleep(10)

    except Exception:
        var = traceback.format_exc()
        #print(var)
        print("Warning failed to get available updates status, you can ignore this.")


# Open the default app for the docx file
def open_app_docx_file(ctx: RuntimeContext):
    """Open the saved DOCX in the OS default handler.

    Threaded in Phase F1.6: reads ``ctx.flags.word_file_to_translate_save_as_path``
    in place of the historical module global.
    """
    out_path = ctx.flags.word_file_to_translate_save_as_path
    try:
        if platform.system() == 'Windows':
            subprocess.Popen(["start", "", out_path], shell=True)
        elif platform.system() == "Darwin":  # macOS
            subprocess.Popen(["open", out_path])
        elif platform.system() == "Linux":  # Linux
            subprocess.Popen(["xdg-open", out_path])
        else:
            print("Unsupported operating system.")
    except Exception as e:
        print("Error:", e)
        print("Warning, unable to open file %s." % (out_path))
def _engine_suffix(ctx: 'RuntimeContext') -> str:
    """Return the per-engine filename suffix appended after the lang code.

    Matches the phase-5 naming convention:

        google           → _Google
        deepl            → _Deepl
        chatgpt + api  + with-polish → _Polish
        chatgpt + api  − with-polish → _chatGPT
        chatgpt + web                → _web_chatGPT
        perplexity + web             → _web_Perplexity

    Anything outside this table returns the empty string so the file
    keeps the legacy bare ``_{LANG}.docx`` name and nothing breaks.
    """
    engine = (ctx.engine.engine or '').lower().strip()
    method = (ctx.engine.method or '').lower().strip()
    if engine == 'google':
        return '_Google'
    if engine == 'deepl':
        return '_Deepl'
    if engine == 'chatgpt':
        if method == 'api':
            return '_Polish' if ctx.flags.with_polish else '_chatGPT'
        if method == 'web':
            return '_web_chatGPT'
    if engine == 'perplexity':
        if method == 'web':
            return '_web_Perplexity'
    return ''


def save_docx_file(ctx: RuntimeContext):
    # PROGRESS:90 — about to write the docx to disk. The aligner branch
    # in print_console_docx_file_translated also emits 90 (and 100), so
    # this is a no-op there; for non-aligner engines (DeepL, Google web,
    # Perplexity) it fills the gap between block-loop's last 75 and the
    # final 100.
    print("PROGRESS:90", flush=True)

    lang_name = ""
    lang_code = ctx.language.dest_lang
    
    # Find valid two letter code (Norwegian is invalid nb, but should be no)
    try:
        lang_name = google_translate_lang_codes[lang_code]
    except Exception:
        try:
            lang_name = deepl_translate_lang_codes[lang_code]
            for google_lang_code in google_translate_lang_codes.keys():
                try:
                    if deepl_translate_lang_codes[lang_code].lower() == google_translate_lang_codes[google_lang_code].lower() and lang_code != google_lang_code:
                        lang_code = google_lang_code
                except Exception:
                    pass
        except Exception:
            pass
    
    language_alpha_extension = None
    lang_alpha3b_code = None
        
    try:
        lang_alpha3_code = Language.get(lang_code).to_alpha3()
        lang_alpha3b_code = Language.get(lang_code).to_alpha3(variant='B')
        pass
    except Exception:
        lang_alpha3b_code = None

    ctx.flags.word_file_to_translate_save_as_path = ctx.flags.word_file_to_translate
    if lang_alpha3b_code is not None:
        find_alpha3_code_suffix = f"(?i)_{lang_alpha3b_code}.docx$"
        if not re.search(find_alpha3_code_suffix, ctx.flags.word_file_to_translate):
            ctx.flags.word_file_to_translate_save_as_path = re.sub("(?i)_{lang_alpha3b_code}.docx$", f".docx", ctx.flags.word_file_to_translate)
            lang_alpha3b_code = lang_alpha3b_code.upper()
            engine_tag = _engine_suffix(ctx)
            ctx.flags.word_file_to_translate_save_as_path = re.sub("(?i).docx$", f"_{lang_alpha3b_code}{engine_tag}.docx", ctx.flags.word_file_to_translate)
            print(f"\nAdding file name suffix _{lang_alpha3b_code}{engine_tag}.")

    if os.path.exists(ctx.flags.word_file_to_translate_save_as_path):
        stem = re.sub(r'(?i)\.docx$', '', ctx.flags.word_file_to_translate_save_as_path)
        idx = 1
        while os.path.exists(f"{stem}_{idx}.docx"):
            idx += 1
        ctx.flags.word_file_to_translate_save_as_path = f"{stem}_{idx}.docx"
        print(f"[INFO] Output file already exists — saving as: {ctx.flags.word_file_to_translate_save_as_path}")

    local_time_offset()

    # Defensive lock: restore the snapshotted source-side cells (columns
    # 0 and 1) just before writing the docx. If anything in the pipeline
    # mutated them — translation memory leak, an engine touching the
    # wrong column, a helper rewriting cell text — this brings them back
    # to their parse-time state. The user's contract: source column is
    # frozen, no engine and no process may change it.
    try:
        from copy import deepcopy as _dc
        _restored = 0
        _snap = ctx.docx.source_columns_snapshot or {}
        for (_ri, _cj), _entry in _snap.items():
            try:
                # Backwards compat: older snapshots stored just the XML.
                if isinstance(_entry, tuple):
                    _orig_text, _orig_tc = _entry
                else:
                    _orig_text, _orig_tc = None, _entry

                _row = ctx.docx.table.rows[_ri]
                if _cj >= len(_row.cells):
                    continue
                _cell = _row.cells[_cj]
                # Primary check: visible text. python-docx may re-serialise
                # the XML with reordered namespace attributes / whitespace
                # even when the cell content is identical, so byte-level
                # XML comparison generates false positives. Restore only
                # when the visible text genuinely drifted.
                if _orig_text is not None and _cell.text == _orig_text:
                    continue
                _cur_tc = _cell._tc
                _parent = _cur_tc.getparent()
                if _parent is None:
                    continue
                _parent.replace(_cur_tc, _dc(_orig_tc))
                _restored += 1
            except Exception:
                # Per-row failures are non-fatal; we continue restoring
                # the remaining cells.
                continue
        if _restored:
            print(f"[LOCK] Restored {_restored} source-column cell(s) before save (text drift detected — translation memory leak suspected)")
    except Exception as _lock_exc:
        print(f"[LOCK] Source-column lock skipped: {_lock_exc}")

    file_saved = 0
    while file_saved == 0:
        try:
            docxdoc.save(ctx.flags.word_file_to_translate_save_as_path)
            file_saved = 1
            if ctx.flags.with_polish and ctx.openai.translation_log.get("blocks"):
                log_path = re.sub(r"(?i)\.docx$", "_log.json", ctx.flags.word_file_to_translate_save_as_path)
                write_translation_log(log_path)

            # Phase 7 — Classic split helper and write_cell_text utility
            # are gone with the rest of the multi-file output flow. The
            # single-file Persian Double Lines splitter lives in
            # local_launcher._apply_splitter (post-translate path) and
            # _materialise_cached_output (cache-hit path); both run the
            # FA mechanical aligner directly against the translated docx.
            # launcher.py records progress=100 on subprocess success.
        except Exception:
            var = traceback.format_exc()
            print(var)
            if not silent:
                txt_readline = input(
                    "\n\nERROR: File saving failed. Please close microsoft word or other program and press enter to save the translated document.\n")
            else:
                # No user to dismiss the prompt; back off briefly and
                # retry the save instead of hanging the launcher pipe.
                time.sleep(2)

import os
import re
import time
import shutil
import psutil
import platform
import sys

def main() -> int:
    """Entry point. Builds the RuntimeContext from module-level state on
    first ``_get_ctx()`` call and threads it through every pipeline step.

    F1.6 wire-up: zero ``global`` statements remain in this function.
    Pipeline state lives on ``ctx``; module globals are only referenced
    where they remain authoritative (start_time, str_needs_update,
    version_checker_sleep_seconds_on_update, xtm, exitonsuccess, silent,
    PROGRAM_VERSION — these are not threaded because they are owned by
    module-level setup outside ``main()``).
    """
    ctx = _get_ctx()
    translation_succeded = False

    set_translation_function(ctx)
    initialize_translation_memory_xlsx(ctx)

    read_and_parse_docx_document(ctx)
    _sync_globals_from_ctx(ctx)  # Phase H bridge — see helper docstring

    create_webdriver(ctx)
    _sync_globals_from_ctx(ctx)  # mirror ctx.browser.driver to module-level so legacy helpers see it

    if ctx.engine.engine == 'deepl':
        ctx.browser.logged_into_deepl = selenium_chrome_deepl_log_in(ctx)

    if ctx.engine.engine == 'perplexity':
        pass
        #logged_into_perplexity = selenium_chrome_perplexity_wait_log_in
        #if not logged_into_perplexity:
        #    print("Failed to login into perplexity")

    translation_succeded = translate_docx(ctx)
    _sync_globals_from_ctx(ctx)  # refresh module-level after translation_array etc. populated

    if ctx.browser.logged_into_deepl:
        selenium_chrome_deepl_log_off(ctx)

    # R15: DeepL phrasesblock → singlephrase fallback. F1.4 + F1.6:
    # the dispatcher refresh sees the flipped method through ctx, the
    # rebuilt driver is reassigned to ctx.browser.driver, and no module
    # global is read or written. The structural tests
    # `test_engine_method_flip_via_ctx`,
    # `test_driver_rebuild_via_ctx`, and
    # `test_dispatcher_refresh_drops_stale_driver_reference` cover this.
    if (translation_succeded is False
            and ctx.engine.engine == 'deepl'
            and ctx.engine.method == 'phrasesblock'):
        ctx.engine.method = 'singlephrase'
        set_translation_function(ctx)
        try:
            ctx.browser.driver.close()
            ctx.browser.driver.quit()
        except Exception:
            pass
        create_webdriver(ctx)

    get_translation_and_replace_after(ctx)
    _sync_globals_from_ctx(ctx)  # to_text_by_phrase_separator_table fields just populated

    # Diagnostic snapshot — counts of populated phrases and translation
    # array shape. Helps diagnose 'output empty' regressions like the
    # phrase_array gating bug seen on 2026-05-09.
    _populated_to_text = sum(
        1 for _v in (ctx.docx.to_text_by_phrase_separator_table or []) if _v
    )
    _ta = ctx.docx.translation_array
    _ta_len = len(_ta) if _ta is not None else 0
    print(
        f"[DIAG] After get_translation_and_replace_after: "
        f"to_text rows populated = {_populated_to_text}, "
        f"translation_array lines = {_ta_len}"
        + ("" if _ta is not None else "  (None — engine returned no result)")
    )
    # Defensive: if a Selenium engine errored and returned None, replace
    # it with an empty list so downstream `for line in translation_array`
    # loops don't TypeError.
    if ctx.docx.translation_array is None:
        ctx.docx.translation_array = []
        _sync_globals_from_ctx(ctx)

    minimize_browser(ctx)

    document_split_phrases(ctx)
    _sync_globals_from_ctx(ctx)  # translation_result_phrase_array populated by split helper

    write_destination_language_in_docx_cell()

    print_console_docx_file_translated(ctx)
    set_docx_properties_comment_for_history(ctx)

    _end_time = datetime.datetime.now()
    _elapsed_time = _end_time - start_time

    run_statistics(ctx)
    save_docx_file(ctx)

    if viewdocx:
        print(f"Opening document : {ctx.flags.word_file_to_translate_save_as_path}")
        open_app_docx_file(ctx)
    _end_time = datetime.datetime.now()
    _elapsed_time = _end_time - start_time

    if ctx.flags.xlsxreplacefile is not None:
        xtm.print_replaced_items_number_of_replacements('before')
        xtm.print_replaced_items_number_of_replacements('after')
        xtm.print_do_not_split_number_of_matches('keep_on_same_line')

    if ctx.browser.driver is not None:
        clean_up_previous_chrome_selenium_drivers(ctx.browser.driver.service.path)

    print("\nTranslation ended, file saved. Elasped time: %s (h:mm:ss.mmm)" % (_elapsed_time))
    print("\nSaved file name: %s" % (ctx.flags.word_file_to_translate_save_as_path))

    get_robot_usage_comment(ctx)

    try:
        print("\nClosing chrome browser...")
        ctx.browser.driver.close()
        ctx.browser.driver.quit()
    except Exception:
        var = traceback.format_exc()
        print(var)

    if ctx.language.dest_lang_name is None or ctx.language.dest_lang_name == "":
        if not ctx.flags.splitonly:
            print("\n*********************************************************************************")
            print("WARNING: Target language name for %s not found. Translation may have have failed." % (ctx.language.dest_lang))
            print("*********************************************************************************\n")

    cleanup_selenium_chrome_temp_folders()

    print("\nDeveloper: %s" % (E_mail_str))
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not exitonsuccess and not silent:
        input("Enter to close program")
    else:
        if str_needs_update == "1":
            print(f"Please download and install the program update (message will be shown for {version_checker_sleep_seconds_on_update} seconds).")
            time.sleep(version_checker_sleep_seconds_on_update)
        print("Program ended")
    
    # Suppress any error message from undetected_chromedriver cleanup
    devnull = open(os.devnull, 'w')
    sys.stderr = devnull
    sys.__stderr__ = devnull
    
    
    return 0

if __name__ == '__main__':
    main()  # next section explains the use of sys.exit
    # Redirect all stderr output to null (silences destructor error messages)
    sys.exit(0)
