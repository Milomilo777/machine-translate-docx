"""Per-cell write helpers for the destination column.

Extracted from the entry script in the 2026-05-10 docx_io extraction
pass. Three functions:

  * :func:`change_cell_font` — apply a font name to every run in a cell.
  * :func:`set_first_paragraph` — write text into ``cell.paragraphs[0]``
    with RTL handling for Persian / Arabic / Hebrew etc.
  * :func:`add_paragraph` — append a new paragraph to the cell with
    the same RTL handling.

The functions take their dependencies (``dest_lang``, ``dest_font``,
``rtlstyle``) as explicit arguments rather than reading them from
module globals. The thin shim wrappers in the entry script
(``cell_set_1st_paragraph``, ``cell_add_paragraph``,
``change_cell_font``) read the entry-script globals and pass them
through, so callers' signatures stay unchanged.

This is the seam that lets the per-cell write path be unit-tested
in isolation, without spinning up the whole RuntimeContext.
"""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import right_to_left_languages_list


__all__ = [
    "change_cell_font",
    "set_first_paragraph",
    "add_paragraph",
]


def change_cell_font(cell, dest_font: str) -> None:
    """Apply ``dest_font`` to every run in every paragraph of ``cell``.

    No-op if ``dest_font`` is empty (the entry-script shim already
    short-circuits but we double-check here so direct callers behave
    the same).
    """
    if not dest_font:
        return
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = dest_font


def _write_into_paragraph(
    cell_paragraph,
    paragraph_text: str,
    dest_lang: str,
    rtlstyle,
) -> None:
    """Shared body of :func:`set_first_paragraph` and :func:`add_paragraph`.

    For RTL targets, builds an RTL run with the per-document
    ``rtlstyle`` style and right-aligns the paragraph. For LTR
    targets, writes the text directly and applies ``Normal`` (or
    ``Default Paragraph Font`` if ``Normal`` is missing).
    """
    if dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(paragraph_text, style="rtlstyle")
        run.style = rtlstyle
        run.font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = paragraph_text
        try:
            cell_paragraph.style = "Normal"
        except KeyError:
            try:
                cell_paragraph.style = "Default Paragraph Font"
            except KeyError:
                # No usable default style — proceed without
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_first_paragraph(
    cell,
    paragraph_text: str,
    *,
    dest_lang: str,
    dest_font: str,
    rtlstyle,
) -> None:
    """Replace the text of ``cell.paragraphs[0]`` with ``paragraph_text``.

    RTL handling per :data:`right_to_left_languages_list`. Optionally
    applies ``dest_font`` afterwards.
    """
    cell_paragraph = cell.paragraphs[0]
    _write_into_paragraph(cell_paragraph, paragraph_text, dest_lang, rtlstyle)
    if dest_font:
        change_cell_font(cell, dest_font)


def add_paragraph(
    cell,
    paragraph_text: str,
    *,
    dest_lang: str,
    dest_font: str,
    rtlstyle,
) -> None:
    """Append a new paragraph carrying ``paragraph_text`` to ``cell``.

    RTL handling per :data:`right_to_left_languages_list`. Optionally
    applies ``dest_font`` afterwards.
    """
    cell_paragraph = cell.add_paragraph("")
    _write_into_paragraph(cell_paragraph, paragraph_text, dest_lang, rtlstyle)
    if dest_font:
        change_cell_font(cell, dest_font)
