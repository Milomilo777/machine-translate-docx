"""Per-cell read + write helpers for the destination column.

Extracted from the entry script across two passes:

  * 2026-05-10 docx_io initial extraction — three write helpers
    (``change_cell_font``, ``set_first_paragraph``, ``add_paragraph``).
  * 2026-05-10 G2 thread-globals pass — :func:`get_cell_data`, the read
    helper that walks every paragraph + run in a source cell, plus its
    private shading-detection helpers (:func:`_paragraph_shading_color`,
    :func:`_run_shading_color`).

The write helpers take their dependencies (``dest_lang``, ``dest_font``,
``rtlstyle``) as explicit arguments rather than reading them from
module globals. The thin shim wrappers in the entry script
(``cell_set_1st_paragraph``, ``cell_add_paragraph``,
``change_cell_font``) read the entry-script globals and pass them
through, so callers' signatures stay unchanged.

The read helper :func:`get_cell_data` takes ``ctx`` (already its
historical signature) and reads the colour-ignore list from
``ctx.config.shading_color_ignore_text``, so it depends only on
``RuntimeContext``.

This is the seam that lets the per-cell read + write paths be
unit-tested in isolation, without spinning up the whole pipeline.
"""
from __future__ import annotations

import re
from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from lxml import etree

from ..config import right_to_left_languages_list
from ..docx_io.runs import _iter_paragraph_runs

if TYPE_CHECKING:
    from ..runtime import RuntimeContext


__all__ = [
    "change_cell_font",
    "set_first_paragraph",
    "add_paragraph",
    "get_cell_data",
]


# ── shading colour detection (private) ───────────────────────────────────────

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_FILL_ATTR = f"{{{_W_NS}}}fill"


def _shading_fill_color(xml_str: str | bytes, *, xpath: str) -> str | None:
    """Return the ``w:fill`` attribute of the first matching ``<w:shd>``.

    ``xpath`` is the search path under the parsed element — either
    ``.//w:pPr/w:shd`` (paragraph shading) or ``.//w:rPr/w:shd``
    (run shading). Returns ``None`` if no shading element exists.
    """
    xml = etree.fromstring(xml_str)
    namespaces = {"w": _W_NS}
    try:
        # Prefer the prefix the document actually declares — Word docs
        # in the wild have used `w14` and other variants in addition to
        # the canonical `w`.
        if xml.prefix:
            namespaces = {xml.prefix: xml.nsmap[xml.prefix]}
    except Exception:
        pass
    fill = None
    for element in xml.findall(xpath, namespaces):
        try:
            fill = element.attrib.get(_W_FILL_ATTR)
        except Exception:
            pass
    return fill


def _paragraph_shading_color(xml_paragraph_str: str | bytes) -> str | None:
    """Return the paragraph-level shading colour, or ``None``."""
    return _shading_fill_color(xml_paragraph_str, xpath=".//w:pPr/w:shd")


def _run_shading_color(xml_run_str: str | bytes) -> str | None:
    """Return the run-level shading colour, or ``None``."""
    return _shading_fill_color(xml_run_str, xpath=".//w:rPr/w:shd")


def change_cell_font(cell, dest_font: str) -> None:
    """Apply ``dest_font`` to every run in every paragraph of ``cell``.

    No-op if ``dest_font`` is empty (the entry-script shim already
    short-circuits but we double-check here so direct callers behave
    the same).
    """
    if not dest_font:
        return
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = dest_font


def _write_into_paragraph(
    cell_paragraph,
    paragraph_text: str,
    dest_lang: str,
    rtlstyle,
) -> None:
    """Shared body of :func:`set_first_paragraph` and :func:`add_paragraph`.

    For RTL targets, builds an RTL run with the per-document
    ``rtlstyle`` style and right-aligns the paragraph. For LTR
    targets, writes the text directly and applies ``Normal`` (or
    ``Default Paragraph Font`` if ``Normal`` is missing).
    """
    if dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(paragraph_text, style="rtlstyle")
        run.style = rtlstyle
        run.font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = paragraph_text
        try:
            cell_paragraph.style = "Normal"
        except KeyError:
            try:
                cell_paragraph.style = "Default Paragraph Font"
            except KeyError:
                # No usable default style — proceed without
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_first_paragraph(
    cell,
    paragraph_text: str,
    *,
    dest_lang: str,
    dest_font: str,
    rtlstyle,
) -> None:
    """Replace the text of ``cell.paragraphs[0]`` with ``paragraph_text``.

    RTL handling per :data:`right_to_left_languages_list`. Optionally
    applies ``dest_font`` afterwards.
    """
    cell_paragraph = cell.paragraphs[0]
    _write_into_paragraph(cell_paragraph, paragraph_text, dest_lang, rtlstyle)
    if dest_font:
        change_cell_font(cell, dest_font)


def add_paragraph(
    cell,
    paragraph_text: str,
    *,
    dest_lang: str,
    dest_font: str,
    rtlstyle,
) -> None:
    """Append a new paragraph carrying ``paragraph_text`` to ``cell``.

    RTL handling per :data:`right_to_left_languages_list`. Optionally
    applies ``dest_font`` afterwards.
    """
    cell_paragraph = cell.add_paragraph("")
    _write_into_paragraph(cell_paragraph, paragraph_text, dest_lang, rtlstyle)
    if dest_font:
        change_cell_font(cell, dest_font)


# ── source-cell read ─────────────────────────────────────────────────────────

_RE_PAUSE = re.compile(r"(?i)<pause>")
_RE_ENTER = re.compile(r"(?i)<enter>")
_RE_LINE_BREAKS = re.compile(r"[\r\n  ]+")
_RE_MULTI_SPACE = re.compile(r" +")


def get_cell_data(ctx: "RuntimeContext", cell, row_n: int):
    """Walk every paragraph + run of ``cell`` and return the cleaned text.

    Returns a 3-tuple ``(cell_non_greyed_text, cell_is_gray, cell_is_red)``
    where the two flags are ``int | None`` — ``None`` if no run carried
    a colour signal at all, ``0`` for "not flagged" and ``1`` for
    "flagged". Side-effect: updates
    ``ctx.docx.from_text_nb_lines_in_cell[row_n - 1]`` with the line
    count derived from ``<pause>`` / ``<enter>`` markers + paragraph
    count.

    Reads :attr:`ctx.config.shading_color_ignore_text` to decide which
    paragraph / run shadings count as "ignore this run". The colour
    list is populated in the entry script after the JSON configuration
    merge.
    """
    shading_color_ignore_text = ctx.config.shading_color_ignore_text or []

    cell_is_gray = None
    cell_is_red = None
    cell_non_greyed_text = ""

    n_paragraph = 0
    n_cell_lines = 1

    for paragraph in cell.paragraphs:
        n_paragraph += 1

        p_shading_color = _paragraph_shading_color(paragraph._p.xml)

        # Materialise the run list once so both p_text (used for
        # <pause> / <enter> counting) and the run-by-run loop below
        # see the same source of truth — including hyperlinked text.
        # Relying on paragraph.text would tie us to a python-docx
        # implementation detail that has flipped behaviour across
        # versions on whether <w:hyperlink> contents are included.
        paragraph_runs = list(_iter_paragraph_runs(paragraph))
        p_text = "".join(r.text for r in paragraph_runs)
        nb_pause = len(_RE_PAUSE.findall(p_text))
        nb_enter = len(_RE_ENTER.findall(p_text))

        n_cell_lines += nb_pause + nb_enter

        if p_shading_color is not None and p_shading_color in shading_color_ignore_text:
            continue

        # Walk every <w:r> below the paragraph, including those nested
        # inside <w:hyperlink>. Using paragraph.runs alone drops the
        # text of every clickable link in the document.
        for run in paragraph_runs:
            current_run_text = run.text

            run_shading_color = _run_shading_color(run.element.xml)
            # Note: the historical code only used run_shading_color in
            # the highlight check below — it did not skip on shaded
            # runs the way it skipped on shaded paragraphs.

            if str(run.font.color.rgb) == "FF0000":
                if cell_is_red is None:
                    cell_is_red = 1
            else:
                if current_run_text != "":
                    if cell_is_red is None:
                        cell_is_red = 0
                    else:
                        cell_is_red = cell_is_red * 0

            if run.font.highlight_color == WD_COLOR_INDEX.RED:
                pass

            if (
                run.font.highlight_color == WD_COLOR_INDEX.GRAY_25
                or run.font.highlight_color == WD_COLOR_INDEX.GRAY_50
                or run.font.strike
                or run.font.double_strike
                or run.font.highlight_color == WD_COLOR_INDEX.PINK
                or run.font.highlight_color == WD_COLOR_INDEX.RED
                or run_shading_color in shading_color_ignore_text
            ):
                cell_non_greyed_text += " "
                if cell_is_gray is None:
                    cell_is_gray = 1
            else:
                if current_run_text != "":
                    cell_non_greyed_text += current_run_text
                    if cell_is_gray is None:
                        cell_is_gray = 0
                    else:
                        cell_is_gray = cell_is_gray * 0

    ctx.docx.from_text_nb_lines_in_cell[row_n - 1] = n_cell_lines

    cell_non_greyed_text = cell_non_greyed_text.replace("’", "'")
    cell_non_greyed_text = cell_non_greyed_text.replace("\n", " ")
    cell_non_greyed_text = cell_non_greyed_text.replace("\r", " ")
    cell_non_greyed_text = _RE_LINE_BREAKS.sub(" ", cell_non_greyed_text)

    cell_non_greyed_text = _RE_PAUSE.sub("", cell_non_greyed_text)
    cell_non_greyed_text = _RE_ENTER.sub("", cell_non_greyed_text)

    cell_non_greyed_text = _RE_MULTI_SPACE.sub(" ", cell_non_greyed_text)
    cell_non_greyed_text = cell_non_greyed_text.strip()

    return cell_non_greyed_text, cell_is_gray, cell_is_red
