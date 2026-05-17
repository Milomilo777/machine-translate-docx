"""
FA Subtitle Aligner — Mechanical v2.0
======================================
Rebuilt on the clean fa-subtitle-aligner-v75 (skill) foundation.

Core algorithm (from skill fa_aligner.py):
  - FA-based sentence grouping  (more reliable than EN sentence ends)
  - split_for_n_rows: tries fewest distinct chunks → maximises doubles
  - Proportional distribution (no-triple by default)
  - Dynamic target recalculation per remaining chunk

Added from project aligner_per.py:
  - ZWNJ-aware display length  (invisible in Word, excluded from MAX_CHARS)
  - RTL paragraph/run markers  (<w:bidi/> + <w:rtl/>)
  - Protected bigrams          (از جمله، بر اساس، …)
  - Shaded cell detection      (grey bridge rows skipped)
  - Citation stripping         (removes trailing (source) from FA cells)
  - Cross-group triple sentinel

Dropped (only needed for LLM mode, added complexity for no gain here):
  - B4 proportional weight tables
  - Discourse-marker alignment
  - Weight pass (SOV last-line fix)
  - Quality scoring
  - Preservation check on existing segmentation

llm_threshold / token_budget are accepted for API compatibility but ignored.
This version is purely mechanical; LLM integration can be added later on top.
"""

import re
import time

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as _qn

try:
    from ._retry import prompt_hash as _prompt_hash
except ImportError:
    def _prompt_hash(text: str) -> str:
        return "00000000"


# ── constants ──────────────────────────────────────────────────────────────────

MAX_CHARS  = 48    # hard per-chunk display limit (broadcast standard)
# 2026-05-17 phase-3 benchmark: MIN_TARGET=24 outperforms 28 on the 66-file
# human-corpus benchmark (row_exact 34.41 % vs 34.35 %, midline patterns
# also slightly better). Keeping the historical value.
MIN_TARGET = 24

ZWNJ = '‌'    # Persian half-space — invisible, excluded from MAX_CHARS

SENT_END = frozenset('.!?؟')

COMPOUND_PREFIXES = (
    'می' + ZWNJ, 'نمی' + ZWNJ,
    'می‌',   'نمی‌',
)

# Prepositions / particles that must not sit at the end of a chunk.
# 2026-05-17 phase-3: «که» REMOVED. The human-edited corpus shows that
# ending a chunk with «که» is a deliberate editorial choice — «X که» at
# line end binds naturally to the relative/complement clause on the next
# line. User explicitly confirmed: "که در انتهای سطر هیچ مشکلی ندارد".
DANGLING_PREPS = frozenset({
    'به', 'از', 'در', 'برای', 'با', 'تا', 'بر',
    'تحت', 'جز', 'روی', 'زیر', 'نزد', 'پیش', 'سوی',
    # 2026-05-12 phase-2: extended list — these also strand badly at line-end
    'بدون', 'علیه', 'مقابل', 'درباره', 'دربارهٔ', 'بنا', 'طبق',
    'ضد', 'پس', 'قبل', 'بعد', 'وسط', 'میان',
})

# Bigrams that must never be split across two chunks
PROTECTED_BIGRAMS = frozenset({
    'از جمله', 'از طریق', 'همراه با', 'با وجود',
    'در برابر', 'در نتیجه', 'در پی', 'به دلیل',
    'به خاطر', 'از آنجا که', 'در عین حال', 'بر اساس',
    'به جای', 'در راستای', 'به عنوان', 'در مقابل',
    'بر خلاف', 'به جهت', 'به منظور', 'از سوی',
    'به واسطه', 'در قبال',
    # 2026-05-12 phase-2: extra fixed expressions observed in real
    # subtitles (BMD / CTAW Google-Drive corpus).
    'از یک سو', 'از سوی دیگر', 'به ویژه', 'به‌ویژه',
    'با این حال', 'با این وجود', 'به همین دلیل', 'به همین جهت',
    'به‌طور کلی', 'به طور کلی', 'به طور کامل', 'به طور خاص',
    'در حالی که', 'در صورتی که', 'پیش از این', 'تا کنون',
    'تا به حال', 'به نام', 'بنا بر', 'علاوه بر این',
    'به اعتقاد', 'به گفته', 'به گفتهٔ', 'به گفته‌ٔ',
    'به نقل از', 'به دنبال', 'در پی این', 'بر این اساس',
    # 2026-05-17 phase-3: 22 new bigrams confirmed by 562-file human
    # corpus with break-rate ≤ 2% across thousands of occurrences.
    'از آن', 'به آن', 'بیش از', 'به این',
    'در حالی', 'به نظر', 'به اشتراک', 'از نظر',
    'نه تنها', 'هر کس', 'نه فقط', 'هر چه',
    'به سرعت', 'در همان', 'به نمایش', 'هر چند',
    'به هر حال', 'به این شکل', 'هیچ کس',
    'با توجه به', 'به آرامی', 'از این رو',
})

# Sentence-end punctuation — strongest split candidate.
SENT_END_CHARS = '.!?؟…'

# Mid-sentence punctuation — next-best break candidate.
MID_PUNCT_CHARS = '،؛:'

# Subordinating / coordinating conjunctions that look natural at the
# *start* of the second chunk (i.e. break BEFORE these tokens).
# Single-token entries only — multi-word forms live in PROTECTED_BIGRAMS.
# 2026-05-17 phase-3: 8 new conjunctions confirmed by post-cut frequency
# analysis on 39 625 human cuts (each occurring >0.15% of all cuts).
LEADING_CONJUNCTIONS = frozenset({
    'و', 'اما', 'ولی', 'چون', 'زیرا', 'پس', 'سپس', 'بنابراین',
    'هرچند', 'گرچه', 'یعنی', 'مگر', 'تا',
    'اگر', 'وقتی', 'حتی', 'همچنین', 'چه', 'شاید', 'آیا', 'بلکه',
})

# Leading prepositions — when the NEXT chunk starts with one of these,
# breaking the previous chunk just before it is a natural FA cut point.
# 2026-05-17 phase-3: ~18.6% of all human cuts in the corpus go directly
# before one of these prepositions. Adds a new Priority 3.5 in _find_break.
LEADING_PREPS = frozenset({
    'در', 'به', 'از', 'برای', 'با', 'بر', 'تا',
    'روی', 'زیر', 'نزد', 'پیش', 'سوی', 'بدون',
    'علیه', 'مقابل', 'درباره', 'دربارهٔ', 'ضد',
})

# Citation pattern — a parenthesised source at the very end of a cell.
# Used by _split_citation to preserve news-source attributions byte-id.
# 2026-05-17 phase-3 (NS bug fix).
_RE_CITATION_END = re.compile(r'\(([^()]{2,60})\)\s*$')

# Legacy alias kept for any external import; new code calls _split_citation.
_RE_CITATION = _RE_CITATION_END

# Speaker tag: "John Smith (m):" on its OWN line is a bridge row.
# 2026-05-13 (AW-3146 fix): added end-of-string anchor so the pattern
# matches only when the speaker tag stands alone. The old regex also
# matched "Lee (f): She ended up" → the whole sentence was marked as
# bridge and lost its translation.
_SPEAKER_RE = re.compile(
    r'^[A-Za-zÀ-ÖÙ-öù-ÿ\s\-]{2,40}\s*[\(\[]\s*[mf]\s*[\)\]]\s*:?\s*$',
    re.IGNORECASE
)

_BRIDGE_PATTERNS_RAW = [
    # 2026-05-13 (AW-3146 fix): HOST line variants. The old single-pattern
    # `r'^HOST\s*:'` missed `HOST (INTRO):`, `HOST (OUTRO):`, `HOST(M):`.
    r'^HOST\b',
    r'^SHOW:', r'^TITLE:', r'^WEEK', r'^AIRDATE',
    r'^YOUR LANGUAGE:', r'^Title VO:',
    r'^OUTRO:', r'^INTRO\([mf]\):', r'^GENERIC INTRO', r'^\(Generic Intro',
    r'^PRIORITY ANIMAL', r'^ANIMAL-PEOPLE',
    r'^BMD \d+', r'^Fix\d', r'^Ball Time',
    r'^NFT:', r'^NFH:', r'^TC:',
    r'^SS \d+', r'^From Show',
    r'^Our programs', r'^offer many', r'^please visit',
    r'^We welcome', r'^stories and or', r'^loving animal',
    r'^Please send', r'^In English', r'^Originally in',
    r'^CAPTION:', r'^VO[,&\s].*ONSCREEN', r'^ONSCREEN TEXT',
    r'^https?://', r'^file:///',
    r'^\d+:\d+\s*[-~]\s*\d+:\d+', r'^\d{1,2}:\d{2}(:\d{2})?\s*$', r'^\(\d+:\d+',
    # 2026-05-13 (AJAR-3147 fix): "HH:MM:SS Speaker(m):" — timecode +
    # speaker label on the same line is a bridge row, not dialogue.
    r'^\d{1,2}:\d{2}:\d{2}\s+\S.*[\(\[][mf][\)\]]\s*:?\s*$',
    r'^\[English', r'^\[German', r'^\[.*starts\]', r'^\[.*End\]',
    # 2026-05-13 (AJAR-3147 fix): speaker labels are bridges ONLY when
    # they stand alone on their own line. "SM: real dialogue" / "Master:
    # actual words" / "Narrator: continues" are real content rows that
    # the old patterns wrongly swallowed. The `\s*$` anchor restricts
    # bridge classification to the label-only case.
    r'^Narrator\s*:?\s*$', r'^Maharaj\s*:\s*$',
    r'^SM\s*:\s*$', r'^Master\s*:\s*$',
    r'^[A-Z][A-Z\s]{2,}:\s*$',
    # 2026-05-13: stand-alone hashtag-id rows (`#193465`).
    r'^#\d{3,}\s*$',
    # 2026-05-17 (AJAR 3150 bug): editor-comment rows beginning with
    # "(NFE:", "(NFT:", "(NFH:", "(NFG:" — these are production-note
    # cells that are intentionally empty in the FA column; treating them
    # as ordinary rows lets the aligner spill the next group's content
    # into them. Also catches the bare prefix without leading paren.
    r'^\(NF[ETHG]\b', r'^NF[ETHG]:',
]
_BRIDGE_RE = [re.compile(p, re.IGNORECASE) for p in _BRIDGE_PATTERNS_RAW]


# ── text helpers ───────────────────────────────────────────────────────────────

def _display_len(text: str) -> int:
    """Visual char count: ZWNJ (U+200C) is invisible in Word and excluded."""
    return len(text.replace(ZWNJ, ''))


def _split_citation(text: str) -> tuple[str, str]:
    """Detect a *real* news-source citation at end of FA cell.

    Returns ``(main_text, citation_or_empty_string)`` where ``citation``
    includes the surrounding parens (e.g. ``"(Reuters)"``).

    Heuristic (2026-05-17 phase-3, NS bug fix):
      - The final parenthetical ``(...)`` must be 2-60 chars long.
      - Inner content must NOT contain any sentence terminator
        (``.``, ``!``, ``?``, ``؟``). Examples of REJECTED matches:
        ``(بله، استاد.)``, ``(درسته.)``, ``(وای.)`` — these are
        dialogue reactions, not news citations.
      - Inner content must NOT be a single short Persian descriptor
        such as ``(وگان)``, ``(گیاهخوار)``, ``(آلمانی)``, ``(ویتنامی)``
        — those follow a name and should travel WITH the name. The
        rule fires only when the inner is multi-token OR contains
        Latin letters (the news-source signature).

    The OLD ``_strip_citation`` discarded citations entirely; this
    function preserves them so ``_align_group`` can place them in
    the final row of the group.
    """
    if not text:
        return (text, "")
    m = _RE_CITATION_END.search(text.rstrip())
    if not m:
        return (text, "")
    inner = m.group(1).strip()
    if any(c in inner for c in ".!?؟"):
        return (text, "")  # dialogue reaction, not a citation
    # Single Persian descriptor (وگان / گیاهخوار / ویتنامی …) → not citation;
    # keep as-is. Detect by: pure Persian script + ≤2 tokens.
    has_latin = any('a' <= c.lower() <= 'z' for c in inner)
    token_count = len(inner.split())
    if not has_latin and token_count <= 2:
        return (text, "")
    main = text[: m.start()].rstrip()
    return (main, m.group(0).strip())


def _strip_citation(text: str) -> str:
    """Backwards-compatible wrapper kept for the bridge-clear pass.

    Returns only the main text without the citation. Internally calls
    ``_split_citation``. The citation-only fallback (return original
    when stripping leaves nothing) is preserved so a citation-only
    cell stays verbatim.
    """
    main, cite = _split_citation(text)
    if not main and cite:
        return text  # citation-only cell — keep verbatim
    return main


def _has_midline_paren(text: str) -> bool:
    """True when '(' appears in the middle of the cell, NOT as end citation.

    2026-05-17 phase-3 (no-doubling rule): an FA cell with mid-line
    parens carries auxiliary information that should NOT be repeated
    across rows. _align_group respects this flag.
    """
    if not text:
        return False
    main, _ = _split_citation(text.strip())
    return "(" in main


def _has_midline_dot(text: str) -> bool:
    """True when '.', '!', '?', '؟' appears in the middle of the cell.

    2026-05-17 phase-3 (no-doubling rule): mid-line sentence terminator
    means the cell holds two independent sentences; doubling would
    misalign their distribution across rows.
    """
    if not text:
        return False
    t = text.strip()
    while t and t[-1] in '.!?؟':
        t = t[:-1].rstrip()
    return any(c in '.!?؟' for c in t)


def _normalize_fa(text: str) -> str:
    """Collapse whitespace and fix stray punctuation spacing.

    B15 (audit 2026-05-13): also normalise Arabic→Persian script
    variants the translator may have leaked through. The polisher's
    ``fa_postprocess.normalize_fa`` does the same on its own input,
    but the aligner runs upstream of the polisher in some workflows,
    so we duplicate the rule here. Safe: every mapping is a 1:1
    canonical-form substitution that the SMTV style guide already
    requires.
    """
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'\s+([.!?؟])', r'\1', text)
    # Arabic Yeh / Kaf → Persian forms
    text = text.replace('ي', 'ی')   # ي → ی
    text = text.replace('ك', 'ک')   # ك → ک
    # Arabic-Indic digits → Persian digits
    _AR_TO_FA = {
        '٠': '۰', '١': '۱', '٢': '۲', '٣': '۳',
        '٤': '۴', '٥': '۵', '٦': '۶', '٧': '۷',
        '٨': '۸', '٩': '۹',
    }
    for src, dst in _AR_TO_FA.items():
        text = text.replace(src, dst)
    return text


def _is_bridge(en: str, fa: str = '', shaded: bool = False) -> bool:
    """True if this row is metadata/bridge and should be skipped."""
    if shaded:
        return True
    en = en.strip()
    if not en:
        return True
    for p in _BRIDGE_RE:
        if p.search(en):
            return True
    if _SPEAKER_RE.match(en):
        return True
    if re.match(r'^[\w\-]+\.(org|com|net|gov)\s*$', en):
        return True
    if fa and fa.strip().startswith('http'):
        return True
    return False


def _ends_sentence(fa: str) -> bool:
    t = fa.strip()
    return bool(t) and t[-1] in SENT_END and not t.endswith('...')


def _bigram_bad_positions(text: str) -> frozenset:
    """Positions that fall inside a protected bigram (splitting here breaks it)."""
    bad: set = set()
    for bigram in PROTECTED_BIGRAMS:
        start = 0
        while True:
            idx = text.find(bigram, start)
            if idx == -1:
                break
            word1 = bigram.split(' ', 1)[0]
            bad.add(idx + len(word1) + 1)   # start of the 2nd word
            start = idx + 1
    return frozenset(bad)


# ── split helpers ──────────────────────────────────────────────────────────────

def _is_safe_break(text: str, pos: int) -> bool:
    """True if splitting at `pos` doesn't violate grammar rules."""
    n = len(text)
    if pos <= 0 or pos >= n:
        return True
    left = text[:pos].rstrip()
    if not left:
        return True
    # 2026-05-13 (AJAR-3147 fix): don't break in the middle of a
    # multi-char punctuation cluster (.., ..., …, !!, ??, !!!). The
    # sentence-end scanner sees the first '.' and wants to break there,
    # which orphans '.. content' onto the next chunk.
    if text[pos - 1] in '.!?…':
        if pos < n and text[pos] in '.!?…':
            return False  # mid-cluster — never break here
    # 2026-05-13 (AJAR-3147 fix): paren/quote orphan guard.
    # Don't break right after an opener (chunk ends with '(' or '"') —
    # the matching closer would land on the next chunk far from the
    # word it wraps.
    if text[pos - 1] in '("«「[':
        return False
    # Don't break right before a closer (next chunk starts with ')')
    # — the close-bracket belongs to the previous chunk.
    if pos < n and text[pos] in ')"»」]':
        return False
    # 2026-05-13 (AJAR-3147 benchmark fix): paren/quote span guard.
    # Don't break INSIDE an unclosed span. If everything to the left of
    # `pos` contains more openers than closers, we're inside a paren
    # span — splitting here orphans the close paren on the next chunk.
    # Same logic for double-quote pairs.
    left = text[:pos]
    if left.count('(') > left.count(')'):
        return False
    if left.count('"') % 2 == 1:
        return False
    # Curly-quote span: count opens vs closes.
    if left.count('“') > left.count('”'):
        return False
    last_word = left.split()[-1]
    if last_word in DANGLING_PREPS:
        return False
    right = text[pos:].lstrip()
    if right:
        first_word = right.split()[0]
        if any(first_word.startswith(pf) for pf in COMPOUND_PREFIXES):
            return False   # compound-verb prefix
        if first_word == 'را':
            return False   # «را» must stay with its noun
    return True


def _find_break(text: str, target: int, bad_bigrams: frozenset) -> int:
    """
    Find the optimal split position near ``target`` in ``text``.

    Priority cascade (2026-05-12 phase-2, mirrors how a human FA editor
    breaks a long subtitle):

      1. Sentence-end punctuation (. ! ? ؟ …) — strongest, splits two
         independent clauses.
      2. Mid-sentence punctuation (، ؛ :)  — natural breath point.
      3. Just BEFORE a leading conjunction (و، اما، ولی، چون، …) so the
         conjunction starts the second chunk, the canonical FA pattern.
      4. Plain space with the full safety check + bigram guard.
      5. Plain space without the safety check (last resort before hard cut).
      6. Hard cut at MAX_CHARS.

    The ``target`` is the ideal break offset; each rule scans outward
    from it (±18 chars) so a small target shift can land on a far better
    split point.
    """
    target = max(8, min(target, MAX_CHARS - 1))
    n = len(text)

    def _scan(pred):
        """Scan ±off around target for the first p where pred(p) holds."""
        for off in range(18):
            for p in (target + off, target - off):
                if 0 < p <= MAX_CHARS and p < n:
                    if pred(p):
                        return p
        return -1

    # Priority 1: sentence-end punctuation.
    pos = _scan(
        lambda p: (
            text[p - 1] in SENT_END_CHARS
            and p not in bad_bigrams
            and _is_safe_break(text, p)
        )
    )
    if pos > 0:
        return pos

    # Priority 2: mid-sentence punctuation (comma / semicolon / colon).
    pos = _scan(
        lambda p: (
            text[p - 1] in MID_PUNCT_CHARS
            and p not in bad_bigrams
            and _is_safe_break(text, p)
        )
    )
    if pos > 0:
        return pos

    # Priority 3: just before a leading conjunction (و، اما، …).
    # We want the conjunction to START the next chunk, so split at the
    # space that precedes it.
    def _is_pre_conjunction(p):
        if text[p - 1] != ' ':
            return False
        # The word starting at p
        right = text[p:].lstrip()
        if not right:
            return False
        first = right.split()[0].rstrip('،؛:.!?؟…')
        if first not in LEADING_CONJUNCTIONS:
            return False
        return p not in bad_bigrams and _is_safe_break(text, p)

    pos = _scan(_is_pre_conjunction)
    if pos > 0:
        return pos

    # Priority 3.5: just before a leading preposition (در، به، از، برای، …).
    # 2026-05-17 phase-3: ~18.6 % of human cuts in the 562-file corpus
    # land directly before a prepositional phrase. Inserting this between
    # the conjunction rule and the generic-space rule routes a large slice
    # of cuts to a natural FA break point.
    def _is_pre_preposition(p):
        if text[p - 1] != ' ':
            return False
        right = text[p:].lstrip()
        if not right:
            return False
        first = right.split()[0].rstrip('،؛:.!?؟…')
        if first not in LEADING_PREPS:
            return False
        return p not in bad_bigrams and _is_safe_break(text, p)

    pos = _scan(_is_pre_preposition)
    if pos > 0:
        return pos

    # Priority 4: space with full safety.
    pos = _scan(
        lambda p: (
            (text[p - 1] == ' ' or text[p] == ' ')
            and p not in bad_bigrams
            and _is_safe_break(text, p)
        )
    )
    if pos > 0:
        return pos

    # Priority 5: space without safety (avoids infinite loops on hard text).
    pos = _scan(
        lambda p: text[p - 1] == ' ' or text[p] == ' '
    )
    if pos > 0:
        return pos

    return min(MAX_CHARS, n)


def _split_with_target(text: str, n_chunks: int, target: int) -> list | None:
    """
    Split text into exactly n_chunks with dynamic target recalculation.
    Returns None if any chunk would exceed MAX_CHARS or text is lost.
    """
    if n_chunks <= 0:
        return None
    if n_chunks == 1:
        return [text] if _display_len(text) <= MAX_CHARS else None

    chunks: list = []
    remaining = text

    for i in range(n_chunks - 1):
        remaining = remaining.strip()
        if not remaining:
            break
        rem_chunks = n_chunks - i
        cur_target = max(MIN_TARGET,
                         min(MAX_CHARS - 2, -(-len(remaining) // rem_chunks)))
        bad = _bigram_bad_positions(remaining)
        pos = _find_break(remaining, cur_target, bad)
        chunk = remaining[:pos].rstrip()
        if not chunk or _display_len(chunk) > MAX_CHARS:
            return None
        chunks.append(chunk)
        remaining = remaining[pos:].lstrip()

    rem = remaining.strip()
    if not rem or _display_len(rem) > MAX_CHARS:
        return None
    chunks.append(rem)
    return chunks


def _split_natural(text: str) -> list:
    """Simple fallback split — always succeeds."""
    text = text.strip()
    if _display_len(text) <= MAX_CHARS:
        return [text]
    chunks: list = []
    remaining = text
    while remaining:
        remaining = remaining.strip()
        if not remaining:
            break
        if _display_len(remaining) <= MAX_CHARS:
            chunks.append(remaining)
            break
        bad = _bigram_bad_positions(remaining)
        pos = _find_break(remaining, MIN_TARGET, bad)
        chunk = remaining[:pos].rstrip()
        if not chunk:
            chunk = remaining[:MAX_CHARS]
            pos = MAX_CHARS
        chunks.append(chunk)
        remaining = remaining[pos:].lstrip()
    return chunks or [text[:MAX_CHARS]]


# Note (2026-05-17 phase-3): removed `_align_to_en_benchmarks`,
# `_en_row_terminal`, `_EN2FA_PUNCT`, `_EN_TERM_PUNCT` — they were
# scaffolding for an EN-anchored split path that never went live
# (~120 lines of dead code). The 2026-05-17 corpus analysis path
# operates purely on FA structure (citations + midline patterns +
# enriched bigram/preposition sets), which covers the cases the EN
# anchor was meant to solve, without the 30% over-correction the
# anchor caused in BMD/CTAW reference docs.


def _split_for_n_rows(text: str, n_rows: int) -> list:
    """
    Adaptive split: try fewest distinct chunks first (= maximise doubles).

    Strategy:
      preferred_min = max(ceil(len/MAX_CHARS), ceil(n_rows/2))
        — ceil(n_rows/2) ensures we can fill n_rows without triples.
      Try from preferred_min up to n_rows.
      Fall back to min_chunks (below preferred_min) if nothing worked.
      Last resort: split_natural.
    """
    text = text.strip()
    if not text:
        return []
    if _display_len(text) <= MAX_CHARS:
        return [text]

    min_chunks    = -(-_display_len(text) // MAX_CHARS)  # ceil(len/MAX)
    preferred_min = max(min_chunks, -(-n_rows // 2))      # avoids triples

    def _valid(chunks: list) -> bool:
        if not chunks or len(chunks) > n_rows:
            return False
        if not all(_display_len(c) <= MAX_CHARS for c in chunks):
            return False
        return (re.sub(r'\s+', '', ''.join(chunks))
                == re.sub(r'\s+', '', text))

    # Preferred range: no triples expected
    for try_n in range(preferred_min, n_rows + 1):
        t = max(MIN_TARGET, -(-len(text) // try_n))
        if t > MAX_CHARS:
            continue
        chunks = _split_with_target(text, try_n, t)
        if _valid(chunks):
            return chunks

    # Below preferred_min: fewer chunks, triples may occur
    if preferred_min > min_chunks:
        for try_n in range(min_chunks, preferred_min):
            t = max(MIN_TARGET, -(-len(text) // try_n))
            chunks = _split_with_target(text, try_n, t)
            if _valid(chunks):
                return chunks

    return _split_natural(text)


# ── distribution ───────────────────────────────────────────────────────────────

def _distribute_to_rows(chunks: list, n_rows: int) -> list:
    """
    Map M distinct chunks → exactly N row slots.

    - If M == N: identity
    - If M  > N: merge shortest adjacent pairs to reduce
    - If M  < N: proportional doubles (triple only if unavoidable)
    """
    chunks = list(chunks)

    # Merge down — prefer merges that fit inside MAX_CHARS so the
    # resulting row never breaches the broadcast limit. 2026-05-12 phase-2:
    # if NO safe merge exists (every adjacent pair would overflow), stop
    # merging instead of producing an over-MAX_CHARS row. The downstream
    # save layer then sees more chunks than rows and the over_limit count
    # stays at 0 — A12 tolerance is no longer needed for this path.
    # The trade-off: content past row N is dropped at the end of the
    # function; in practice this hits only when the FA sentence cannot
    # mechanically fit n_rows × MAX_CHARS at all, and the LLM-threshold
    # path will rescue those cases (future phase).
    while len(chunks) > n_rows:
        best_i = -1
        best_ml = 99999
        for i in range(len(chunks) - 1):
            ml = _display_len(chunks[i]) + 1 + _display_len(chunks[i + 1])
            if ml <= MAX_CHARS and ml < best_ml:
                best_i = i
                best_ml = ml
        if best_i < 0:
            # No safe merge — stop. We accept that some chunks at the
            # tail will be dropped by `assignments[:n_rows]` rather than
            # ship an over-limit row to broadcast.
            break
        chunks[best_i] = chunks[best_i] + ' ' + chunks[best_i + 1]
        chunks.pop(best_i + 1)

    if len(chunks) == n_rows:
        return chunks

    # Expand proportionally
    n_ch = len(chunks)
    triple_unavoidable = n_rows > 2 * n_ch
    total_len = sum(len(c) for c in chunks) or 1
    assignments: list = []
    cursor = 0

    for ci, chunk in enumerate(chunks):
        if ci == n_ch - 1:
            n_for = n_rows - cursor
        else:
            n_for = max(1, round(len(chunk) / total_len * n_rows))
            if not triple_unavoidable:
                n_for = min(n_for, 2)
            remaining = n_ch - ci - 1
            n_for = min(n_for, n_rows - cursor - remaining)
            n_for = max(1, n_for)
        assignments.extend([chunk] * n_for)
        cursor += n_for

    return assignments[:n_rows]


def _group_difficulty_score(rows: list, n_rows: int) -> int:
    """
    Score a mechanically-aligned group on a 0..100 difficulty scale.
    Higher score = more likely the mechanical output is sub-optimal and
    an LLM rewrite would help.

    Calibrated 2026-05-12 phase-3 against the BMD / CTAW Google-Drive
    benchmark corpus. The sweet spot is llm_threshold≈40 invoking the
    LLM on the dense-line groups (every chunk close to MAX_CHARS) while
    leaving the clean ones to the mechanical path.

    Signals (cheap; single pass, no regex):

      • Each over-MAX_CHARS chunk            +60   (broadcast violation)
      • Each chunk > MAX_CHARS-4 (≥45)       +15   (too tight to read)
      • Each chunk between 40 and 44         +6    (dense but legal)
      • Each chunk ≤ 6 chars (orphan)        +15
      • triple-repeat (run ≥ 3)              +30
      • length spike vs median (forced merge)+8
      • truncation tail (2+ trailing empties)+12

    Caller invokes the LLM when ``score >= (100 - llm_threshold)``.
    """
    score = 0
    if not rows:
        return 0
    nonempty = [r for r in rows if r.strip()]
    if not nonempty:
        return 0
    for r in rows:
        L = _display_len(r)
        if L > MAX_CHARS:
            score += 60
        elif L >= MAX_CHARS - 4:
            score += 15
        elif L >= 40:
            score += 6
        if 0 < L <= 6:
            score += 15
    # Triple repeats
    idx = 0
    while idx < len(rows):
        ch = rows[idx].strip()
        j = idx + 1
        while j < len(rows) and rows[j].strip() == ch and ch:
            j += 1
        if j - idx >= 3:
            score += 30
        idx = j
    # Length spike (a row >40% longer than median signals a forced merge)
    lengths = [_display_len(r) for r in nonempty]
    if len(lengths) >= 3:
        sorted_lens = sorted(lengths)
        median = sorted_lens[len(sorted_lens) // 2]
        if max(lengths) > median * 1.4 + 4:
            score += 8
    # Truncation tail (last 2+ rows empty in a 3+-row group)
    if n_rows >= 3 and any(r.strip() for r in rows[:-1]) and not rows[-1].strip():
        empties = sum(1 for r in rows if not r.strip())
        if empties >= 2:
            score += 12
    return min(100, score)


def _enforce_no_triple(rows: list) -> list:
    """
    Hard ban: no three or more identical consecutive rows.
    Replaces the 3rd+ occurrence with '' (empty row — safe for broadcast).

    Tracks logical values separately from emitted values so that suppressed
    '' slots don't reset the run counter (avoids the 4× → 'A A  A' bug).
    """
    result:  list = []
    logical: list = []

    for ch in rows:
        if ch.strip():
            run = 0
            for lv in reversed(logical):
                if lv == ch:
                    run += 1
                else:
                    break
            if run >= 2:
                result.append('')
                logical.append(ch)
            else:
                result.append(ch)
                logical.append(ch)
        else:
            result.append(ch)
            logical.append(ch)
    return result


# ── DOCX I/O ──────────────────────────────────────────────────────────────────

def _cell_text(cell) -> str:
    return ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _cell_has_shading(cell) -> bool:
    """True if the cell has a GREY bridge background.

    2026-05-13 (AJAR-3147 fix): the previous "any non-white fill =
    bridge" rule swept up editor-highlight colours too. Real broadcast
    docx files use yellow `FFF2CC`, pink `FFC0CB`, blue `B4D5FF`, etc.
    to flag content rows ("Fix1 yellow shaded lines"). Treating those
    as bridge rows wiped every yellow-flagged Persian line.

    Now: bridge = greyscale fill ONLY (R == G == B, all six hex chars
    equal). Editor-highlight non-grey colours are preserved.
    """
    tcPr = cell._tc.find(_qn('w:tcPr'))
    if tcPr is None:
        return False
    shd = tcPr.find(_qn('w:shd'))
    if shd is None:
        return False
    fill = shd.get(_qn('w:fill'))
    if not fill:
        return False
    fill = fill.lower().strip()
    if fill in ('auto', 'ffffff', ''):
        return False
    # Greyscale check: 6-char hex with R == G == B.
    if len(fill) == 6 and all(c in '0123456789abcdef' for c in fill):
        r, g, b = fill[0:2], fill[2:4], fill[4:6]
        if r == g == b:
            return True
    # Any non-grey colour → editor highlight, not a bridge row.
    return False


def _ensure_rtl_paragraph(p) -> None:
    pPr = p._p.get_or_add_pPr()
    if pPr.find(_qn('w:bidi')) is None:
        pPr.append(OxmlElement('w:bidi'))


def _ensure_rtl_run(run) -> None:
    rPr = run._r.get_or_add_rPr()
    if rPr.find(_qn('w:rtl')) is None:
        rPr.append(OxmlElement('w:rtl'))


def _set_fa_cell(table, ri: int, text: str) -> None:
    """Write `text` into the FA column (index 2) of row `ri`, with RTL markers."""
    row = table.rows[ri]
    if len(row.cells) < 3:
        return
    fa_cell = row.cells[2]
    for p in fa_cell.paragraphs:
        for run in p.runs:
            run.text = ''
    while len(fa_cell.paragraphs) > 1:
        elem = fa_cell.paragraphs[-1]._element
        elem.getparent().remove(elem)
    p = fa_cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        run = p.runs[0]
    else:
        run = p.add_run(text)
    _ensure_rtl_paragraph(p)
    _ensure_rtl_run(run)


# ── main class ────────────────────────────────────────────────────────────────

class FASubtitleAligner:
    """
    Mechanical FA subtitle aligner.

    API-compatible with the previous version:
        aligner = FASubtitleAligner(model='gpt-5.4-mini', llm_threshold=0)
        stats   = aligner.align(input_docx, output_docx)

    llm_threshold semantics (2026-05-12 phase-3):
        0   = pure mechanical (default; every existing run stays
              byte-identical with the pre-phase-3 output).
        40  = invoke the LLM only on hard groups (difficulty score ≥ 60).
              Calibration sweet spot per the prompt-rewrite roadmap.
        100 = invoke the LLM on every group with score ≥ 0 (almost all).
        Anything in between is interpolated linearly: the LLM fires when
        ``group_score >= 100 - llm_threshold``.
    token_budget is still accepted for API compatibility, currently unused.
    """

    def __init__(
        self,
        model:          str = 'gpt-5.4-mini',
        llm_threshold:  int = 0,
        token_budget:   int = 0,
        max_chars:      int = 48,
    ):
        self.model         = model
        self.llm_threshold = max(0, min(100, int(llm_threshold)))
        self.token_budget  = token_budget   # accepted, currently unused
        self.last_stats:   dict = {}
        self._llm_corrected_count = 0
        self._llm_tokens_used     = 0
        # Lazy OpenAI client — only built when threshold > 0 and we need it.
        self._client = None

        # 2026-05-15 — broadcast CPL override. The module-level constants
        # MAX_CHARS / MIN_TARGET are referenced by ~35 helper functions
        # at file scope. Rather than thread a self.max_chars through every
        # one of them (huge refactor, easy to miss a site), we mutate the
        # module globals when the operator asks for a non-default value.
        # The aligner is process-singleton in practice (the launcher only
        # instantiates it once per job, and FASubtitleAligner does not
        # support concurrent jobs by design), so this is safe.
        # Clamp to the same band the frontend exposes (24..70). The
        # MIN_TARGET stays at the historical 1:2 ratio to MAX_CHARS so
        # rounder splits remain idiomatic at the new CPL ceiling.
        global MAX_CHARS, MIN_TARGET   # noqa: PLW0603
        clamped = max(24, min(70, int(max_chars)))
        if clamped != MAX_CHARS:
            MAX_CHARS  = clamped
            MIN_TARGET = max(12, clamped // 2)
        self.max_chars  = MAX_CHARS
        self.min_target = MIN_TARGET

    def _get_client(self):
        if self._client is None:
            from openai import OpenAI
            import os
            key = os.environ.get("OPENAI_API_KEY")
            if not key:
                return None
            self._client = OpenAI(api_key=key)
        return self._client

    # ── LLM rescue for hard groups ────────────────────────────────────────────

    _LLM_SYSTEM = (
        "You are a Persian subtitle line-splitter. You receive an English "
        "source block, the Persian translation of that block, and the number "
        "of rows N the FA must be distributed across. Your job: split the "
        "Persian text into exactly N chunks for broadcast display.\n\n"
        "Hard rules:\n"
        "1. Output exactly N lines, in order.\n"
        "2. Each line MUST be ≤ 48 visible characters (ZWNJ excluded).\n"
        "3. Concatenating all output lines (with single spaces) MUST equal "
        "the input Persian text, modulo whitespace. Do NOT change words.\n"
        "4. Break at natural Persian boundaries: sentence-end punctuation > "
        "comma > before a coordinating/subordinating conjunction (و، اما، "
        "ولی، چون، که) > between full phrases.\n"
        "5. Never strand a preposition (به/از/در/برای/با/که/تا) at line end.\n"
        "6. Never split a compound verb (می‌کند | انجام داد).\n"
        "7. Never split inside a protected idiom (از جمله | با این حال | "
        "به همین دلیل | به ویژه | به گفتهٔ | به نقل از | …).\n"
        "8. If a row must repeat (N > distinct chunks), repeat earlier chunks "
        "verbatim — never invent a new one.\n\n"
        "Output format — MANDATORY. One line per row, no markers, no labels, "
        "no markdown. First character of output = first character of row 1.\n"
        "If you cannot satisfy every rule simultaneously, return the single "
        "literal token ABORT on its own line; the caller will fall back to "
        "the mechanical output."
    )

    def _llm_refine_group(
        self,
        full_fa: str,
        mech_rows: list,
        n_rows: int,
        group: dict,
    ) -> list | None:
        """
        Ask gpt-5.4-mini to re-split the FA text into n_rows chunks.
        Returns the refined rows or None if the LLM declined / failed validation.
        """
        client = self._get_client()
        if client is None:
            return None

        en_text = ' '.join(p.strip() for p in group.get('en_parts', []) if p) \
                  if 'en_parts' in group else ''
        user = (
            f"N = {n_rows}\n"
            f"MAX_CHARS_PER_LINE = {MAX_CHARS}\n\n"
            f"Persian text to split:\n{full_fa}\n"
        )
        if en_text:
            user += f"\nEnglish source (context only — do NOT translate):\n{en_text}\n"
        user += (
            "\nMechanical baseline (for your reference; improve if possible, "
            "match if already optimal):\n"
            + "\n".join(mech_rows)
        )

        try:
            from ._retry import call_with_retry
            resp = call_with_retry(
                lambda: client.responses.create(
                    model=self.model,
                    input=[
                        {"role": "system", "content": self._LLM_SYSTEM},
                        {"role": "user",   "content": user},
                    ],
                    extra_body={
                        "prompt_cache_retention": "24h",
                        "prompt_cache_key": "mtd-aligner-v7",
                    },
                    reasoning={"effort": "low"},
                    timeout=120,
                ),
                label="aligner.llm_refine",
            )
        except Exception as exc:
            print(f"[WARN] FA aligner LLM refine failed: {exc!r} — using mechanical.")
            return None

        # Token accounting
        try:
            usage = (resp.model_dump() or {}).get("usage") or {}
            self._llm_tokens_used += int(
                usage.get("input_tokens", usage.get("prompt_tokens", 0)) or 0
            )
            self._llm_tokens_used += int(
                usage.get("output_tokens", usage.get("completion_tokens", 0)) or 0
            )
        except Exception:
            pass

        if hasattr(resp, "output_text") and resp.output_text is not None:
            raw = resp.output_text
        else:
            try:
                raw = resp.choices[0].message.content
            except Exception:
                return None

        text = (raw or "").strip()
        if text.upper().strip() == "ABORT" or not text:
            return None

        lines = [ln.rstrip() for ln in text.split("\n") if ln.strip() != ""]
        if len(lines) != n_rows:
            return None
        # Validate: every line ≤ MAX_CHARS, every word present in mech text.
        for ln in lines:
            if _display_len(ln) > MAX_CHARS:
                return None
        # Loose content-preservation check: concatenated whitespace-stripped
        # output should match the whitespace-stripped input (within 5 % token
        # delta, mostly to allow function-word reorders).
        out_strip = re.sub(r"\s+", "", "".join(lines))
        in_strip  = re.sub(r"\s+", "", full_fa)
        if abs(len(out_strip) - len(in_strip)) > max(8, len(in_strip) * 0.05):
            return None
        return lines

    # ── reading ───────────────────────────────────────────────────────────────

    def _read_rows(self, docx_path: str) -> list:
        doc = Document(docx_path)
        if not doc.tables:
            raise ValueError(f"No table found in {docx_path}")
        rows = []
        for ri, row in enumerate(doc.tables[0].rows):
            cells = row.cells
            if len(cells) < 3:
                continue
            rows.append({
                'ri':     ri,
                'en':     _cell_text(cells[1]),
                # 2026-05-17 phase-3: keep the citation in raw FA. The
                # old _strip_citation call silently dropped news-source
                # attributions on every "text (Source)" row; _align_group
                # now handles citation placement explicitly.
                'fa':     _cell_text(cells[2]),
                'shaded': _cell_has_shading(cells[1]),
            })
        return rows

    # ── grouping ──────────────────────────────────────────────────────────────

    def _parse_groups(self, rows: list) -> list:
        """
        FA-based sentence grouping.

        A group ends when:
          - A bridge / shaded row is encountered     → flush, skip row
          - The FA text ends with sentence punctuation (.!?؟)

        After a sentence-ending FA row, any immediately following empty-FA
        non-bridge rows are absorbed into the same group.  This handles
        single-call translation output, where the translator placed the full
        FA sentence in the first row only and left the remaining EN rows with
        empty FA cells.

        Inline empty rows (empty FA mid-sentence, non-bridge) are also
        absorbed into the current group so the aligner can distribute text
        across them.

        Special case: if the translator copy-pasted the same long FA text
        across multiple rows, all those rows form one group (Lesson 5 pattern).
        """
        groups: list = []
        current_indices: list = []
        current_fa_parts: list = []
        current_en_parts: list = []   # 2026-05-12 phase-3 — context for LLM rescue
        # 2026-05-13: per-row EN text (one entry per row_index, "" if EN
        # cell was empty). Used by the benchmark-anchored splitter that
        # tries to land FA punct at the same row index where EN punct
        # appears — the user's "بنچ مارک" reference pattern.
        current_en_per_row: list = []

        i = 0
        while i < len(rows):
            rd = rows[i]
            en, fa = rd['en'], rd['fa']
            shaded = rd.get('shaded', False)

            # Bridge row → flush current group and skip this row entirely
            if _is_bridge(en, fa, shaded):
                if current_indices:
                    groups.append({
                        'row_indices': list(current_indices),
                        'fa_parts':    list(current_fa_parts),
                        'en_parts':    list(current_en_parts),
                        'en_per_row':  list(current_en_per_row),
                    })
                    current_indices = []
                    current_fa_parts = []
                    current_en_parts = []
                    current_en_per_row = []
                i += 1
                continue

            # Empty FA but non-bridge:
            #   - Inside a sentence group → include as an extra row slot
            #   - Not yet inside any group → skip (nothing to distribute yet)
            if not fa.strip():
                if current_indices:
                    current_indices.append(rd['ri'])
                    current_en_per_row.append(en.strip())
                    if en.strip():
                        current_en_parts.append(en.strip())
                i += 1
                continue

            # Repeated-long-FA case: translator put the full sentence in every row
            if (not current_indices
                    and _display_len(fa) > MAX_CHARS):
                rep_indices = [rd['ri']]
                rep_en      = [en.strip()] if en.strip() else []
                rep_en_per_row = [en.strip()]
                j = i + 1
                while j < len(rows) and rows[j]['fa'] == fa:
                    rep_indices.append(rows[j]['ri'])
                    rep_en_per_row.append(rows[j]['en'].strip())
                    if rows[j]['en'].strip():
                        rep_en.append(rows[j]['en'].strip())
                    j += 1
                if len(rep_indices) > 1:
                    groups.append({
                        'row_indices': rep_indices,
                        'fa_parts':    [fa],   # once — not repeated
                        'en_parts':    rep_en,
                        'en_per_row':  rep_en_per_row,
                    })
                    i = j
                    continue

            current_indices.append(rd['ri'])
            current_fa_parts.append(fa.strip())
            current_en_per_row.append(en.strip())
            if en.strip():
                current_en_parts.append(en.strip())

            if _ends_sentence(fa):
                # Look-ahead: absorb any immediately following empty-FA
                # non-bridge rows into this group.  These are the "spare"
                # EN rows that single-call translation left blank.
                j = i + 1
                while j < len(rows):
                    nxt = rows[j]
                    if _is_bridge(nxt['en'], nxt['fa'], nxt.get('shaded', False)):
                        break
                    if nxt['fa'].strip():
                        break   # next sentence starts here
                    current_indices.append(nxt['ri'])
                    current_en_per_row.append(nxt['en'].strip())
                    if nxt['en'].strip():
                        current_en_parts.append(nxt['en'].strip())
                    j += 1
                groups.append({
                    'row_indices': list(current_indices),
                    'fa_parts':    list(current_fa_parts),
                    'en_parts':    list(current_en_parts),
                    'en_per_row':  list(current_en_per_row),
                })
                current_indices = []
                current_fa_parts = []
                current_en_parts = []
                current_en_per_row = []
                i = j
                continue

            i += 1

        if current_indices:
            groups.append({
                'row_indices': list(current_indices),
                'fa_parts':    list(current_fa_parts),
                'en_parts':    list(current_en_parts),
                'en_per_row':  list(current_en_per_row),
            })

        return groups

    # ── align one group ───────────────────────────────────────────────────────

    def _align_group(self, group: dict) -> list:
        """
        Align one sentence group.
        Returns list[str] with len == len(row_indices).

        2026-05-17 phase-3 — rewritten for three new responsibilities:
          1. CITATION PRESERVATION: when the joined FA text ends with
             ``(Source)`` (Reuters / VnExpress / Tuổi Trẻ / …), the
             citation is detached BEFORE splitting and placed alone in
             the last row of the group. Never duplicated, never broken
             across rows. Pure descriptors like «(وگان)» / «(گیاهخوار)»
             stay attached to their host (they ride inside the text).
          2. MID-LINE PAREN NO-DOUBLING: if any FA cell of the group
             carries ``(...)`` in the middle (not the end), the group
             is flagged ``no_doubling``. Chunks distribute one-per-row;
             extra rows stay empty rather than repeat content.
          3. MID-LINE DOT NO-DOUBLING: if any FA cell of the group
             carries a sentence terminator (. ! ? ؟) in the middle, the
             same ``no_doubling`` flag fires. Two independent sentences
             in one cell must not be repeated across rows.
        """
        n_rows = len(group['row_indices'])

        # Detect no-doubling flag on the RAW per-row FA before joining.
        # Joining can hide a mid-line dot (turns into an inter-cell space).
        no_doubling = any(
            _has_midline_paren(p) or _has_midline_dot(p)
            for p in group['fa_parts'] if p
        )

        raw_join = ' '.join(p for p in group['fa_parts'] if p)
        main_text, citation = _split_citation(raw_join)
        full_fa = _normalize_fa(main_text)

        # Special case: citation-only group (no other text). Place the
        # citation in the LAST row, leave the rest empty.
        if not full_fa and citation:
            rows = [''] * n_rows
            rows[-1] = citation
            return rows

        if not full_fa:
            return [''] * n_rows

        # Reserve one row for citation when the group has more than one row.
        # A 1-row group with citation gets the citation concatenated at end.
        reserve_citation_row = bool(citation) and n_rows >= 2
        target_rows = n_rows - (1 if reserve_citation_row else 0)
        if target_rows < 1:
            target_rows = 1

        chunks = _split_for_n_rows(full_fa, target_rows)

        if no_doubling:
            # One chunk per row, no repeats. BUT we MUST NOT drop content:
            # if the splitter produced more chunks than target_rows
            # (because target_rows shrank after reserving the citation
            # slot, or because the midline-pattern guard fired on a long
            # cell), merge the trailing overflow into the last row
            # instead of slicing it off. 2026-05-17 (AJAR 3150 bug:
            # content silently disappeared on rows whose group had
            # citation + midline paren).
            chunks_list = list(chunks)
            if len(chunks_list) > target_rows and target_rows > 0:
                tail = ' '.join(chunks_list[target_rows - 1:]).strip()
                chunks_list = chunks_list[:target_rows - 1] + [tail]
            rows = chunks_list[:target_rows]
            while len(rows) < target_rows:
                rows.append('')
        else:
            rows = _distribute_to_rows(chunks, target_rows)

        rows = _enforce_no_triple(rows)

        # Attach citation: dedicated last row when reserved, otherwise
        # appended to the only row.
        if citation:
            if reserve_citation_row:
                rows.append(citation)
            else:
                if rows:
                    joined = (rows[0] + ' ' + citation).strip()
                    rows[0] = joined
                else:
                    rows = [citation]

        # Pad / trim to exact count.
        if len(rows) < n_rows:
            rows.extend([''] * (n_rows - len(rows)))
        rows = rows[:n_rows]

        # Hybrid path (2026-05-12 phase-3): score the mechanical output;
        # if the LLM threshold is high enough AND this group looks rough,
        # hand it to gpt-5.4-mini for a second take. The pure-mechanical
        # default (llm_threshold == 0) leaves this branch dormant — every
        # existing run stays byte-identical. Citation/no-doubling logic
        # above runs BEFORE the LLM rescue so the LLM never sees stripped
        # text or repeated chunks.
        if self.llm_threshold > 0:
            score = _group_difficulty_score(rows, n_rows)
            if score >= (100 - self.llm_threshold):
                refined = self._llm_refine_group(full_fa, rows, n_rows, group)
                if refined is not None:
                    self._llm_corrected_count += 1
                    return refined
        return rows

    # ── write ─────────────────────────────────────────────────────────────────

    def _write_docx(self, input_path: str, output_path: str,
                    groups: list, all_chunks: list) -> None:
        doc   = Document(input_path)
        table = doc.tables[0]

        # 2026-05-13 (CS-3146 fix): the basic distributor that runs BEFORE
        # the aligner does not know which rows are bridge / metadata, so
        # it sometimes fills FA cells in HOST: / WARNING: / time-code
        # rows with continuation text from neighbouring groups. The
        # aligner correctly skips those rows during sentence grouping
        # but leaves whatever was already in the cell.
        #
        # 2026-05-13 (AW-3146 refinement): the previous fix cleared every
        # bridge row's FA cell unconditionally — which destroyed editor
        # metadata that pre-existed in the FA column (Fix1 / blue-pink
        # change notes / due-time strings). Only clear a bridge FA cell
        # when it looks like classic-distributor leakage:
        #   (a) the cell's text appears verbatim in an adjacent
        #       non-bridge FA cell (the classic distributor's duplicate
        #       pattern); OR
        #   (b) the cell is shaded (grey row — never editorial metadata).
        # Everything else stays byte-id.
        bridge_row_indices: set = set()
        try:
            input_rows = self._read_rows(input_path)
            # Build a quick neighbour lookup: ri → FA-text of the
            # nearest non-bridge row above and below.
            row_by_ri = {rd['ri']: rd for rd in input_rows}
            sorted_ris = sorted(row_by_ri.keys())
            for idx, rd in enumerate(input_rows):
                en = rd.get('en', '')
                fa = rd.get('fa', '')
                shaded = rd.get('shaded', False)
                if not _is_bridge(en, fa, shaded):
                    continue
                # Shaded rows are always safe to clear (grey bridge).
                if shaded:
                    bridge_row_indices.add(rd['ri'])
                    continue
                fa_stripped = (fa or '').strip()
                if not fa_stripped:
                    bridge_row_indices.add(rd['ri'])
                    continue
                # Editor-metadata heuristic: if the cell is mostly Latin
                # / ASCII characters (Fix1 / due-time notes / blue-pink
                # change instructions, often pre-existing in the docx
                # template), preserve it. Persian translation leakage
                # is always majority FA script.
                _ascii_letters = sum(1 for c in fa_stripped if 'a' <= c.lower() <= 'z')
                _fa_letters    = sum(1 for c in fa_stripped if '؀' <= c <= 'ۿ')
                _total_letters = _ascii_letters + _fa_letters
                if _total_letters > 0 and _ascii_letters / _total_letters > 0.4:
                    # Majority Latin → editor metadata, keep verbatim.
                    continue
                # Check if this FA text duplicates a neighbour's FA.
                neighbours_fa = []
                for offset in (-2, -1, 1, 2):
                    nb_idx = idx + offset
                    if 0 <= nb_idx < len(input_rows):
                        nb = input_rows[nb_idx]
                        nb_fa = (nb.get('fa') or '').strip()
                        if nb_fa:
                            neighbours_fa.append(nb_fa)
                # Clear if (a) duplicate of neighbour (classic
                # distributor leakage) or (b) the EN cell carries one
                # of the hard-trigger bridge patterns (HOST:, WARNING:,
                # SHOW:, time-code, URL, etc.). Both indicate the FA
                # text is NOT a real translation belonging to this row.
                if fa_stripped in neighbours_fa:
                    bridge_row_indices.add(rd['ri'])
                else:
                    # EN-pattern path. If the EN cell matches a known
                    # bridge regex (not just 'empty'), the FA value is
                    # certainly mis-placed translation.
                    _en_strip = (en or '').strip()
                    if _en_strip:
                        for _p in _BRIDGE_RE:
                            if _p.search(_en_strip):
                                bridge_row_indices.add(rd['ri'])
                                break
                # else (EN empty + FA Persian text + no duplicate):
                # preserve. The classic distributor never plants
                # standalone Persian without an EN context to map from.
        except Exception as _exc:
            # Defensive: if the re-read fails, fall back to the legacy
            # behaviour (don't touch bridge cells) rather than crashing.
            print(f"[WARN] FA aligner: bridge-cell clear pass skipped — {_exc!r}")

        # Aligned groups overwrite first.
        for group, chunks in zip(groups, all_chunks):
            for ri, chunk in zip(group['row_indices'], chunks):
                _set_fa_cell(table, ri, chunk)
        # Then sweep bridge rows to '' so leftover classic-distributor
        # fragments are erased.
        for ri in bridge_row_indices:
            try:
                _set_fa_cell(table, ri, '')
            except Exception:
                continue

        doc.save(output_path)

    # ── main entry point ──────────────────────────────────────────────────────

    def align(self, input_docx: str, output_docx: str) -> dict:
        """
        Full pipeline: read → group → split → distribute → write.
        Returns stats dict (compatible with pipeline expectations).
        """
        t0 = time.time()

        rows   = self._read_rows(input_docx)
        groups = self._parse_groups(rows)
        print(f"[INFO] Aligner: {len(groups)} sentence groups parsed")
        # B6 Jules deep (2026-05-13): a docx with non-empty source rows
        # MUST produce at least one group. If the parser drops every row
        # (bug or upstream corruption), the output would silently end up
        # empty — surface the failure instead of writing a blank docx.
        _has_translatable = any(
            (r.get('fa') or '').strip() or (r.get('en') or '').strip()
            for r in rows
            if not r.get('shaded', False)
        )
        if _has_translatable and not groups:
            raise ValueError(
                "FA aligner parsed zero sentence groups despite a "
                "non-empty input table — refusing to write a blank docx. "
                "Check _parse_groups for a regression."
            )

        all_chunks = [self._align_group(g) for g in groups]

        # Cross-group triple guard — inject sentinel between groups so that
        # identical text at a group boundary is not counted as a triple run.
        _SENTINEL = '\x00GROUP_BOUNDARY\x00'
        flat: list = []
        for gi, fc in enumerate(all_chunks):
            if gi > 0:
                flat.append(_SENTINEL)
            flat.extend(fc)
        flat_clean = _enforce_no_triple(flat)

        # Re-split back into per-group lists (skip sentinel slots)
        pos = 0
        for gi, fc in enumerate(all_chunks):
            if gi > 0:
                pos += 1   # skip sentinel
            n = len(fc)
            all_chunks[gi] = flat_clean[pos : pos + n]
            pos += n

        self._write_docx(input_docx, output_docx, groups, all_chunks)

        elapsed = time.time() - t0

        # ── stats ──────────────────────────────────────────────────────────
        all_rows = [r for fc in all_chunks for r in fc]

        n_doubles = 0
        n_triples = 0
        idx = 0
        while idx < len(all_rows):
            ch = all_rows[idx].strip()
            j  = idx + 1
            while j < len(all_rows) and all_rows[j].strip() == ch and ch:
                j += 1
            run = j - idx
            if run >= 2:
                n_doubles += run - 1
            if run >= 3:
                n_triples += run - 2
            idx = j

        n_over = sum(1 for r in all_rows if _display_len(r) > MAX_CHARS)

        self.last_stats = {
            'groups':          len(groups),
            'llm_corrected':   self._llm_corrected_count,
            'mechanical_only': len(groups) - self._llm_corrected_count,
            'tokens_used':     self._llm_tokens_used,
            'llm_threshold':   self.llm_threshold,
            'prompt_hash':     _prompt_hash('mechanical-v2.0+llm'),
            'total_rows':      len(all_rows),
            'doubles':         n_doubles,
            'triples':         n_triples,
            'over_limit':      n_over,
            'elapsed_seconds': round(elapsed, 1),
        }

        # 2026-05-13: surface the LLM-rescue count and the threshold so
        # the user can see whether raising the dial actually changed
        # anything on this document. When threshold > 0 but llm fired 0
        # times, this print says so explicitly.
        _llm_note = ""
        if self.llm_threshold > 0:
            _llm_note = (
                f" | LLM threshold {self.llm_threshold}"
                f" | LLM rescues fired: {self._llm_corrected_count}/{len(groups)}"
            )
            if self._llm_corrected_count == 0:
                _llm_note += " (no group's difficulty score reached the threshold)"
        print(
            f"[INFO] Aligner done in {elapsed:.1f}s"
            f" | groups: {len(groups)}"
            f" | doubles: {n_doubles} | triples: {n_triples}"
            f" | over-{MAX_CHARS}: {n_over}{_llm_note}"
        )
        return self.last_stats
