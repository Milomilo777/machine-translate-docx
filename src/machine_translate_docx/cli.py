#!/usr/bin/python3
# - *- coding: utf- 8 - *-
PROGRAM_VERSION="2026-03-29"

import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pkg_resources")
# Suppress deprecation warning from pyobjc using pkg_resources
warnings.filterwarnings(
    "ignore",
    category=UserWarning,
    module="objc._bridgesupport"
)

import sys
import io

# If all these flags appear anywhere on the command line, exit quietly.
UNWANTED_FLAGS = {"-B", "-S", "-E", "-s", "-c"}

if UNWANTED_FLAGS.issubset(set(sys.argv[1:])):
    # Optional extra safety: ensure there's something after -c (the inline code)
    try:
        c_index = sys.argv.index("-c")
        if c_index + 1 < len(sys.argv):
            # There *is* an argument after -c; likely the resource_tracker helper
            sys.exit(0)
        # If -c is the last token (rare/invalid), still exit if you want:
        # sys.exit(0)
    except ValueError:
        # -c not found (shouldn't happen when issubset passed), but just in case:
        sys.exit(0)
    
# For bidirectional text display right to left and left to right
from bidi.algorithm import get_display

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
sys.stdout.reconfigure(encoding='utf-8')
import platform

json_configuration_url='https://raw.githubusercontent.com/translation-robot/machine-translate-docx/main/src/configuration/configuration.json'
# Day 0 is October 3rd 2017


print("*********************************************************")
print("*  machine-translate-docx program version : %s" % (PROGRAM_VERSION))
print("*********************************************************")

print("Python programming language %s\n" % (platform.python_version()))

from pprint import pprint                          # 2026-05-18 cleanup: drop bare `import gc`, `import pprint`, `import shlex`, `import codecs` — unused at module scope.
import traceback
import subprocess
import os
import re
import time
import urllib
import urllib.request
import requests
import json
import json5 # json 5 with the ability to have comments


from inspect import currentframe, getframeinfo
import chardet
import getpass
import datetime

import zipfile
import xml.dom.minidom
# used to get elements in XML, shading in docx for example
from lxml import etree

# This library automatically downloads chrome driver
# pyderman was replaced with webdriver_manager
# then selenium 4.11.2 managed downloading the drivers
# For selenium 3

# When translation engine is deepl or chatgpt : use undetected_chromedriver
# Else, use standard selenium webdriver
# drivers are loaded after getting the engine value

from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.remote.remote_connection import LOGGER
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import StaleElementReferenceException, TimeoutException

from screeninfo import get_monitors


from time import sleep
import argparse
import clipboard

import psutil


import docx
from docx import Document
from docx import oxml
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT,WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

# For japanese

# For Thai

import timeit
import datetime
import progressbar


from timeit import default_timer as timer

import re
import inspect

from .xlsx_translation_memory import xlsx_translation_memory

# Module-level translation-memory handle. Pre-refactor code assumed this
# existed as a module-level global, but no top-level binding was ever
# committed — `initialize_translation_memory_xlsx` only created a local
# `xtm`, so every later read raised NameError once the code path ran
# live (silent in tests because the F-012 mid-layer wasn't exercised).
# Initialise as None here so `if xtm is not None` reads correctly when
# `--xlsxreplacefile` is not provided.
xtm = None

import html

from urllib.parse import urlencode, quote_plus

from openpyxl import load_workbook
from openpyxl import Workbook

from bs4 import BeautifulSoup

# pip install pycryptodome
# used for passwords (deepl, etc)

# Load configuration from a json file on internet (github for example)

import json

from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.text.run import Run as _DocxRun

import glob
from langcodes import *
import math
import shutil
import signal
import atexit

import random

from .config import (
    DefaultJsonConfiguration,
    validate_json_string,
    get_nested_value_from_json_array,
    google_translate_lang_codes,
    deepl_translate_lang_codes,
    office_language_tags,
    right_to_left_languages_list,
    eol_array,
    eol_conditional_array,
    bol_array,
    MAX_LINE_SIZE,
    COUNTRY_QUERY_HTTP_TIMEOUT,
)

from .runtime import RuntimeContext
from .selenium_utils import (
    safe_click,
    browser_fill_form_field_value,
    set_chrome_window_2_3_screen,
    create_webdriver,
    minimize_browser,
    clean_up_previous_chrome_selenium_drivers,
    cleanup_selenium_chrome_temp_folders,
)
from .engines.google import (
    selenium_chrome_google_translate,
    selenium_chrome_google_click_cookies_consent_button,
)
from .engines.deepl import (
    selenium_chrome_deepl_log_in,
    selenium_chrome_deepl_log_off,
    deepl_close_messages,
    selenium_chrome_deepl_translate,
)
# Google file-mode workers extracted to engines/google_file_modes.py in
# Sprint D-B (2026-05-16). cli.translate_docx calls the three top-level
# dispatchers; the seven internal helpers are now private to that module.
from .engines import (
    google_translate_from_text_file,
    google_translate_from_html_javascript,
    google_translate_from_html_xlsxfile,
)
from .runner import selenium_chrome_translate_maxchar_blocks as _runner_translate_maxchar_blocks

# Module-level RuntimeContext singleton — Phase F1 transition shim.
# Lazily built by `_get_ctx()` on first call from any function that has
# already been threaded with a `ctx` parameter. Phase F1.6 collapses this
# helper into an explicit `ctx = RuntimeContext.empty()` at the top of
# main(); until then, threaded functions can be called from un-threaded
# callers without a signature cascade.
_ctx: RuntimeContext | None = None


def _get_ctx() -> RuntimeContext:
    """Return the singleton RuntimeContext, snapshotting current module-level
    globals on first call. Threaded callers should propagate their own
    parameter where possible — this helper exists only to bridge un-threaded
    callers to threaded callees during the F1 transition."""
    global _ctx
    if _ctx is None:
        _ctx = RuntimeContext.empty()
        # Snapshot whatever module-level globals are already populated by
        # import time. The dataclass defaults cover anything that is not.
        # Each F1.x batch may extend this snapshot as it threads more
        # globals.
        try:
            _ctx.browser.driver = driver
        except NameError:
            pass
        try:
            _ctx.language.src_lang_name  = src_lang_name
            _ctx.language.dest_lang_name = dest_lang_name
        except NameError:
            pass
        # F1.2 — configuration snapshot
        try:
            _ctx.config.json_configuration_array = json_configuration_array
        except NameError:
            pass
        try:
            _ctx.config.max_translation_block_size = MAX_TRANSLATION_BLOCK_SIZE
        except NameError:
            pass
        # F1.4 — engine + flags snapshot
        try:
            _ctx.engine.engine = translation_engine
        except NameError:
            pass
        try:
            _ctx.engine.method = engine_method
        except NameError:
            pass
        try:
            _ctx.flags.splitonly = splitonly
        except NameError:
            pass
        # 2026-05-17 (AJAR 3150 basic-split bug): `split_translation` flag
        # had no snapshot here. The module-level global is set to True in
        # splitonly mode (line ~799) but ``ctx.flags.split_translation``
        # stayed at the RuntimeContext default (False). Downstream
        # ``print_console_docx_file_translated`` then took the
        # ``if not ctx.flags.split_translation:`` branch, which writes the
        # phrase's full FA blob into the FIRST row of the phrase only and
        # leaves the rest empty. Adding the snapshot fixes the
        # basic-split distribution.
        try:
            _ctx.flags.split_translation = split_translation
        except NameError:
            pass
        try:
            _ctx.docx.translation_errors_count = translation_errors_count
        except NameError:
            pass
        try:
            _ctx.browser.deepl_sleep_wait_translation_seconds = deepl_sleep_wait_translation_seconds
        except NameError:
            pass
        # F1.5 — DeepL session flags + OpenAI handles
        try:
            _ctx.browser.closed_cookies_accept_message_bool = closed_cookies_accept_message_bool
        except NameError:
            pass
        try:
            _ctx.browser.close_install_extension_message_bool = close_install_extension_message_bool
        except NameError:
            pass
        try:
            _ctx.browser.deepl_nb_clear_cached_times = deepl_nb_clear_cached_times
        except NameError:
            pass
        try:
            _ctx.browser.logged_into_deepl = logged_into_deepl
        except NameError:
            pass
        try:
            _ctx.openai.translator = oai_translator
        except NameError:
            pass
        try:
            _ctx.openai.polisher = oai_polisher
        except NameError:
            pass
        try:
            _ctx.openai.translation_log = translation_log
        except NameError:
            pass
        # F1.5 — language + paths + translation_array
        try:
            _ctx.language.src_lang = src_lang
        except NameError:
            pass
        try:
            _ctx.language.dest_lang = dest_lang
        except NameError:
            pass
        try:
            _ctx.flags.word_file_to_translate = word_file_to_translate
        except NameError:
            pass
        try:
            _ctx.flags.use_api = use_api
        except NameError:
            pass
        try:
            _ctx.docx.translation_array = translation_array
        except NameError:
            pass
        try:
            _ctx.docx.blocks_nchar_max_to_translate_array = blocks_nchar_max_to_translate_array
        except NameError:
            pass
        # G1 — webdriver module + paths + remaining flags
        try:
            _ctx.browser.webdriver_module = webdriver
        except NameError:
            pass
        try:
            _ctx.browser.chrome_options = chrome_options
        except NameError:
            pass
        try:
            _ctx.flags.exitonsuccess = exitonsuccess
        except NameError:
            pass
        try:
            _ctx.flags.silent = silent
        except NameError:
            pass
        try:
            _ctx.flags.viewdocx = viewdocx
        except NameError:
            pass
        try:
            _ctx.flags.xlsxreplacefile = xlsxreplacefile
        except NameError:
            pass
        # G4 — aimodel + with_polish snapshots
        try:
            _ctx.flags.aimodel = args.aimodel
        except (NameError, AttributeError):
            pass
        try:
            _ctx.flags.with_polish = args.with_polish
        except (NameError, AttributeError):
            pass
        # 2026-05-11 — split_engine snapshot. Module-level `split_engine`
        # is normalised at line ~776 (lower-cased, whitelisted). Mirror it
        # onto ctx.flags so engine_suffix() and the aligner-invocation
        # branches read the same value.
        try:
            _ctx.flags.split_engine = split_engine
        except (NameError, AttributeError):
            pass
        # 2026-05-12 — aligner_llm_threshold snapshot. 0..100 slider from
        # the legacy frontend. Stored on ctx.flags so save.py can pass it
        # to FASubtitleAligner. Currently a no-op (aligner is mechanical).
        try:
            _ctx.flags.aligner_llm_threshold = getattr(args, "alignerllmthreshold", 0)
        except (NameError, AttributeError):
            pass
        # 2026-05-10 G1 — docxdoc + use_html snapshot for the upcoming
        # docx_io.parse / docx_io.cells extraction.
        try:
            _ctx.docx.docxdoc = docxdoc
        except NameError:
            pass
        try:
            _ctx.docx.use_html = use_html
        except NameError:
            pass
        # 2026-05-10 G2 — shading colour list snapshot for the
        # docx_io.cells.get_cell_data extraction.
        try:
            _ctx.config.shading_color_ignore_text = shading_color_ignore_text
        except NameError:
            pass
        # 2026-05-16 Sprint D-C — xtm + rtlstyle snapshot. These were
        # historically module globals (xtm initialised by
        # initialize_translation_memory_xlsx; rtlstyle by the docx-load
        # block at ~1146). Snapshotting them onto ctx.docx lets the
        # cell-write helpers, generate_char_blocks_array_from_phrases,
        # get_translation_and_replace_after and main() read via ctx,
        # which paves the way for the bridge deletion in slice 6.
        try:
            _ctx.docx.xtm = xtm
        except NameError:
            pass
        try:
            _ctx.docx.rtlstyle = rtlstyle
        except NameError:
            pass
    return _ctx


def _atexit_cleanup_driver() -> None:
    """Best-effort browser shutdown on interpreter exit.

    The happy-path quit lives at the bottom of ``main()``; if anything
    above it raises, the child Chrome process gets orphaned. Registering
    this with ``atexit`` makes sure the driver is closed on any normal
    termination — including crashes — so the launcher's job pool doesn't
    accumulate zombie Chrome processes between failed jobs.

    B5 Jules deep (2026-05-13): also scans `_spawned_driver_pids` so a
    driver that was spawned BEFORE `_ctx` was initialised (e.g. an
    exception fired during `_get_ctx()` itself) still gets cleaned up.
    """
    # First the ctx-tracked driver (normal path)
    try:
        if _ctx is not None and _ctx.browser.driver is not None:
            try:
                _ctx.browser.driver.quit()
            except Exception:
                pass
    except Exception:
        pass
    # Then any PIDs we recorded on spawn before ctx existed
    try:
        for pid in list(_spawned_driver_pids):
            try:
                import os as _os
                import signal as _signal
                _os.kill(pid, _signal.SIGTERM)
            except (ProcessLookupError, PermissionError, OSError):
                pass
        _spawned_driver_pids.clear()
    except Exception:
        pass


# Driver PIDs recorded immediately on spawn so the atexit hook can
# still clean them up if `_ctx` failed to fully initialise.
_spawned_driver_pids: set = set()


def _track_spawned_driver(driver) -> None:
    """Record a driver process so atexit can SIGTERM it on crash.

    Safe to call multiple times; idempotent. Failures are silent —
    tracking is best-effort defence-in-depth, not a hard invariant.
    """
    try:
        svc = getattr(driver, "service", None)
        proc = getattr(svc, "process", None)
        pid = getattr(proc, "pid", None)
        if isinstance(pid, int) and pid > 0:
            _spawned_driver_pids.add(pid)
    except Exception:
        pass


import atexit as _atexit
_atexit.register(_atexit_cleanup_driver)


# _sync_globals_from_ctx (Phase-H mirror bridge) was deleted on
# 2026-05-16 in Sprint D-C slice 6. Every previously-bare-name read in
# cli.py's helper functions has been threaded through ctx; the
# `RuntimeContext` dataclass is now the SOLE canonical state surface for
# downstream pipeline steps. Module-level globals at the top of this
# file remain authoritative only for argparse-time CLI inputs (silent,
# splitonly, viewdocx, etc.) and one-shot import-time setup (rtlstyle,
# docxdoc); these are snapshotted into ctx by `_get_ctx()` and never
# read by name from any function body.


# Track the child processes
def kill_child_process():
    import time
    parent = psutil.Process(os.getpid())
    children = parent.children(recursive=True)

    #print(f"[CLEANUP] Found {len(children)} child process(es).")

    for child in children:
        try:
            #print(f"[CLEANUP] Terminating PID {child.pid} ({' '.join(child.cmdline())})")
            child.terminate()
        except psutil.NoSuchProcess:
            print(f"[CLEANUP] PID {child.pid} already gone.")
        except Exception as e:
            print(f"[CLEANUP] Could not terminate PID {child.pid}: {e}")

    # Give processes some time to exit gracefully
    _, alive = psutil.wait_procs(children, timeout=3)

    for child in alive:
        try:
            print(f"[CLEANUP] Forcing kill on PID {child.pid} ({' '.join(child.cmdline())})")
            child.kill()
        except psutil.NoSuchProcess:
            print(f"[CLEANUP] PID {child.pid} disappeared before kill().")
        except Exception as e:
            print(f"[CLEANUP] Could not kill PID {child.pid}: {e}")

    # Re-check if any children remain
    still_alive = parent.children(recursive=True)
    if still_alive:
        print(f"[CLEANUP] WARNING: Still alive: {[p.pid for p in still_alive]}")
    else:
        #print("[CLEANUP] All child processes terminated.")
        pass

# Run cleanup on normal exit
atexit.register(kill_child_process)

# Trap termination signals
def handle_signal(signum, frame):
    kill_child_process()
    raise SystemExit(0)

signal.signal(signal.SIGINT, handle_signal)
signal.signal(signal.SIGTERM, handle_signal)

my_hazm_normalizer = None


# validate_json_string + get_nested_value_from_json_array + DefaultJsonConfiguration
# now live in src/config.py.


# Network helpers were extracted to ``src/machine_translate_docx/network_utils.py``
# in the 2026-05-16 cli.py shrink phase 2. ``probe_internet`` was renamed
# from the historical ``test_internet`` so pytest doesn't collect it as
# a test function (Sprint B, 2026-05-16).
from .network_utils import (
    probe_internet,
    fetch_country_data,
    check_mirror_url,
    set_se_driver_mirror_url_if_needed,
)

try:
    json_online_configuration = requests.get(json_configuration_url).content
except Exception:
    print("Warning, unable to get configuration from internet at {json_configuration_url}")
    if not probe_internet():
        print("Warning, internet connection seems to be down, google name servers don't respond")
        time.sleep(1)
        
    json_online_configuration = "{}"


# Find default configuration file name from other configuration files
local_configuration_json_path_key = ["local_configuration", "json_filename_path"]
local_configuration_json_path = get_nested_value_from_json_array([json_online_configuration,DefaultJsonConfiguration], local_configuration_json_path_key)

# determine if application is a script file or frozen exe
if getattr(sys, 'frozen', False):
    application_path = os.path.dirname(sys.executable)
elif __file__:
    application_path = os.path.dirname(__file__)

configuration_file_full_path = os.path.join(application_path, local_configuration_json_path)

try:
    if os.path.isfile(configuration_file_full_path):
        with open(configuration_file_full_path) as configuration_file:
          local_json_contents = configuration_file.read()
    else:
        #print(f"Optional local json configuration file not found at {configuration_file_full_path}, ignoring")
        local_json_contents = None
except Exception:
    local_json_contents = None
          
json_configuration_array = [local_json_contents,json_online_configuration,DefaultJsonConfiguration]

# Find default sleep time for update message
version_checker_sleep_seconds_on_update_key = ["version_checker", "sleep_seconds_on_update"]
version_checker_sleep_seconds_on_update = get_nested_value_from_json_array(json_configuration_array,
    version_checker_sleep_seconds_on_update_key, default_when_none=30)
# We assume the program does not need update. Value of 1 is for update needed.
str_needs_update = "0"


# Colors : grey and pink backgroud to ignore
# https://learn.microsoft.com/en-us/office/vba/api/word.wdcolor
shading_color_ignore_text_update_key = ["document", "shading_color_ignore_text"]
shading_color_ignore_text = get_nested_value_from_json_array(json_configuration_array,
    shading_color_ignore_text_update_key)
# G2 (2026-05-10): the colour list is mirrored onto ctx.config inside
# `_get_ctx()`'s lazy snapshot — see the snapshot block above. Calling
# `_get_ctx()` here would force the snapshot to fire before the engine /
# CLI args are parsed below, which would leave ctx.engine.engine empty.

process_platform = platform.system()
if platform.system() == 'Windows':
    from win32com.client import Dispatch

tried_login_in_deepl = False
logged_into_deepl = False

from_text_table = [''] *1
from_text_is_greyed_table = [0] *1
from_text_is_red_color_table = [0] *1
from_text_is_end_of_line_table = [0] *1
from_text_is_beginning_of_line_table = [0] *1
from_text_is_empty_line_table = [0] *1
from_text_is_conditional_end_of_line_table = [0] *1
from_text_by_phrase_separator_table = [''] *1
from_text_by_phrase_table = [''] *1
#number of lines in per phrase
from_text_nb_lines_in_phrase = [0] *1
from_text_nb_lines_in_cell = [0] *1
#
to_text_by_phrase_separator_table = [''] *1
to_text_by_phrase_separator_removed_table = [''] *1
to_text_splited_table1 = [''] *1
to_text_by_phrase_table = [''] *1
to_text_table = [''] *1
to_raw_translated_table = [''] *1
to_text_removed_line_separator = [''] *1
translation_result_using_separator = [''] *1
translation_result_phrase_array = [[]] *1
translation_result = [''] *1
from_text_is_read = [0] *1
word_translation_table_length = 0
table = None

table_cells = [['' for i in range(1)] for j in range(1)]

docxfile_table_number_of_phrases = 0

selenium_chrome_machine_translate_once = None

numerrors_deepl = 0

# We have found zero phrase up to now
docxfile_table_number_of_characters = 0
docxfile_table_number_of_phrases = 0
docxfile_table_number_of_words = 0

numrows = 0
numcols = 0

# 2026-05-16 (P3.2): canonical name is ``config.SUPPORT_EMAIL``. This
# module-level alias is kept for the three local prints + any operator
# scripts that grep for the legacy spelling.
from .config import SUPPORT_EMAIL as E_mail_str


cf = currentframe()
filename = getframeinfo(cf).filename

start_time = datetime.datetime.now()

for m in get_monitors():
    #print(str(m))
    break

parser = argparse.ArgumentParser()

#parser = argparse.ArgumentParser(description = "Translate everything!")
#parser.add_argument('--source-language', required = True, choices = Languages, help="Specify the source language!")
parser.add_argument('--srclang', '-sl', required = False, help="Specify the default source language, en is default (hi,ja,ru,de,ru,hi,ja,in, etc)", default='en')
parser.add_argument('--destlang', '--dl', required = False, help="Specify the destination language with 2 letter code (hi,ja,ru,de,ru,hi,ja,in, etc)")
parser.add_argument('--engine', '-e', required = False, help="Specify the translation engine (google, deepl, chatgpt)")
parser.add_argument('--enginemethod', '-m', required = False, help="Specify the method (javascript, phrasesblock, singlephrase, xlsxfile, textfile )")
parser.add_argument('--aimodel', '-am', required = False, help="Specify the ai model when applicable")
parser.add_argument('--docxfile', '-d', required = False, help="Input file name")
parser.add_argument('--xlsxreplacefile', '-x', required = False, help="Excel xlsx search and replace file")
parser.add_argument('--destfont', '-f', required = False, help="Destination font name")
parser.add_argument('--useapi', '-a', required = False, help="Use api to get translation, lower quality but faster", action='store_true')
parser.add_argument('--split', '-s', required = False, help="Split web translation into cells", action='store_true')
parser.add_argument('--splitengine', '-p', required = False, help="Specify split engine (openai | persian_double_lines)")
parser.add_argument('--alignerllmthreshold', required=False, type=int, default=0,
    help="Persian Double Lines aligner LLM threshold (0..100). 0=mechanical-only (default, current behaviour); 100=fully model-driven. Currently a no-op; reserved for hybrid aligner.")
parser.add_argument('--splitonly', required = False, help="Split translation into lines only, do not translate.", action='store_true')
parser.add_argument('--showbrowser', '-b', required = False, help="Show browser", action='store_true')
parser.add_argument('--exitonsuccess', '-t', required = False, help="Exit progream on success", action='store_true')
parser.add_argument('--viewdocx', '-l', required = False, help="Open the docx file with the default application after completion.", action='store_true')
parser.add_argument('--silent', '-q', required = False, help="Silent, do not ask question and exit silently", action='store_true')
parser.add_argument("--verbose", '-v', help="increase output verbosity", action="store_true")
parser.add_argument("--clientip", '-i', help="Client IP for statistics")
#parser.add_argument('--destination-file', required = True, help="Output file name")
#args = parser.parse_args()
parser.add_argument('--version', required = False, help="Show program version", action='store_true')
parser.add_argument('--with-polish', required = False, help="Run a Persian polish pass after translation and before splitting (API only)", action='store_true', dest='with_polish')

try:
    args = parser.parse_args()
except Exception:  #print("Waiting for the input_element...")
    var = traceback.format_exc()
    print(var)
    #input ("Type enter to continue")

show_version = args.version
silent = args.silent
if show_version:

    print("\nDeveloper: %s\n" %(E_mail_str))
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not silent:
        input("\nEnter to close program")
    sys.exit(0)

if args.docxfile is None:
    parser.print_help()
    print("\nDeveloper: smtv.bot@gmail.com\n")
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not silent:
        input("\nEnter to close program")
    sys.exit(1)

# B-004 / W-3: reject unknown OpenAI model identifiers at CLI parse
# time. Without this guard, `--aimodel gpt-5.5-mini` (a stale dropdown
# value or typo) used to be accepted here and only surfaced as a 400
# BadRequestError deep inside the API call, after the docx had been
# parsed and Chrome had been launched. The whitelist lives in
# `config.py` so the v2 frontend dropdown can pull from the same list.
from .config import VALID_AI_MODELS as _VALID_AI_MODELS, is_valid_ai_model as _is_valid_ai_model
if args.aimodel is not None and not _is_valid_ai_model(args.aimodel):
    print(
        f"ERROR: --aimodel '{args.aimodel}' is not a recognised OpenAI "
        f"model identifier. Allowed values: {', '.join(_VALID_AI_MODELS)}."
    )
    if not silent:
        input("\nEnter to close program")
    else:
        print("Program ended with errors")
    sys.exit(1)

use_html = False


script_folder = os.path.dirname(os.path.realpath(__file__))
os_path = os.environ["PATH"]
new_os_path = "%s;%s" %(script_folder, os_path)

#print("\nAdding %s to path.\n" % (script_folder))
os.environ["PATH"] = new_os_path

use_translation_api = False


#line_separator_str = ' () '
line_separator_str = ' '
#line_separator_nospace_str = '()'
line_separator_nospace_str = '()'
line_separator_regex_str = ' ?\(\) ?'

#pp = pprint.PrettyPrinter(indent=4)


html_file_path = ''

nb_character_total = 0

# Maximum 5000 characters on the free version
# but only 1500 if not logged on to deepl with free version
deepl_max_char_bloc_size_key = ['deepl', 'no_account','maximum_character_block']
deepl_maximum_character_block = get_nested_value_from_json_array(json_configuration_array, deepl_max_char_bloc_size_key)

deepl_sleep_wait_translation_seconds = 0.1
translation_errors_count = 0

word_file_to_translate = args.docxfile

viewdocx = args.viewdocx
client_ip = args.clientip

xlsxreplacefile = args.xlsxreplacefile
dest_font = args.destfont
split_translation = args.split

split_engine = args.splitengine
if split_engine is not None:
    split_engine = split_engine.lower()
    if split_engine not in ('openai', 'persian_double_lines'):
        print("Unknown split engine, accepted values : openai, persian_double_lines. Defaulting to non AI line splitting.")
        split_engine = None

use_api = args.useapi
#use_browser = args.useapi

showbrowser = args.showbrowser
exitonsuccess = args.exitonsuccess
splitonly = args.splitonly
with_polish = args.with_polish
if splitonly:
    split_translation = True

driver = None

src_lang = args.srclang
dest_lang = args.destlang
if dest_lang is not None:
    dest_lang = dest_lang.lower()
else:
    dest_lang = ""
    if not splitonly:
        # C16 + P1-2: respect --silent. The launcher always passes
        # --silent --destlang, so this guard never triggers in production —
        # but any caller that forgets --destlang would otherwise hang
        # forever waiting on stdin. Fail fast with the structured marker.
        if silent:
            print("[FAIL] reason=missing_destlang message=--destlang is required in silent mode", flush=True)
            sys.exit(20)
        dest_lang = input("Please enter language translation code (fr,de,ru,hi,etc.)")

# cjk_segmenter = None 
# if dest_lang == 'zh':
# jieba.enable_paddle()# 启动paddle模式。 0.40版之后开始支持，早期版本不支持
# if dest_lang == 'zh-cn':
# dest_lang = 'zh-CN'
# jieba.enable_paddle()# 启动paddle模式。 0.40版之后开始支持，早期版本不支持
# if dest_lang == 'zh-tw':
# dest_lang = 'zh-TW'
# jieba.enable_paddle()# 启动paddle模式。 0.40版之后开始支持，早期版本不支持
# if dest_lang == 'th':
# if dest_lang == 'ko':
# if dest_lang == 'ja':
# cjk_segmenter = TinySegmenter()
# if dest_lang == 'fa':
# my_hazm_normalizer = Normalizer()


cjk_segmenter = None 
if dest_lang == 'zh-cn':
    dest_lang = 'zh-CN'
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang == 'zh-tw':
    dest_lang = 'zh-TW'
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang.lower() == 'zh-hant' or dest_lang == 'zh-hans':
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
    
if dest_lang == 'th':
    from newmm_tokenizer.tokenizer import word_tokenize
if dest_lang == 'zh' or dest_lang == 'ja' or dest_lang == 'ko':
    from tinysegmenter import TinySegmenter
    cjk_segmenter = TinySegmenter()
if dest_lang == 'fa':
    # 2026-05-13 (feat/exe-packaging): hazm is heavy and only used on
    # the Selenium-scraping FA path. Make it best-effort so a packaged
    # CLI build (OpenAI API only) doesn't fail to start when hazm is
    # absent. Falls back to a passthrough Normalizer that returns text
    # unchanged — fine because the API path never calls .normalize().
    try:
        from hazm import Normalizer
        my_hazm_normalizer = Normalizer()
    except Exception as _hazm_exc:
        class _PassthroughNormalizer:
            def normalize(self, text=""):
                return text
        my_hazm_normalizer = _PassthroughNormalizer()



valid_online_json = validate_json_string(json_online_configuration)
if not valid_online_json:
    print(f"json_online_configuration={json_online_configuration}")
    print(f"Warning: Json file at {json_configuration_url} is not valid. Ignoring this configuration file.")
else:
    #print(f"Using JSON configuration file at {json_configuration_url} : OK")
    pass

valid_local_json = validate_json_string(local_json_contents)
if os.path.isfile(configuration_file_full_path):
    if not valid_local_json:
        print(f"Warning: Json file at {configuration_file_full_path} is not valid. Ignoring this configuration file.")
    else:
        print(f"Using JSON configuration file at {configuration_file_full_path}")

print("")


src_lang_name = (google_translate_lang_codes.get(src_lang))
if src_lang_name is None:
    src_lang_name = ""
    if not splitonly:
        print("Source language name for %s not found. Continuing as it is." % (dest_lang))
else:
    print("Source language name for '%s' : %s" % (src_lang, src_lang_name))

dest_lang_name = (google_translate_lang_codes.get(dest_lang))

if dest_lang_name is None:
    dest_lang_name = deepl_translate_lang_codes.get(dest_lang)
    if not splitonly and dest_lang_name is None:
        print("Target language name for %s not found. Continuing as it is." % (dest_lang))
else:
    print("Target language name for '%s' : %s" % (dest_lang, dest_lang_name))

dest_lang_tag = ""
try:
    dest_lang_tag = office_language_tags[dest_lang]
except Exception:
    pass

translation_engine = args.engine

if translation_engine is not None:
    translation_engine = translation_engine.lower()
else:
    translation_engine = ""

if translation_engine in ['chatgpt', 'deepl']:
    showbrowser = True
else:
    translation_engine = 'google'

if use_api and translation_engine != 'chatgpt':
    use_api = False


engine_method = args.enginemethod
engine_method = "%s" % engine_method
engine_method = engine_method.strip().lower()

translation_array = []

if splitonly:
    engine_method = ''
elif translation_engine == 'google':
    if engine_method == 'api' or use_api == True:
        engine_method = 'api'
    elif engine_method  == 'singlephrase':
        engine_method = 'singlephrase'
    elif engine_method  == 'phrasesblock':
        engine_method = 'phrasesblock'
    elif engine_method =='xlsxfile':
        engine_method = 'xlsxfile'
        # There is a bug on xlsxfile method, show browser for debugging purpose
        showbrowser = True
    elif engine_method  == 'textfile':
        engine_method = 'textfile'
        # There is a bug on textfile method, show browser for debugging purpose
        showbrowser = True
    elif engine_method  == 'javascript':
        engine_method = 'javascript'
    else:
        # Default for `--engine google` was historically `javascript`,
        # which uploads a local HTML file to translate.google.com. That
        # path stopped working when modern Chrome started blocking the
        # widget on file:// URLs (~2022); it now fails fast with an
        # error message but produces an empty docx. Switch the default
        # to `phrasesblock`, which uses the textarea URL — fast and
        # actually working. Users who genuinely want the old file-mode
        # path can still pass `--enginemethod javascript` explicitly.
        engine_method = 'phrasesblock'
elif translation_engine == 'deepl':
    if engine_method == 'singlephrase' or use_api == True:
        engine_method = 'singlephrase'
    elif engine_method  == 'phrasesblock':
        engine_method = 'phrasesblock'
    else:
        engine_method = 'phrasesblock'
elif translation_engine == 'chatgpt':
    # chatgpt-web was removed in the 2026-05-10 cleanup; only the API
    # path remains for chatgpt.
    if engine_method == 'api' or use_api == True:
        engine_method = 'api'
        showbrowser = False
    else:
        engine_method = 'api'
        showbrowser = False

else:
    engine_method = "web"

if engine_method == 'webservice':
    showbrowser = False

if translation_engine == 'chatgpt' and engine_method == 'api':
    from .openai_tools import OpenAITranslator, OpenAIPolisher
    # FASubtitleAligner is no longer imported here (Phase 1): the aligner is
    # decoupled from chatgpt-polish and reached only via the Persian Double
    # Lines Split Method, which performs its own local import on demand.
    chatgpt_max_char_bloc_size_key = ['chatgpt', 'api','maximum_character_block']
else:
    chatgpt_max_char_bloc_size_key = ['chatgpt', 'no_account','maximum_character_block']
chatgpt_maximum_character_block = get_nested_value_from_json_array(json_configuration_array, chatgpt_max_char_bloc_size_key)

# Load openai line splitting package
from .openai_tools import OpenAISubtitleSplitter
    
if translation_engine == 'chatgpt':
    MAX_TRANSLATION_BLOCK_SIZE = chatgpt_maximum_character_block
else:
    MAX_TRANSLATION_BLOCK_SIZE = deepl_maximum_character_block
# Override MAX_TRANSLATION_BLOCK_SIZE value after logging on Deepl

# When translation engine is deepl or chatgpt : use undetected_chromedriver
# Else, use standard selenium webdriver

if translation_engine == 'chatgpt' and engine_method != "webservice":
    # 2026-05-13 (feat/exe-packaging): undetected_chromedriver is only
    # needed for the legacy chatgpt-WEB path (Selenium browser
    # automation). The OpenAI-API engine doesn't touch a browser at
    # all. Treat the import as best-effort so packaged builds that
    # only ship the API path don't fail to start.
    try:
        import undetected_chromedriver as webdriver
    except Exception:
        from selenium import webdriver
else:
    from selenium import webdriver  # regular selenium webdriver

# OpenAI handles + translation log dict must exist BEFORE the first
# `_get_ctx()` runtime call below, otherwise the lazy snapshot inside
# `_get_ctx()` hits NameError on these names and ctx.openai.translation_log
# ends up as the dataclass empty {} default instead of pointing at this
# module-global dict. (Until 2026-05-16 the Phase-H mirror
# `_sync_globals_from_ctx` papered over this; with the mirror deleted
# in Sprint D-C slice 6 the seed dict MUST exist before _get_ctx() runs,
# and the runner mutates it by reference so subsequent reads on
# ctx.openai.translation_log stay in sync.)
oai_translator = None
oai_polisher = None
translation_log = {"run_info": {}, "blocks": [], "summary": {}}

# Mirror the webdriver module onto ctx now that it has been chosen.
# `_get_ctx()`'s lazy snapshot may have fired earlier (e.g. from the
# G2 shading-color mirror at line ~570) when this name did not yet
# exist; re-set it explicitly so create_webdriver(ctx) downstream
# sees the right module.
_get_ctx().browser.webdriver_module = webdriver


# write_translation_log was extracted to
# ``src/machine_translate_docx/translation_log_writer.py`` in the
# 2026-05-16 cli.py shrink phase 3. The injected callback in
# ``docx_io/save.py`` still gets the historical 1-arg signature, so a
# thin shim here forwards ``log_path`` plus the live ctx.
from .translation_log_writer import write_translation_log as _write_translation_log_impl  # noqa: E402


def write_translation_log(log_path: str):
    _write_translation_log_impl(_get_ctx(), log_path)


if not os.path.exists(word_file_to_translate) :
    print("ERROR: File not found: %s" % (word_file_to_translate))
    sys.exit(1)

splitted_filename = os.path.splitext(os.path.basename(word_file_to_translate))

# number of segment separated by dot in the docx filename
splitted_filename_size = len(splitted_filename)

docx_file_name =  "%s%s" % (splitted_filename[splitted_filename_size-2], splitted_filename[splitted_filename_size-1])

if splitted_filename_size > 1:
    word_file_to_translate_extension = splitted_filename[splitted_filename_size-1].lower()

if word_file_to_translate_extension == ".docx":
    try:
        docxdoc = docx.Document(word_file_to_translate)
    except Exception:
        print(f"Error, file {word_file_to_translate} does not appear to be a valid Microsoft Word docx file.")
        print("Please check that the file is a valid document and rerun on a valid Microsoft docx document.\n")

        print("\nDeveloper: %s" % (E_mail_str))
        print("Program version: %s\n" % (PROGRAM_VERSION))
        if not silent:
            input("Enter to close program")
        else:
            print("Program ended with errors")
        sys.exit(2)
    # G1: thread the freshly opened Document and the use_html flag onto
    # the shared RuntimeContext so the upcoming docx_io.parse extraction
    # can read them from ctx instead of module globals.
    _get_ctx().docx.docxdoc  = docxdoc
    _get_ctx().docx.use_html = use_html
    styles = docxdoc.styles
    
    if dest_lang_tag != '':
        styles_element = docxdoc.styles.element
        try:
            # Some office suite like WPS does not handle language tag in a document, ignore it
            rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
            lang_default = rpr_default.xpath('w:lang')[0]
            lang_default.set(docx.oxml.shared.qn('w:val'),dest_lang_tag)
        except Exception:
            # Ignore the language tag of the document, it is not supported by some office suites
            pass

    # Create Right to Left Style if it is not found
    try:
        rtlstyle = styles['rtlstyle']
    except Exception:
        rtlstyle = docxdoc.styles.add_style('rtlstyle', WD_STYLE_TYPE.CHARACTER)
    if dest_lang == "" or dest_lang is None:
        dest_lang_name_from_cell = docxdoc.tables[0].cell(1, 2).text
        print("Lang cell: %s" % (dest_lang_name_from_cell))
        for lang_code, lang_name in google_translate_lang_codes.items():
            #print("%s : %s" % (lang_code, lang_name))
            if dest_lang_name_from_cell == lang_name:
                dest_lang = lang_code

print("File: %s" %(args.docxfile))
print("Language code: %s" %(dest_lang))
print("Language name: %s" %(dest_lang_name))

print("Destination font: %s" %(dest_font))
print("Split: %s" %(split_translation))
print("Splitonly: %s" %(splitonly))


tmx_file_path = "%s\%s_%s.tmx" % (os.path.dirname(word_file_to_translate), os.path.splitext(os.path.basename(word_file_to_translate))[0],dest_lang)
#print(tmx_file_path)


print("Extension: %s" % (word_file_to_translate_extension))

if not os.path.exists(word_file_to_translate):
    print("ERROR: docxfile '%s' not found, exiting." % (word_file_to_translate))


if word_file_to_translate_extension != ".docx":
    print("ERROR: not a docx file. Please convert to docx first then run on docx file. Exiting.")
    if not silent:
        input("Enter to close program")
    else:
        print("Program ended with errors")
    os._exit(3)

print("")


location_primary_country_checker_url_key = ["location", "primary_country_checker_url"]
location_primary_country_checker_url = get_nested_value_from_json_array(json_configuration_array, location_primary_country_checker_url_key)

location_secondary_country_checker_url_key = ["location", "secondary_country_checker_url"]
location_secondary_country_checker_url = get_nested_value_from_json_array(json_configuration_array, location_secondary_country_checker_url_key)

location_http_query_timeout_key = ["location", "http_query_timeout"]
location_http_query_timeout = get_nested_value_from_json_array(json_configuration_array, location_http_query_timeout_key)

# Check if location_http_query_timeout is not an integer > 0
if not isinstance(location_http_query_timeout, int) or location_http_query_timeout <= 0:
    location_http_query_timeout = COUNTRY_QUERY_HTTP_TIMEOUT  # Set to 3 if the condition is not met

chrome_driver_restricted_countries_key = ["chrome_driver", "restricted_countries"]
chrome_driver_restricted_countries = get_nested_value_from_json_array(json_configuration_array, chrome_driver_restricted_countries_key)

chrome_driver_mirror_url_key = ["chrome_driver", "mirror_url"]
chrome_driver_mirror_url = get_nested_value_from_json_array(json_configuration_array, chrome_driver_mirror_url_key)

#print(f"location_primary_country_checker_url = {location_primary_country_checker_url}")
#print(f"location_secondary_country_checker_url = {location_secondary_country_checker_url}")
#print(f"chrome_driver_restricted_countries = {chrome_driver_restricted_countries}")
#print(f"chrome_driver_mirror_url = {chrome_driver_mirror_url}")

# fetch_country_data / check_mirror_url / set_se_driver_mirror_url_if_needed
# were extracted to ``src/machine_translate_docx/network_utils.py``
# in the 2026-05-16 cli.py shrink phase 2. The thin module-level startup
# sequence below drives them via the already-loaded JSON config values.

# Set chrome driver download proxy URL for restricted countries
country_name = fetch_country_data(
    location_primary_country_checker_url, http_timeout=location_http_query_timeout,
)

# If primary URL fails or does not return a valid country name, fallback to the secondary URL
if not country_name:
    print("Falling back to secondary URL...")
    country_name = fetch_country_data(
        location_secondary_country_checker_url, http_timeout=location_http_query_timeout,
    )

# Set environment variable if needed
set_se_driver_mirror_url_if_needed(
    country_name,
    chrome_driver_mirror_url,
    restricted_countries=chrome_driver_restricted_countries or [],
    http_timeout=location_http_query_timeout,
)

# Set up Chrome options.
chrome_options = Options()
chrome_options.add_argument("--disable-web-security")
chrome_options.add_argument("--disable-xss-auditor")
chrome_options.add_argument("--lang=en-GB")
chrome_options.add_argument("--log-level=3")  # fatal
chrome_options.add_argument("--password-store=basic")
# Mirror onto ctx so create_webdriver(ctx) sees the populated options.
# `_get_ctx()`'s lazy snapshot may have fired earlier (e.g. from the
# G1 docxdoc mirror at line ~1091) when this name did not yet exist.
_get_ctx().browser.chrome_options = chrome_options


if not showbrowser :
    chrome_options.add_argument("--headless")
    if platform.system() == "Linux":  # Linux
        chrome_options.add_argument("--disable-gpu")         # remove GPU fallback flutters
        chrome_options.add_argument("--disable-software-rasterizer")
        chrome_options.add_argument("--enable-features=UseOzonePlatform")
        chrome_options.add_argument("--use-system-clipboard")
        chrome_options.add_argument("--disable-features=site-per-process")
        chrome_options.add_argument("--allow-running-insecure-content")
        chrome_options.add_argument("--disable-extensions")
        chrome_options.add_argument("--remote-allow-origins=*")
        chrome_options.add_argument("--window-size=1920,1080")
elif platform.system() == "Linux":  # Linux    
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-backgrounding-occluded-windows")
    chrome_options.add_argument("--disable-features=OptimizationGuideModelDownloading")
    chrome_options.add_argument("--disable-infobars")

# Used to tokenize thai
#thai_segmenter = thai_tokenizer_tokenizer()
#word_tokenize(text)

#translator = Translator(service_urls=['translate.google.com'], user_agent='Mozilla/5.0 (Windows NT 6.1; Win64; x64; rv:47.0) Gecko/20100101 Firefox/47.0')
#user_agent='Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36')

#driver = webdriver.Chrome()
#driver.set_window_position(0, 350)
#driver.set_window_size(400, 650)

#driver.get("https://translate.google.com/#%s/%s/" % (src_lang,dest_lang))
#driver = uc.Firefox()

def get_translated_cells_content(ctx: RuntimeContext, lineno, to_translate):
    print("get_translation for line %d" % (lineno))
    print("from_text_nb_lines_in_phrase %d" % (ctx.docx.from_text_nb_lines_in_phrase[lineno]))
    translation = ""

    _dest_lang = ctx.language.dest_lang
    if _dest_lang.lower() == 'ja' or _dest_lang.lower() == 'zh-cn' or _dest_lang.lower() == 'zh-tw' or _dest_lang.lower() == 'ko':
        cell_space = ''
    else:
        cell_space = ' '

    last_row_n = ctx.docx.from_text_nb_lines_in_phrase[lineno] + lineno
    for row_n in range(lineno, last_row_n):
        #cell_text = docxdoc.tables[0].cell(row_n, 2).text
        cell_text = ctx.docx.table_cells[row_n][2].text
        cell_text = cell_text.strip()
        if cell_text != "":
            print("adding cell %s " % (cell_text))
            if row_n == lineno:
                translation = cell_text
            else:
                translation = translation + cell_space + cell_text
    return translation

found_google_cookies_consent_button = False
google_translate_first_page_loaded = False

def selenium_chrome_translate_get_from_text_array(to_translate, index):
    # Guard: pre-refactor code raised IndexError here whenever the block
    # loop produced fewer translated lines than the document has phrases
    # (off-by-one or alignment mismatch). Returning '' instead lets the
    # caller log "Error translating='…'" and continue, rather than
    # killing the whole job.
    #
    # 2026-05-16 Sprint D-C: signature is dictated by the dispatcher
    # contract (`ctx.engine.dispatcher(to_translate, index)`), so we
    # cannot take ``ctx`` as a positional arg. Resolve the singleton
    # via ``_get_ctx()`` and read the canonical array off it.
    ta = _get_ctx().docx.translation_array
    if 0 < index <= len(ta):
        return ta[index - 1]
    print(f"[WARN] translation_array index out of range: index={index}, len={len(ta)}")
    return ""
    
# method to get the downloaded file name
# function to wait for download to finish and then rename the latest downloaded file
# deleted in the same pass. The only remaining mentions in the tree
# are historical CHANGELOG and PROJECT_MEMORY notes.


# set_translation_function was extracted to ``src/dispatch.py`` in the
# 2026-05-10 architecture cleanup. The dispatch module is the single
# source of truth for engine routing — both ``set_translation_function``
# AND the ``use_phrasesblock`` predicate live there, so future drift
# between them is impossible.
from .dispatch import set_translation_function, use_phrasesblock as _dispatch_use_phrasesblock
from . import dispatch as _dispatch_module
_dispatch_module.set_array_dispatcher(selenium_chrome_translate_get_from_text_array)


def selenium_chrome_machine_translate(ctx: RuntimeContext, to_translate, index):
    """Per-call wrapper around ``ctx.engine.dispatcher`` with retry logic.

    Threaded in Phase F1.4: reads the dispatcher pointer, the active
    engine + method, the retry counter, and the inter-retry sleep
    factor from ``ctx`` instead of module globals.
    """
    translation = ""
    translation_try_count = 1
    max_try_count = 15
    try:
        while translation_try_count < max_try_count and translation == "":
            if translation_try_count > 1:
                print("Retrying to translate again (%d)..." % (translation_try_count))
                ctx.docx.translation_errors_count += 1
                ctx.browser.deepl_sleep_wait_translation_seconds *= 1.1
                print("%d translation retry so far..." % (ctx.docx.translation_errors_count))
            if ctx.engine.engine == 'deepl':
                if ctx.engine.method == 'phrasesblock':
                    translation = ctx.engine.dispatcher(to_translate, index)
                else:
                    translation = ctx.engine.dispatcher(to_translate, translation_try_count - 1)
            elif ctx.engine.method in (
                'textfile', 'xlsxfile', 'phrasesblock',
                'javascript', 'webservice', 'api',
            ):
                translation = ctx.engine.dispatcher(to_translate, index)
            else:
                translation = ctx.engine.dispatcher(to_translate)
            translation_try_count += 1
    except Exception:
        print("Error in selenium_chrome_machine_translate function.")
    return translation
    
def initialize_translation_memory_xlsx(ctx: RuntimeContext):
    global xtm
    # If --xlsxreplacefile was provided in the command line
    if xlsxreplacefile is not None:
        print("xlsxreplacefile: %s" % (xlsxreplacefile))
        xtm = xlsx_translation_memory.xlsx_translation_memory(xlsxreplacefile)
        print("")
    else:
        xtm = xlsx_translation_memory.xlsx_translation_memory(None)
    # 2026-05-16 Sprint D-C — mirror the newly-created XTM instance onto
    # ctx.docx. The module-level `xtm` is kept in sync so the
    # `print_replaced_items_*` calls in main() that go through ctx.docx.xtm
    # see the populated handle. (Pre-slice-6 the Phase-H mirror would have
    # stomped this from the snapshot's default None; the mirror is gone now,
    # but threading every reader through ctx.docx still requires this explicit
    # write.)
    ctx.docx.xtm = xtm


def is_end_of_line(line):
    for eol in eol_array:
        #print("Testing is_end_of_line '%s' on string '%s'" %(eol, line))
        if re.search(eol, line):
            #print("Found is_end_of_line '%s' on string '%s'" %(eol, line))
            return 1
    return 0


def is_conditional_end_of_line(line):
    for ceol in eol_conditional_array:
        #print("Testing is_conditional_end_of_line '%s' on string '%s'" %(ceol, line))
        if re.search(ceol, line):
            return 1
    return 0

def is_beginning_of_line(line):
    for bol in bol_array:
        if re.search(bol, line):
            return 1
    return 0

def is_empty_line(line):
    line_trimmed = re.sub(' +', '', line)
    length = len(line_trimmed)
    if length == 0:
        return 1
    return 0

# Per-cell read + write helpers were extracted to ``src/docx_io/cells.py``
# in the 2026-05-10 docx_io extraction pass.
#
# - The two shading-detection helpers (formerly
#   ``get_paragraph_shading_color`` and ``get_run_shading_color``) are now
#   private to ``docx_io.cells`` (``_paragraph_shading_color`` /
#   ``_run_shading_color``).
# - ``get_cell_data(ctx, cell, row_n)`` lives in ``docx_io.cells`` and
#   reads the colour-ignore list from ``ctx.config.shading_color_ignore_text``.
# - The write shims (``change_cell_font``, ``cell_set_1st_paragraph``,
#   ``cell_add_paragraph``) below thread the entry-script globals into
#   the new explicit-kwarg implementations.
from .docx_io import (
    _iter_paragraph_runs,
    _cell_add_paragraph_impl,
    _change_cell_font_impl,
    _cell_set_first_paragraph_impl,
    get_cell_data,
)



# change_cell_font was extracted to ``src/docx_io/cells.py`` in the
# 2026-05-10 docx_io extraction pass. Thin shim — reads the entry-script
# global ``dest_font`` and delegates to the new implementation.
def change_cell_font(ctx: RuntimeContext, cell):
    _change_cell_font_impl(cell, ctx.language.dest_font)

def tokenize_text_to_array(ctx: RuntimeContext, text, lang_code):
    lang_code = lang_code + ""
    lang_code = lang_code.lower()

    words = []
    # In japanese tokenize words
    if lang_code == 'ja' or lang_code== 'zh-cn' or lang_code == 'zh' or lang_code == 'zh-tw' or lang_code == 'ko' or lang_code== 'zh-hans' or lang_code == 'zh-hant':
        words = cjk_segmenter.tokenize(text)
    # In other languages, just use spaces
    elif lang_code == 'th':
        #words = thai_segmenter(text)
        words = word_tokenize(text)
    # In other languages, just use spaces
    else:
        #ctx.docx.xtm.tokenize_phrase(text, dest_lang)

        # search do not split here
        #ctx.docx.xtm.pprint_translation_memory_list()

        # Old simple split method replaced by tokenize_phrase method having do not split
        # words = text.split()

        words = ctx.docx.xtm.tokenize_phrase(text, lang_code)
        #input("Wait, remove tokenize_phrase here..")


    return words

def divide_array(words_array, dest_lang, width):
    dest_lang = dest_lang.lower()

    print("Divide into max %d size lines" % (width))
    count = len(words_array)
    offsets = [0]
    for w in words_array:
        offsets.append(offsets[-1] + len(w))

    minima = [0] + [10 ** 20] * count
    breaks = [0] * (count + 1)

    def cost(i, j):
        w = offsets[j] - offsets[i] + j - i - 1
        if w > width:
            return 10 ** 10
        return minima[i] + (width - w) ** 2

    def search(i0, j0, i1, j1):
        stack = [(i0, j0, i1, j1)]
        while stack:
            i0, j0, i1, j1 = stack.pop()
            if j0 < j1:
                j = (j0 + j1) // 2
                for i in range(i0, i1):
                    c = cost(i, j)
                    if c <= minima[j]:
                        minima[j] = c
                        breaks[j] = i
                stack.append((breaks[j], j+1, i1, j1))
                stack.append((i0, j0, breaks[j]+1, j))

    n = count + 1
    i = 0
    offset = 0
    while True:
        r = min(n, 2 ** (i + 1))
        edge = 2 ** i + offset
        search(0 + offset, edge, edge, r + offset)
        x = minima[r - 1 + offset]
        for j in range(2 ** i, r - 1):
            y = cost(j + offset, r - 1 + offset)
            if y <= x:
                n -= j
                i = 0
                offset += j
                break
        else:
            if r == n:
                break
            i = i + 1

    lines = []
    j = count
    while j > 0:
        i = breaks[j]
        # In japanese just join words_array without adding any spaces
        if dest_lang == 'ja' or dest_lang == 'zh-cn' or dest_lang == 'zh-tw' \
            or dest_lang.lower() == 'zh-hans' or dest_lang.lower() == 'zh-hant' or dest_lang == 'ko' \
                or dest_lang == 'th':
            lines.append(''.join(words_array[i:j]))
        # In other languages, join words_array using a space
        else:
            lines.append(' '.join(words_array[i:j]))

        j = i
    lines.reverse()
    return lines


def split_phrases(ctx: RuntimeContext):
    """Group consecutive cells into a single phrase keyed by the
    first cell of the phrase.

    Phase H bridge: the historical version wrote to module-level
    globals (`from_text_by_phrase_separator_table`,
    `from_text_by_phrase_table`, `from_text_nb_lines_in_phrase`)
    that were sized `[''] * 1`. Every row beyond index 0 raised
    IndexError, which the enclosing `try/except` in
    `read_and_parse_docx_document` swallowed silently — leaving the
    arrays empty so every translation block downstream became
    empty too. All reads + writes now go through ``ctx.docx`` which
    is properly sized at parse time.
    """
    docx = ctx.docx
    n_last_row_phrase = 3
    last_table_row = docx.word_translation_table_length
    cur_row_n = 2
    while cur_row_n < (last_table_row):
        if docx.from_text_nb_lines_in_cell[cur_row_n] > 1:
            pass
        if docx.from_text_is_beginning_of_line_table[cur_row_n] == 1:
            n_last_row_phrase = cur_row_n
            nb_lines_in_phrase = 1
            docx.from_text_nb_lines_in_phrase[cur_row_n] = docx.from_text_nb_lines_in_cell[cur_row_n]
            while docx.from_text_is_end_of_line_table[n_last_row_phrase] != 1 \
                and n_last_row_phrase < (last_table_row - 1):
                if docx.from_text_by_phrase_separator_table[cur_row_n] == "":
                    docx.from_text_by_phrase_separator_table[cur_row_n] = docx.from_text_table[n_last_row_phrase]
                    docx.from_text_by_phrase_table[cur_row_n] = docx.from_text_table[n_last_row_phrase]
                else:
                    docx.from_text_by_phrase_separator_table[cur_row_n] = docx.from_text_by_phrase_separator_table[cur_row_n] + line_separator_str + docx.from_text_table[n_last_row_phrase]
                    docx.from_text_by_phrase_table[cur_row_n] = docx.from_text_by_phrase_table[cur_row_n] + ' ' + docx.from_text_table[n_last_row_phrase]
                    nb_lines_in_phrase += 1
                    if docx.from_text_nb_lines_in_cell[cur_row_n] > 1:
                        pass
                    docx.from_text_nb_lines_in_phrase[cur_row_n] += docx.from_text_nb_lines_in_cell[n_last_row_phrase]
                n_last_row_phrase += 1
            if docx.from_text_by_phrase_separator_table[cur_row_n] == "":
                docx.from_text_by_phrase_separator_table[cur_row_n] = docx.from_text_table[n_last_row_phrase]
                docx.from_text_by_phrase_table[cur_row_n] = docx.from_text_table[n_last_row_phrase]
            else:
                docx.from_text_by_phrase_separator_table[cur_row_n] = docx.from_text_by_phrase_separator_table[cur_row_n] + line_separator_str + docx.from_text_table[n_last_row_phrase]
                nb_lines_in_phrase += 1
                docx.from_text_nb_lines_in_phrase[cur_row_n] += docx.from_text_nb_lines_in_cell[n_last_row_phrase]
                docx.from_text_by_phrase_table[cur_row_n] = docx.from_text_by_phrase_table[cur_row_n] + ' ' + docx.from_text_table[n_last_row_phrase]
            if docx.use_html:
                print("(%d)from_text_by_phrase_table[%d]=%s<br>" % (n_last_row_phrase, cur_row_n, docx.from_text_by_phrase_table[cur_row_n]))
            nb_lines_in_phrase_str = "[%s]" % (nb_lines_in_phrase)

            cur_row_n = n_last_row_phrase + 1
        else:
            cur_row_n += 1

    return 0

# delete_paragraph was extracted to ``docx_io/cells.py`` in the
# 2026-05-16 cli.py shrink phase 2.
from .docx_io.cells import delete_paragraph  # noqa: E402


def prepare_and_clear_cell_for_writing(ctx: RuntimeContext, row_n, translation_cell_text):
    """Clear and re-init a target-language cell.

    Threaded in Phase F1.3 + Sprint D-C slice 2 (2026-05-16): reads
    ``ctx.docx.table_cells``, ``ctx.language.dest_lang`` /
    ``ctx.language.dest_font``, and ``ctx.docx.rtlstyle`` in place of
    the historical module-level globals. All four are mirrored back to
    module scope by the Phase-H bridge until it is deleted in slice 6.
    """
    paragraph_no = 0
    # Skip rows that don't have a third cell (footer / single-column notes
    # at the end of subtitle DOCX files). Without this guard the legacy
    # code would IndexError on every short row and spam the log.
    if row_n >= len(ctx.docx.table_cells) or len(ctx.docx.table_cells[row_n]) < 3:
        return
    current_cell = ctx.docx.table_cells[row_n][2]

    current_cell._element.clear_content()
    
    # Clear paragraphs in the cell
    for paragraph in current_cell.paragraphs:
        if paragraph_no != 0:
            delete_paragraph(paragraph)
        else:
            paragraph.text = ''
        paragraph_no += 1

    # Ensure there's at least one paragraph in the cell
    if len(current_cell.paragraphs) == 0:
        cell_paragraph = current_cell.add_paragraph("")
    else:
        cell_paragraph = current_cell.paragraphs[0]

    # Add orientation for Right-to-Left (RTL) languages
    if ctx.language.dest_lang in right_to_left_languages_list.keys():
        run = cell_paragraph.add_run(translation_cell_text)
        run.style = ctx.docx.rtlstyle  # Ensure `rtlstyle` exists in the document
        font = run.font
        font.rtl = True
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.text = translation_cell_text
        # Check if the 'Normal' style exists before applying
        try:
            cell_paragraph.style = 'Normal'
        except KeyError:
            #print("Warning: 'Normal' style not found. Falling back to 'Default Paragraph Font'.")
            try:
                cell_paragraph.style = 'Default Paragraph Font'
            except KeyError:
                #print("Error: No usable default style found. Proceeding without style assignment.")
                pass
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Apply font changes if necessary
    if ctx.language.dest_font != "":
        change_cell_font(ctx, current_cell)

    ctx.docx.table_cells[row_n][2] = current_cell


# cell_set_1st_paragraph and cell_add_paragraph were extracted to
# ``src/docx_io/cells.py`` in the 2026-05-10 docx_io extraction pass.
# Thin shims here — read the entry-script globals and delegate. The
# new implementations take their dependencies as explicit kwargs so
# they can be unit-tested without RuntimeContext.
def cell_set_1st_paragraph(ctx: RuntimeContext, row_n, paragraph_text):
    # Docx Cell objects are mutable, so we don't need to extract+reassign:
    # the impl mutates the cell in place and ctx.docx.table_cells already
    # references the same object.
    _cell_set_first_paragraph_impl(
        ctx.docx.table_cells[row_n][2],
        paragraph_text,
        dest_lang=ctx.language.dest_lang,
        dest_font=ctx.language.dest_font,
        rtlstyle=ctx.docx.rtlstyle,
    )


def cell_add_paragraph(ctx: RuntimeContext, row_n, paragraph_text):
    _cell_add_paragraph_impl(
        ctx.docx.table_cells[row_n][2],
        paragraph_text,
        dest_lang=ctx.language.dest_lang,
        dest_font=ctx.language.dest_font,
        rtlstyle=ctx.docx.rtlstyle,
    )


# read_and_parse_docx_document was extracted to ``src/docx_io/parse.py``
# in the 2026-05-10 G3 thread-globals pass. The function is re-exported
# from the package root so existing callers (`main()` in this module)
# keep working without a signature change. The new implementation reads
# `docxdoc` and `use_html` from `ctx.docx`, `silent` / `splitonly` /
# `word_file_to_translate` from `ctx.flags`, and lazy-imports the four
# `is_*_line` predicates plus `prepare_and_clear_cell_for_writing` and
# `split_phrases` from this module to avoid an import cycle.
from .docx_io.parse import read_and_parse_docx_document  # noqa: E402,F401


# deepl_double_linefeed_between_phrases was extracted to
# ``engines/deepl.py`` in the 2026-05-16 cli.py shrink phase 2.
from .engines.deepl import deepl_double_linefeed_between_phrases  # noqa: E402


def generate_char_blocks_array_from_phrases(ctx: RuntimeContext, text_file_path):
    ctx.docx.docxfile_table_number_of_phrases = 0
    print("Generating %d character blocks for translation..." % (ctx.config.max_translation_block_size))
    #if ctx.docx.xtm.wb is not None:
    if ctx.docx.xtm is not None:
        print("Replacing text before using excel file...\n")
    text_to_translate = ''
    text_to_translate_array = []
    ctx.docx.blocks_nchar_max_to_translate_array = []
    
    double_linefeed_between_phrases = False
    phrase_separator = "\n"
    phrase_separator_len = 1
    if ctx.engine.engine == 'deepl':
        if deepl_double_linefeed_between_phrases(ctx.language.dest_lang):
            double_linefeed_between_phrases = False
            phrase_separator = "\n\n"
            phrase_separator_len = 2
    
    for i, line in enumerate(ctx.docx.from_text_table):
        item = ctx.docx.from_text_by_phrase_separator_table[i]
        item = item.strip()

        item_searched_and_replaced_before = item

        if item_searched_and_replaced_before != '':
            if xlsxreplacefile is not None:
                #if ctx.docx.xtm.wb is not None:
                if ctx.docx.xtm.wb is not None:
                    #print("%d/%d" % (i, word_translation_table_length))
                    #print("Phrase to translate :'%s'\n" % (item.strip()))
                    item_searched_and_replaced_before, nb_searched_and_replaced_before = ctx.docx.xtm.search_and_replace_text('before', item)
                    if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                        continue
        
        if item_searched_and_replaced_before != '':
            #text_to_translate = text_to_translate + '''%s
#''' % (item)
            text_to_translate_array.append(item_searched_and_replaced_before)
            ctx.docx.docxfile_table_number_of_phrases = ctx.docx.docxfile_table_number_of_phrases + 1
                
    #print (text_to_translate)
    #print (text_to_translate_array)
    
    len_text_to_translate_array = len(text_to_translate_array)
    #print("len(text_to_translate_array)=%d" % (len(text_to_translate_array)))
    
    
    current_text_block = ""
    current_text_block_len = len(current_text_block)
    
    
    #print(text_to_translate_array)
    #print(len_text_to_translate_array)
    #input("len_text_to_translate_array")
    #print("Before current_text_block generation")
    for index in range(0, len(text_to_translate_array)):
        #print("%d : '%s'" % (index, text_to_translate_array[index]))
        current_phrase_str = text_to_translate_array[index]
        
        if current_text_block_len + len(current_phrase_str) <= ctx.config.max_translation_block_size + phrase_separator_len:
            if len(current_text_block) == 0:
                current_text_block = current_phrase_str
                current_text_block_len = len(current_text_block)
                #print(current_phrase_str)
                #input("adding first phrase")
            else:
               
                current_text_block = current_text_block + phrase_separator + current_phrase_str
                current_text_block_len = current_text_block_len + len(current_phrase_str) + phrase_separator_len
                #print(current_phrase_str)
                #print("current_text_block")
                #print(current_text_block)
        else:
            ctx.docx.blocks_nchar_max_to_translate_array.append(current_text_block)
            #print("Current block of %d characters:\n-------------------------------------------------" % (ctx.config.max_translation_block_size))
            #print("current_text_block")
            #print(current_text_block)
            #print("end")
            #input("OK Here")
            #input("adding more phrase")
            #print(current_text_block.split("\n"))
            #print("------------------------------------------\nBlock size : %d" % (current_text_block_len))
            
            current_text_block = current_phrase_str
            current_text_block_len = len(current_phrase_str)
            #input("------------------------------------------\nType enter to continue")
        
        if index == (len(text_to_translate_array) - 1):
            ctx.docx.blocks_nchar_max_to_translate_array.append(current_text_block)
            #print("Current block of %d characters:\n-------------------------------------------------"  % (ctx.config.max_translation_block_size))
            #print(current_text_block.split("\n"))
            #print("------------------------------------------\nBlock size : %d" % (current_text_block_len))
            #input("------------------------------------------\nType enter to continue")
    
    #print("blocks_nchar_max_to_translate_array:")
    #print("\n******************************************\n".join(ctx.docx.blocks_nchar_max_to_translate_array))
    #print ("len(text_to_translate_array) = %d " % (len(text_to_translate_array) - 1)) 

    #print("text_to_translate=\n%s" % (text_to_translate))
    try:
        #text_file_path = docx_file_name + '.txt'
        #text_file_to_translate = open(text_file_path, 'w', encoding='utf-8')
        #text_file_to_translate.write(text_to_translate)
        #text_file_to_translate.close()
        pass
    except Exception:
        var = traceback.format_exc()
        print(var)


def translate_from_phrasesblock(ctx: RuntimeContext):
    text_file_path = docx_file_name + '.txt'
    text_file_full_path = os.path.realpath(text_file_path)
    #print("text_file_full_path=%s" % text_file_full_path)
    #generate_text_file_from_phrases(ctx, text_file_full_path)
    generate_char_blocks_array_from_phrases(ctx, text_file_full_path)

    translation_succeeded = True

    #input("phrasesblock")
    print("Starting translation in %s using phrase blocks of %d characters..." % (ctx.engine.engine, ctx.config.max_translation_block_size))

    translation_succeeded, ctx.docx.translation_array = _runner_translate_maxchar_blocks(ctx)
    try:
        os.remove(text_file_path)
        pass
    except Exception:
        pass
    return translation_succeeded

def translate_docx(ctx: RuntimeContext):
    translation_succeeded = True

    # Splitonly mode = "only run document_split_phrases, do NOT translate".
    # Without this guard the chatgpt branch of use_phrasesblock() fires
    # (it always returns True for chatgpt regardless of method), and the
    # runner is invoked with engine_method='' (cleared by the splitonly
    # gate at line ~983) — which fails immediately with
    # `chatgpt method '' not supported (supported: api)`.
    # Added 2026-05-16 to unblock launcher._apply_basic_split (raw-cache
    # refactor): the spec'd `--splitonly --engine chatgpt --enginemethod
    # api` re-spawn relies on this short-circuit.
    if ctx.flags.splitonly:
        return translation_succeeded

    # ------------------------------------------------------------------
    # Engine-method specific translators
    # ------------------------------------------------------------------
    if engine_method == "textfile":
        google_translate_from_text_file(ctx)
        return translation_succeeded

    if engine_method == "javascript":
        google_translate_from_html_javascript(ctx)
        # Fail loudly when Google's web widget refuses to translate the
        # local file (a known modern-Chrome limitation, since ~2022).
        # Silent failure here used to cascade into ~14 retries per
        # phrase × N phrases of `translation_array index out of range`
        # warnings before producing an empty docx — completely
        # unreadable failure mode for the user.
        if not ctx.docx.translation_array:
            print("[ERROR] Google web translate returned 0 lines.")
            print("[ERROR] Modern Chrome blocks the Google translate widget on")
            print("[ERROR] file:// URLs (CORS / sandboxing). This engine path")
            print("[ERROR] cannot complete in current Chrome versions.")
            print("[INFO] Use the OpenAI API engine (chatgpt) or DeepL instead.")
        return translation_succeeded

    if engine_method == "xlsxfile":
        google_translate_from_html_xlsxfile(ctx)
        return translation_succeeded

    # ------------------------------------------------------------------
    # Phrase-block logic — predicate lives in src/dispatch.py so it
    # can never drift from set_translation_function (both share the
    # same engine ↔ method matrix). See dispatch.use_phrasesblock for
    # the per-engine policy.
    # ------------------------------------------------------------------
    if _dispatch_use_phrasesblock(translation_engine, engine_method):
        translation_succeeded = translate_from_phrasesblock(ctx)

    return translation_succeeded



def get_translation_and_replace_after(ctx: RuntimeContext):
    # Phase H: seed local `driver` from ctx so the later reassignment
    # branches don't trigger UnboundLocalError on prior reads.
    driver = ctx.browser.driver

    phrase_no = 0

    # ── DIAGNOSTIC: confirm we enter here with polished ctx.docx.translation_array ─
    if ctx.openai.polisher is not None:
        print(f"[DIAG] get_translation_and_replace_after: ctx.docx.translation_array has "
              f"{len(ctx.docx.translation_array)} lines — distributing into ctx.docx.to_text_by_phrase_separator_table")
    # ─────────────────────────────────────────────────────────────────────

    p_remove_pause = re.compile('(?i)<pause>')
    p_remove_double_spaces = re.compile(' +')
    p_remove_parenthesis_spaces = re.compile('\( +')

    for i, line in enumerate(ctx.docx.from_text_table):
        item = ctx.docx.from_text_by_phrase_separator_table[i]
        item = item.strip()
        from_language = ctx.language.src_lang
        phrase_separator_removed_str = ''

        p_remove_separator = re.compile(line_separator_regex_str)
        p_remove_double_spaces = re.compile(' +')

        # Avec separateurs ()

        try:
            web_translation_separators = ''
            if item.strip() != '':
                phrase_no = phrase_no + 1
                print("\n%d/%d" % (i, ctx.docx.word_translation_table_length))
                print("Phrase to translate :'%s'\n" % (item.strip()))
                item = item.strip()

                item_searched_and_replaced_before = item
                if xlsxreplacefile is not None:
                    if ctx.docx.xtm.wb is not None:
                        item_searched_and_replaced_before, nb_searched_and_replaced_before = ctx.docx.xtm.search_and_replace_text('before', item)
                        if item_searched_and_replaced_before.strip() == '' or item_searched_and_replaced_before is None:
                            continue
                if ctx.flags.splitonly:
                    web_translation_separators = get_translated_cells_content(ctx, i, item_searched_and_replaced_before)
                elif ctx.flags.use_api:
                    try:
                        web_translation_separators = ""
                        #if ctx.flags.use_api:
                        #    translation = ctx.openai.translator.translate(item_searched_and_replaced_before, src=ctx.language.src_lang, dest=ctx.language.dest_lang)
                        #    web_translation_separators = translation.text
                        if not len(web_translation_separators) > 0:
                            ctx.flags.use_api = False
                            # Faster google Chrome translate failed, using Selenium as backup

                            if driver is None:
                                print(f"[Line {inspect.currentframe().f_lineno}] Starting Chrome browser\n")
                                
                                service = Service()                                
                                driver = webdriver.Chrome(service=service, options=ctx.browser.chrome_options)
                                
                                driver.set_window_position(100, 100)
                                driver.set_window_size(800, 700)
                                #driver.set_window_size(400, 650)

                            print("phrase_no=%d" % phrase_no)
                            web_translation_separators = selenium_chrome_machine_translate(ctx, item_searched_and_replaced_before, phrase_no)
                    except Exception:
                        ctx.flags.use_api = False
                        # Faster google Chrome translate failed, using Selenium as backup

                        if driver is not None:
                            print(f"Starting Chrome browser\n")
                            
                            service = Service()                                
                            driver = webdriver.Chrome(service=service, options=ctx.browser.chrome_options)

                        if ctx.engine.engine == 'google' and driver is not None:
                            driver.set_window_position(100, 100)
                            driver.set_window_size(800, 700)

                        print("phrase_no = %d" % phrase_no)
                        web_translation_separators = selenium_chrome_machine_translate(ctx, item_searched_and_replaced_before, phrase_no)
                else:
                    if ctx.engine.method == "singlephrase" and ctx.engine.engine == 'deepl':
                        translation_succeeded, web_translation_separators  = selenium_chrome_machine_translate(ctx, item_searched_and_replaced_before, phrase_no)
                    else:
                        web_translation_separators = selenium_chrome_machine_translate(ctx, item_searched_and_replaced_before, phrase_no)
                        
                #web_translation_separators = translation.text
                phrase_separator_removed_str = p_remove_double_spaces.sub(' ', web_translation_separators)

                #print("Google translation='%s'" % (phrase_separator_removed_str.encode('utf8')))
                if xlsxreplacefile is not None:
                    nb_searched_and_replaced = 0
                    web_translation_separators_searched_and_replaced, nb_searched_and_replaced = ctx.docx.xtm.search_and_replace_text('after', phrase_separator_removed_str)
                    if nb_searched_and_replaced > 0:
                        #print("\nPhrase %d replacements :\n'%s'" % (nb_searched_and_replaced, web_translation_separators))
                        #print("Replaced phrase :\n'%s'" % (web_translation_separators_searched_and_replaced))
                        phrase_separator_removed_str = web_translation_separators_searched_and_replaced

                if ctx.language.dest_lang in right_to_left_languages_list.keys():
                    #phrase_separator_removed_aligned_str = reverse_string (phrase_separator_removed_str)
                    #phrase_separator_removed_aligned_str = "\u202B" + phrase_separator_removed_str + "\u202C"
                    phrase_separator_removed_aligned_str = get_display(phrase_separator_removed_str)
                else:
                    phrase_separator_removed_aligned_str = phrase_separator_removed_str
                try:
                    if ctx.flags.splitonly:
                        print("Translated text :'%s'\n" % (phrase_separator_removed_aligned_str))
                    else:
                        print("%s translation (%s):'%s'" % (ctx.engine.engine.title() ,ctx.language.dest_lang_name, phrase_separator_removed_aligned_str))
                except Exception:
                    print("")
                    print("Google translation='%s'" % (phrase_separator_removed_str.encode('utf8').decode('utf8')))
                if web_translation_separators.strip() == '' and not ctx.flags.splitonly:
                    print("Error translating='%s'" % (item))
                ctx.docx.to_text_by_phrase_separator_table[i] = phrase_separator_removed_str
                phrase_separator_removed_str = p_remove_separator.sub(' ', phrase_separator_removed_str)
                phrase_separator_removed_str.strip()
                ctx.docx.to_text_by_phrase_separator_removed_table[i] = phrase_separator_removed_str
        except Exception:
            var = traceback.format_exc()
            ctx.browser.numerrors_deepl = ctx.browser.numerrors_deepl + 1
            web_translation_separators = var
            print("ERROR:%s" % (var))

        item = ctx.docx.from_text_by_phrase_table[i]
        try:
            web_translation_no_separators = ''
            if item.strip() != '':
                #google_translation_res = translator.translate(item, src=ctx.language.src_lang, dest='fr')
                #time.sleep(5)
                #web_translation_no_separators = pydeepl.translate(item, to_language)
                phrase_separator_removed_str = p_remove_double_spaces.sub(' ', web_translation_no_separators)
                phrase_separator_removed_str = p_remove_parenthesis_spaces.sub('(', phrase_separator_removed_str)
                ctx.docx.to_text_by_phrase_table[i] = phrase_separator_removed_str
        except Exception:
            var = traceback.format_exc()
            numerrors_googletranslate = numerrors_googletranslate + 1
            web_translation_no_separators = var
        Identical_with_without_separators = 'DIFFERENT<BR>'
        if phrase_separator_removed_str == web_translation_no_separators:
            Identical_with_without_separators = 'SAME<BR>'


def document_split_phrases(ctx: RuntimeContext):
    # Split phrases into multiple lines to match source language number of lines

    oai_sub_splitter = None
    if ctx.flags.split_engine == "openai":
        if ctx.openai.translator is not None:
            oai_sub_splitter = OpenAISubtitleSplitter(filename=ctx.flags.word_file_to_translate, doc_id=ctx.openai.translator.get_doc_id())
        else:
            oai_sub_splitter = OpenAISubtitleSplitter()
    
    for i, line in enumerate(ctx.docx.from_text_table):
        if ctx.docx.to_text_by_phrase_separator_table[i] != '':
            #ctx.docx.docxfile_table_number_of_phrases = ctx.docx.docxfile_table_number_of_phrases + 1
            ctx.docx.docxfile_table_number_of_characters = ctx.docx.docxfile_table_number_of_characters + len(ctx.docx.from_text_by_phrase_separator_table[i])
            ctx.docx.phrase_number_of_words = len(ctx.docx.from_text_by_phrase_separator_table[i].strip().split(" "))
            #print("Phrase to split: %s" % (ctx.docx.from_text_by_phrase_separator_table[i]))
            #print("number of words: %d" % (ctx.docx.phrase_number_of_words))
            ctx.docx.docxfile_table_number_of_words = ctx.docx.docxfile_table_number_of_words + ctx.docx.phrase_number_of_words
            input_phrase_lines = ""
            #translation_result_using_separator[i] = []
            #translation_result_phrase_array[i] = []
            
            try:
                current_line = ctx.docx.to_text_by_phrase_separator_table[i]
                # Using () as separator for splitting phrases, not used anymore
                #lines = current_line.split(line_separator_nospace_str)
                str_translation_len = len(current_line)

                try:
                    if str_translation_len <= 0:
                        str_phrase_stats = ""
                    else:
                        str_nb_lines = ctx.docx.from_text_nb_lines_in_phrase[i]
                        if str_nb_lines > 0:
                            str_line_average = str_translation_len / str_nb_lines
                            str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                        else:
                            str_line_average = 0
                            str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                        #print("str_phrase_stats=%s" % (str_phrase_stats))
                except Exception:
                    var = traceback.format_exc()
                    print("  ERROR:%s<br>" % (var))
                
                
                # --- Step 1: Build input safely ---
                if str_nb_lines > 1:
                    input_phrase_lines = "\n".join(
                        ctx.docx.from_text_table[i + idx] for idx in range(str_nb_lines)
                    )
                else:
                    input_phrase_lines = current_line


                # --- Step 2: Split functions ---
                def split_with_openai():
                    def expected_line_count(text):
                        if not text:
                            return 0
                        return text.count("\n") + 1

                    def normalize(lines):
                        # If string → split
                        if isinstance(lines, str):
                            lines = lines.strip().split("\n")

                        # Ensure list of strings
                        if isinstance(lines, list):
                            return [str(l).strip() for l in lines if str(l).strip() != ""]

                        return []

                    def is_valid(lines, expected):
                        return isinstance(lines, list) and len(lines) == expected

                    expected = expected_line_count(input_phrase_lines)

                    oai_sub_splitter.set_model('gpt-5.4-mini')
                    # --- Try 2 times with default model ---
                    for attempt in range(2):
                        response, lines = oai_sub_splitter.split_phrase(
                            ctx.language.src_lang_name, ctx.language.dest_lang_name, input_phrase_lines, current_line
                        )

                        lines = normalize(lines)

                        if is_valid(lines, expected):
                            return lines

                        print(f"[Retry {attempt+1}] Invalid OpenAI split: got {len(lines)} expected {expected}")

                    # --- Switch model ---
                    print("[Switching model to gpt-5.5]")
                    oai_sub_splitter.set_model('gpt-5.5')

                    response, lines = oai_sub_splitter.split_phrase(
                        ctx.language.src_lang_name, ctx.language.dest_lang_name, input_phrase_lines, current_line
                    )

                    # --- Final handling ---
                    if isinstance(lines, list):
                        lines = [str(l).strip() for l in lines if str(l).strip() != ""]
                        if lines:
                            return lines
                        else:
                            return []

                    elif isinstance(lines, str):
                        lines = lines.strip().split("\n")
                        lines = [l.strip() for l in lines if l.strip() != ""]
                        if lines:
                            return [lines[0]]  # return first line only
                        else:
                            return []

                    # fallback safety
                    return []


                def split_with_algorithm():
                    local_avg = str_line_average

                    if local_avg > MAX_LINE_SIZE:
                        local_avg = math.ceil(local_avg)

                    tokens = tokenize_text_to_array(ctx, current_line, ctx.language.dest_lang)
                    lines = divide_array(tokens, ctx.language.dest_lang, local_avg + 4)

                    divide_max_try = MAX_LINE_SIZE
                    while (len(lines) > str_nb_lines) and (divide_max_try > 0):
                        local_avg += 1
                        lines = divide_array(tokens, ctx.language.dest_lang, local_avg + 4)
                        divide_max_try -= 1

                    divide_max_try = MAX_LINE_SIZE
                    local_avg = str_translation_len / str_nb_lines
                    local_avg = math.ceil(local_avg)

                    while (len(lines) <= str_nb_lines) and (len(lines) >= 1) and (divide_max_try > 0):
                        local_avg -= 1
                        attempt_lines = divide_array(tokens, ctx.language.dest_lang, local_avg + 4)

                        if len(attempt_lines) <= str_nb_lines:
                            lines = attempt_lines

                        divide_max_try -= 1

                    return lines


                # --- Only 1 line, no splitting ---
                if str_nb_lines == 1:
                    lines_divided = [current_line]
                if str_nb_lines > 1:
                    if ctx.flags.split_engine == "openai":
                        lines_divided = split_with_openai()
                        
                        number_lines = len(lines_divided)
                        if number_lines != str_nb_lines:
                            print('Fallback to algorythm line splitting')
                            lines_divided = split_with_algorithm()
                    else:
                        lines_divided = split_with_algorithm()

                    # --- Step 4: Normalize ---
                    if not lines_divided:
                        lines_divided = [current_line]

                    lines_divided = [str(line).strip() for line in lines_divided]

                    # --- Step 5: Enforce exact number of lines ---
                    number_lines = len(lines_divided)
                    print(number_lines)
                    print(lines_divided)

                    if number_lines != str_nb_lines:
                        print(f"Adjusting output: {number_lines} -> {str_nb_lines}")

                        if number_lines > str_nb_lines:
                            lines_divided = lines_divided[:str_nb_lines - 1] + [
                                " ".join(lines_divided[str_nb_lines - 1:])
                            ]
                        else:
                            lines_divided += [""] * (str_nb_lines - number_lines)

                        number_lines = len(lines_divided)


                # --- Step 6: Store results ---
                ctx.docx.translation_result_phrase_array[i] = lines_divided

                number_lines = len(lines_divided)
                for line_no in range(number_lines):
                    try:
                        ctx.docx.translation_result_using_separator[line_no + i] = lines_divided[line_no]
                        print(f"{line_no} -> {lines_divided[line_no]}")
                    except Exception:
                        ctx.docx.translation_result_using_separator[line_no + i] = "Error"


                # --- Step 7: Logging ---
                try:
                    print(
                        "Splitting phrase : %s (%d) = %d lines"
                        % (ctx.docx.to_text_by_phrase_separator_table[i], i, str_nb_lines)
                    )
                except Exception:
                    try:
                        print(
                            "%s (%d): %d "
                            % (ctx.docx.to_text_by_phrase_separator_table[i].encode("utf-8"), i, str_nb_lines)
                        )
                    except Exception:
                        print("(unable to print content to screen) (%d): %d" % (i, str_nb_lines))


                # --- Final check ---
                if number_lines != str_nb_lines:
                    print("Error in number of line %d, expected %d." % (number_lines, str_nb_lines))
                
            except Exception:
                var = traceback.format_exc()
                print("  ERROR:%s<br>" % (var))


# write_destination_language_in_docx_cell was extracted to
# ``docx_io/metadata.py`` in the 2026-05-16 cli.py shrink phase 2.
# Thin shim — reads the entry-script globals and delegates.
from .docx_io.metadata import (
    write_destination_language_in_docx_cell as _write_dest_lang_impl,
    set_docx_properties_comment_for_history as _set_docx_history_impl,
)  # noqa: E402


def write_destination_language_in_docx_cell(ctx: RuntimeContext):
    _write_dest_lang_impl(
        ctx.docx.docxdoc,
        splitonly=ctx.flags.splitonly,
        dest_lang_name=ctx.language.dest_lang_name,
        dest_lang=ctx.language.dest_lang,
    )


def print_console_docx_file_translated(ctx: RuntimeContext):
    print("\nTranslated text:\n")
    numrows = len(ctx.docx.table.rows)
    numcols = len(ctx.docx.table.columns)
    current_cell_row = 2
    for row_n in range(2, (numrows)):

        str_translation_len = len(ctx.docx.translation_result_using_separator[row_n])
        translation_phrase_lines_len = len(ctx.docx.translation_result_phrase_array[row_n])
        if translation_phrase_lines_len == 0 and current_cell_row < row_n:
            print("%d :" % row_n)
        #print("row_n = %d" %  row_n)

        # Non-split path: write the translation regardless of whether
        # document_split_phrases populated translation_result_phrase_array.
        # The original guard `if translation_phrase_lines_len >= 1` was a
        # silent failure mode \u2014 if the split helper ever skipped a row
        # (e.g. its own to_text guard), the translation would never reach
        # the cell. By keying on to_text_by_phrase_separator_table directly
        # we write whatever the translation engine returned.
        if not ctx.flags.split_translation:
            translation_cell_text = ctx.docx.to_text_by_phrase_separator_table[row_n]
            if translation_cell_text:
                prepare_and_clear_cell_for_writing(ctx, row_n, translation_cell_text)
                if ctx.language.dest_lang in right_to_left_languages_list.keys():
                    translation_cell_aligned_text = get_display(translation_cell_text)
                else:
                    translation_cell_aligned_text = translation_cell_text
                print("%d : %s" % (row_n, translation_cell_aligned_text))
            continue

        if translation_phrase_lines_len >= 1 :
            #print("%d : %s" % (row_n,' '.join(ctx.docx.translation_result_phrase_array[row_n])))

            if False:  # legacy non-split branch above already handled by `if not split_translation`
                pass
            else:
                #translation_cell_text = ctx.docx.translation_result_using_separator[row_n]
                #print("len array: %d" % (translation_phrase_lines_len))
                #print("translation_result_phrase_array[%d] : %s" % (row_n,'\n'.join(ctx.docx.translation_result_phrase_array[row_n])))

                translation_phrase_line_pos = 0
                translation_phrase_cell_pos = 0

                while translation_phrase_line_pos < translation_phrase_lines_len:
                    current_cell_row = row_n + translation_phrase_cell_pos
                    cell_lines_len = ctx.docx.from_text_nb_lines_in_cell[row_n + translation_phrase_cell_pos]
                    cell_line_pos = 0
                    current_cell = ctx.docx.table_cells[current_cell_row][2]
                    while cell_line_pos < cell_lines_len \
                        and translation_phrase_line_pos < translation_phrase_lines_len:

                        translation_phrase_line_str = ctx.docx.translation_result_phrase_array[row_n][translation_phrase_line_pos]
                        if ctx.language.dest_lang in right_to_left_languages_list.keys():
                            #translation_cell_aligned_text = reverse_string (translation_phrase_line_str)
                            #translation_cell_aligned_text = "\u202B" + translation_phrase_line_str + "\u202C"
                            translation_cell_aligned_text = get_display(translation_phrase_line_str)
                        else:
                            translation_cell_aligned_text = translation_phrase_line_str
                        if cell_lines_len > 1:
                            print("%d-%d : %s" % (current_cell_row, cell_line_pos + 1, translation_cell_aligned_text))
                        else:
                            print("%d : %s" % (current_cell_row, translation_cell_aligned_text))
                        if cell_line_pos == 0:
                            #print("cell_line_pos=%d" % cell_line_pos)
                            if ctx.flags.splitonly:
                                prepare_and_clear_cell_for_writing(ctx, current_cell_row, translation_phrase_line_str)
                            else:
                            #prepare_and_clear_cell_for_writing(current_cell_row, translation_phrase_line_str)
                                cell_set_1st_paragraph(ctx, current_cell_row, translation_phrase_line_str)
                            # Not needed
                            #current_cell.paragraphs[0].text = translation_phrase_line_str
                        else:
                            # Add empty paragraph between translation lines
                            cell_add_paragraph(ctx, current_cell_row, "")
                            # Add the translation line
                            cell_add_paragraph(ctx, current_cell_row, translation_phrase_line_str)
                        cell_line_pos = cell_line_pos + 1
                        translation_phrase_line_pos = translation_phrase_line_pos + 1
                        #input("press enter")
                    translation_phrase_cell_pos = translation_phrase_cell_pos + 1

        try:
            if str_translation_len <= 0:
                str_phrase_stats = ""
            else:
                str_translation_len = len(ctx.docx.translation_result_using_separator[row_n])
                str_nb_lines = ctx.docx.from_text_nb_lines_in_phrase[row_n]
                if str_nb_lines > 0:
                    str_line_average = str_translation_len / str_nb_lines
                    str_phrase_stats = "[%d/%d=%d] " % (str_translation_len, str_nb_lines, str_line_average)
                else:
                    str_line_average = 0
                    str_phrase_stats = ""

        except Exception:
            var = traceback.format_exc()
            print("  ERROR:%s<br>" % (var))


def set_docx_properties_comment_for_history(ctx: RuntimeContext):
    _set_docx_history_impl(
        ctx.docx.docxdoc, program_version=PROGRAM_VERSION, engine=ctx.engine.engine,
    )



# local_time_offset extracted to ``statistics.py`` in Sprint D (2026-05-16).
# run_statistics extracted to ``statistics.py`` in Sprint D-A.4 (2026-05-16).
# get_robot_usage_comment extracted to ``statistics.py`` in Sprint D-A.5 (2026-05-16).
from .statistics import (  # noqa: E402
    local_time_offset,
    run_statistics as _run_statistics_impl,
    get_robot_usage_comment as _get_robot_usage_comment_impl,
)


def run_statistics(ctx: RuntimeContext):
    """Stats reporter — extracted to statistics.py in Sprint D-A.4 (2026-05-16)."""
    return _run_statistics_impl(ctx)


def get_robot_usage_comment(ctx: RuntimeContext):
    """Available-updates check — extracted to statistics.py in Sprint D-A.5 (2026-05-16)."""
    return _get_robot_usage_comment_impl(ctx)


# Open the default app for the docx file
def open_app_docx_file(ctx: RuntimeContext):
    """Open the saved DOCX in the OS default handler.

    Threaded in Phase F1.6: reads ``ctx.flags.word_file_to_translate_save_as_path``
    in place of the historical module global.
    """
    out_path = ctx.flags.word_file_to_translate_save_as_path
    try:
        if platform.system() == 'Windows':
            # A2 (2026-05-12): `Popen(["start", "", path], shell=True)` runs
            # the cmd parser on `path`, so a docx filename containing `&`,
            # `|`, `^`, `(`, `)`, double quotes, or other cmd metacharacters
            # could execute attacker-chosen commands when the user clicks
            # "Open file". `os.startfile` is the documented Windows
            # shell-open API: no cmd parser, no shell, the filename is
            # passed verbatim to ShellExecuteEx.
            os.startfile(out_path)
        elif platform.system() == "Darwin":  # macOS
            subprocess.Popen(["open", out_path])
        elif platform.system() == "Linux":  # Linux
            subprocess.Popen(["xdg-open", out_path])
        else:
            print("Unsupported operating system.")
    except Exception as e:
        print("Error:", e)
        print("Warning, unable to open file %s." % (out_path))
# _engine_suffix was extracted to src/docx_io/save.py in the
# 2026-05-10 docx_io extraction pass. Re-imported under the legacy
# private name so any other call site keeps resolving. The
# implementation save_docx_file is wrapped by the same-named shim
# below this comment block.
from .docx_io.save import engine_suffix as _engine_suffix
from .docx_io.save import save_docx_file as _save_docx_file_impl


# save_docx_file was extracted to ``src/docx_io/save.py`` in the
# 2026-05-10 docx_io extraction pass. Thin shim — emits the
# PROGRESS:90 marker the launcher watches, then delegates to the
# implementation, passing the entry-script globals (docxdoc, silent,
# write_translation_log) through.
def save_docx_file(ctx: RuntimeContext):
    # PROGRESS:90 — about to write the docx to disk. The aligner branch
    # in print_console_docx_file_translated also emits 90 (and 100), so
    # this is a no-op there; for non-aligner engines (DeepL, Google web,
    # chatgpt) it fills the gap before the save-time final 100.
    print("PROGRESS:90", flush=True)
    _save_docx_file_impl(
        ctx,
        ctx.docx.docxdoc,
        silent=ctx.flags.silent,
        write_translation_log_fn=write_translation_log,
    )

import os
import re
import time
import shutil
import psutil
import platform
import sys

def _print_structured_failure(reason: str, message: str) -> None:
    """Emit a structured ``[FAIL] reason=...`` line for the launcher to parse.

    Added 2026-05-10 (B-001 fix). The launcher's stdout watcher will
    pick up this line, set ``jobs[id].status = "error"`` with the
    ``reason`` token, and (if B-002 alerting is wired up) trigger the
    failure-archive hook.
    """
    print(f"[FAIL] reason={reason} message={message}", flush=True)


def main() -> int:
    """Entry point. Builds the RuntimeContext from module-level state on
    first ``_get_ctx()`` call and threads it through every pipeline step.

    F1.6 wire-up: zero ``global`` statements remain in this function.
    Pipeline state lives on ``ctx``; module globals are only referenced
    where they remain authoritative (start_time, str_needs_update,
    version_checker_sleep_seconds_on_update, xtm, exitonsuccess, silent,
    PROGRAM_VERSION — these are not threaded because they are owned by
    module-level setup outside ``main()``).

    B-001 (2026-05-10): wraps the body in a ``try`` that catches
    :class:`exceptions.TranslationFailure` so empty-source docx and
    engine-empty runs produce a structured ``[FAIL] reason=...`` line
    plus a non-zero exit instead of "Translation ended, file saved".
    """
    from .exceptions import TranslationFailure
    from .translation_health import (
        assert_source_has_content,
        assert_translation_present,
    )

    # 2026-05-13: sweep old sidecars > 10 days from `Log json file/`.
    try:
        from .log_paths import cleanup_old_logs
        _removed = cleanup_old_logs(retention_days=10)
        if _removed:
            print(f"[INFO] Log retention swept {_removed} file(s) older than 10 days.")
    except Exception as _exc:
        print(f"[WARN] Log retention sweep failed: {_exc!r}")

    ctx = _get_ctx()
    translation_succeeded = False

    set_translation_function(ctx)
    initialize_translation_memory_xlsx(ctx)

    read_and_parse_docx_document(ctx)

    # B-001: fail fast on a docx with no translatable text. Without this
    # guard, the pipeline finishes "successfully" with an empty output
    # and the user thinks the run worked.
    assert_source_has_content(ctx)

    create_webdriver(ctx)

    if ctx.engine.engine == 'deepl':
        ctx.browser.logged_into_deepl = selenium_chrome_deepl_log_in(ctx)

    translation_succeeded = translate_docx(ctx)

    if ctx.browser.logged_into_deepl:
        selenium_chrome_deepl_log_off(ctx)

    # R15: DeepL phrasesblock → singlephrase fallback. F1.4 + F1.6:
    # the dispatcher refresh sees the flipped method through ctx, the
    # rebuilt driver is reassigned to ctx.browser.driver, and no module
    # global is read or written. The structural tests
    # `test_engine_method_flip_via_ctx`,
    # `test_driver_rebuild_via_ctx`, and
    # `test_dispatcher_refresh_drops_stale_driver_reference` cover this.
    if (translation_succeeded is False
            and ctx.engine.engine == 'deepl'
            and ctx.engine.method == 'phrasesblock'):
        ctx.engine.method = 'singlephrase'
        set_translation_function(ctx)
        try:
            ctx.browser.driver.close()
            ctx.browser.driver.quit()
        except Exception:
            pass
        create_webdriver(ctx)

    get_translation_and_replace_after(ctx)

    # B-001: fail with a structured reason if the engine returned
    # nothing usable. Catches the "all rows empty" + "single-row dump"
    # failure modes the user reported pre-pass.
    assert_translation_present(ctx)

    # Diagnostic snapshot — counts of populated phrases and translation
    # array shape. Helps diagnose 'output empty' regressions like the
    # phrase_array gating bug seen on 2026-05-09.
    _populated_to_text = sum(
        1 for _v in (ctx.docx.to_text_by_phrase_separator_table or []) if _v
    )
    _ta = ctx.docx.translation_array
    _ta_len = len(_ta) if _ta is not None else 0
    print(
        f"[DIAG] After get_translation_and_replace_after: "
        f"to_text rows populated = {_populated_to_text}, "
        f"translation_array lines = {_ta_len}"
        + ("" if _ta is not None else "  (None — engine returned no result)")
    )
    # Defensive: if a Selenium engine errored and returned None, replace
    # it with an empty list so downstream `for line in translation_array`
    # loops don't TypeError.
    if ctx.docx.translation_array is None:
        ctx.docx.translation_array = []

    minimize_browser(ctx)

    # 2026-05-17 (AJAR 3150 basic-split bug fix): in splitonly mode the
    # translation phase is skipped, so
    # ``ctx.docx.to_text_by_phrase_separator_table`` is never populated
    # by get_translation_and_replace_after. document_split_phrases then
    # iterates with every entry == '' and does nothing — leaving the
    # raw "first-row-of-phrase has all the FA, rest empty" shape
    # unchanged. Populate the array directly from the docx's FA column
    # so the distributor has something to split.
    if ctx.flags.splitonly:
        # In splitonly mode the translation phase is skipped, so
        # ``to_text_by_phrase_separator_table`` is never filled by
        # get_translation_and_replace_after. Read the FA column of the
        # source docx directly so document_split_phrases has phrase-level
        # text to distribute across the phrase's rows.
        try:
            _table = ctx.docx.docxdoc.tables[0]
            _last_row = ctx.docx.word_translation_table_length
            _populated = 0
            for _ri in range(_last_row):
                if _ri >= len(_table.rows):
                    break
                _cells = _table.rows[_ri].cells
                if len(_cells) < 3:
                    continue
                _fa = " ".join(
                    p.text.strip() for p in _cells[2].paragraphs if p.text.strip()
                )
                if _fa:
                    ctx.docx.to_text_by_phrase_separator_table[_ri] = _fa
                    _populated += 1
            print(f"[INFO] splitonly: populated {_populated} FA rows for distribution")
        except Exception as _exc:
            print(f"[WARN] splitonly populate-from-docx failed: {_exc}")

    document_split_phrases(ctx)

    write_destination_language_in_docx_cell(ctx)

    print_console_docx_file_translated(ctx)
    set_docx_properties_comment_for_history(ctx)

    _end_time = datetime.datetime.now()
    _elapsed_time = _end_time - start_time

    run_statistics(ctx)
    save_docx_file(ctx)

    if ctx.flags.viewdocx:
        print(f"Opening document : {ctx.flags.word_file_to_translate_save_as_path}")
        open_app_docx_file(ctx)
    _end_time = datetime.datetime.now()
    _elapsed_time = _end_time - start_time

    if ctx.flags.xlsxreplacefile is not None:
        ctx.docx.xtm.print_replaced_items_number_of_replacements('before')
        ctx.docx.xtm.print_replaced_items_number_of_replacements('after')
        ctx.docx.xtm.print_do_not_split_number_of_matches('keep_on_same_line')

    if ctx.browser.driver is not None:
        clean_up_previous_chrome_selenium_drivers(ctx.browser.driver.service.path)

    print("\nTranslation ended, file saved. Elasped time: %s (h:mm:ss.mmm)" % (_elapsed_time))
    print("\nSaved file name: %s" % (ctx.flags.word_file_to_translate_save_as_path))

    get_robot_usage_comment(ctx)

    # 2026-05-13 (feat/exe-packaging): chatgpt-api path never opened a
    # browser, so ctx.browser.driver is None. Skip the cleanup
    # gracefully instead of stamping a noisy traceback.
    if ctx.browser.driver is not None:
        try:
            print("\nClosing chrome browser...")
            ctx.browser.driver.close()
            ctx.browser.driver.quit()
        except Exception:
            var = traceback.format_exc()
            print(var)

    if ctx.language.dest_lang_name is None or ctx.language.dest_lang_name == "":
        if not ctx.flags.splitonly:
            print("\n*********************************************************************************")
            print("WARNING: Target language name for %s not found. Translation may have have failed." % (ctx.language.dest_lang))
            print("*********************************************************************************\n")

    cleanup_selenium_chrome_temp_folders()

    print("\nDeveloper: %s" % (E_mail_str))
    print("Program version: %s\n" % (PROGRAM_VERSION))
    if not ctx.flags.exitonsuccess and not ctx.flags.silent:
        input("Enter to close program")
    else:
        if str_needs_update == "1":
            print(f"Please download and install the program update (message will be shown for {version_checker_sleep_seconds_on_update} seconds).")
            time.sleep(version_checker_sleep_seconds_on_update)
        print("Program ended")
    
    # Suppress any error message from undetected_chromedriver's
    # __del__ — it writes to stderr AFTER main() returns (during
    # process teardown), so a normal context manager wouldn't cover
    # the window. The legacy approach opened ``os.devnull`` and
    # leaked the file descriptor across process exit, plus
    # reassigned ``sys.__stderr__`` (a frozen reference debuggers
    # rely on) — both anti-patterns. Use an in-memory discard
    # stream instead: no fd allocated, no immutable-reference
    # mutation, GC reclaims it cleanly on process exit. Chrome
    # itself is already silenced via ``--log-level=3`` (set in
    # chrome_options at module setup) — this only catches the
    # post-main destructor noise.
    sys.stderr = io.StringIO()
    return 0

if __name__ == '__main__':
    # B-001: catch structured pipeline failures (empty docx / engine
    # returned empty) and exit with a non-zero status so the launcher
    # flags the job as ``status=error`` instead of inheriting our
    # default exit-zero "success".
    try:
        from .exceptions import TranslationFailure as _TranslationFailure
    except Exception:
        _TranslationFailure = None
    try:
        main()
    except Exception as _exc:
        if _TranslationFailure is not None and isinstance(_exc, _TranslationFailure):
            _print_structured_failure(_exc.reason, str(_exc))
            try:
                ctx_for_cleanup = _get_ctx()
                if ctx_for_cleanup.browser.driver is not None:
                    ctx_for_cleanup.browser.driver.quit()
            except Exception:
                pass
            sys.exit(20)
        # Anything else: propagate (existing global atexit hooks still
        # run, traceback is preserved in stdout).
        raise
    # Redirect all stderr output to null (silences destructor error messages)
    sys.exit(0)
