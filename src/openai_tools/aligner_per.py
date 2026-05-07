"""
FA Subtitle Aligner — Score-then-LLM pipeline
==============================================
Reads a bilingual EN|FA subtitle DOCX, re-distributes Persian text
across TV rows using:
  Pass 1 — Mechanical split (no API, ~80% quality)
  Pass 2 — Quality scoring per group (no API)
  Pass 3 — LLM batch re-split for low-quality groups (gpt-5.4-mini)

Hard limits enforced in both passes:
  - Each FA chunk ≤ 48 characters
  - Triple (3 identical rows) is FORBIDDEN — max 2 identical (double only)
  - Token budget cap (default 40 000) — excess groups keep mechanical result

Mechanical pass improvements (2026-05):
  - B4 proportional distribution (EN word-count → FA char budget per row)
  - Content-type rules (DIALOGUE/SPIRITUAL/NEWS_ATTR/INGREDIENT skip doubling)
  - Multi-candidate split scoring with balance bonus
  - Tiered validation: FATAL blocks, WARN accepted with log
  - CONTINUATION_STARTERS: prevents premature group splits at clause boundaries
  - DANGEROUS_SPLITS: detects light-verb compound pairs beyond می‌/نمی‌ prefix
  - Weight pass: fixes heavy-last-line from Persian SOV (verb-final) structure
  - Modulo-cycle distribution: evenly spaces doubles (not longest-chunk-first)
  - Existing-segmentation preservation: skips re-split when FA is already balanced
  - 5-part alignment score: discourse(0.30) + numbers(0.20) + punctuation(0.10)
                             + length_ratio(0.20) + base(0.10)
  - BREAK_RATIO_MEDIAN=0.45: empirical split target from broadcast data
  - Citation stripping: removes trailing (source) attribution from FA cells

Usage from pipeline:
    from openai_tools.aligner_per import FASubtitleAligner
    aligner = FASubtitleAligner(model="gpt-5.4-mini")
    stats = aligner.align(input_docx, output_docx)
"""

import os
import re
import json
import time
import unicodedata
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as _qn

try:
    from openai import OpenAI as _OpenAI
    _HAS_OPENAI = True
except ImportError:
    _HAS_OPENAI = False

try:
    from ._retry import call_with_retry as _call_with_retry
    from ._retry import prompt_hash as _prompt_hash
except ImportError:
    def _call_with_retry(fn, *, label="openai"):  # fallback when openai missing
        return fn()
    def _prompt_hash(text: str) -> str:
        return "00000000"

try:
    import tiktoken as _tiktoken
    _HAS_TIKTOKEN = True
except ImportError:
    _HAS_TIKTOKEN = False


# ── constants ─────────────────────────────────────────────────────────────────

MAX_CHARS  = 48   # hard limit per FA chunk
MIN_TARGET = 24   # minimum target length when splitting

BREAK_RATIO_MEDIAN = 0.45
# Empirical: Persian splits better at 45% of text length rather than 50%.
# Persian SOV structure means verbs (heavier content) cluster at the end,
# so splitting slightly before the midpoint produces more balanced halves.

# Per-content-type break ratios. NEWS leans front-loaded (event/subject up
# front), so the split target moves to roughly the midpoint. SAGE / spiritual
# narration keeps the legacy verb-final 0.45. Dialogue tends to balance speaker
# turn vs. content, so 0.50 is a safe middle.
_BREAK_RATIO_BY_TYPE: dict = {
    # constants are referenced before the _CT_* names below;
    # we use the literal strings _CT_* will be assigned to.
    'narration':  0.45,
    'spiritual':  0.45,
    'news_attr':  0.55,
    'dialogue':   0.50,
    'ingredient': 0.50,
}

# Density-based minimum double recommendation (technique C from para_bridge).
# Based on empirical broadcast data: FA char length → minimum doubles needed.
# Used as a hint in _distribute() and passed to LLM payload.
_DOUBLE_DENSITY: list[tuple[int, int]] = [
    (20, 0),   # FA ≤ 20 chars → 0 doubles
    (35, 1),   # FA ≤ 35 chars → 1 double
    (48, 2),   # FA ≤ 48 chars → 2 doubles
    # > 48 chars → split first (already handled by MAX_CHARS enforcement)
]


def _min_doubles_for(fa_len: int) -> int:
    """Return minimum recommended doubles for a given FA text length."""
    for threshold, min_d in _DOUBLE_DENSITY:
        if fa_len <= threshold:
            return min_d
    return 2


# B4 proportional weight table — empirical EN word-count → FA display weight.
# Derived from 3,036 real broadcast rows (para_bridge research).
# 0-1 words = meta/bridge row (weight 0); 7+ words → weight 1.01
_B4_WEIGHTS: dict = {0: 0.0, 1: 0.0, 2: 2.86, 3: 2.05, 4: 1.57, 5: 1.34, 6: 1.13}
_B4_DEFAULT = 1.01  # 7+ words


def _fa_budget_per_row(en_rows: list, total_fa_len: int) -> list:
    """Compute target FA character count per EN row using B4 weight table.

    Rows with 0-1 EN words (likely meta/bridge) receive proportionally
    less budget. Result sums to approximately total_fa_len.
    """
    weights = [_B4_WEIGHTS.get(len(en.split()), _B4_DEFAULT) for en in en_rows]
    total_w = sum(weights)
    if total_w == 0:
        each = max(1, total_fa_len // len(en_rows)) if en_rows else 0
        return [each] * len(en_rows)
    return [max(1, round(total_fa_len * w / total_w)) for w in weights]


# ── content type constants ────────────────────────────────────────────────────

_CT_NARRATION  = 'narration'    # default — normal doubling rules
_CT_DIALOGUE   = 'dialogue'     # speaker-tagged rows — conservative doubling
_CT_SPIRITUAL  = 'spiritual'    # SM:/Master: — conservative doubling
_CT_NEWS_ATTR  = 'news_attr'    # (Reuters)/(AP) attribution lines — skip doubling
_CT_INGREDIENT = 'ingredient'   # recipe ingredient list row — skip doubling
# NOTE: general cooking narrative (methods, descriptions) uses _CT_NARRATION
# and gets normal doubling. Only isolated ingredient-list rows skip doubling.

_RE_DIALOGUE   = re.compile(r'^(?:Q\s*\([mf]\)|[A-Z]{1,4})\s*:', re.I)
_RE_SPIRITUAL  = re.compile(r'^(?:SM|Master)\s*:', re.I)
_RE_NEWS_ATTR  = re.compile(
    r'\((?:Reuters|AP|BBC|CNN|AFP|Al\s*Jazeera|Xinhua|Fox\s*News|Euronews|euronews)\)',
    re.I
)

# Ingredient list row: starts with measurement unit or fraction/number.
# Examples: "1 cup of vegan cream cheese", "¼ cup of sugar", "2 tablespoons salt"
_RE_INGREDIENT = re.compile(
    r'(?:'
    r'^[¼½¾⅓⅔⅛⅜⅝⅞]\s'
    r'|^\d+\s*(?:cup|tbsp|tsp|tablespoon|teaspoon'
    r'|oz|ounce|gram|g\b|ml|kg|lb|pound'
    r'|inch|cm|pinch|clove|slice|piece|handful)s?'
    r'|^\d+\s+\w'
    r')',
    re.I
)

# Citation pattern: trailing "(source)" at end of FA cell.
# Strips e.g. "این خبر مهم بود. (یورونیوز)" → "این خبر مهم بود."
_RE_CITATION = re.compile(r'\s*\([^()]{2,40}\)\s*$')


def _classify_content(en: str, fa: str = '') -> str:
    """Return content type for a row based on EN (and optionally FA) text."""
    if _RE_SPIRITUAL.match(en):
        return _CT_SPIRITUAL
    if _RE_DIALOGUE.match(en):
        return _CT_DIALOGUE
    if _RE_NEWS_ATTR.search(en) or _RE_NEWS_ATTR.search(fa):
        return _CT_NEWS_ATTR
    if _RE_INGREDIENT.match(en.strip()):
        return _CT_INGREDIENT
    return _CT_NARRATION


# ── Persian linguistic constants ──────────────────────────────────────────────

# Protected bigrams — never split between these two words (technique A).
PROTECTED_BIGRAMS: frozenset = frozenset({
    'از جمله', 'از طریق', 'همراه با', 'با وجود',
    'در برابر', 'در نتیجه', 'در پی', 'به دلیل',
    'به خاطر', 'از آنجا که', 'در عین حال', 'بر اساس',
    'به جای', 'در راستای', 'به عنوان', 'در مقابل',
    'بر خلاف', 'به جهت', 'به منظور', 'از سوی',
    'به واسطه', 'در قبال',
})

ZWNJ = '‌'  # Persian half-space

SENT_END = frozenset('.!?؟')

COMPOUND_PREFIXES = {
    'می' + ZWNJ, 'نمی' + ZWNJ,
    'می ',       'نمی ',
}

# Continuation starters — Persian conjunctions that open dependent clauses.
# When the NEXT row's FA begins with one of these, the current group should
# remain open to prevent premature sentence-boundary splitting.
CONTINUATION_STARTERS: frozenset = frozenset({
    'که', 'و', 'تا', 'با', 'اما', 'ولی', 'یا', 'چون',
    'زیرا', 'هرچند', 'بلکه', 'پس', 'سپس', 'بنابراین', 'چراکه',
})

# Dangerous split patterns — light-verb compound constructions that span
# a word boundary beyond the می‌/نمی‌ prefix check.
# Splitting between these components produces grammatically broken subtitles.
DANGEROUS_SPLITS: list = [
    re.compile(r'(انجام)\s+(می[‌ ]|نمی[‌ ]|داد\b|دادند|دهد\b|دهند)'),
    re.compile(r'(استفاده)\s+(می[‌ ]|نمی[‌ ]|کرد\b|کردند|کردیم|کند\b)'),
    re.compile(r'(صحبت|بیان|اعلام|تصمیم)\s+(کرد\b|کردند|کردیم|گرفت\b|گرفتند|می[‌ ]|نمی[‌ ])'),
    re.compile(r'(نگاه)\s+(کرد\b|کردند|می[‌ ]|نمی[‌ ])'),
    re.compile(r'(دسترسی|آمادگی|توانایی)\s+(داشت\b|داشتند|ندارد\b|ندارند|دارد\b|دارند)'),
    re.compile(r'(موفق|قادر)\s+(شد\b|شدند|می[‌ ]|نمی[‌ ])'),
    re.compile(r'(به\s+وجود|راه\s*اندازی)\s+(آمد\b|آمدند|کرد\b|کردند|می[‌ ])'),
]

# Meaningful short FA words that should never be treated as bridges.
# E.g. "نه" (no) is 2 chars but is real dialogue content.
_MEANINGFUL_SHORT_FA: frozenset = frozenset({'نه', 'من', 'تو', 'آب', 'ما', 'او'})

# NOTE: dangling-preposition rule intentionally removed per project requirements.

BRIDGE_PATTERNS = [
    re.compile(r'^file:///'),
    re.compile(r'^\([^()]{2,40}\)\s*$'),     # entire-row citation: "(euronews)"
    re.compile(r'^https?://'),
    re.compile(r'^\d+:\d+'),                 # timecodes: "0:34 ~ 0:44"
    re.compile(r'^\d+[\.\)]?\s*$'),
    re.compile(r'^[A-Z][A-Z\s]{2,}:\s*$'),   # ALL-CAPS label: "ONSCREEN TEXT:"
    re.compile(r'^[A-Z][a-z]+\s*:$'),
    re.compile(r'^Narrator'),
    re.compile(r'^HOST:', re.I),
    re.compile(r'^SHOW:', re.I),
    re.compile(r'^CAPTION:', re.I),
    re.compile(r'^ONSCREEN', re.I),
    re.compile(r'^VO\s*[&:]', re.I),
    re.compile(r'^BMD\s'),
    re.compile(r'^Fix[12]'),
]

# Minimal built-in discourse cues (project can supply alignment_cues.json)
_BUILTIN_CUES = {
    'cause':      {'weight': 0.8, 'cues': {
        'because': ['چون', 'زیرا', 'چراکه'],
        'since':   ['چون', 'از آنجا که'],
    }},
    'result':     {'weight': 0.9, 'cues': {
        'therefore': ['بنابراین', 'ازاین‌رو'],
        'so':        ['پس', 'بنابراین'],
    }},
    'contrast':   {'weight': 0.7, 'cues': {
        'but':     ['اما', 'ولی'],
        'however': ['اما', 'با این حال'],
    }},
    'concession': {'weight': 0.8, 'cues': {
        'although':    ['اگرچه', 'هرچند'],
        'even though': ['با اینکه', 'هرچند'],
    }},
    'condition':  {'weight': 0.8, 'cues': {
        'if': ['اگر', 'در صورتی که'],
    }},
    'time':       {'weight': 0.6, 'cues': {
        'when':  ['وقتی', 'هنگامی که'],
        'while': ['در حالی که'],
    }},
    # Categories below added 2026-05-09 from `docs/aligner-research.md`
    # — common in SMTV news / educational content.
    'addition':   {'weight': 0.6, 'cues': {
        'also':           ['همچنین', 'نیز'],
        'moreover':       ['علاوه بر این', 'افزون بر این'],
        'in addition':    ['علاوه بر این', 'افزون بر این'],
        'furthermore':    ['افزون بر این', 'همچنین'],
    }},
    'sequence':   {'weight': 0.5, 'cues': {
        'then':    ['سپس', 'پس از آن'],
        'next':    ['سپس', 'پس از آن'],
        'finally': ['در نهایت', 'سرانجام'],
        'first':   ['نخست', 'ابتدا'],
    }},
    'example':    {'weight': 0.5, 'cues': {
        'for instance': ['برای مثال', 'به‌عنوان مثال'],
        'such as':      ['مانند', 'از جمله'],
        'for example':  ['برای مثال', 'به‌عنوان مثال'],
        'e.g.':         ['مانند', 'از جمله'],
    }},
    'emphasis':   {'weight': 0.5, 'cues': {
        'in fact':  ['در واقع'],
        'indeed':   ['به‌راستی', 'در واقع'],
        'actually': ['در واقع', 'به‌راستی'],
    }},
}


# ── standalone helpers ────────────────────────────────────────────────────────

def _cell_text(cell) -> str:
    return ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _cell_has_shading(cell) -> bool:
    """Return True if the cell has a non-white / non-auto background fill."""
    tcPr = cell._tc.find(_qn('w:tcPr'))
    if tcPr is None:
        return False
    shd = tcPr.find(_qn('w:shd'))
    if shd is None:
        return False
    fill = shd.get(_qn('w:fill'))
    if not fill:
        return False
    return fill.lower() not in ('auto', 'ffffff', '')


def _strip_citation(text: str) -> str:
    """Remove trailing source attribution like '(euronews)' or '(رویترز)' from FA."""
    return _RE_CITATION.sub('', text).strip()


def _is_bridge(en: str, fa: str = '') -> bool:
    """Return True if this row should be skipped (bridge row).

    EN-side: pattern-matched (timecodes, labels, URLs, etc.).
    FA override: if FA is a known meaningful short word (نه، من، تو…),
    never treat the row as a bridge regardless of EN appearance.
    """
    en = en.strip()
    if not en:
        return True
    # Meaningful short FA overrides bridge detection
    if fa and fa.strip() in _MEANINGFUL_SHORT_FA:
        return False
    return any(p.search(en) for p in BRIDGE_PATTERNS)


def _ends_sentence(text: str) -> bool:
    t = text.strip()
    return bool(t) and t[-1] in SENT_END and not t.endswith('...')


def _normalize_fa(text: str) -> str:
    """Normalize Persian text for comparison (not for output)."""
    text = text.replace('|', ' ')                  # normalize pipe character
    text = text.replace('ي', 'ی').replace('ك', 'ک')
    text = re.sub(r'[ًٌٍَُِّْ]', '', text)
    return re.sub(r'\s+', ' ', text).strip()


def _normalize_text(text: str) -> str:
    """Collapse all whitespace — used for text-preservation checks."""
    return re.sub(r'\s+', '', text)


def _display_len(text: str) -> int:
    """Return the visual character count of `text`.

    ZWNJ (U+200C, "نیم‌فاصله") is invisible — Word renders it with zero
    width — so it must NOT count toward MAX_CHARS. Counting it would
    waste display budget on glyphs the viewer never sees, producing
    chunks that look short but pass the limit.
    """
    return len(text.replace(ZWNJ, ''))


def _estimate_tokens(text: str) -> int:
    if _HAS_TIKTOKEN:
        try:
            enc = _tiktoken.get_encoding('cl100k_base')
            return len(enc.encode(text))
        except Exception:
            pass
    return max(1, len(text) // 3)


# ── main class ────────────────────────────────────────────────────────────────

class FASubtitleAligner:
    """Score-then-LLM subtitle aligner for Persian/English DOCX files."""

    _SYSTEM_PROMPT = (
        "You are a bilingual Persian/English subtitle aligner for TV broadcast.\n"
        "You receive a JSON array of sentence groups. Each group has:\n"
        "  id          — batch index (integer)\n"
        "  n_rows      — number of TV display rows to fill (integer)\n"
        "  en_rows     — English text for each row (array of strings)\n"
        "  full_fa     — complete Persian sentence to distribute (string)\n"
        "  min_doubles — minimum recommended doubles based on FA length (integer)\n\n"
        "Task: split full_fa into exactly n_rows Persian chunks.\n\n"
        "HARD RULES — never violate:\n"
        "  1. Each chunk ≤ 48 characters (count carefully).\n"
        "  2. Chunks joined with single space must equal full_fa exactly.\n"
        "  3. Consecutive identical chunks = double (allowed, max 2 identical).\n"
        "  4. TRIPLE IS FORBIDDEN — never use 3 or more identical chunks.\n"
        "  5. Never split compound verbs:\n"
        "     می‌کند، نمی‌دهد، انجام می‌دهد، استفاده کرد، صحبت کرد،\n"
        "     نگاه کرد، تصمیم گرفت، بیان کرد، اعلام کرد،\n"
        "     دسترسی دارد، موفق شد، به وجود آمد.\n"
        "  6. Never start a chunk with standalone 'را'.\n"
        "  7. Never split protected bigrams across chunks:\n"
        "     از جمله، از طریق، همراه با، با وجود، در برابر، در نتیجه، در پی،\n"
        "     به دلیل، به خاطر، از آنجا که، در عین حال، بر اساس، به جای،\n"
        "     در راستای، به عنوان، در مقابل، بر خلاف، به جهت، به منظور،\n"
        "     از سوی، به واسطه، در قبال\n\n"
        "GUIDANCE:\n"
        "  • Use at least min_doubles double rows when FA length warrants it.\n"
        "  • Split at clause boundaries: که، اما، ولی، زیرا، چون، بنابراین\n"
        "  • Align FA chunk to EN row: numbers and named entities in same row\n"
        "  • Discourse markers (because→چون, however→اما) aligned where possible\n"
        "  • Chunks roughly balanced in length\n\n"
        "Respond ONLY with JSON (no explanation):\n"
        '{"results": [{"id": <id>, "fa_rows": ["chunk1", "chunk2", ...]}, ...]}\n'
        "fa_rows must have exactly n_rows strings. Use identical strings for doubles."
    )

    def __init__(
        self,
        model:         str = 'gpt-5.4-mini',
        llm_threshold: int = 90,
        token_budget:  int = 40_000,
    ):
        self.model         = model
        self.llm_threshold = llm_threshold
        self.token_budget  = token_budget
        self.tokens_used   = 0
        self.last_stats    = {}
        self.client        = None
        self._cues         = _BUILTIN_CUES

        if _HAS_OPENAI:
            api_key = os.environ.get('OPENAI_API_KEY')
            if api_key:
                self.client = _OpenAI(api_key=api_key)

        self._load_cues()

    def _load_cues(self):
        """Load alignment_cues.json if available alongside this file."""
        cues_path = os.path.join(
            os.path.dirname(__file__), '..', '..', 'data', 'alignment_cues.json'
        )
        if os.path.exists(cues_path):
            try:
                with open(cues_path, encoding='utf-8') as f:
                    data = json.load(f)
                self._cues = data.get('en_to_fa', _BUILTIN_CUES)
            except Exception:
                pass

    # ── DOCX I/O ──────────────────────────────────────────────────────────────

    def _read_rows(self, docx_path: str) -> list:
        doc = Document(docx_path)
        if not doc.tables:
            raise ValueError(f"No table in {docx_path}")
        rows = []
        for ri, row in enumerate(doc.tables[0].rows):
            cells = row.cells
            if len(cells) < 3:
                continue
            rows.append({
                'ri':     ri,
                'en':     _cell_text(cells[1]),
                'fa':     _cell_text(cells[2]),
                'shaded': _cell_has_shading(cells[1]),
            })
        return rows

    def _parse_groups(self, rows: list) -> list:
        """
        Group consecutive non-bridge rows into sentence groups.

        Key improvements over naive grouping:
        - CONTINUATION_STARTERS: if the next row's FA begins with a Persian
          conjunction (که، چون، اما…), keep the current group open instead of
          flushing at the sentence-end punctuation.
        - orig_fa_rows: stored per group to enable preservation check.
        - citation stripping: trailing (source) attributions removed from FA.
        - Safety cap: groups with ≥10 rows are always flushed to avoid runaway
          groups from cascading continuation checks.
        """
        groups  = []
        current = []
        orig_fa = []   # original FA text per row (before dedup/combining)

        def flush():
            if not current:
                return
            unique_fa = []
            seen: set = set()
            for r in current:
                fa = _strip_citation(r['fa'].strip())
                if fa and fa not in seen:
                    unique_fa.append(fa)
                    seen.add(fa)
            full_fa = ' '.join(unique_fa)
            groups.append({
                'row_indices':  [r['ri'] for r in current],
                'en_rows':      [r['en'] for r in current],
                'orig_fa_rows': list(orig_fa),
                'full_fa':      full_fa,
                'n_rows':       len(current),
            })
            current.clear()
            orig_fa.clear()

        for idx, rd in enumerate(rows):
            if _is_bridge(rd['en'], rd.get('fa', '')) or rd.get('shaded', False):
                flush()
                continue

            current.append(rd)
            orig_fa.append(_strip_citation(rd['fa'].strip()))

            if _ends_sentence(rd['en']):
                # Safety cap: very long groups always flush.
                if len(current) >= 10:
                    flush()
                    continue

                # Continuation lookahead: check the next non-bridge row's FA.
                # If its first word is a continuation starter, keep group open.
                continuation = False
                for look in range(idx + 1, min(idx + 4, len(rows))):
                    nr = rows[look]
                    if _is_bridge(nr['en'], nr.get('fa', '')) or nr.get('shaded', False):
                        continue  # skip bridge rows in lookahead
                    next_fa = nr['fa'].strip()
                    if next_fa:
                        first_word = next_fa.split()[0].replace(ZWNJ, '')
                        if first_word in CONTINUATION_STARTERS:
                            continuation = True
                    break   # only examine the first qualifying next row

                if not continuation:
                    flush()

        flush()
        return groups

    @staticmethod
    def _ensure_rtl_paragraph(p):
        """Add <w:bidi/> to paragraph properties so Word renders RTL.

        Uses python-docx's built-in `get_or_add_pPr` so element ordering
        respects the OOXML schema rather than relying on manual insert(0).
        """
        pPr = p._p.get_or_add_pPr()
        if pPr.find(_qn('w:bidi')) is None:
            pPr.append(OxmlElement('w:bidi'))

    @staticmethod
    def _ensure_rtl_run(run):
        """Add <w:rtl/> to run properties so glyph order is right-to-left."""
        rPr = run._r.get_or_add_rPr()
        if rPr.find(_qn('w:rtl')) is None:
            rPr.append(OxmlElement('w:rtl'))

    def _set_fa_cell(self, table, ri: int, text: str):
        row = table.rows[ri]
        if len(row.cells) < 3:
            return
        fa_cell = row.cells[2]
        for p in fa_cell.paragraphs:
            for run in p.runs:
                run.text = ''
        while len(fa_cell.paragraphs) > 1:
            elem = fa_cell.paragraphs[-1]._element
            elem.getparent().remove(elem)
        p = fa_cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            run = p.runs[0]
        else:
            run = p.add_run(text)
        # Ensure Persian text renders RTL even if the source DOCX cell
        # template lacked w:bidi/w:rtl markers (caused mirrored display).
        self._ensure_rtl_paragraph(p)
        self._ensure_rtl_run(run)

    def _write_docx(self, input_path: str, output_path: str,
                    groups: list, all_chunks: list):
        doc   = Document(input_path)
        table = doc.tables[0]
        for group, chunks in zip(groups, all_chunks):
            for ri, chunk in zip(group['row_indices'], chunks):
                self._set_fa_cell(table, ri, chunk)
        doc.save(output_path)

    # ── split position helpers ────────────────────────────────────────────────

    @staticmethod
    def _bigram_bad_positions(text: str) -> frozenset:
        """Character positions inside a protected bigram (splitting here breaks it)."""
        bad: set = set()
        for bigram in PROTECTED_BIGRAMS:
            start = 0
            while True:
                idx = text.find(bigram, start)
                if idx == -1:
                    break
                word1 = bigram.split(' ', 1)[0]
                bad.add(idx + len(word1) + 1)
                start = idx + 1
        return frozenset(bad)

    @staticmethod
    def _dangerous_split_positions(text: str) -> frozenset:
        """
        Character positions at the START of the second component of a
        dangerous compound verb pair (e.g. 'می‌کند' in 'انجام می‌کند').
        Splitting here breaks a light-verb construction.
        """
        bad: set = set()
        for pattern in DANGEROUS_SPLITS:
            for m in pattern.finditer(text):
                bad.add(m.start(2))
        return frozenset(bad)

    def _find_split_points(self, text: str) -> list:
        """Return list of (position, quality) sorted by quality descending."""
        candidates = []
        words = text.split()
        pos = 0
        boundaries = []
        for w in words:
            idx = text.find(w, pos)
            boundaries.append((idx, idx + len(w), w))
            pos = idx + len(w)

        bad_bigrams   = self._bigram_bad_positions(text)
        bad_dangerous = self._dangerous_split_positions(text)

        for i, (start, end, word) in enumerate(boundaries):
            if i == 0:
                continue
            quality = 50
            prev = boundaries[i - 1][2]

            if prev and prev[-1] in SENT_END and not prev.endswith('...'):
                quality = 100
            elif prev and prev[-1] in ',;،؛':
                quality = 85

            # Discourse marker alignment bonus
            w_norm = _normalize_fa(word)
            for cat in self._cues.values():
                for fa_list in cat.get('cues', {}).values():
                    for fa_eq in fa_list:
                        if w_norm.startswith(_normalize_fa(fa_eq)):
                            quality = max(quality, 80)

            # Balance bonus: positions near text midpoint
            text_mid      = len(text) / 2
            proximity     = 1.0 - abs(start - text_mid) / max(text_mid, 1)
            balance_bonus = int(proximity * 18)  # up to +18
            if quality < 80:
                quality = min(79, quality + balance_bonus)

            # Compound verb prefix: never split before می‌ / نمی‌
            if any(word.startswith(pf) for pf in COMPOUND_PREFIXES):
                quality = 5

            # «را» stays with its noun phrase
            if word.strip() == 'را':
                quality = 5

            # Protected bigram: penalize
            if start in bad_bigrams:
                quality = 2

            # Dangerous split (light-verb compound): strongly penalize
            if start in bad_dangerous:
                quality = max(quality, 5)   # same level as compound prefix
                quality = min(quality, 5)

            candidates.append((start, quality))

        candidates.sort(key=lambda x: -x[1])
        return candidates

    # ── recursive split ───────────────────────────────────────────────────────

    def _recursive_split(self, text: str, n_parts: int, target: int) -> list:
        text = text.strip()
        if n_parts <= 1 or not text:
            return [text] if text else ['']

        candidates = self._find_split_points(text)
        best = None
        best_score = -1

        for pos, quality in candidates:
            left  = text[:pos].rstrip()
            right = text[pos:].lstrip()
            if not left or not right or _display_len(left) > MAX_CHARS:
                continue
            ideal   = len(text) / n_parts
            balance = 1.0 - abs(len(left) - ideal) / max(len(text), 1)
            score   = quality * 0.6 + balance * 100 * 0.4
            if score > best_score:
                best_score = score
                best = (left, right)

        if best is None:
            sp = text.rfind(' ', 0, min(target, MAX_CHARS))
            if sp <= 0:
                sp = text.find(' ', min(target, MAX_CHARS))
            if sp <= 0:
                sp = MAX_CHARS
            best = (text[:sp].rstrip(), text[sp:].lstrip())

        left, right = best
        rest = self._recursive_split(right, n_parts - 1, target)
        return [left] + rest

    def _emergency_split(self, text: str, n_parts: int) -> list:
        chunks = []
        rem = text.strip()
        for i in range(n_parts):
            if i == n_parts - 1:
                chunks.append(rem)
                break
            if _display_len(rem) <= MAX_CHARS:
                chunks.append(rem)
                rem = ''
                continue
            sp = rem.rfind(' ', 0, MAX_CHARS)
            if sp <= 0:
                sp = MAX_CHARS
            chunks.append(rem[:sp].rstrip())
            rem = rem[sp:].lstrip()
        while len(chunks) < n_parts:
            chunks.append('')
        return chunks

    def _split_distinct(self, text: str, n_distinct: int,
                        content_type: str | None = None) -> list:
        """Split text into n_distinct chunks each ≤ MAX_CHARS.

        For 2-part splits the break ratio is content-type aware:
        SAGE/narration keeps 0.45 (Persian verb-final bias);
        NEWS shifts to 0.55 (front-loaded event/subject);
        DIALOGUE / INGREDIENT use 0.50 (neutral balance).
        See `_BREAK_RATIO_BY_TYPE`. When `content_type` is None the legacy
        `BREAK_RATIO_MEDIAN` is used so existing call sites keep behaviour.
        """
        text = text.strip()
        if not text:
            return [''] * n_distinct
        if n_distinct <= 1:
            return [text]

        # 2-part split — pick break ratio per content type.
        if n_distinct == 2:
            ratio = _BREAK_RATIO_BY_TYPE.get(content_type, BREAK_RATIO_MEDIAN)
            target = max(MIN_TARGET, int(len(text) * ratio))
        else:
            target = max(MIN_TARGET, len(text) // n_distinct)

        chunks = self._recursive_split(text, n_distinct, target)

        for c in chunks:
            if _display_len(c) > MAX_CHARS:
                return self._emergency_split(text, n_distinct)
        return chunks

    # ── distribute ────────────────────────────────────────────────────────────

    def _distribute(self, chunks: list, n_rows: int, fa_len: int = 0) -> list:
        """
        Map M distinct chunks → exactly N row slots.

        Strategy: modulo-cycle doubles for even spatial distribution.
        This avoids clustering all doubles at the beginning of a group
        (which was the flaw in the previous longest-chunk-first approach).
        TRIPLES ARE FORBIDDEN — max 2 identical rows.

        fa_len: total FA char count, used for density-based double hint.
        """
        m = len(chunks)
        if m >= n_rows:
            return list(chunks[:n_rows])

        extra = n_rows - m
        min_d = _min_doubles_for(fa_len) if fa_len else 0
        # Effective doubles to apply: honour density hint but stay within budget
        n_doubles = max(min_d, extra)
        n_doubles = min(n_doubles, m, extra)

        # Modulo cycling: evenly space double positions across all chunks.
        # e.g. m=5, n_doubles=2 → step=2.5 → positions int(1.25)=1, int(3.75)=3
        double_set: set = set()
        if n_doubles > 0 and m > 0:
            if n_doubles >= m:
                double_set = set(range(m))
            else:
                step = m / n_doubles
                for k in range(n_doubles):
                    pos = int(k * step + step / 2)
                    pos = min(pos, m - 1)
                    double_set.add(pos)
                # Fill any collisions from remaining positions
                remaining = n_doubles - len(double_set)
                if remaining > 0:
                    for i in range(m):
                        if remaining <= 0:
                            break
                        if i not in double_set:
                            double_set.add(i)
                            remaining -= 1

        rows = []
        for i, ch in enumerate(chunks):
            rows.append(ch)
            if i in double_set:
                rows.append(ch)

        while len(rows) < n_rows:
            rows.append(rows[-1] if rows else '')
        return rows[:n_rows]

    def _enforce_no_triple(self, rows: list) -> list:
        """Remove ALL triples — hard ban. Max 2 identical consecutive rows.

        The previous implementation used a global occurrence counter which
        failed to detect triples: when counts[ch] > 2 it appended result[-1]
        which was still the same value ch, leaving the triple intact.

        This version tracks the *logical* consecutive run of each non-empty
        value by scanning backwards over both real and suppressed ('') slots
        that were emitted while the same value was active.  When a 3rd
        identical consecutive chunk would be appended, it is replaced with ''
        (empty row) so the row count stays constant.
        Empty FA rows are safe for broadcast — the subtitle slot shows blank.

        Bug fixed: the previous scan stopped at '' entries, so a quadruple
        ['A','A','A','A'] became ['A','A','','A'] instead of ['A','A','',''].
        The corrected scan skips over '' slots to count the true logical run.
        """
        result: list[str] = []
        # Parallel list: for each result slot, store the *logical* value it
        # represents (the original ch before any suppression to '').
        logical: list[str] = []

        for ch in rows:
            if ch.strip():
                # Count consecutive logical slots at the tail that equal ch
                run = 0
                for lv in reversed(logical):
                    if lv == ch:
                        run += 1
                    else:
                        break
                if run >= 2:
                    result.append('')   # suppress 3rd+ consecutive duplicate
                    logical.append(ch)  # record logical value for future scan
                else:
                    result.append(ch)
                    logical.append(ch)
            else:
                # Empty string — never suppress; reset logical run
                result.append(ch)
                logical.append(ch)
        return result

    # ── weight pass ───────────────────────────────────────────────────────────

    def _weight_pass(self, rows: list, en_rows: list) -> list:
        """
        Weight pass: fix the 'heavy last line' problem from Persian SOV structure.

        Persian verb-final order means the last subtitle row often carries heavy
        content (the main verb cluster). If the corresponding EN row is short
        (≤4 words, little display time), viewers can't read the FA text.

        This pass forces the NEXT row to display the same chunk (doubling) when:
          - EN row has ≤ 4 words  (short display time)
          - FA chunk has ≥ 28 chars  (heavy content)
          - The next row is not already a double of this chunk
          - Making it a double would not create a triple

        Changes are collected first (non-cascading), then applied.
        Caller must call _enforce_no_triple() afterward (done in _mechanical_align).
        """
        if len(rows) < 2 or len(rows) != len(en_rows):
            return rows

        result  = list(rows)
        changes: dict = {}

        for i in range(len(result) - 1):
            en_words = len(en_rows[i].split()) if i < len(en_rows) else 99
            fa_chars = len(result[i].strip())
            if en_words <= 4 and fa_chars >= 28:
                if result[i + 1] != result[i]:
                    # Would making [i+1] a double create a triple at [i+2]?
                    if i + 2 >= len(result) or result[i + 2] != result[i]:
                        changes[i + 1] = result[i]

        for idx, val in changes.items():
            result[idx] = val
        return result

    # ── best split near target ────────────────────────────────────────────────

    def _best_split_near(self, text: str, target: int) -> int:
        """Find best word-boundary split position near `target` chars from start.

        Respects compound verbs, dangerous splits, protected bigrams, and
        sentence punctuation. Returns a character index into `text`.
        """
        if len(text) <= target:
            return len(text)

        lo = max(1, int(target * 0.65))
        hi = min(MAX_CHARS, int(target * 1.35), len(text) - 1)

        bad_bigrams   = self._bigram_bad_positions(text)
        bad_dangerous = self._dangerous_split_positions(text)

        words  = text.split()
        pos    = 0
        best_p = None
        best_q = -1

        for w in words:
            idx = text.find(w, pos)
            pos = idx + len(w)
            if idx == 0:
                continue

            if not (lo <= idx <= hi):
                continue

            proximity = 1.0 - abs(idx - target) / max(target, 1)
            quality   = proximity * 60

            left = text[:idx].rstrip()

            if left and left[-1] in '.!?؟' and not left.endswith('...'):
                quality = 100
            elif left and left[-1] in '،؛,;':
                quality = min(100, quality + 25)

            if any(w.startswith(pf) for pf in COMPOUND_PREFIXES):
                continue

            if w.strip() == 'را':
                continue

            if idx in bad_bigrams:
                quality = max(0, quality - 55)

            if idx in bad_dangerous:
                quality = max(0, quality - 55)  # same penalty as bigram

            if quality > best_q:
                best_q = quality
                best_p = idx

        if best_p is None:
            sp = text.rfind(' ', 0, min(target + 10, MAX_CHARS))
            best_p = sp if sp > 0 else min(target, MAX_CHARS, len(text))

        return best_p

    def _split_by_budget(self, text: str, budgets: list) -> list:
        """Split text into len(budgets) chunks guided by per-row char budgets.

        Each budget[i] is the target character count from B4 weighting.
        Falls back to equal splitting when budget guidance is unhelpful.
        """
        n = len(budgets)
        if n <= 1 or not text:
            return [text[:MAX_CHARS]] if text else ['']

        chunks    = []
        remaining = text.strip()

        for i, budget in enumerate(budgets[:-1]):
            if not remaining:
                chunks.append('')
                continue
            remaining_budget = sum(budgets[i:])
            if remaining_budget > 0:
                scaled = max(1, round(len(remaining) * budget / remaining_budget))
            else:
                scaled = len(remaining) // max(1, n - i)

            split_pos = self._best_split_near(remaining, min(scaled, MAX_CHARS))
            left      = remaining[:split_pos].rstrip()
            remaining = remaining[split_pos:].lstrip()

            if _display_len(left) > MAX_CHARS:
                sp = left.rfind(' ', 0, MAX_CHARS)
                if sp > 0:
                    remaining = left[sp:].lstrip() + (' ' if remaining else '') + remaining
                    left = left[:sp].rstrip()
                else:
                    left = left[:MAX_CHARS]
            chunks.append(left)

        chunks.append(remaining[:MAX_CHARS])
        return [c[:MAX_CHARS] for c in chunks]

    # ── preservation check ────────────────────────────────────────────────────

    def _should_preserve_existing_segmentation(self, orig_fa_rows: list) -> bool:
        """
        Return True if the existing FA segmentation is already balanced and
        should be kept as-is (skip mechanical re-split).

        Criteria (from legacy v9 module, empirical broadcast data):
          - Mean chunk length between 18 and 42 characters
          - Short-chunk ratio (chunks < 10 chars) below 34 %
          - No chunk exceeds MAX_CHARS (48 chars)

        Preserving good translator output avoids introducing re-split noise.
        """
        nonempty = [r.strip() for r in orig_fa_rows if r.strip()]
        if len(nonempty) < 2:
            return False
        if any(_display_len(r) > MAX_CHARS for r in nonempty):
            return False
        mean_len    = sum(len(r) for r in nonempty) / len(nonempty)
        short_ratio = sum(1 for r in nonempty if len(r) < 10) / len(nonempty)
        return 18.0 <= mean_len <= 42.0 and short_ratio < 0.34

    # ── mechanical align ──────────────────────────────────────────────────────

    def _mechanical_align(self, group: dict) -> list:
        """Mechanical pass for one group. Returns list of len == n_rows."""
        full_fa = group['full_fa'].strip()
        n_rows  = group['n_rows']
        en_rows = group['en_rows']

        if not full_fa:
            return [''] * n_rows

        # ── Preservation check ─────────────────────────────────────────────
        # If the existing segmentation is already balanced, use it unchanged.
        orig_fa_rows = group.get('orig_fa_rows', [])
        if orig_fa_rows and len(orig_fa_rows) == n_rows:
            if self._should_preserve_existing_segmentation(orig_fa_rows):
                return list(orig_fa_rows)

        # ── Compute split dimensions ───────────────────────────────────────
        min_by_len  = max(1, -(-len(full_fa) // MAX_CHARS))   # ceil(len/MAX)
        min_by_rows = -(-n_rows // 2)                         # ceil(n_rows/2)
        min_dist    = max(min_by_len, min_by_rows)

        # Content type for this group (use first non-empty EN row)
        first_en     = next((r for r in en_rows if r.strip()), '')
        content_type = _classify_content(first_en, full_fa)

        # Skip doubling for attribution and ingredient rows
        no_double_types = {_CT_NEWS_ATTR, _CT_INGREDIENT}

        if content_type in no_double_types:
            # Single allocation: one distinct chunk per row, no doubles
            chunks = self._split_distinct(full_fa, n_rows, content_type=content_type)
            chunks = self._try_marker_align(full_fa, en_rows, chunks)
            rows   = list(chunks[:n_rows])
        else:
            # ── B4-guided split ────────────────────────────────────────────
            budgets = _fa_budget_per_row(en_rows, len(full_fa))

            if min_dist < n_rows:
                # Aggregate budgets for the distinct-chunk stage
                split_budgets = budgets[:min_dist] if min_dist <= len(budgets) else budgets
            else:
                split_budgets = budgets

            chunks = self._split_by_budget(full_fa, split_budgets)

            # Safety: if budget split violated MAX_CHARS, fall back
            if any(_display_len(c) > MAX_CHARS for c in chunks):
                chunks = self._split_distinct(full_fa, min_dist, content_type=content_type)

            # Discourse-marker alignment to improve EN↔FA correspondence
            chunks = self._try_marker_align(full_fa, en_rows, chunks)

            # Distribute into n_rows with modulo-cycled doubles
            rows = self._distribute(chunks, n_rows, fa_len=len(full_fa))

            # Weight pass: give extra display time to heavy SOV last-lines
            rows = self._weight_pass(rows, en_rows)

        # Hard ban: remove any triples that slipped through
        rows = self._enforce_no_triple(rows)

        # Length guard (pad/trim to exact n_rows)
        if len(rows) < n_rows:
            rows.extend([''] * (n_rows - len(rows)))
        return rows[:n_rows]

    def _try_marker_align(self, full_fa: str, en_rows: list, chunks: list) -> list:
        """Try to shift split points to align discourse markers EN↔FA.

        Improved search window: ±(max(20, len/6)) around the boundary,
        rather than the fixed ±15/+25 of the previous version.
        """
        if len(chunks) < 2 or not en_rows:
            return chunks

        improved = list(chunks)
        for i, en in enumerate(en_rows):
            if i == 0 or i >= len(improved):
                continue
            en_lower = en.strip().lower()
            for cat in self._cues.values():
                for en_marker, fa_equivs in cat.get('cues', {}).items():
                    if not en_lower.startswith(en_marker.lower()):
                        continue
                    # Check if chunk[i] already starts with a FA equivalent
                    chunk_norm = _normalize_fa(improved[i])
                    if any(chunk_norm.startswith(_normalize_fa(eq)) for eq in fa_equivs):
                        break  # already aligned

                    # Search for FA equivalent near the current boundary
                    preceding    = ' '.join(improved[:i])
                    search_near  = len(preceding)
                    search_range = max(20, len(full_fa) // 6)
                    fa_norm      = _normalize_fa(full_fa)

                    for fa_eq in fa_equivs:
                        eq_norm = _normalize_fa(fa_eq)
                        lo  = max(0, search_near - search_range)
                        hi  = search_near + search_range + 10
                        pos = fa_norm.find(eq_norm, lo)
                        if pos < 0 or pos > hi:
                            continue
                        orig_pos = self._norm_to_orig(full_fa, pos)
                        if orig_pos is None:
                            continue
                        left  = full_fa[:orig_pos].rstrip()
                        right = full_fa[orig_pos:].lstrip()
                        if not left or not right or _display_len(left) > MAX_CHARS:
                            continue
                        if not _normalize_fa(right).startswith(eq_norm):
                            continue
                        left_ch  = self._split_distinct(left, i)
                        right_ch = self._split_distinct(right, len(improved) - i)
                        if (all(_display_len(c) <= MAX_CHARS for c in left_ch) and
                                all(_display_len(c) <= MAX_CHARS for c in right_ch)):
                            improved = left_ch + right_ch
                        break
        return improved

    def _norm_to_orig(self, orig: str, norm_pos: int) -> int:
        """Map position in normalized text back to original string."""
        n = 0
        for i, ch in enumerate(orig):
            if unicodedata.category(ch) in ('Mn', 'Me'):
                continue
            if n >= norm_pos:
                return i
            if ch in ' \t\n':
                if i + 1 < len(orig) and orig[i + 1] in ' \t\n':
                    continue
            n += 1
        return len(orig)

    # ── quality scoring ───────────────────────────────────────────────────────

    def _quality_score(self, rows: list, group: dict) -> int:
        """
        Score a row assignment 0–100.
        Lower score → more benefit from LLM re-split.

        Double-ratio penalty is content-type aware:
          NARRATION         → warn above 55 %
          DIALOGUE/SPIRITUAL→ warn above 30 %
          NEWS_ATTR/INGREDIENT→ warn above  5 %
        """
        score    = 100
        nonempty = [r for r in rows if r.strip()]

        if not nonempty:
            return 100

        # Hard violation: any chunk over limit (visible length, ZWNJ-aware)
        for r in rows:
            if _display_len(r) > MAX_CHARS:
                score -= 50
                break

        # Compound verb split: chunk starts with می‌/نمی‌
        for i in range(len(rows) - 1):
            b = rows[i + 1].strip()
            if b and any(b.startswith(pf) for pf in COMPOUND_PREFIXES):
                score -= 40
                break

        # «را» orphan at start of chunk
        for r in rows:
            if r.strip().startswith('را ') or r.strip() == 'را':
                score -= 20
                break

        # Triple → heavy penalty (triples are forbidden)
        cnt: dict = {}
        for r in rows:
            cnt[r] = cnt.get(r, 0) + 1
        for ch, c in cnt.items():
            if c >= 3 and ch.strip():
                score -= 60
                break

        # Length imbalance
        lengths = [len(r) for r in nonempty]
        if len(lengths) >= 2:
            rng = max(lengths) - min(lengths)
            if rng > 35:
                score -= 25
            elif rng > 22:
                score -= 12

        # 5-part alignment score
        align = self._alignment_score(group['en_rows'], rows)
        if align < 0.3:
            score -= 20
        elif align < 0.5:
            score -= 8

        # Content-type-aware double ratio check
        first_en     = next((r for r in group['en_rows'] if r.strip()), '')
        content_type = _classify_content(first_en)
        if content_type in (_CT_NEWS_ATTR, _CT_INGREDIENT):
            high_ratio = 0.05
        elif content_type in (_CT_DIALOGUE, _CT_SPIRITUAL):
            high_ratio = 0.30
        else:
            high_ratio = 0.55   # NARRATION — wider tolerance

        n_total   = len(rows)
        n_doubles = sum(max(0, c - 1) for c in cnt.values() if c >= 2)
        if n_total > 2 and n_doubles / n_total > high_ratio:
            score -= 15

        return max(0, score)

    def _alignment_score(self, en_rows: list, fa_chunks: list) -> float:
        """
        5-part calibrated EN↔FA alignment confidence (0–1).

        Parts:
          1. Discourse marker alignment  — weight 0.30
          2. Per-row number alignment    — weight 0.20
          3. Punctuation alignment       — weight 0.10
          4. Length ratio FA/EN          — weight 0.20
          5. Base                        — weight 0.10
          Total maximum                           0.90 + 0.10 neutral fills = 1.0
        """
        if not en_rows or not fa_chunks:
            return 0.5

        score = 0.0
        n = min(len(en_rows), len(fa_chunks))

        # ── Part 1: Discourse marker alignment (0.30) ─────────────────────
        matched = total = 0
        for i in range(n):
            en_lower = en_rows[i].strip().lower()
            for cat in self._cues.values():
                for en_m, fa_list in cat.get('cues', {}).items():
                    if en_lower.startswith(en_m.lower()):
                        total += 1
                        chunk_norm = _normalize_fa(fa_chunks[i])
                        if any(chunk_norm.startswith(_normalize_fa(eq)) for eq in fa_list):
                            matched += 1
        score += (0.30 * matched / total) if total else 0.15

        # ── Part 2: Per-row number alignment (0.20) ───────────────────────
        # Check: numbers in EN row i appear in FA chunk i (window ±1).
        num_hits = num_total = 0
        for i in range(n):
            en_nums = set(re.findall(r'\d+', en_rows[i]))
            if not en_nums:
                continue
            num_total += 1
            # Search in a ±1 window to handle slight row shifts
            window   = fa_chunks[max(0, i - 1) : min(n, i + 2)]
            fa_text  = ' '.join(window)
            fa_nums  = set()
            for raw in re.findall(r'[\d۰-۹]+', fa_text):
                conv = raw
                for fa_d, en_d in zip('۰۱۲۳۴۵۶۷۸۹', '0123456789'):
                    conv = conv.replace(fa_d, en_d)
                fa_nums.add(conv)
            if en_nums & fa_nums:
                num_hits += 1
        score += (0.20 * num_hits / num_total) if num_total else 0.10

        # ── Part 3: Punctuation alignment (0.10) ──────────────────────────
        # EN row ends with sentence punctuation → FA chunk should too.
        punc_hits = punc_total = 0
        for i in range(n):
            en_end = en_rows[i].strip()
            if en_end and en_end[-1] in SENT_END:
                punc_total += 1
                fa_chunk = fa_chunks[i].strip()
                if fa_chunk and fa_chunk[-1] in SENT_END:
                    punc_hits += 1
        score += (0.10 * punc_hits / punc_total) if punc_total else 0.05

        # ── Part 4: Length ratio FA/EN (0.20) ─────────────────────────────
        en_len = sum(len(r) for r in en_rows)
        fa_len = sum(len(c) for c in fa_chunks)
        if en_len and fa_len:
            ratio  = fa_len / en_len
            score += 0.20 if 0.5 <= ratio <= 2.5 else 0.05
        else:
            score += 0.10

        # ── Part 5: Base (0.10) ───────────────────────────────────────────
        score += 0.10

        return min(1.0, score)

    # ── LLM batch ─────────────────────────────────────────────────────────────

    def _llm_batch(self, candidates: list) -> dict:
        """
        Re-split low-quality groups via LLM.
        candidates: list of (group_idx, group, mech_rows)
        Returns: {group_idx: new_rows}
        """
        if not self.client or not candidates:
            return {}

        corrections = {}
        BATCH_SIZE  = 20

        for b_start in range(0, len(candidates), BATCH_SIZE):
            batch = candidates[b_start : b_start + BATCH_SIZE]

            payload = [
                {
                    'id':          local_i,
                    'n_rows':      grp['n_rows'],
                    'en_rows':     grp['en_rows'],
                    'full_fa':     grp['full_fa'],
                    'min_doubles': _min_doubles_for(len(grp['full_fa'].strip())),
                }
                for local_i, (_, grp, _) in enumerate(batch)
            ]

            user_msg   = json.dumps(payload, ensure_ascii=False)
            est_tokens = _estimate_tokens(self._SYSTEM_PROMPT + user_msg)

            if self.tokens_used + est_tokens > self.token_budget:
                remaining = len(candidates) - b_start
                print(
                    f"[WARN] Aligner: token budget {self.token_budget} reached — "
                    f"{remaining} group(s) will keep mechanical result."
                )
                break

            try:
                t0   = time.time()
                resp = _call_with_retry(
                    lambda: self.client.chat.completions.create(
                        model=self.model,
                        messages=[
                            {'role': 'system', 'content': self._SYSTEM_PROMPT},
                            {'role': 'user',   'content': user_msg},
                        ],
                        response_format={'type': 'json_object'},
                        temperature=0,
                        timeout=120,
                        extra_body={'prompt_cache_retention': '24h'},
                    ),
                    label=f"aligner.batch_{b_start // BATCH_SIZE + 1}",
                )
                elapsed           = time.time() - t0
                used              = resp.usage.total_tokens
                self.tokens_used += used
                print(
                    f"[INFO] Aligner LLM batch {b_start // BATCH_SIZE + 1}: "
                    f"{len(batch)} groups — {used} tokens — {elapsed:.1f}s"
                )

                parsed = json.loads(resp.choices[0].message.content)
                for item in parsed.get('results', []):
                    local_i = item.get('id')
                    if local_i is None or local_i >= len(batch):
                        continue
                    group_idx, grp, mech_rows = batch[local_i]
                    fa_rows   = item.get('fa_rows', [])
                    validated = self._validate(fa_rows, grp)
                    if validated:
                        corrections[group_idx] = validated
                    else:
                        print(
                            f"[WARN] Aligner: LLM output for group {group_idx} "
                            f"failed validation — keeping mechanical."
                        )

            except Exception as e:
                print(
                    f"[ERROR] Aligner LLM batch {b_start // BATCH_SIZE + 1} "
                    f"failed: {e} — keeping mechanical for this batch."
                )

        return corrections

    def _validate(self, fa_rows: list, group: dict) -> list | None:
        """Validate LLM output using tiered rules.

        Returns fa_rows if FATAL-free (WARNs accepted with log).
        Returns None on any FATAL violation.

        FATAL (reject): wrong count, text modified, over MAX_CHARS, triple.
        WARN  (accept): doubling ratio outside ideal range.
        """
        n = group['n_rows']

        if not isinstance(fa_rows, list) or len(fa_rows) != n:
            return None

        for ch in fa_rows:
            if not isinstance(ch, str) or _display_len(ch) > MAX_CHARS:
                return None

        # FATAL: triple ban
        cnt: dict = {}
        for ch in fa_rows:
            cnt[ch] = cnt.get(ch, 0) + 1
        for ch, c in cnt.items():
            if c >= 3 and ch.strip():
                return None

        # FATAL: text preservation
        joined = ' '.join(ch for ch in fa_rows if ch.strip())
        if _normalize_text(joined) != _normalize_text(group['full_fa']):
            return None

        # WARN: doubling ratio outside [0.10, 0.65]
        n_doubles = sum(max(0, c - 1) for c in cnt.values() if c >= 2)
        if n > 2:
            ratio = n_doubles / n
            if ratio > 0.65:
                print(f"[WARN] Aligner: high double ratio {ratio:.0%} — accepted")

        return fa_rows

    # ── main entry point ──────────────────────────────────────────────────────

    def align(self, input_docx: str, output_docx: str) -> dict:
        """
        Full pipeline: read → mechanical → score → LLM → write.
        Returns stats dict.
        """
        self.tokens_used = 0
        t0 = time.time()

        rows   = self._read_rows(input_docx)
        groups = self._parse_groups(rows)
        print(f"[INFO] Aligner: {len(groups)} sentence groups parsed")

        # Pass 1 — mechanical
        mech_all = [self._mechanical_align(g) for g in groups]

        # Pass 2 — score
        scores = [
            self._quality_score(rows_out, grp)
            for rows_out, grp in zip(mech_all, groups)
        ]
        needs_llm = [
            (i, groups[i], mech_all[i])
            for i, s in enumerate(scores)
            if s < self.llm_threshold
        ]
        print(
            f"[INFO] Aligner: {len(needs_llm)}/{len(groups)} groups "
            f"below threshold {self.llm_threshold} -> LLM"
        )

        # Pass 3 — LLM batch (with budget guard)
        corrections = self._llm_batch(needs_llm)

        # Merge
        final_chunks = list(mech_all)
        for idx, corrected in corrections.items():
            final_chunks[idx] = corrected

        # Global triple guard: _enforce_no_triple runs per-group, but adjacent
        # groups can share the same boundary text and create cross-group triples.
        # Insert a sentinel between groups so the run scan in _enforce_no_triple
        # treats each group boundary as a break — without it, two groups whose
        # last/first chunks coincide would be counted as a real consecutive run.
        # The sentinel is a U+0000 NUL string that no real FA chunk can equal.
        _SENTINEL = '\x00GROUP_BOUNDARY\x00'
        flat_all: list[str] = []
        for gi, fc in enumerate(final_chunks):
            if gi > 0:
                flat_all.append(_SENTINEL)
            flat_all.extend(fc)
        flat_clean = self._enforce_no_triple(flat_all)

        # Re-chunk back into per-group lists, dropping sentinels.
        pos = 0
        for gi, fc in enumerate(final_chunks):
            if gi > 0:
                # Skip the sentinel slot for this boundary.
                pos += 1
            n = len(fc)
            final_chunks[gi] = flat_clean[pos : pos + n]
            pos += n

        # Write
        self._write_docx(input_docx, output_docx, groups, final_chunks)

        elapsed = time.time() - t0

        # Stats
        all_rows  = [r for fc in final_chunks for r in fc]
        n_total   = len(all_rows)

        # Count doubles and triples using CONSECUTIVE runs (not global frequency).
        # Global frequency overcounts: "است" appearing 5× across the file would
        # be reported as 3 triples even when no three consecutive rows match.
        n_doubles = 0
        n_triples = 0
        i = 0
        while i < len(all_rows):
            ch = all_rows[i].strip()
            j = i + 1
            while j < len(all_rows) and all_rows[j].strip() == ch and ch:
                j += 1
            run = j - i
            if run >= 2:
                n_doubles += run - 1       # each pair beyond the first is a double
            if run >= 3:
                n_triples += run - 2       # each 3rd+ occurrence is a triple violation
            i = j

        n_over    = sum(1 for r in all_rows if _display_len(r) > MAX_CHARS)

        self.last_stats = {
            'groups':          len(groups),
            'llm_corrected':   len(corrections),
            'mechanical_only': len(groups) - len(corrections),
            'tokens_used':     self.tokens_used,
            'prompt_hash':     _prompt_hash(self._SYSTEM_PROMPT),
            'total_rows':      n_total,
            'doubles':         n_doubles,
            'triples':         n_triples,
            'over_limit':      n_over,
            'elapsed_seconds': round(elapsed, 1),
        }

        print(
            f"[INFO] Aligner done in {elapsed:.1f}s"
            f" | groups: {len(groups)} | LLM: {len(corrections)}"
            f" | tokens: {self.tokens_used}"
            f" | doubles: {n_doubles} | triples: {n_triples}"
            f" | over-{MAX_CHARS}: {n_over}"
        )
        return self.last_stats
