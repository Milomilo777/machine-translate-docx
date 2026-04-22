"""DOCX reading utilities."""
import re
from typing import List, Dict, Optional
from lxml import etree
from docx import Document
from ..core.config import (
    TABLE_INDEX, SOURCE_COL, TRANSLATION_COL,
    shading_color_ignore_text, eol_array, bol_array
)

class DocxReader:
    """Handles opening DOCX files and extracting cell data from tables."""

    def __init__(self):
        self.doc = None

    def load(self, filepath: str) -> Document:
        """Opens docx and validates that tables[TABLE_INDEX] has >= 3 columns."""
        self.doc = Document(filepath)
        if len(self.doc.tables) <= TABLE_INDEX:
            raise ValueError(f"Document has no table at index {TABLE_INDEX}")

        table = self.doc.tables[TABLE_INDEX]
        if len(table.columns) < 3:
            raise ValueError(f"Table at index {TABLE_INDEX} must have at least 3 columns.")

        return self.doc

    def extract_cells(self) -> List[Dict]:
        """Iterates rows from index 2 and returns a list of cell data dicts."""
        if not self.doc:
            raise RuntimeError("No document loaded. Call load() first.")

        table = self.doc.tables[TABLE_INDEX]
        cells_data = []

        for row_n, row in enumerate(table.rows):
            if row_n < 2:  # Skip header rows
                continue

            source_cell = row.cells[SOURCE_COL]
            target_cell = row.cells[TRANSLATION_COL]

            # Logic for is_already_translated
            is_already_translated = bool(target_cell.text.strip())

            # Use a NoiseFilter-like logic for other properties (simplified here,
            # delegating actual filtering to processing/noise_filter.py)
            cells_data.append({
                "row_n": row_n,
                "source_text": source_cell.text,
                "cell_obj": source_cell,
                "is_already_translated": is_already_translated,
                # These will be populated by NoiseFilter in the pipeline
                "clean_text": "",
                "is_gray": False,
                "is_red": False,
                "nb_lines_in_cell": len(source_cell.text.split('\n'))
            })

        return cells_data

    @staticmethod
    def is_end_of_line(line: str) -> bool:
        """Checks if a line ends with sentence-ending punctuation."""
        for eol in eol_array:
            if re.search(eol, line):
                return True
        return False

    @staticmethod
    def is_beginning_of_line(line: str) -> bool:
        """Checks if a line begins with a capitalized letter."""
        for bol in bol_array:
            if re.search(bol, line):
                return True
        return False

    @staticmethod
    def is_empty_line(line: str) -> bool:
        """Checks if a line contains only whitespace."""
        return not line.strip()

    def _get_paragraph_shading_color(self, xml_paragraph_str: str) -> Optional[str]:
        """Extracts shading fill color from paragraph XML."""
        try:
            paragraph_xml = etree.fromstring(xml_paragraph_str)
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            for e in paragraph_xml.findall('.//w:pPr/w:shd', namespaces):
                return e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        except:
            pass
        return None

    def _get_run_shading_color(self, xml_run_str: str) -> Optional[str]:
        """Extracts shading fill color from run XML."""
        try:
            run_xml = etree.fromstring(xml_run_str)
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            for e in run_xml.findall('.//w:rPr/w:shd', namespaces):
                return e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        except:
            pass
        return None
