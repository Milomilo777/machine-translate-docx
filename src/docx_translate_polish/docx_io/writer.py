"""DOCX writing utilities."""
import os
from typing import List
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from ..core.config import right_to_left_languages_list, TRANSLATION_COL, TABLE_INDEX

class DocxWriter:
    """Handles writing translations back to DOCX files."""

    def __init__(self, doc):
        self.doc = doc
        self.table = self.doc.tables[TABLE_INDEX]
        self._ensure_rtl_style()

    def _ensure_rtl_style(self):
        """Creates 'rtlstyle' if it doesn't exist in the document."""
        styles = self.doc.styles
        if 'rtlstyle' not in styles:
            style = styles.add_style('rtlstyle', WD_STYLE_TYPE.CHARACTER)
            style.font.rtl = True

    def _prepare_and_clear_cell(self, cell):
        """Clears all paragraphs in a cell except the first one."""
        for i in range(len(cell.paragraphs) - 1, 0, -1):
            p = cell.paragraphs[i]._element
            p.getparent().remove(p)
        cell.paragraphs[0].text = ""
        return cell.paragraphs[0]

    def write_translation(self, row_n: int, text: str, dest_lang: str, dest_font: str = ""):
        """Writes a single-line translation to the target cell with RTL/LTR support."""
        cell = self.table.rows[row_n].cells[TRANSLATION_COL]
        paragraph = self._prepare_and_clear_cell(cell)

        if dest_lang in right_to_left_languages_list:
            run = paragraph.add_run(text)
            run.style = 'rtlstyle'
            run.font.rtl = True
            if dest_font:
                run.font.name = dest_font
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.text = text
            try:
                paragraph.style = 'Normal'
            except KeyError:
                pass
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def write_multiline(self, row_n: int, lines: List[str], dest_lang: str, dest_font: str = ""):
        """Writes multiple lines of translation to the target cell."""
        cell = self.table.rows[row_n].cells[TRANSLATION_COL]
        self._prepare_and_clear_cell(cell)

        for i, line in enumerate(lines):
            if i == 0:
                paragraph = cell.paragraphs[0]
            else:
                paragraph = cell.add_paragraph()

            if dest_lang in right_to_left_languages_list:
                run = paragraph.add_run(line)
                run.style = 'rtlstyle'
                run.font.rtl = True
                if dest_font:
                    run.font.name = dest_font
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.text = line
                try:
                    paragraph.style = 'Normal'
                except KeyError:
                    pass
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def save(self, output_path: str):
        """Saves the document to the specified path."""
        self.doc.save(output_path)
